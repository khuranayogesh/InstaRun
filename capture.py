import sys
import json
import os
import re # <-- Make sure this is present
import re
import copy
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QPushButton,
    QLabel, QDockWidget, QTabWidget, QTabBar, QFrame, QMenuBar,
    QToolBar, QStatusBar, QTreeWidget, QTreeWidgetItem, QHeaderView,
    QSplitter, QTableWidget, QTableWidgetItem, QMessageBox, QStyle, QMenu,
    QFileDialog, QTextEdit, QSizePolicy, QDialog, QLineEdit, QFormLayout, QDialogButtonBox,
    QSpacerItem, QComboBox, QLineEdit, QListWidget, QListWidgetItem,
    QCheckBox, QRadioButton, QToolButton, QSlider, QStackedWidget, QInputDialog,QSpinBox
)
from PyQt6.QtCore import Qt, QSize, QByteArray, QPoint, QTimer, QPropertyAnimation, QEasingCurve, pyqtSignal
from PyQt6.QtGui import QPixmap, QIcon, QAction, QFont, QFontMetrics, QTextCursor, QIntValidator, QPalette, QColor, QTextTableFormat, QTextFrameFormat
import pyautogui
import pygetwindow as gw
import pyperclip
import time

# --- New Imports for Win32 API ---
import win32gui
import win32ui
import win32con
from ctypes import windll
from PIL import Image
import win32com.client
import pythoncom
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
# ---------------------------------

# --- NEW: Smart PCOMM Wait Functions ---
def get_screen_content(autECLPS, rows=24, cols=80):
    """
    Gets the current screen content as a string.
    
    Args:
        autECLPS: The PCOMM session object
        rows: Number of rows (default 24)
        cols: Number of columns (default 80)
    
    Returns:
        str: Screen content or None if error
    """
    try:
        screen_size = rows * cols
        screen_text = autECLPS.GetText(1, screen_size)
        return screen_text
    except Exception as e:
        print(f"Error reading screen: {e}")
        return None

def wait_for_screen_change(autECLPS, initial_screen, timeout=30, min_wait=0.3, check_interval=0.1):
    """
    Waits for the screen content to change (indicates PCOMM finished processing).
    
    Args:
        autECLPS: The PCOMM session object
        initial_screen: The screen content before sending the command
        timeout: Maximum seconds to wait (default 30)
        min_wait: Minimum seconds to wait before checking (default 0.3)
        check_interval: How often to check in seconds (default 0.1)
    
    Returns:
        tuple: (bool: success, float: elapsed_time)
    """
    from datetime import datetime
    import time
    
    start_time = time.time()
    
    # Wait minimum time first (let PCOMM start processing)
    time.sleep(min_wait)
    
    while time.time() - start_time < timeout:
        current_screen = get_screen_content(autECLPS)
        
        if current_screen is None:
            return False, time.time() - start_time
        
        # Check if screen has changed
        if current_screen != initial_screen:
            elapsed = time.time() - start_time
            return True, elapsed
        
        time.sleep(check_interval)
    
    # Timeout
    return False, time.time() - start_time

def wait_for_pcomm_ready_smart(autECLPS, action_description="", timeout=30):
    """
    Smart wait that captures screen before/after and waits for changes.
    Use this BEFORE sending keys that will change the screen.
    
    Args:
        autECLPS: The PCOMM session object
        action_description: Description for logging (e.g., "Enter Key")
        timeout: Maximum seconds to wait
    
    Returns:
        tuple: (before_screen, success, elapsed_time)
    """
    # Capture screen before action
    before_screen = get_screen_content(autECLPS)
    return before_screen

def complete_pcomm_wait(autECLPS, before_screen, action_description="", timeout=30):
    """
    Completes the wait after an action was performed.
    Use this AFTER sending keys.
    
    Args:
        autECLPS: The PCOMM session object
        before_screen: Screen content captured before the action
        action_description: Description for logging  # ✅ Must be named parameter
        timeout: Maximum seconds to wait
    
    Returns:
        tuple: (bool: success, float: elapsed_time)
    """
    if before_screen is None:
        print(f"⚠️ Could not capture screen before {action_description}, using fixed wait")
        time.sleep(2.0)
        return True, 2.0
    
    success, elapsed = wait_for_screen_change(autECLPS, before_screen, timeout)
    
    if success:
        print(f"✅ {action_description} completed in {elapsed:.2f}s")
    else:
        print(f"⚠️ {action_description} timeout after {elapsed:.2f}s")
    
    return success, elapsed

class AddLabelDialog(QDialog):
    """
    A dialog box to manually add a new label with all its properties.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add New Label")
        self.setFixedSize(300, 180)  # ✅ CHANGED: Increased height from 150 to 180
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)

        self.layout = QFormLayout(self)
        self.layout.setVerticalSpacing(12)  # ✅ NEW: Add vertical spacing between rows
        
        self.name_input = QLineEdit()
        self.row_input = QLineEdit()
        self.column_input = QLineEdit()
        self.length_input = QLineEdit()
        
        # Add validators to ensure only integers can be entered for numeric fields
        self.row_input.setValidator(QIntValidator(1, 999))
        self.column_input.setValidator(QIntValidator(1, 999))
        self.length_input.setValidator(QIntValidator(0, 999))

        self.layout.addRow("Field Name:", self.name_input)
        self.layout.addRow("Row:", self.row_input)
        self.layout.addRow("Column:", self.column_input)
        self.layout.addRow("Length:", self.length_input)

        self.button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        self.button_box.accepted.connect(self.accept)
        self.button_box.rejected.connect(self.reject)
        
        self.layout.addWidget(self.button_box)

    def get_data(self):
        """Returns the data entered by the user."""
        return {
            "name": self.name_input.text(),
            "row": int(self.row_input.text() or 0),
            "column": int(self.column_input.text() or 0),
            "length": int(self.length_input.text() or 0),
        }
        
class DocumentConfigDialog(QDialog):
    """
    Dialog to configure the layout of DOCX documents for screenshots.
    Users can add multiple text configurations with alignment, font type, and size.
    """
    def __init__(self, parent=None, existing_config=None, existing_highlight_color='Yellow'):
        super().__init__(parent)
        self.setWindowTitle("Configure Document Layout")
        self.setMinimumSize(600, 500)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        
        # ✅ CHANGED: Create a deep copy to avoid modifying the original config
        self.config_items = copy.deepcopy(existing_config) if existing_config else []
        
        self.setup_ui()
        self.populate_existing_config()
        
        # ✅ NEW: Set the highlight color after setup_ui creates the combo box
        self.highlight_color_combo.setCurrentText(existing_highlight_color)
        self.update_color_preview(existing_highlight_color)
    
    def setup_ui(self):
        main_layout = QVBoxLayout(self)
        
        # Title
        title_label = QLabel("Document Layout Configuration")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        main_layout.addWidget(title_label)
        
        # Description
        desc_label = QLabel("Configure text elements that will appear in the DOCX screenshot document.")
        desc_label.setWordWrap(True)
        main_layout.addWidget(desc_label)
        
        main_layout.addSpacing(10)
        
        # ✅ NEW: Highlight Color Configuration Section
        highlight_group = QWidget()
        highlight_layout = QHBoxLayout(highlight_group)
        highlight_layout.setContentsMargins(0, 0, 0, 0)
        
        highlight_label = QLabel("Screenshot Highlight Color:")
        highlight_label.setStyleSheet("font-weight: bold;")
        highlight_layout.addWidget(highlight_label)
        
        self.highlight_color_combo = QComboBox()
        self.highlight_color_combo.setFixedWidth(150)
        
        # Add color options with their Word color index values
        self.color_options = {
            'Yellow': 7,
            'Bright Green': 4,
            'Turquoise': 3,
            'Pink': 5,
            'Blue': 2,
            'Red': 6,
            'Dark Blue': 9,
            'Dark Cyan': 10,
            'Dark Green': 11,
            'Dark Magenta': 12,
            'Dark Red': 13,
            'Dark Yellow': 14,
            'Gray 25%': 16,
            'Gray 50%': 15
        }
        
        for color_name in self.color_options.keys():
            self.highlight_color_combo.addItem(color_name)
        
        # Set default to Yellow
        self.highlight_color_combo.setCurrentText('Yellow')
        
        highlight_layout.addWidget(self.highlight_color_combo)
        
        # Color preview box
        self.color_preview = QLabel("  Preview  ")
        self.color_preview.setStyleSheet("background-color: yellow; border: 1px solid black; padding: 5px;")
        self.color_preview.setFixedWidth(80)
        highlight_layout.addWidget(self.color_preview)
        
        highlight_layout.addStretch()
        
        # Connect color change to preview update
        self.highlight_color_combo.currentTextChanged.connect(self.update_color_preview)
        
        main_layout.addWidget(highlight_group)
        main_layout.addSpacing(10)
        
        # Separator line
        separator = QFrame()
        separator.setFrameShape(QFrame.Shape.HLine)
        separator.setFrameShadow(QFrame.Shadow.Sunken)
        main_layout.addWidget(separator)
        main_layout.addSpacing(10)
        
        # Configuration items list
        list_label = QLabel("Text Elements:")
        main_layout.addWidget(list_label)
        
        self.config_list = QListWidget()
        self.config_list.setSelectionMode(QListWidget.SelectionMode.SingleSelection)
        main_layout.addWidget(self.config_list)
        
        # Buttons for managing config items
        button_layout = QHBoxLayout()
        
        self.add_item_button = QPushButton("Add Text Element")
        self.add_item_button.clicked.connect(self.add_config_item)
        button_layout.addWidget(self.add_item_button)
        
        self.add_blank_line_button = QPushButton("Add Blank Line")
        self.add_blank_line_button.clicked.connect(self.add_blank_line)
        button_layout.addWidget(self.add_blank_line_button)
        
        self.edit_item_button = QPushButton("Edit Selected")
        self.edit_item_button.clicked.connect(self.edit_config_item)
        button_layout.addWidget(self.edit_item_button)
        
        self.delete_item_button = QPushButton("Delete Selected")
        self.delete_item_button.clicked.connect(self.delete_config_item)
        button_layout.addWidget(self.delete_item_button)
        
        self.move_up_button = QPushButton("↑ Move Up")
        self.move_up_button.clicked.connect(self.move_item_up)
        button_layout.addWidget(self.move_up_button)
        
        self.move_down_button = QPushButton("↓ Move Down")
        self.move_down_button.clicked.connect(self.move_item_down)
        button_layout.addWidget(self.move_down_button)
        
        button_layout.addStretch()
        main_layout.addLayout(button_layout)
        
        # Dialog buttons
        dialog_buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        dialog_buttons.accepted.connect(self.accept)
        dialog_buttons.rejected.connect(self.reject)
        main_layout.addWidget(dialog_buttons)
        
    def update_color_preview(self, color_name):
        """Updates the color preview box based on selected color."""
        color_map = {
            'Yellow': '#FFFF00',
            'Bright Green': '#00FF00',
            'Turquoise': '#00FFFF',
            'Pink': '#FF00FF',
            'Blue': '#0000FF',
            'Red': '#FF0000',
            'Dark Blue': '#00008B',
            'Dark Cyan': '#008B8B',
            'Dark Green': '#006400',
            'Dark Magenta': '#8B008B',
            'Dark Red': '#8B0000',
            'Dark Yellow': '#808000',
            'Gray 25%': '#C0C0C0',
            'Gray 50%': '#808080'
        }
        
        hex_color = color_map.get(color_name, '#FFFF00')
        self.color_preview.setStyleSheet(
            f"background-color: {hex_color}; border: 1px solid black; padding: 5px;"
        )
        
    def add_blank_line(self):
        """Adds a blank line element to the configuration."""
        blank_line_item = {
            'type': 'blank_line',
            'text': '',
            'font_name': 'Arial',
            'font_size': 12,
            'alignment': 'Left',
            'bold': False,
            'italic': False
        }
        self.config_items.append(blank_line_item)
        self.populate_existing_config()
        self.update_step_combo_options()
        
    def populate_existing_config(self):
        """Populate the list with existing configuration items."""
        self.config_list.clear()
        for item in self.config_items:
            display_text = self.format_config_item_display(item)
            self.config_list.addItem(display_text)
    
    def format_config_item_display(self, item):
        """Format a config item for display in the list."""
        # Check if it's a blank line
        if item.get('type') == 'blank_line':
            return "[Blank Line]"
        
        text = item.get('text', 'No text')
        font = item.get('font_name', 'Arial')
        size = item.get('font_size', 12)
        align = item.get('alignment', 'Left')
        
        # Truncate text if too long
        display_text = text[:30] + "..." if len(text) > 30 else text
        return f"{display_text} | {font} {size}pt | {align}"
    
    def add_config_item(self):
        """Open dialog to add a new text element."""
        dialog = TextElementDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            item_data = dialog.get_data()
            self.config_items.append(item_data)
            self.populate_existing_config()
    
    def edit_config_item(self):
        """Edit the selected configuration item."""
        current_row = self.config_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "No Selection", "Please select a text element to edit.")
            return
        
        item_data = self.config_items[current_row]
        
        # Check if it's a blank line - can't edit blank lines
        if item_data.get('type') == 'blank_line':
            QMessageBox.information(self, "Blank Line", "Blank lines cannot be edited. You can only move or delete them.")
            return
        
        dialog = TextElementDialog(self, existing_data=item_data)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            updated_data = dialog.get_data()
            self.config_items[current_row] = updated_data
            self.populate_existing_config()
            self.config_list.setCurrentRow(current_row)
    
    def delete_config_item(self):
        """Delete the selected configuration item."""
        current_row = self.config_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "No Selection", "Please select a text element to delete.")
            return
        
        reply = QMessageBox.question(
            self, "Confirm Delete",
            "Are you sure you want to delete this text element?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.config_items.pop(current_row)
            self.populate_existing_config()
    
    def move_item_up(self):
        """Move the selected item up in the list."""
        current_row = self.config_list.currentRow()
        if current_row <= 0:
            return
        
        self.config_items[current_row], self.config_items[current_row - 1] = \
            self.config_items[current_row - 1], self.config_items[current_row]
        self.populate_existing_config()
        self.config_list.setCurrentRow(current_row - 1)
    
    def move_item_down(self):
        """Move the selected item down in the list."""
        current_row = self.config_list.currentRow()
        if current_row < 0 or current_row >= len(self.config_items) - 1:
            return
        
        self.config_items[current_row], self.config_items[current_row + 1] = \
            self.config_items[current_row + 1], self.config_items[current_row]
        self.populate_existing_config()
        self.config_list.setCurrentRow(current_row + 1)
    
    def get_configuration(self):
        """Return the configuration items including highlight color."""
        return {
            'text_elements': self.config_items,
            'highlight_color': self.highlight_color_combo.currentText()
        }

class MaskingConfigDialog(QDialog):
    """
    Dialog to configure text masking rules.
    """
    def __init__(self, parent=None, existing_patterns=None):
        super().__init__(parent)
        self.setWindowTitle("Configure Masking Rules")
        self.setMinimumSize(800, 600)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        
        self.pattern_objects = copy.deepcopy(existing_patterns) if existing_patterns else []
        
        self.setup_ui()
        self.populate_patterns_list()
    
    def setup_ui(self):
        main_layout = QVBoxLayout(self)
        
        # Title
        title_label = QLabel("Text Masking Configuration")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        main_layout.addWidget(title_label)
        
        # Description
        desc_label = QLabel("Configure patterns to mask sensitive data in screenshots.")
        desc_label.setWordWrap(True)
        main_layout.addWidget(desc_label)
        
        main_layout.addSpacing(10)
        
        # Input Section
        input_group = QWidget()
        input_layout = QVBoxLayout(input_group)
        input_group.setStyleSheet("QWidget { background-color: #f3f4f6; border-radius: 4px; }")
        
        # Step 1: Sample Structure
        step1_label = QLabel("Step 1: Enter Sample Structure (e.g., 491100-5530000-15-6)")
        step1_label.setStyleSheet("font-weight: bold; padding: 8px;")
        input_layout.addWidget(step1_label)
        
        self.sample_input = QLineEdit()
        self.sample_input.setPlaceholderText("Enter sample pattern here...")
        input_layout.addWidget(self.sample_input)
        
        # Step 2: Masking Positions
        step2_label = QLabel("Step 2: Masking Positions (1-based, e.g., 8-10, 19)")
        step2_label.setStyleSheet("font-weight: bold; padding: 8px; padding-top: 16px;")
        input_layout.addWidget(step2_label)
        
        positions_layout = QHBoxLayout()
        self.position_inputs = []
        for i in range(4):
            pos_input = QLineEdit()
            pos_input.setPlaceholderText(f"Position {i+1}")
            pos_input.setMinimumHeight(30)  # ✅ NEW: Set minimum height
            positions_layout.addWidget(pos_input)
            self.position_inputs.append(pos_input)

        input_layout.addLayout(positions_layout)

        input_layout.addSpacing(5)  # ✅ CHANGED: Reduced spacing

        # Add Rule Button
        add_button = QPushButton("Add Rule")
        add_button.clicked.connect(self.add_pattern)
        input_layout.addWidget(add_button)
        
        main_layout.addWidget(input_group)
        main_layout.addSpacing(10)
        
        # Patterns List Section
        # Patterns List Section
        list_label = QLabel("Active Masking Rules")
        list_label.setStyleSheet("font-weight: bold;")
        main_layout.addWidget(list_label)

        self.patterns_list = QListWidget()
        self.patterns_list.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        self.patterns_list.setMinimumHeight(200)
        main_layout.addWidget(self.patterns_list)
        
        # List Buttons
        # List Buttons
        list_buttons_layout = QHBoxLayout()
        
        remove_button = QPushButton("Remove Selected")
        remove_button.clicked.connect(self.remove_selected_patterns)
        list_buttons_layout.addWidget(remove_button)
        
        list_buttons_layout.addStretch()
        
        main_layout.addLayout(list_buttons_layout)
        
        # Dialog Buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        main_layout.addWidget(button_box)
    
    def convert_to_regex(self, sample_string):
        """Converts a sample string into a structural regex."""
        if not sample_string:
            return ""
        
        regex = ""
        digit_count = 0
        
        for char in sample_string:
            if char.isdigit():
                digit_count += 1
            else:
                if digit_count > 0:
                    regex += f"\\d{{{digit_count}}}"
                    digit_count = 0
                regex += re.escape(char)
        
        if digit_count > 0:
            regex += f"\\d{{{digit_count}}}"
        
        return r"\b" + regex + r"\b"
    
    def parse_mask_positions(self, position_strings):
        """Parses position strings into 0-based indices."""
        indices = set()
        
        for pos_str in position_strings:
            if not pos_str.strip():
                continue
            
            parts = [p.strip() for p in pos_str.split(',')]
            
            for part in parts:
                if '-' in part:
                    try:
                        start, end = map(int, part.split('-'))
                        if start < 1 or end < 1 or start > end:
                            raise ValueError(f"Invalid range: {part}")
                        for i in range(start - 1, end):
                            indices.add(i)
                    except ValueError as e:
                        raise ValueError(f"Invalid range format: {part}")
                else:
                    try:
                        single_pos = int(part)
                        if single_pos < 1:
                            raise ValueError(f"Position must be >= 1: {part}")
                        indices.add(single_pos - 1)
                    except ValueError as e:
                        raise ValueError(f"Invalid position: {part}")
        
        return sorted(list(indices))
    
    def apply_mask_by_indices(self, text, indices):
        """Applies 'x' masking to text at given indices."""
        text_list = list(text)
        max_len = len(text_list)
        
        for index in indices:
            if 0 <= index < max_len:
                text_list[index] = 'x'
        
        return "".join(text_list)
    
    def add_pattern(self):
        """Adds a new masking pattern."""
        sample_string = self.sample_input.text().strip()
        position_strings = [inp.text().strip() for inp in self.position_inputs]
        
        if not sample_string:
            QMessageBox.warning(self, "Warning", "Please enter a sample structure.")
            return
        
        if not any(position_strings):
            QMessageBox.warning(self, "Warning", "Please enter at least one position to mask.")
            return
        
        try:
            mask_indices = self.parse_mask_positions(position_strings)
        except ValueError as e:
            QMessageBox.critical(self, "Invalid Input", str(e))
            return
        
        if mask_indices and max(mask_indices) >= len(sample_string):
            QMessageBox.critical(self, "Invalid Position", 
                               f"Position {max(mask_indices) + 1} exceeds sample length ({len(sample_string)}).")
            return
        
        generated_regex = self.convert_to_regex(sample_string)
        masked_sample = self.apply_mask_by_indices(sample_string, mask_indices)
        
        new_pattern = {
            "regex": generated_regex,
            "sample": sample_string,
            "mask_indices": mask_indices,
            "masked_sample": masked_sample
        }
        
        # Check for duplicates
        if any(p.get('regex') == generated_regex for p in self.pattern_objects):
            QMessageBox.information(self, "Duplicate", "This pattern already exists.")
            return
        
        self.pattern_objects.append(new_pattern)
        self.populate_patterns_list()
        
        # Clear inputs
        self.sample_input.clear()
        for inp in self.position_inputs:
            inp.clear()
    
    def remove_selected_patterns(self):
        """Removes selected patterns from the list."""
        selected_items = self.patterns_list.selectedItems()
        if not selected_items:
            return
        
        selected_indices = [self.patterns_list.row(item) for item in selected_items]
        
        for index in sorted(selected_indices, reverse=True):
            if 0 <= index < len(self.pattern_objects):
                del self.pattern_objects[index]
        
        self.populate_patterns_list()
    
    def populate_patterns_list(self):
        """Populates the patterns list widget."""
        self.patterns_list.clear()
        
        for pattern in self.pattern_objects:
            regex = pattern.get('regex', '')
            sample = pattern.get('sample', '')
            masked = pattern.get('masked_sample', '')
            
            display_text = f"{sample} → {masked}  |  Pattern: {regex}"
            self.patterns_list.addItem(display_text)
    
    def get_patterns(self):
        """Returns the configured patterns."""
        return self.pattern_objects

class SettingsDialog(QDialog):
    """
    Unified settings dialog with vertical tabs on the left side.
    """
    def __init__(self, parent=None):
        super().__init__(parent)
        self.main_window = parent
        self.setWindowTitle("Configure")
        self.setMinimumSize(1000, 700)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        
        self.setup_ui()
    
    def setup_ui(self):
        main_layout = QHBoxLayout(self)
        main_layout.setContentsMargins(0, 0, 0, 0)
        main_layout.setSpacing(0)
        
        # Left side - Vertical Tab Bar
        self.tab_list = QListWidget()
        self.tab_list.setFixedWidth(200)
        self.tab_list.setStyleSheet("""
            QListWidget {
                background-color: #f3f4f6;
                border: none;
                border-right: 2px solid #e5e7eb;
                font-size: 11pt;
                outline: none;
            }
            QListWidget::item {
                padding: 15px 10px;
                border-bottom: 1px solid #e5e7eb;
            }
            QListWidget::item:selected {
                background-color: #6B2C91;
                color: white;
                font-weight: bold;
            }
            QListWidget::item:hover:!selected {
                background-color: #e5e7eb;
            }
        """)
        
        self.tab_list.addItem("Default Location")
        self.tab_list.addItem("Document Layout")
        self.tab_list.addItem("Text Masking")
        self.tab_list.addItem("PCOMM Window")
        
        self.tab_list.setCurrentRow(0)
        self.tab_list.currentRowChanged.connect(self.change_tab)
        
        main_layout.addWidget(self.tab_list)
        
        # Right side - Stacked Widget for content
        self.content_stack = QWidget()
        content_layout = QVBoxLayout(self.content_stack)
        content_layout.setContentsMargins(20, 20, 20, 20)
        
        # Create all tabs
        # Create all tabs
        self.default_location_widget = self.create_default_location_tab()
        self.document_layout_widget = self.create_document_layout_tab()
        self.masking_widget = self.create_masking_tab()
        self.pcomm_widget = self.create_pcomm_tab()

        # Stack widget to hold all tabs
        self.stack = QStackedWidget()
        self.stack.addWidget(self.default_location_widget)
        self.stack.addWidget(self.document_layout_widget)
        self.stack.addWidget(self.masking_widget)
        self.stack.addWidget(self.pcomm_widget)
        
        content_layout.addWidget(self.stack)
        
        # Bottom buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        save_button = QPushButton("Save && Close")
        save_button.clicked.connect(self.save_and_close)
        button_layout.addWidget(save_button)
        
        cancel_button = QPushButton("Cancel")
        cancel_button.clicked.connect(self.reject)
        button_layout.addWidget(cancel_button)
        button_layout.addSpacing(16)
        
        content_layout.addLayout(button_layout)
        
        main_layout.addWidget(self.content_stack)
    
    def change_tab(self, index):
        """Changes the displayed tab."""
        self.stack.setCurrentIndex(index)
    
    def create_document_layout_tab(self):
        """Creates the Document Layout configuration tab with inline controls."""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Title
        title_label = QLabel("Document Layout Configuration")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        layout.addSpacing(10)
        
        # Description
        desc_label = QLabel("Configure text elements that will appear in the DOCX screenshot document.")
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)
        
        layout.addSpacing(15)

        # Generate Documentation Toggle
        doc_gen_layout = QHBoxLayout()

        self.generate_documentation_checkbox = QCheckBox("Generate DOCX Documentation for Screenshots")
        self.generate_documentation_checkbox.setChecked(self.main_window.document_config.get('generate_documentation', True))
        self.generate_documentation_checkbox.setStyleSheet("font-weight: bold; font-size: 11pt; color: #6B2C91;")
        doc_gen_layout.addWidget(self.generate_documentation_checkbox)
        doc_gen_layout.addStretch()

        layout.addWidget(self.generate_documentation_checkbox)

        layout.addSpacing(10)

        # Info label
        info_label = QLabel("When disabled, 'Capture Screenshot' steps will be skipped during test execution.")
        info_label.setStyleSheet("color: #6b7280; font-size: 9pt; font-style: italic;")
        info_label.setWordWrap(True)
        layout.addWidget(info_label)

        layout.addSpacing(15)

        # ✅ NEW: Capture Screen Flow Toggle
        self.capture_screen_flow_checkbox = QCheckBox("Capture Screen Flow (Before/After JPEG for Module Steps)")
        self.capture_screen_flow_checkbox.setChecked(self.main_window.document_config.get('capture_screen_flow', False))
        self.capture_screen_flow_checkbox.setStyleSheet("font-weight: bold; font-size: 11pt; color: #6B2C91;")
        layout.addWidget(self.capture_screen_flow_checkbox)

        layout.addSpacing(10)

        # Info label for screen flow
        flow_info_label = QLabel("When enabled, automatically captures before/after JPEG screenshots for all module steps.")
        flow_info_label.setStyleSheet("color: #6b7280; font-size: 9pt; font-style: italic;")
        flow_info_label.setWordWrap(True)
        layout.addWidget(flow_info_label)

        layout.addSpacing(15)

        # Separator
        separator = QFrame()
        separator.setFrameShape(QFrame.Shape.HLine)
        separator.setFrameShadow(QFrame.Shadow.Sunken)
        layout.addWidget(separator)
        layout.addSpacing(10)

        # Highlight Color Configuration
        
        # Highlight Color Configuration
        highlight_group = QWidget()
        highlight_layout = QHBoxLayout(highlight_group)
        highlight_layout.setContentsMargins(0, 0, 0, 0)
        
        highlight_label = QLabel("Screenshot Highlight Color:")
        highlight_label.setStyleSheet("font-weight: bold;")
        highlight_layout.addWidget(highlight_label)
        
        self.highlight_color_combo = QComboBox()
        self.highlight_color_combo.setFixedWidth(150)
        
        color_options = {
            'Yellow': 7, 'Bright Green': 4, 'Turquoise': 3, 'Pink': 5,
            'Blue': 2, 'Red': 6, 'Dark Blue': 9, 'Dark Cyan': 10,
            'Dark Green': 11, 'Dark Magenta': 12, 'Dark Red': 13,
            'Dark Yellow': 14, 'Gray 25%': 16, 'Gray 50%': 15
        }
        
        for color_name in color_options.keys():
            self.highlight_color_combo.addItem(color_name)
        
        current_color = self.main_window.document_config.get('highlight_color', 'Yellow')
        self.highlight_color_combo.setCurrentText(current_color)
        
        highlight_layout.addWidget(self.highlight_color_combo)
        
        # Color preview
        self.color_preview = QLabel("  Preview  ")
        self.color_preview.setStyleSheet("background-color: yellow; border: 1px solid black; padding: 5px;")
        self.color_preview.setFixedWidth(80)
        highlight_layout.addWidget(self.color_preview)
        
        highlight_layout.addStretch()
        
        self.highlight_color_combo.currentTextChanged.connect(self.update_color_preview)
        self.update_color_preview(current_color)
        
        layout.addWidget(highlight_group)
        layout.addSpacing(15)
        
        # Separator
        separator = QFrame()
        separator.setFrameShape(QFrame.Shape.HLine)
        separator.setFrameShadow(QFrame.Shadow.Sunken)
        layout.addWidget(separator)
        layout.addSpacing(10)
        
        # Text Elements List
        list_label = QLabel("Text Elements:")
        list_label.setStyleSheet("font-weight: bold;")
        layout.addWidget(list_label)
        
        self.config_list = QListWidget()
        self.config_list.setSelectionMode(QListWidget.SelectionMode.SingleSelection)
        layout.addWidget(self.config_list)
        
        # Populate existing config
        self.populate_config_list()
        
        # Buttons
        button_layout = QHBoxLayout()
        
        self.add_item_button = QPushButton("Add Text Element")
        self.add_item_button.clicked.connect(self.add_config_item)
        button_layout.addWidget(self.add_item_button)
        
        self.add_blank_line_button = QPushButton("Add Blank Line")
        self.add_blank_line_button.clicked.connect(self.add_blank_line)
        button_layout.addWidget(self.add_blank_line_button)
        
        self.edit_item_button = QPushButton("Edit Selected")
        self.edit_item_button.clicked.connect(self.edit_config_item)
        button_layout.addWidget(self.edit_item_button)
        
        self.delete_item_button = QPushButton("Delete Selected")
        self.delete_item_button.clicked.connect(self.delete_config_item)
        button_layout.addWidget(self.delete_item_button)
        
        self.move_up_button = QPushButton("↑ Move Up")
        self.move_up_button.setMinimumWidth(100)
        self.move_up_button.clicked.connect(self.move_item_up)
        button_layout.addWidget(self.move_up_button)
        
        self.move_down_button = QPushButton("↓ Move Down")
        self.move_down_button.clicked.connect(self.move_item_down)
        button_layout.addWidget(self.move_down_button)
        
        button_layout.addStretch()
        layout.addLayout(button_layout)
        
        return tab
    
    def update_color_preview(self, color_name):
        """Updates the color preview box."""
        color_map = {
            'Yellow': '#FFFF00', 'Bright Green': '#00FF00', 'Turquoise': '#00FFFF',
            'Pink': '#FF00FF', 'Blue': '#0000FF', 'Red': '#FF0000',
            'Dark Blue': '#00008B', 'Dark Cyan': '#008B8B', 'Dark Green': '#006400',
            'Dark Magenta': '#8B008B', 'Dark Red': '#8B0000', 'Dark Yellow': '#808000',
            'Gray 25%': '#C0C0C0', 'Gray 50%': '#808080'
        }
        
        hex_color = color_map.get(color_name, '#FFFF00')
        self.color_preview.setStyleSheet(
            f"background-color: {hex_color}; border: 1px solid black; padding: 5px;"
        )
    
    def populate_config_list(self):
        """Populates the config list with existing items."""
        self.config_list.clear()
        text_elements = self.main_window.document_config.get('text_elements', [])
        
        for item in text_elements:
            if item.get('type') == 'blank_line':
                self.config_list.addItem("[Blank Line]")
            else:
                text = item.get('text', 'No text')
                font = item.get('font_name', 'Arial')
                size = item.get('font_size', 12)
                align = item.get('alignment', 'Left')
                display_text = text[:30] + "..." if len(text) > 30 else text
                self.config_list.addItem(f"{display_text} | {font} {size}pt | {align}")
    
    def add_config_item(self):
        """Adds a new text element."""
        dialog = TextElementDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            item_data = dialog.get_data()
            if 'text_elements' not in self.main_window.document_config:
                self.main_window.document_config['text_elements'] = []
            self.main_window.document_config['text_elements'].append(item_data)
            self.populate_config_list()
    
    def add_blank_line(self):
        """Adds a blank line element."""
        blank_line_item = {
            'type': 'blank_line',
            'text': '',
            'font_name': 'Arial',
            'font_size': 12,
            'alignment': 'Left',
            'bold': False,
            'italic': False
        }
        if 'text_elements' not in self.main_window.document_config:
            self.main_window.document_config['text_elements'] = []
        self.main_window.document_config['text_elements'].append(blank_line_item)
        self.populate_config_list()
    
    def edit_config_item(self):
        """Edits the selected config item."""
        current_row = self.config_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "No Selection", "Please select a text element to edit.")
            return
        
        text_elements = self.main_window.document_config.get('text_elements', [])
        if current_row >= len(text_elements):
            return
        
        item_data = text_elements[current_row]
        
        if item_data.get('type') == 'blank_line':
            QMessageBox.information(self, "Blank Line", "Blank lines cannot be edited.")
            return
        
        dialog = TextElementDialog(self, existing_data=item_data)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            updated_data = dialog.get_data()
            text_elements[current_row] = updated_data
            self.populate_config_list()
            self.config_list.setCurrentRow(current_row)
    
    def delete_config_item(self):
        """Deletes the selected config item."""
        current_row = self.config_list.currentRow()
        if current_row < 0:
            QMessageBox.warning(self, "No Selection", "Please select a text element to delete.")
            return
        
        reply = QMessageBox.question(
            self, "Confirm Delete",
            "Are you sure you want to delete this text element?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            text_elements = self.main_window.document_config.get('text_elements', [])
            if current_row < len(text_elements):
                text_elements.pop(current_row)
                self.populate_config_list()
    
    def move_item_up(self):
        """Moves the selected item up."""
        current_row = self.config_list.currentRow()
        if current_row <= 0:
            return
        
        text_elements = self.main_window.document_config.get('text_elements', [])
        text_elements[current_row], text_elements[current_row - 1] = \
            text_elements[current_row - 1], text_elements[current_row]
        self.populate_config_list()
        self.config_list.setCurrentRow(current_row - 1)
    
    def move_item_down(self):
        """Moves the selected item down."""
        current_row = self.config_list.currentRow()
        text_elements = self.main_window.document_config.get('text_elements', [])
        if current_row < 0 or current_row >= len(text_elements) - 1:
            return
        
        text_elements[current_row], text_elements[current_row + 1] = \
            text_elements[current_row + 1], text_elements[current_row]
        self.populate_config_list()
        self.config_list.setCurrentRow(current_row + 1)
    
    def create_masking_tab(self):
        """Creates the Text Masking configuration tab with inline controls."""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Title
        title_label = QLabel("Text Masking Configuration")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        layout.addSpacing(10)
        
        # Enable/Disable Toggle
        self.masking_enabled_checkbox = QCheckBox("Enable Text Masking in Screenshots")
        self.masking_enabled_checkbox.setChecked(self.main_window.masking_enabled)
        self.masking_enabled_checkbox.setStyleSheet("font-weight: bold; font-size: 12pt; color: #6B2C91;")
        layout.addWidget(self.masking_enabled_checkbox)
        
        layout.addSpacing(10)
        
        # Description
        desc_label = QLabel(
            "When enabled, sensitive data in screenshots will be masked based on the configured rules below."
        )
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)
        
        layout.addSpacing(15)
        
        # Separator
        separator = QFrame()
        separator.setFrameShape(QFrame.Shape.HLine)
        separator.setFrameShadow(QFrame.Shadow.Sunken)
        layout.addWidget(separator)
        layout.addSpacing(10)
        
        # Input Section
        # Input Section
        input_group = QWidget()
        input_layout = QVBoxLayout(input_group)
        input_layout.setContentsMargins(10, 10, 10, 10)
        input_group.setStyleSheet("background-color: #f3f4f6; border-radius: 4px;")
        
        # Step 1
        step1_label = QLabel("Step 1: Enter Sample Structure (e.g., 491100-5530000-15-6)")
        step1_label.setStyleSheet("font-weight: bold; padding: 5px; background-color: transparent;")
        input_layout.addWidget(step1_label)
        
        self.sample_input = QLineEdit()
        self.sample_input.setPlaceholderText("Enter sample pattern here...")
        input_layout.addWidget(self.sample_input)
        
        input_layout.addSpacing(10)
        
        # Step 2
        step2_label = QLabel("Step 2: Masking Positions (1-based, e.g., 8-10, 19)")
        step2_label.setStyleSheet("font-weight: bold; padding: 5px; background-color: transparent;")
        input_layout.addWidget(step2_label)
        
        positions_layout = QHBoxLayout()
        self.position_inputs = []
        for i in range(4):
            pos_input = QLineEdit()
            pos_input.setPlaceholderText(f"Position {i+1}")
            pos_input.setMinimumHeight(30)
            positions_layout.addWidget(pos_input)
            self.position_inputs.append(pos_input)

        input_layout.addLayout(positions_layout)

        input_layout.addSpacing(15)  # Increased spacing

        # Add Rule Button
        add_button = QPushButton("Add Rule")
        add_button.setMinimumWidth(120)
        add_button.setMinimumHeight(35)  # Increased height
        add_button.setStyleSheet("""
            QPushButton {
                background-color: #6B2C91;
                color: white;
                border-radius: 4px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #7C3AA3;
            }
        """)
        add_button.clicked.connect(self.add_masking_pattern)
        input_layout.addWidget(add_button)
        
        layout.addWidget(input_group)
        layout.addSpacing(15)
        
        # Patterns List Section
        list_label = QLabel("Active Masking Rules:")
        list_label.setStyleSheet("font-weight: bold;")
        layout.addWidget(list_label)

        self.patterns_list = QListWidget()
        self.patterns_list.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        self.patterns_list.setMinimumHeight(200)
        layout.addWidget(self.patterns_list)
        
        # Populate existing patterns
        self.populate_patterns_list()
        
        # List Buttons
        list_buttons_layout = QHBoxLayout()
        
        remove_button = QPushButton("Remove Selected")
        remove_button.clicked.connect(self.remove_selected_patterns)
        list_buttons_layout.addWidget(remove_button)
        
        list_buttons_layout.addStretch()
        
        layout.addLayout(list_buttons_layout)
        
        return tab
    
    def populate_patterns_list(self):
        """Populates the patterns list."""
        self.patterns_list.clear()
        
        for pattern in self.main_window.masking_patterns:
            sample = pattern.get('sample', '')
            masked = pattern.get('masked_sample', '')
            regex = pattern.get('regex', '')
            
            # Truncate regex if too long
            display_regex = regex[:40] + "..." if len(regex) > 40 else regex
            display_text = f"{sample} → {masked}  |  Pattern: {display_regex}"
            self.patterns_list.addItem(display_text)
    
    def add_masking_pattern(self):
        """Adds a new masking pattern."""
        sample_string = self.sample_input.text().strip()
        position_strings = [inp.text().strip() for inp in self.position_inputs]
        
        if not sample_string:
            QMessageBox.warning(self, "Warning", "Please enter a sample structure.")
            return
        
        if not any(position_strings):
            QMessageBox.warning(self, "Warning", "Please enter at least one position to mask.")
            return
        
        try:
            mask_indices = self.parse_mask_positions(position_strings)
        except ValueError as e:
            QMessageBox.critical(self, "Invalid Input", str(e))
            return
        
        if mask_indices and max(mask_indices) >= len(sample_string):
            QMessageBox.critical(self, "Invalid Position", 
                               f"Position {max(mask_indices) + 1} exceeds sample length ({len(sample_string)}).")
            return
        
        generated_regex = self.convert_to_regex(sample_string)
        masked_sample = self.apply_mask_by_indices(sample_string, mask_indices)
        
        new_pattern = {
            "regex": generated_regex,
            "sample": sample_string,
            "mask_indices": mask_indices,
            "masked_sample": masked_sample
        }
        
        # Check for duplicates
        if any(p.get('regex') == generated_regex for p in self.main_window.masking_patterns):
            QMessageBox.information(self, "Duplicate", "This pattern already exists.")
            return
        
        self.main_window.masking_patterns.append(new_pattern)
        self.populate_patterns_list()
        
        # Clear inputs
        self.sample_input.clear()
        for inp in self.position_inputs:
            inp.clear()
    
    def remove_selected_patterns(self):
        """Removes selected patterns."""
        selected_items = self.patterns_list.selectedItems()
        if not selected_items:
            return
        
        selected_indices = [self.patterns_list.row(item) for item in selected_items]
        
        for index in sorted(selected_indices, reverse=True):
            if 0 <= index < len(self.main_window.masking_patterns):
                del self.main_window.masking_patterns[index]
        
        self.populate_patterns_list()
    
    def convert_to_regex(self, sample_string):
        """Converts a sample string into a structural regex."""
        if not sample_string:
            return ""
        
        regex = ""
        digit_count = 0
        
        for char in sample_string:
            if char.isdigit():
                digit_count += 1
            else:
                if digit_count > 0:
                    regex += f"\\d{{{digit_count}}}"
                    digit_count = 0
                regex += re.escape(char)
        
        if digit_count > 0:
            regex += f"\\d{{{digit_count}}}"
        
        return r"\b" + regex + r"\b"
    
    def parse_mask_positions(self, position_strings):
        """Parses position strings into 0-based indices."""
        indices = set()
        
        for pos_str in position_strings:
            if not pos_str.strip():
                continue
            
            parts = [p.strip() for p in pos_str.split(',')]
            
            for part in parts:
                if '-' in part:
                    try:
                        start, end = map(int, part.split('-'))
                        if start < 1 or end < 1 or start > end:
                            raise ValueError(f"Invalid range: {part}")
                        for i in range(start - 1, end):
                            indices.add(i)
                    except ValueError:
                        raise ValueError(f"Invalid range format: {part}")
                else:
                    try:
                        single_pos = int(part)
                        if single_pos < 1:
                            raise ValueError(f"Position must be >= 1: {part}")
                        indices.add(single_pos - 1)
                    except ValueError:
                        raise ValueError(f"Invalid position: {part}")
        
        return sorted(list(indices))
    
    def apply_mask_by_indices(self, text, indices):
        """Applies 'x' masking to text at given indices."""
        text_list = list(text)
        max_len = len(text_list)
        
        for index in indices:
            if 0 <= index < max_len:
                text_list[index] = 'x'
        
        return "".join(text_list)
    
    def create_pcomm_tab(self):
        """Creates the PCOMM Window configuration tab."""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Title
        title_label = QLabel("PCOMM Window Configuration")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        layout.addSpacing(10)
        
        # Description
        desc_label = QLabel("Enter the exact window title as it appears in your PCOMM application.")
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)
        
        layout.addSpacing(20)
        
        # Window Title Input
        form_layout = QFormLayout()
        
        self.pcomm_title_input = QLineEdit()
        self.pcomm_title_input.setText(self.main_window.pcomm_window_title)
        self.pcomm_title_input.setPlaceholderText("e.g., SessionA or Session A")
        
        form_layout.addRow("Window Title:", self.pcomm_title_input)
        
        layout.addLayout(form_layout)
        
        layout.addStretch()
        
        return tab
    
    def save_and_close(self):
        """Saves all settings and closes the dialog."""
        # Save Default Location
        new_location = self.location_path_input.text().strip()
        if new_location and os.path.isdir(new_location):
            self.main_window.default_results_location = new_location
            self.main_window.save_default_location_config()
        
        # Save Document Layout
        self.main_window.document_config['generate_documentation'] = self.generate_documentation_checkbox.isChecked()
        self.main_window.document_config['highlight_color'] = self.highlight_color_combo.currentText()
        self.main_window.document_config['capture_screen_flow'] = self.capture_screen_flow_checkbox.isChecked()  # ✅ NEW
        self.main_window.save_document_config()
        
        # Save Masking Settings
        self.main_window.masking_enabled = self.masking_enabled_checkbox.isChecked()
        self.main_window.save_masking_config()
        
        # Save PCOMM Settings
        new_title = self.pcomm_title_input.text().strip()
        if new_title:
            self.main_window.pcomm_window_title = new_title
            self.main_window.save_pcomm_window_config()
        
        QMessageBox.information(self, "Success", "All settings saved successfully.")
        self.accept()
        
    def create_default_location_tab(self):
        """Creates the Default Location configuration tab."""
        tab = QWidget()
        layout = QVBoxLayout(tab)
        
        # Title
        title_label = QLabel("Default Results Location")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        layout.addSpacing(10)
        
        # Description
        desc_label = QLabel("Configure where test results and execution summaries will be saved.")
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)
        
        layout.addSpacing(20)
        
        # Current location display
        location_group = QWidget()
        location_layout = QHBoxLayout(location_group)
        location_layout.setContentsMargins(0, 0, 0, 0)
        
        location_label = QLabel("Save Location:")
        location_label.setFixedWidth(105)
        location_label.setStyleSheet("font-weight: bold;")
        location_layout.addWidget(location_label)
        
        self.location_path_input = QLineEdit()
        self.location_path_input.setText(self.main_window.default_results_location)
        self.location_path_input.setReadOnly(True)
        self.location_path_input.setMinimumWidth(400)
        location_layout.addWidget(self.location_path_input)
        
        browse_button = QPushButton("Browse...")
        browse_button.setFixedWidth(100)
        browse_button.clicked.connect(self.browse_location)
        location_layout.addWidget(browse_button)
        
        location_layout.addStretch()
        
        layout.addWidget(location_group)
        
        layout.addSpacing(15)
        
        # Info label
        info_label = QLabel(
            "Results will be saved in:\n"
            "• [Location]/Results/[Project or Master]/\n"
            "• [Location]/Test Execution Summary/"
        )
        info_label.setStyleSheet("color: #6b7280; font-size: 9pt; padding: 10px; background-color: #f9fafb; border-radius: 4px;")
        info_label.setWordWrap(True)
        layout.addWidget(info_label)
        
        layout.addSpacing(15)

        layout.addStretch()

        return tab

    def browse_location(self):
        """Opens a directory browser to select save location."""
        directory = QFileDialog.getExistingDirectory(
            self,
            "Select Results Save Location",
            self.location_path_input.text(),
            QFileDialog.Option.ShowDirsOnly
        )
        
        if directory:
            self.location_path_input.setText(directory)

       

class TextElementDialog(QDialog):
    """
    Dialog to add or edit a single text element configuration.
    """
    def __init__(self, parent=None, existing_data=None):
        super().__init__(parent)
        self.setWindowTitle("Text Element Configuration")
        self.setFixedSize(500, 300)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        
        self.existing_data = existing_data
        self.setup_ui()
        
        if existing_data:
            self.populate_existing_data()
    
    def setup_ui(self):
        layout = QFormLayout(self)
        
        # Text input
        self.text_input = QTextEdit()
        self.text_input.setPlaceholderText("Enter text to appear in document...")
        self.text_input.setMaximumHeight(80)
        layout.addRow("Text:", self.text_input)
        
        # NEW: Add a label showing available variables
        variables_label = QLabel(
            "<b>Available Variables:</b> {test_case_id}, {test_description}, {date}, {time}, {datetime}, {total_screenshots}, {space}"
        )      
        variables_label.setWordWrap(True)
        variables_label.setStyleSheet("color: #555; font-size: 9pt; padding: 5px;")
        layout.addRow("", variables_label)
        
        # Font selection
        self.font_combo = QComboBox()
        self.font_combo.addItems([
            "Arial", "Times New Roman", "Calibri", "Courier New", 
            "Georgia", "Verdana", "Tahoma", "Comic Sans MS"
        ])
        layout.addRow("Font Name:", self.font_combo)
        
        # Font size
        self.font_size_spin = QLineEdit()
        self.font_size_spin.setText("12")
        self.font_size_spin.setValidator(QIntValidator(6, 72))
        layout.addRow("Font Size (pt):", self.font_size_spin)
        
        # Alignment
        self.alignment_combo = QComboBox()
        self.alignment_combo.addItems(["Left", "Center", "Right", "Justify"])
        layout.addRow("Alignment:", self.alignment_combo)
        
        # Bold checkbox
        self.bold_checkbox = QCheckBox("Bold")
        layout.addRow("", self.bold_checkbox)
        
        # Italic checkbox
        self.italic_checkbox = QCheckBox("Italic")
        layout.addRow("", self.italic_checkbox)
        
        # Dialog buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
    
    def populate_existing_data(self):
        """Fill the form with existing data."""
        self.text_input.setText(self.existing_data.get('text', ''))
        
        font_name = self.existing_data.get('font_name', 'Arial')
        index = self.font_combo.findText(font_name)
        if index >= 0:
            self.font_combo.setCurrentIndex(index)
        
        self.font_size_spin.setText(str(self.existing_data.get('font_size', 12)))
        
        alignment = self.existing_data.get('alignment', 'Left')
        align_index = self.alignment_combo.findText(alignment)
        if align_index >= 0:
            self.alignment_combo.setCurrentIndex(align_index)
        
        self.bold_checkbox.setChecked(self.existing_data.get('bold', False))
        self.italic_checkbox.setChecked(self.existing_data.get('italic', False))
    
    def get_data(self):
        """Return the configured text element data."""
        return {
            'text': self.text_input.toPlainText(),
            'font_name': self.font_combo.currentText(),
            'font_size': int(self.font_size_spin.text() or 12),
            'alignment': self.alignment_combo.currentText(),
            'bold': self.bold_checkbox.isChecked(),
            'italic': self.italic_checkbox.isChecked()
        }

class TestExecutionDialog(QDialog):

    def __init__(self, parent=None, test_cases_data=None):
        super().__init__(parent)
        self.test_cases = test_cases_data or []
        self.setWindowTitle("Test Execution")
        self.setMinimumSize(700, 500)
        self.main_window = parent
        self.displayed_test_cases = {}
        self.previously_imported = set()
        self.execution_data_file = 'test_execution_data.json'
        self.execution_times = {}
        self.stop_execution = False
        
        # ✅ NEW: Projects dictionary to store project structure
        self.projects = {}  # {project_name: {test_cases: {}, expanded: True}}
        
        self.setup_ui()
        self.load_execution_data()

    def validate_field_value(self, actual_value, expected_value):
        """
        Validates a field value, with special handling for {blank} validation.
        
        Args:
            actual_value: The actual value from PCOMM screen
            expected_value: The expected value from test case (can be '{blank}')
        
        Returns:
            bool: True if validation passes, False otherwise
        """
        # Strip both values for comparison
        actual_stripped = actual_value.strip()
        expected_stripped = expected_value.strip()
        
        # Special handling for {blank} validation
        if expected_stripped.lower() == '{blank}':
            # Check if actual value is empty/blank
            return actual_stripped == '' or actual_stripped.isspace() or len(actual_stripped) == 0
        
        # Normal validation (existing behavior)
        return actual_stripped == expected_stripped

    def setup_ui(self):
        main_layout = QVBoxLayout(self)

        # --- Top Control Section ---
        top_layout = QHBoxLayout()
        self.select_all_checkbox = QCheckBox("Select All")
        self.select_all_checkbox.stateChanged.connect(self.select_all_test_cases)  # ✅ Keep this connection
        top_layout.addWidget(self.select_all_checkbox)

        # Status filter dropdown
        filter_label = QLabel("Filter by Status:")
        top_layout.addWidget(filter_label)

        self.status_filter_combo = QComboBox()
        self.status_filter_combo.addItems(["All", "Not Run", "Running", "Passed", "Failed", "Stopped"])
        self.status_filter_combo.setFixedWidth(120)
        self.status_filter_combo.currentTextChanged.connect(self.filter_by_status)
        top_layout.addWidget(self.status_filter_combo)

        top_layout.addStretch()

        # ✅ NEW: Add Project button
        add_project_button = QPushButton("➕ Add Project")
        add_project_button.clicked.connect(self.add_project)
        top_layout.addWidget(add_project_button)

        # Import Test Cases button with dropdown menu
        import_button = QPushButton("Import Test Cases")
        import_button.setStyleSheet("""
            QPushButton::menu-indicator {
                subcontrol-position: center right;
                subcontrol-origin: padding;
                left: -2px;
            }
        """)
        import_menu = QMenu(self)
        import_from_library_action = QAction("Import from Library", self)
        import_from_library_action.triggered.connect(self.import_from_library)
        import_from_device_action = QAction("Import from Device", self)
        import_from_device_action.triggered.connect(self.import_from_device)
        import_menu.addAction(import_from_library_action)
        import_menu.addAction(import_from_device_action)
        import_button.setMenu(import_menu)
        top_layout.addWidget(import_button)
        main_layout.addLayout(top_layout)

        # --- Test Case List Section ---
        self.test_case_list = QListWidget()
        main_layout.addWidget(self.test_case_list)
        
        self.test_case_list.setDragEnabled(True)
        self.test_case_list.setAcceptDrops(True)
        self.test_case_list.setDragDropMode(QListWidget.DragDropMode.InternalMove)
        self.test_case_list.setDefaultDropAction(Qt.DropAction.MoveAction)

        # --- Bottom Control Section ---
        bottom_layout = QHBoxLayout()
        bottom_layout.addStretch()
        
        self.execute_button = QPushButton("Execute Tests")
        self.execute_button.clicked.connect(self.execute_selected_tests)
        bottom_layout.addWidget(self.execute_button)
        
        self.clear_button = QPushButton("Clear")
        self.clear_button.clicked.connect(self.clear_all_test_cases)
        bottom_layout.addWidget(self.clear_button)
        
        self.close_button = QPushButton("Close")
        self.close_button.clicked.connect(self.reject)
        bottom_layout.addWidget(self.close_button)
        
        self.test_case_list.model().rowsMoved.connect(self.on_rows_moved)
        main_layout.addLayout(bottom_layout)

    def add_project(self):
        """Opens a dialog to add a new project."""
        from PyQt6.QtWidgets import QInputDialog
        
        project_name, ok = QInputDialog.getText(
            self,
            "Add Project",
            "Enter project name:"
        )
        
        if ok and project_name.strip():
            project_name = project_name.strip()
            
            # Check if project already exists
            if project_name in self.projects:
                QMessageBox.warning(self, "Duplicate Project", 
                                  f"Project '{project_name}' already exists.")
                return
            
            # Create new project
            self.projects[project_name] = {
                'test_cases': {},
                'expanded': True
            }
            
            # Add project header to list
            self.add_project_header(project_name)
            self.save_execution_data()
            
            QMessageBox.information(self, "Success", 
                                  f"Project '{project_name}' created successfully.")



    def add_project_header(self, project_name):
        """Adds a collapsible project header to the list."""
        item = QListWidgetItem(self.test_case_list)
        
        # ✅ NEW: Disable dragging for project headers
        item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsDragEnabled)
        
        item_widget = QWidget()
        item_layout = QHBoxLayout(item_widget)
        item_layout.setContentsMargins(0, 4, 8, 4)
        item_layout.setSpacing(3)
        
        # Expand/Collapse button
        expand_button = QPushButton("▼" if self.projects[project_name]['expanded'] else "▶")
        expand_button.setFixedSize(QSize(16, 16))
        expand_button.setObjectName("expand_button")
        expand_button.setStyleSheet("""
            QPushButton {
                border: none;
                background-color: transparent;
                font-size: 8pt;
                padding: 0px;
            }
            QPushButton:hover {
                background-color: #e5e7eb;
                border-radius: 2px;
            }
        """)
        expand_button.clicked.connect(lambda: self.toggle_project_expansion(project_name))
        item_layout.addWidget(expand_button, 0, Qt.AlignmentFlag.AlignVCenter)
    

        
        # Project name label (bold and slightly larger)
        name_label = QLabel(f"📁 {project_name}")
        name_label.setStyleSheet("font-weight: bold; font-size: 9pt; color: #6B2C91; padding: 0px 0px;")  # ✅ CHANGED: Added padding
        name_label.setMinimumWidth(200)
        name_label.setMinimumHeight(24)  # ✅ NEW: Set minimum height for the label
        name_label.setWordWrap(True)
        name_label.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        item_layout.addWidget(name_label, 1, Qt.AlignmentFlag.AlignVCenter)
        
        # Import button for this project
        import_icon_svg = """
        <svg width="16" height="16" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M14 2H6C5.44772 2 5 2.44772 5 3V19C5 19.5523 5.44772 20 6 20H18C18.5523 20 19 19.5523 19 19V8L14 2Z" stroke="#6B2C91" stroke-width="2"/>
            <path d="M12 11V18M9 14L12 11L15 14" stroke="#6B2C91" stroke-width="2" stroke-linecap="round"/>
        </svg>
        """
        import_icon = QIcon()
        import_pixmap = QPixmap()
        import_pixmap.loadFromData(QByteArray(import_icon_svg.encode('utf-8')))
        import_icon.addPixmap(import_pixmap, QIcon.Mode.Normal, QIcon.State.Off)
        
        import_button = QPushButton()
        import_button.setIcon(import_icon)
        import_button.setFixedSize(QSize(24, 24))
        import_button.setToolTip(f"Import test cases to '{project_name}'")
        import_button.setStyleSheet("""
            QPushButton {
                border: none;
                border-radius: 3px;
                background-color: transparent;
                padding: 2px;
            }
            QPushButton:hover {
                background-color: #e5e7eb;
            }
        """)
        import_button.clicked.connect(lambda: self.import_to_project(project_name))
        item_layout.addWidget(import_button, 0, Qt.AlignmentFlag.AlignVCenter)
        
        # Delete project button
        delete_button = QPushButton("✕")
        delete_button.setFixedSize(20, 20)
        delete_font = QFont()
        delete_font.setBold(True)
        delete_font.setPointSize(13)
        delete_button.setFont(delete_font)
        delete_button.setStyleSheet("""
            QPushButton {
                color: #dc2626;
                border: none;
                background-color: transparent;
                padding: 0px;
            }
            QPushButton:hover {
                color: #991b1b;
                background-color: #fee2e2;
                border-radius: 3px;
            }
        """)
        delete_button.clicked.connect(lambda: self.delete_project(project_name))
        item_layout.addWidget(delete_button, 0, Qt.AlignmentFlag.AlignVCenter)
        
        item_widget.setLayout(item_layout)
        item_widget.setStyleSheet("background-color: #f3f4f6; border-radius: 4px;")
        item_widget.setMinimumHeight(35)  # ✅ NEW: Set minimum height for the entire widget
        
        self.test_case_list.addItem(item)
        self.test_case_list.setItemWidget(item, item_widget)
        
        # ✅ NEW: Calculate proper size hint based on content
        size_hint = item_widget.sizeHint()
        size_hint.setHeight(max(40, size_hint.height()))  # Ensure at least 40px height
        item.setSizeHint(size_hint)
        
        # Store reference to project header item
        item.setData(Qt.ItemDataRole.UserRole, {'type': 'project', 'name': project_name})

    def toggle_project_expansion(self, project_name):
        """Toggles the expansion state of a project."""
        self.projects[project_name]['expanded'] = not self.projects[project_name]['expanded']
        
        # ✅ CHANGED: Don't call refresh_test_case_list, manually handle expansion
        is_expanded = self.projects[project_name]['expanded']
        
        # Find the project header item
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'project' and item_data.get('name') == project_name:
                # Update the expand/collapse button icon
                widget = self.test_case_list.itemWidget(item)
                expand_button = widget.findChild(QPushButton, "expand_button")
                if expand_button:
                    expand_button.setText("▼" if is_expanded else "▶")
                
                # Show/hide test cases under this project
                j = i + 1
                while j < self.test_case_list.count():
                    next_item = self.test_case_list.item(j)
                    next_item_data = next_item.data(Qt.ItemDataRole.UserRole)
                    
                    # Stop if we hit another project header
                    if next_item_data and next_item_data.get('type') == 'project':
                        break
                    
                    # Show/hide test case items belonging to this project
                    if next_item_data and next_item_data.get('type') == 'test_case' and next_item_data.get('project') == project_name:
                        next_item.setHidden(not is_expanded)
                    
                    j += 1
                
                break
        
        self.save_execution_data()

    def delete_project(self, project_name):
        """Deletes a project and all its test cases."""
        reply = QMessageBox.question(
            self,
            "Delete Project",
            f"Are you sure you want to delete project '{project_name}' and all its test cases?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            del self.projects[project_name]
            self.refresh_test_case_list()
            self.save_execution_data()
            QMessageBox.information(self, "Deleted", f"Project '{project_name}' deleted.")

    def import_to_project(self, project_name):
        """Shows import options for a specific project."""
        menu = QMenu(self)
        
        from_library = QAction("Import from Library", self)
        from_library.triggered.connect(lambda: self.import_from_library(project_name))
        
        from_device = QAction("Import from Device", self)
        from_device.triggered.connect(lambda: self.import_from_device(project_name))
        
        menu.addAction(from_library)
        menu.addAction(from_device)
        
        # Show menu at cursor position
        menu.exec(self.cursor().pos())

    def filter_by_status(self, status_filter):
        """Filters test cases by status."""
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            # Don't hide project headers
            if item_data and item_data.get('type') == 'project':
                item.setHidden(False)
                continue
            
            # Filter test cases
            if item_data and item_data.get('type') == 'test_case':
                widget = self.test_case_list.itemWidget(item)
                if widget:
                    status_label = widget.findChild(QLabel, "status_label")
                    
                    if status_label:
                        current_status = status_label.text()
                        if status_filter == "All":
                            item.setHidden(False)
                        else:
                            item.setHidden(current_status != status_filter)

    def populate_test_cases(self, new_cases=None):
        """Adds test cases to the dialog's list, handling duplicates."""
        
        if new_cases is None:
            cases_to_add = self.main_window.test_cases
        else:
            cases_to_add = new_cases

        for name, data in cases_to_add.items():
            if name not in self.displayed_test_cases:
                self.displayed_test_cases[name] = data
                self.add_list_item(name)
    
        self.test_case_list.sortItems()

    def add_list_item(self, name, project_name=None):
        """Creates a test case item - updated to support indentation for project tests."""
        item = QListWidgetItem(self.test_case_list)  # ✅ CHANGED: Use test_case_list
        item_widget = QWidget()
        item_layout = QHBoxLayout(item_widget)
        
        # Add left margin for project test cases
        if project_name:
            item_layout.setContentsMargins(30, 1, 8, 6)
        else:
            item_layout.setContentsMargins(0, 1, 8, 6)
        
        item_layout.setSpacing(5)
        
        checkbox = QCheckBox()
        # ✅ FIXED: Use blockSignals to prevent triggering update_select_all_state during initialization
        checkbox.blockSignals(True)
        checkbox.setCheckState(Qt.CheckState.Unchecked)
        checkbox.blockSignals(False)
        
        # ✅ FIXED: Connect directly to stateChanged, not the select_all method
        checkbox.stateChanged.connect(self.update_select_all_state)
        item_layout.addWidget(checkbox, 0, Qt.AlignmentFlag.AlignVCenter)

        name_label = QLabel(name)
        name_label.setMinimumWidth(150)
        name_label.setAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)

        # ✅ NEW: Add tooltip showing assumptions
        if project_name:
            test_case_data = self.projects[project_name]['test_cases'].get(name)
        else:
            test_case_data = self.displayed_test_cases.get(name)

        if test_case_data:
            assumptions_html = test_case_data.get('assumptions', '')
            if assumptions_html:
                # Keep HTML formatting for rich text tooltip (tables will render properly)
                name_label.setToolTip(f"<b>Assumptions:</b><br>{assumptions_html}")
                # Enable rich text rendering for the label
                name_label.setTextFormat(Qt.TextFormat.RichText)
            else:
                name_label.setToolTip("No assumptions defined")

        item_layout.addWidget(name_label, 0, Qt.AlignmentFlag.AlignVCenter)
        
        item_layout.addStretch()

        # Get test case data
        if project_name:
            test_case_data = self.projects[project_name]['test_cases'].get(name)
        else:
            test_case_data = self.displayed_test_cases.get(name)

        # ✅ NEW: Show execution time if test has been executed
        execution_time_label = QLabel("")
        execution_time_label.setObjectName("execution_time_label")
        execution_time_label.setStyleSheet("font-size: 8pt; color: #dc2626; font-style: italic; font-weight: bold;")
        execution_time_label.setFixedWidth(120)
        execution_time_label.setAlignment(Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)

        if name in self.execution_times and 'duration' in self.execution_times[name]:
            duration = self.execution_times[name]['duration']
            execution_time_label.setText(duration)

        item_layout.addWidget(execution_time_label, 0, Qt.AlignmentFlag.AlignVCenter)

        # Step selector
        step_label = QLabel("Start from:")
        step_label.setStyleSheet("font-size: 9pt;")
        item_layout.addWidget(step_label, 0, Qt.AlignmentFlag.AlignVCenter)
        
        step_combo = QComboBox()
        step_combo.setObjectName("step_combo")
        step_combo.setFixedWidth(100)
        step_combo.setFixedHeight(22)
        step_combo.setStyleSheet("font-size: 9pt; margin-bottom: 8px;")
        
        if test_case_data and 'steps' in test_case_data:
            num_steps = len(test_case_data['steps'])
            for i in range(1, num_steps + 1):
                step_combo.addItem(f"Step {i}")
        else:
            step_combo.addItem("Step 1")
        
        step_combo.setCurrentIndex(0)
        item_layout.addWidget(step_combo, 0, Qt.AlignmentFlag.AlignVCenter)
        
        # Status Label
        status_label = QLabel("Not Run")
        status_label.setFixedWidth(60)
        status_label.setFixedHeight(24)
        status_label.setAlignment(Qt.AlignmentFlag.AlignCenter | Qt.AlignmentFlag.AlignVCenter)
        status_label.setObjectName("status_label")
        status_label.setStyleSheet("color: #6B2C91; font-weight: normal;")
        item_layout.addWidget(status_label, 0, Qt.AlignmentFlag.AlignVCenter)

        # Play/Stop button
        play_stop_button = QPushButton()
        play_stop_button.setIcon(self.main_window.style().standardIcon(QStyle.StandardPixmap.SP_MediaPlay))
        play_stop_button.setFixedSize(QSize(30, 30))
        play_stop_button.setObjectName("play_stop_button")
        play_stop_button.clicked.connect(lambda: self.toggle_play_stop(name))
        play_stop_button.setStyleSheet("margin-bottom: 0px;")
        item_layout.addWidget(play_stop_button, 0, Qt.AlignmentFlag.AlignVCenter)

        # Refresh button
        refresh_icon_svg = """<svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M21 10C21 10 18.995 7.26822 17.3662 5.63824C15.7373 4.00827 13.4864 3 11 3C6.02944 3 2 7.02944 2 12C2 16.9706 6.02944 21 11 21C15.1031 21 18.5649 18.2543 19.6482 14.5M21 10V4M21 10H15" stroke="#6B2C91" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>"""
        refresh_icon = QIcon()
        refresh_pixmap = QPixmap()
        refresh_pixmap.loadFromData(QByteArray(refresh_icon_svg.encode('utf-8')))
        refresh_icon.addPixmap(refresh_pixmap, QIcon.Mode.Normal, QIcon.State.Off)
        
        refresh_button = QPushButton()
        refresh_button.setIcon(refresh_icon)
        refresh_button.setFixedSize(QSize(30, 30))
        refresh_button.setStyleSheet("""
            QPushButton { border: none; border-radius: 4px; background-color: transparent; margin-bottom: 0px; }
            QPushButton:hover { background-color: #e5e7eb; }
        """)
        refresh_button.setToolTip("Refresh test case")
        refresh_button.clicked.connect(lambda: self.refresh_test_case(name))
        item_layout.addWidget(refresh_button, 0, Qt.AlignmentFlag.AlignVCenter)

        # Delete button
        delete_button = QPushButton("✕")
        delete_button.setFixedSize(20, 20)
        delete_font = QFont()
        delete_font.setBold(True)
        delete_font.setPointSize(13)
        delete_button.setFont(delete_font)
        delete_button.setStyleSheet("""
            QPushButton { color: #dc2626; border: none; background-color: transparent; padding: 0px; }
            QPushButton:hover { color: #991b1b; background-color: #fee2e2; border-radius: 3px; }
        """)
        delete_button.clicked.connect(lambda: self.delete_single_test(name))
        item_layout.addWidget(delete_button, 0, Qt.AlignmentFlag.AlignVCenter)

        item_widget.setLayout(item_layout)
        self.test_case_list.addItem(item)  # ✅ CHANGED: Use test_case_list
        self.test_case_list.setItemWidget(item, item_widget)  # ✅ CHANGED: Use test_case_list
        item.setSizeHint(item_widget.sizeHint())
        
        # Store metadata
        item.setData(Qt.ItemDataRole.UserRole, {
            'type': 'test_case',
            'name': name,
            'project': project_name
        })
        
    def get_test_case_data(self, test_case_name):
        """Helper method to get test case data from either projects or standalone."""
        # Check in projects first
        for project_data in self.projects.values():
            if test_case_name in project_data['test_cases']:
                return project_data['test_cases'][test_case_name]
        
        # Check in standalone
        return self.displayed_test_cases.get(test_case_name)
        
    def get_start_step_index(self, test_case_name):
        """Gets the selected start step index."""
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case' and item_data.get('name') == test_case_name:
                widget = self.test_case_list.itemWidget(item)
                step_combo = widget.findChild(QComboBox, "step_combo")
                if step_combo:
                    return step_combo.currentIndex()
                break
        return 0
    
    def update_status(self, test_case_name, status):
        """Updates status for a test case."""
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case' and item_data.get('name') == test_case_name:
                widget = self.test_case_list.itemWidget(item)
                status_label = widget.findChild(QLabel, "status_label")
                if status_label:
                    status_label.setText(status)
                    if status == "Passed":
                        status_label.setStyleSheet("color: green; font-weight: bold;")
                    elif status == "Failed":
                        status_label.setStyleSheet("color: red; font-weight: bold;")
                    elif status == "Stopped":
                        status_label.setStyleSheet("color: orange; font-weight: bold;")
                    else:
                        status_label.setStyleSheet("color: black;")
                break
        
        self.save_execution_data()

    def delete_single_test(self, test_case_name):
        """Deletes a test case from projects or standalone."""
        # ✅ FIXED: First find which project (if any) this test case belongs to
        belongs_to_project = None
        
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case' and item_data.get('name') == test_case_name:
                belongs_to_project = item_data.get('project')
                break
        
        reply = QMessageBox.question(
            self, "Delete Test Case",
            f"Are you sure you want to delete '{test_case_name}'?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            # ✅ FIXED: Only remove from the correct location
            if belongs_to_project:
                # Remove from specific project only
                if belongs_to_project in self.projects:
                    self.projects[belongs_to_project]['test_cases'].pop(test_case_name, None)
            else:
                # Remove from standalone only
                self.displayed_test_cases.pop(test_case_name, None)
            
            self.refresh_test_case_list()
            self.save_execution_data()
            QMessageBox.information(self, "Deleted", f"'{test_case_name}' deleted.")


    def select_all_test_cases(self, state):
        """Selects/deselects all test case checkboxes."""
        # ✅ FIXED: Block signals during bulk operation to prevent infinite loops
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            # Only process test case items, not project headers
            if item_data and item_data.get('type') == 'test_case':
                widget = self.test_case_list.itemWidget(item)
                checkbox = widget.findChild(QCheckBox)
                if checkbox:
                    checkbox.blockSignals(True)  # ✅ Block signals
                    checkbox.setCheckState(Qt.CheckState(state))
                    checkbox.blockSignals(False)  # ✅ Unblock signals


    def update_select_all_state(self):
        """Updates the Select All checkbox state."""
        # ✅ FIXED: Add guard to prevent processing during bulk operations
        if self.select_all_checkbox.signalsBlocked():
            return
        
        checked_count = 0
        total_count = 0
        
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            # Only count test case items
            if item_data and item_data.get('type') == 'test_case':
                total_count += 1
                widget = self.test_case_list.itemWidget(item)
                checkbox = widget.findChild(QCheckBox)
                if checkbox and checkbox.isChecked():
                    checked_count += 1
        
        if total_count == 0:
            self.select_all_checkbox.setCheckState(Qt.CheckState.Unchecked)
            return

        # ✅ FIXED: Block signals on select_all_checkbox to prevent recursion
        self.select_all_checkbox.blockSignals(True)
        
        if checked_count == total_count:
            self.select_all_checkbox.setCheckState(Qt.CheckState.Checked)
        elif checked_count > 0:
            self.select_all_checkbox.setCheckState(Qt.CheckState.PartiallyChecked)
        else:
            self.select_all_checkbox.setCheckState(Qt.CheckState.Unchecked)
        
        self.select_all_checkbox.blockSignals(False)  # ✅ Unblock signals

    def import_from_library(self, project_name=None):
        """Imports test cases from library - updated to support projects."""
        library_test_cases = self.main_window.test_cases
        if not library_test_cases:
            QMessageBox.information(self, "No Test Cases", "The Test Cases Library is empty.")
            return

        selection_dialog = QDialog(self)
        title = f"Import to '{project_name}'" if project_name else "Import Test Cases"
        selection_dialog.setWindowTitle(title)
        selection_dialog.setMinimumSize(400, 400)
        
        layout = QVBoxLayout(selection_dialog)
        
        search_bar = QLineEdit()
        search_bar.setPlaceholderText("Search test cases...")
        layout.addWidget(search_bar)
        
        list_widget = QListWidget()
        layout.addWidget(list_widget)

        for name in library_test_cases.keys():
            item = QListWidgetItem(name)
            item.setFlags(item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
            item.setCheckState(Qt.CheckState.Unchecked)
            list_widget.addItem(item)
        
        def filter_test_cases(query):
            query = query.strip().lower()
            for i in range(list_widget.count()):
                item = list_widget.item(i)
                item_text = item.text().lower()
                item.setHidden(query not in item_text)
        
        search_bar.textChanged.connect(filter_test_cases)
            
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(selection_dialog.accept)
        button_box.rejected.connect(selection_dialog.reject)
        layout.addWidget(button_box)

        if selection_dialog.exec() == QDialog.DialogCode.Accepted:
            selected_cases = {}
            for i in range(list_widget.count()):
                item = list_widget.item(i)
                if item.checkState() == Qt.CheckState.Checked:
                    test_case_name = item.text()
                    test_case_data = library_test_cases.get(test_case_name)
                    # ✅ FIXED: Deep copy the test case data to avoid reference issues
                    if test_case_data:
                        selected_cases[test_case_name] = copy.deepcopy(test_case_data)
            
            if selected_cases:
                if project_name:
                    # ✅ FIXED: Ensure project structure exists and is properly initialized
                    if project_name not in self.projects:
                        self.projects[project_name] = {
                            'test_cases': {},
                            'expanded': True,
                            'status_data': {}
                        }
                    
                    # ✅ CRITICAL FIX: Ensure test_cases is a dict
                    if not isinstance(self.projects[project_name].get('test_cases'), dict):
                        self.projects[project_name]['test_cases'] = {}
                    
                    # Add to project only - directly update the dict
                    self.projects[project_name]['test_cases'].update(selected_cases)
                else:
                    # Add to standalone only
                    self.displayed_test_cases.update(selected_cases)
                
                self.refresh_test_case_list()
                self.save_execution_data()
                QMessageBox.information(self, "Import Successful", 
                                      f"Successfully imported {len(selected_cases)} test case(s).")

    def import_from_device(self, project_name=None):
        """Imports test cases from device - updated to support projects."""
        options = QFileDialog.Option.ReadOnly
        file_path, _ = QFileDialog.getOpenFileName(self, "Import Test Cases from Device", "",
                                                   "JSON Files (*.json);;All Files (*)", options=options)
        
        if file_path:
            try:
                with open(file_path, 'r') as f:
                    imported_data = json.load(f)
                
                imported_cases = {}
                if isinstance(imported_data, dict) and 'name' in imported_data and 'steps' in imported_data:
                    test_case_name = imported_data['name']
                    # ✅ FIXED: Deep copy to avoid reference issues
                    imported_cases[test_case_name] = copy.deepcopy(imported_data)
                elif isinstance(imported_data, dict):
                    # ✅ FIXED: Deep copy to avoid reference issues
                    imported_cases = copy.deepcopy(imported_data)
                
                if imported_cases:
                    if project_name:
                        # ✅ FIXED: Ensure project structure exists and is properly initialized
                        if project_name not in self.projects:
                            self.projects[project_name] = {
                                'test_cases': {},
                                'expanded': True,
                                'status_data': {}
                            }
                        
                        # ✅ CRITICAL FIX: Ensure test_cases is a dict
                        if not isinstance(self.projects[project_name].get('test_cases'), dict):
                            self.projects[project_name]['test_cases'] = {}
                        
                        # Add to project only - directly update the dict
                        self.projects[project_name]['test_cases'].update(imported_cases)
                    else:
                        # Add to standalone only
                        self.displayed_test_cases.update(imported_cases)
                    
                    self.refresh_test_case_list()
                    self.save_execution_data()
                    QMessageBox.information(self, "Import Successful", 
                                          f"Successfully imported {len(imported_cases)} test case(s).")
                else:
                    QMessageBox.warning(self, "Import Failed", "No valid test cases found.")

            except Exception as e:
                QMessageBox.critical(self, "Import Error", f"Error: {e}")
                
    def refresh_test_case_list(self):
        """Refreshes the entire test case list with projects and standalone tests."""
        # ✅ NEW: Store current status before refresh
        current_statuses = {}
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case':
                test_name = item_data.get('name')
                project = item_data.get('project')
                widget = self.test_case_list.itemWidget(item)
                status_label = widget.findChild(QLabel, "status_label")
                step_combo = widget.findChild(QComboBox, "step_combo")
                
                if status_label and test_name:
                    key = (test_name, project)  # Use tuple of (name, project) as key
                    current_statuses[key] = {
                        'status': status_label.text(),
                        'status_style': status_label.styleSheet(),
                        'selected_step': step_combo.currentIndex() if step_combo else 0
                    }
        
        self.test_case_list.clear()
        
        # Add projects first
        for project_name, project_data in self.projects.items():
            self.add_project_header(project_name)
            
            # ✅ CRITICAL FIX: Always add test cases to the list, even if collapsed
            # They will be hidden by the toggle_project_expansion logic
            is_expanded = project_data.get('expanded', True)
            
            for test_name, test_data in project_data.get('test_cases', {}).items():
                self.add_list_item(test_name, project_name)
                
                # ✅ NEW: If project is collapsed, hide the item immediately after adding it
                if not is_expanded:
                    # Find the item we just added and hide it
                    for i in range(self.test_case_list.count() - 1, -1, -1):
                        item = self.test_case_list.item(i)
                        item_data = item.data(Qt.ItemDataRole.UserRole)
                        
                        if (item_data and item_data.get('type') == 'test_case' and 
                            item_data.get('name') == test_name and 
                            item_data.get('project') == project_name):
                            item.setHidden(True)
                            break
        
        # Add standalone test cases
        for test_name in self.displayed_test_cases.keys():
            self.add_list_item(test_name, None)
        
        # ✅ NEW: Restore status after refresh
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case':
                test_name = item_data.get('name')
                project = item_data.get('project')
                key = (test_name, project)
                
                if key in current_statuses:
                    widget = self.test_case_list.itemWidget(item)
                    status_label = widget.findChild(QLabel, "status_label")
                    step_combo = widget.findChild(QComboBox, "step_combo")
                    
                    if status_label:
                        saved_data = current_statuses[key]
                        status_label.setText(saved_data['status'])
                        status_label.setStyleSheet(saved_data['status_style'])
                    
                    if step_combo:
                        saved_step = saved_data.get('selected_step', 0)
                        if saved_step < step_combo.count():
                            step_combo.setCurrentIndex(saved_step)


    # In the TestExecutionDialog class, update the execute_single_test method:

    def execute_single_test(self, test_case_name):
        """Executes a single test case."""
        try:
            from datetime import datetime
            start_time = datetime.now()
            start_time_str = start_time.strftime('%Y-%m-%d %H:%M:%S')
            self.stop_execution = False
            self.set_play_stop_button_state(test_case_name, True)
            test_was_stopped = False
            
            # ✅ NEW: Initialize test_project at the beginning
            test_project = None
            for project_name, project_data in self.projects.items():
                if test_case_name in project_data['test_cases']:
                    test_project = project_name
                    break
            
            can_run, reason = self.check_prerequisites(test_case_name)
            if not can_run:
                # ✅ CHANGED: Custom message box with Override option
                msg_box = QMessageBox(self)
                msg_box.setIcon(QMessageBox.Icon.Warning)
                msg_box.setWindowTitle("Prerequisites Not Met")
                msg_box.setText(f"Cannot execute '{test_case_name}':\n\n{reason}")
                
                # Add custom buttons
                ok_button = msg_box.addButton("OK", QMessageBox.ButtonRole.RejectRole)
                override_button = msg_box.addButton("Override && Execute", QMessageBox.ButtonRole.AcceptRole)
                
                msg_box.setDefaultButton(ok_button)
                msg_box.exec()
                
                # Check which button was clicked
                if msg_box.clickedButton() == ok_button:
                    self.update_status(test_case_name, "Failed")
                    self.set_play_stop_button_state(test_case_name, False)
                    return
                # If override_button was clicked, continue execution
            
            self.update_status(test_case_name, "Running")
            QApplication.processEvents()
        
            
            # ✅ UPDATED: Use the helper method to get test case data
            test_case_data = self.get_test_case_data(test_case_name)
            if not test_case_data:
                self.update_status(test_case_name, "Failed")
                self.set_play_stop_button_state(test_case_name, False)
                QMessageBox.warning(self, "Error", f"Test case '{test_case_name}' not found.")
                return

            # Initialize COM once for the entire test case
            # NEW CODE:
            pythoncom.CoInitialize()
            autECLSession = win32com.client.Dispatch("PCOMM.autECLSession")
            connection_name = self.main_window.get_connection_name_from_title(self.main_window.pcomm_window_title)
            autECLSession.SetConnectionByName(connection_name)
            autECLPS = autECLSession.autECLPS

            # Track validation results and captured data
            validation_failures = []
            text_captures = []
            docx_screenshots = []

            # Get screen dimensions (default to 24x80)
            screen_rows = 24
            screen_cols = 80

            # Get the selected start step (0-based index)
            start_step_index = self.get_start_step_index(test_case_name)
            
            # Process each step sequentially from the selected start step
            # Process each step sequentially from the selected start step
            all_steps = test_case_data.get("steps", [])

            # ✅ CHANGED: Use while loop instead of for loop to allow dynamic step list updates
            current_step_index = start_step_index
            while current_step_index < len(all_steps):
                step_index = current_step_index + 1  # Display as 1-based
                step = all_steps[current_step_index]
                
                # ✅ Check if stop was requested
                if self.stop_execution:
                    print(f"Execution stopped by user at step {step_index}")
                    break
                
                QApplication.processEvents()  # ✅ Allow UI to update
                
                step_type = step.get("type")

                if step_type == "break":
                    message = step.get("message", "")
                    print(f"Step {step_index}: Break point reached")
                    
                    # Show break dialog
                    break_dialog = BreakExecutionDialog(message, self)
                    
                    # ✅ CHANGED: Use show() instead of exec() to make it non-modal initially
                    break_dialog.show()
                    
                    # Wait for user action in an event loop
                    while break_dialog.result_action is None:
                        QApplication.processEvents()
                        time.sleep(0.1)
                    
                    action = break_dialog.result_action
                    
                    if action == BreakExecutionDialog.STOP:
                        print("User chose to stop execution at break point")
                        self.stop_execution = True
                        break_dialog.close()
                        break
                    
                    elif action == BreakExecutionDialog.EDIT:
                        print("User chose to edit test case at break point")
                        
                        # âœ… NEW: Reset the action so the dialog can be used again
                        break_dialog.result_action = None
                        
                        # Get current test case data
                        existing_steps = test_case_data.get('steps', [])
                        test_case_description = test_case_data.get('description', '')
                        test_case_assumptions = test_case_data.get('assumptions', '')
                        
                        # âœ… CHANGED: Open edit dialog as modal (it will appear in front)
                        edit_dialog = EditTestCaseDialog(
                            existing_steps, 
                            self.main_window.modules, 
                            self.main_window,
                            test_case_name, 
                            test_case_description, 
                            test_case_assumptions
                        )
                        
                        # Set the break dialog as parent so it stays behind
                        edit_dialog.setParent(break_dialog, edit_dialog.windowFlags())
                        edit_dialog.exec()
                        
                        # Get updated data
                        updated_steps = edit_dialog.get_updated_steps()
                        updated_description = edit_dialog.get_test_case_description()
                        updated_assumptions = edit_dialog.get_test_case_assumptions()
                        updated_prerequisites = edit_dialog.get_prerequisites()
                        
                        # Update in-memory test case data
                        test_case_data['steps'] = updated_steps
                        test_case_data['description'] = updated_description
                        test_case_data['assumptions'] = updated_assumptions
                        test_case_data['prerequisites'] = updated_prerequisites
                        
                        # Save to file
                        self.main_window.save_test_cases_to_file()
                        
                        # âœ… CRITICAL FIX: Refresh the all_steps list with updated data
                        # ✅ CRITICAL FIX: Refresh the all_steps list with updated data
                        all_steps = test_case_data.get("steps", [])
                        
                        print(f"Test case '{test_case_name}' updated during execution")
                        print(f"Total steps after update: {len(all_steps)}")
                        
                        # ✅ NEW: Show message in the break dialog instead of separate popup
                        QMessageBox.information(
                            break_dialog,  # Use break_dialog as parent
                            "Test Case Updated",
                            f"Test case '{test_case_name}' has been updated.\n"
                            f"Total steps: {len(all_steps)}\n"
                            f"Press 'Resume Execution' to continue from Step {step_index}."  # ✅ CHANGED: Removed +1
                        )
                        
                        # âœ… NEW: Bring break dialog back to front
                        break_dialog.raise_()
                        break_dialog.activateWindow()
                        
                        # âœ… NEW: Continue waiting for the next action (Resume or Stop)
                        while break_dialog.result_action is None:
                            QApplication.processEvents()
                            time.sleep(0.1)
                        
                        action = break_dialog.result_action
                        
                        if action == BreakExecutionDialog.STOP:
                            self.stop_execution = True
                            break_dialog.close()
                            break
                        elif action == BreakExecutionDialog.RESUME:
                            break_dialog.close()
                            # Continue to next step
                    
                    elif action == BreakExecutionDialog.RESUME:
                        print("User chose to resume execution")
                        break_dialog.close()
                        # Continue to next step
                    
                    
                
                if step_type == "module_import":
                    import os
                    
                    # Process module fields FIRST
                    module_name = step.get("module_name")
                    
                    # ✅ NEW: Capture "Before" screenshot if Screen Flow is enabled
                    capture_screen_flow = self.main_window.document_config.get('capture_screen_flow', False)
                    if capture_screen_flow:
                        try:
                            screen_flow_dir = os.path.join(self.main_window.default_results_location, 'Screen Flows')
                            os.makedirs(screen_flow_dir, exist_ok=True)
                            
                            before_filename = f"{test_case_name}_Step {step_index}_Before.jpg"
                            before_path = os.path.join(screen_flow_dir, before_filename)
                            
                            if self.main_window.capture_pcomm_screen_as_jpeg(before_path):
                                print(f"Step {step_index}: Captured 'Before' screen flow: {before_filename}")
                        except Exception as e:
                            print(f"Error capturing before screenshot: {e}")
                    
                    # Process module fields for Input
                    for field in step.get("fields", []):
                        if self.stop_execution:
                            break
                            
                        action_type = field.get("action_type", "Input")
                        value = str(field.get("value", "")).strip()
                        
                        if action_type == "Input":
                            if not value:
                                continue
                            
                            if module_name in self.main_window.modules:
                                module_data = self.main_window.modules[module_name]
                                labels = module_data.get("labels", [])
                                
                                field_name = field.get("field_name")
                                for label in labels:
                                    label_name = label.get('name') or label.get('label') or label.get('text', '')
                                    if label_name == field_name:
                                        row = int(label.get("row", 1))
                                        col = int(label.get("column", 1))
                                        
                                        # Substitute variables in the value before sending
                                        substituted_value = self.substitute_execution_variables(value, test_case_name)
                                        
                                        autECLPS.SetCursorPos(row, col)
                                        autECLPS.SendKeys(substituted_value)
                                        time.sleep(0.1)
                                        break
                    
                    # ✅ NEW: Capture "After" screenshot with smart wait if Screen Flow is enabled
                    if capture_screen_flow:
                        try:
                            screen_flow_dir = os.path.join(self.main_window.default_results_location, 'Screen Flows')
                            after_filename = f"{test_case_name}_Step {step_index}_After.jpg"
                            after_path = os.path.join(screen_flow_dir, after_filename)
                            
                            # Small delay to ensure screen is updated
                            time.sleep(0.5)
                            
                            if self.main_window.capture_pcomm_screen_as_jpeg(after_path):
                                print(f"Step {step_index}: Captured 'After' screen flow: {after_filename}")
                        except Exception as e:
                            print(f"Error capturing after screenshot: {e}")
                    
                    # Process validation actions for module fields
                    for field in step.get("fields", []):
                        # Check for stop during field processing
                        QApplication.processEvents()
                        if self.stop_execution:
                            test_was_stopped = True
                            break
                        
                        action_type = field.get("action_type", "Input")
                        value = str(field.get("value", "")).strip()
                        
                        if action_type == "Validate":
                            if not value:
                                continue
                            
                            if module_name in self.main_window.modules:
                                module_data = self.main_window.modules[module_name]
                                labels = module_data.get("labels", [])
                                
                                field_name = field.get("field_name")
                                for label in labels:
                                    label_name = label.get('name') or label.get('label') or label.get('text', '')
                                    if label_name == field_name:
                                        row = int(label.get("row", 1))
                                        col = int(label.get("column", 1))
                                        length = int(label.get("length", len(value)))
                                                                          
                                        actual_value = autECLPS.GetText(row, col, length)  # ← Remove .strip() here
                                        
                                        # ✅ NEW: Use validation helper function
                                        validation_passed = self.validate_field_value(actual_value, value)
                                        
                                        if not validation_passed:
                                            # ✅ ENHANCED: Better error message for {blank} validation
                                            if value.lower() == '{blank}':
                                                validation_failures.append({
                                                    "step": step_index,
                                                    "field": field_name,
                                                    "expected": '<blank>',
                                                    "actual": f"'{actual_value.strip()}'" if actual_value.strip() else '<blank>'
                                                })
                                            else:
                                                validation_failures.append({
                                                    "step": step_index,
                                                    "field": field_name,
                                                    "expected": value,
                                                    "actual": actual_value.strip()
                                                })
                                            # Stop execution immediately on validation failure
                                            print(f"❌ Validation failed at Step {step_index} - Field '{field_name}': Expected '{value}', Got '{actual_value}'")
                                            break
                                        
                                        time.sleep(0.1)
                                        break
                            
                            # Check if validation failed and stop test execution
                            if validation_failures:
                                print(f"🛑 Stopping test execution due to validation failure at Step {step_index}")
                                break  # Break out of fields loop

                    # Check if validation failed and stop processing steps
                    if validation_failures:
                        break  # Break out of steps loop
                
                elif step_type == "special_key":
                    key_value = step.get("key_value", "")
                    
                    key_mapping = {
                        "Enter Key": "[enter]",
                        "Clear Key": "[clear]",
                        "End Key": "[EraseEof]",
                        "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                        "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                        "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                        "F13": "[pf13]", "F14": "[pf14]", "F15": "[pf15]", "F16": "[pf16]",
                        "F17": "[pf17]", "F18": "[pf18]", "F19": "[pf19]", "F20": "[pf20]",
                        "F21": "[pf21]", "F22": "[pf22]", "F23": "[pf23]", "F24": "[pf24]",
                    }
                    
                    # Check if it's a combo key (contains + sign)
                    if "+" in key_value:
                        pcomm_key = self.main_window.convert_combo_key_to_pcomm(key_value)
                    else:
                        pcomm_key = key_mapping.get(key_value, key_value)
                    
                    # ✅ FIXED: Smart wait for special keys
                    # Keys that typically cause screen changes
                    action_keys = ["Enter Key"] + [f"F{i}" for i in range(1, 25)]
                    
                    if key_value in action_keys or "+" in key_value:
                        # Capture screen before action
                        before_screen = wait_for_pcomm_ready_smart(autECLPS, key_value)
                        
                        # Send the key
                        autECLPS.SendKeys(pcomm_key)
                        print(f"Step {step_index}: Sent {key_value}")
                        
                        # ✅ FIXED: Pass key_value as a string, not variable reference
                        success, elapsed = complete_pcomm_wait(
                            autECLPS, 
                            before_screen, 
                            action_description=f"Special Key: {key_value}",  # ✅ Fixed here
                            timeout=30
                        )
                        
                        if not success:
                            QMessageBox.warning(self, "Timeout Warning", 
                                f"Step {step_index}: {key_value} did not complete within 30 seconds.\n\n"
                                "The test will continue, but results may be unreliable.")
                    else:
                        # Non-action keys (just send with small delay)
                        autECLPS.SendKeys(pcomm_key)
                        time.sleep(0.5)
                
                elif step_type == "capture_screen_text":
                    if self.stop_execution:
                        break
                        
                    print(f"Step {step_index}: Capturing screen text...")
                    screen_size = screen_rows * screen_cols
                    full_screen_text = autECLPS.GetText(1, screen_size)
                    
                    captured_text_lines = []
                    for row_num in range(screen_rows):
                        start_index_text = row_num * screen_cols
                        line_text = full_screen_text[start_index_text:start_index_text + screen_cols]
                        captured_text_lines.append(line_text)
                    
                    from datetime import datetime
                    header = f"--- Step {step_index}: Screen Text Capture at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ---"
                    full_capture_string = header + "\n" + "\n".join(captured_text_lines)
                    text_captures.append(full_capture_string)
                    print(f"Step {step_index}: Screen text captured successfully")
                
                elif step_type == "capture_screenshot":
                    if self.stop_execution:
                        break
                    
                    # ✅ NEW: Check if documentation generation is enabled
                    if not self.main_window.document_config.get('generate_documentation', True):
                        print(f"Step {step_index}: Screenshot capture skipped (documentation disabled)")
                        continue
                        
                    print(f"Step {step_index}: Capturing screenshot for DOCX...")
                    screen_size = screen_rows * screen_cols
                    full_screen_text = autECLPS.GetText(1, screen_size)
                    
                    # ✅ NEW: Get highlight information for this screenshot
                    highlight_info = {}
                    reference_module = step.get('reference_module')
                    
                    if reference_module and reference_module in self.main_window.modules:
                        module_data = self.main_window.modules[reference_module]
                        labels = module_data.get('labels', [])
                        
                        # Get highlight flags from step fields
                        for field in step.get('fields', []):
                            if field.get('highlight', False):
                                field_name = field.get('field_name')
                                # Find the corresponding label
                                for label in labels:
                                    label_name = label.get('name') or label.get('label') or label.get('text', '')
                                    if label_name == field_name:
                                        highlight_info[field_name] = {
                                            'row': label.get('row', 1),
                                            'column': label.get('column', 1),
                                            'length': label.get('length', 10)
                                        }
                                        break
                    
                    from datetime import datetime
                    docx_screenshots.append({
                        'step': step_index,
                        'screen_text': full_screen_text,
                        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'highlight_info': highlight_info  # ✅ NEW: Add highlight info
                    })
                    print(f"Step {step_index}: Screenshot captured for DOCX")
                    
                elif step_type == "random_input":
                    if self.stop_execution:
                        break
                        
                    row = int(step.get("row", 1))
                    col = int(step.get("column", 1))
                    value = str(step.get("value", "")).strip()
                    is_special_key = step.get("is_special_key", False)
                    
                    if value:
                        autECLPS.SetCursorPos(row, col)
                        
                        if is_special_key:
                            key_mapping = {
                                "Enter Key": "[enter]", "Clear Key": "[clear]", "End Key": "[EraseEof]",
                                "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                                "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                                "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                            }
                            pcomm_key = key_mapping.get(value, value)
                            autECLPS.SendKeys(pcomm_key)
                            print(f"Step {step_index}: Sent special key '{value}' to position ({row}, {col})")
                            time.sleep(1.0)
                        else:
                            # ✅ NEW: Substitute variables before sending
                            substituted_value = self.substitute_execution_variables(value, test_case_name)
                            autECLPS.SendKeys(substituted_value)
                            print(f"Step {step_index}: Sent '{substituted_value}' (from '{value}') to position ({row}, {col})")
                            
                elif step_type == "wait":
                    if self.stop_execution:
                        break
                        
                    seconds = float(step.get("seconds", 0))
                    if seconds > 0:
                        print(f"Step {step_index}: Waiting for {seconds} second(s)...")
                        time.sleep(seconds)
                        print(f"Step {step_index}: Wait completed")
                
                
                if 'utility_steps' in step:
                    for sub_index, utility_step in enumerate(step['utility_steps'], 1):
                        if self.stop_execution:
                            print(f"Execution stopped at utility step {step_index}.{sub_index}")
                            break
                        
                        QApplication.processEvents()
                        
                        utility_type = utility_step.get("type")
                        print(f"Executing utility step {step_index}.{sub_index}: {utility_step.get('name', 'Unknown')}")
                        
                        if utility_type == "special_key":
                            key_value = utility_step.get("key_value", "")
                            
                            key_mapping = {
                                "Enter Key": "[enter]",
                                "Clear Key": "[clear]",
                                "End Key": "[EraseEof]",
                                "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                                "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                                "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                                "F13": "[pf13]", "F14": "[pf14]", "F15": "[pf15]", "F16": "[pf16]",
                                "F17": "[pf17]", "F18": "[pf18]", "F19": "[pf19]", "F20": "[pf20]",
                                "F21": "[pf21]", "F22": "[pf22]", "F23": "[pf23]", "F24": "[pf24]",
                            }
                            
                            if "+" in key_value:
                                pcomm_key = self.main_window.convert_combo_key_to_pcomm(key_value)
                            else:
                                pcomm_key = key_mapping.get(key_value, key_value)
                            
                            # ✅ FIXED: Smart wait for utility special keys
                            action_keys = ["Enter Key"] + [f"F{i}" for i in range(1, 25)]
                            
                            if key_value in action_keys or "+" in key_value:
                                before_screen = wait_for_pcomm_ready_smart(autECLPS, f"Utility {key_value}")
                                autECLPS.SendKeys(pcomm_key)
                                print(f"Step {step_index}.{sub_index}: Sent utility special key '{key_value}'")
                                
                                # ✅ FIXED: Pass action_description as a string
                                success, elapsed = complete_pcomm_wait(
                                    autECLPS, 
                                    before_screen, 
                                    action_description=f"Utility Special Key: {key_value}",  # ✅ Fixed here
                                    timeout=30
                                )
                                
                                if not success:
                                    print(f"⚠️ Utility {key_value} timeout at step {step_index}.{sub_index}")
                            else:
                                autECLPS.SendKeys(pcomm_key)
                                time.sleep(0.5)
                        
                        elif utility_type == "wait":
                            seconds = float(utility_step.get("seconds", 0))
                            if seconds > 0:
                                print(f"Step {step_index}.{sub_index}: Utility wait for {seconds} second(s)...")
                                time.sleep(seconds)
                                print(f"Step {step_index}.{sub_index}: Utility wait completed")
                        
                        elif utility_type == "capture_screenshot":
                            if not self.main_window.document_config.get('generate_documentation', True):
                                print(f"Step {step_index}.{sub_index}: Utility screenshot skipped (documentation disabled)")
                                continue
                            
                            print(f"Step {step_index}.{sub_index}: Capturing utility screenshot for DOCX...")
                            screen_size = screen_rows * screen_cols
                            full_screen_text = autECLPS.GetText(1, screen_size)
                            
                            # ✅ FIXED: Get highlight information for utility screenshot
                            highlight_info = {}
                            reference_module = utility_step.get('reference_module')
                            
                            if reference_module and reference_module in self.main_window.modules:
                                module_data = self.main_window.modules[reference_module]
                                labels = module_data.get('labels', [])
                                
                                # Get highlight flags from utility step fields
                                for field in utility_step.get('fields', []):
                                    if field.get('highlight', False):
                                        field_name = field.get('field_name')
                                        # Find the corresponding label
                                        for label in labels:
                                            label_name = label.get('name') or label.get('label') or label.get('text', '')
                                            if label_name == field_name:
                                                highlight_info[field_name] = {
                                                    'row': label.get('row', 1),
                                                    'column': label.get('column', 1),
                                                    'length': label.get('length', 10)
                                                }
                                                break
                            
                            from datetime import datetime
                            docx_screenshots.append({
                                'step': f"{step_index}.{sub_index}",
                                'screen_text': full_screen_text,
                                'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                                'highlight_info': highlight_info  # ✅ FIXED: Add highlight info
                            })
                            print(f"Step {step_index}.{sub_index}: Utility screenshot captured for DOCX with {len(highlight_info)} highlighted field(s)")
                        
                        elif utility_type == "capture_screen_text":
                            print(f"Step {step_index}.{sub_index}: Capturing utility screen text...")
                            screen_size = screen_rows * screen_cols
                            full_screen_text = autECLPS.GetText(1, screen_size)
                            
                            captured_text_lines = []
                            for row_num in range(screen_rows):
                                start_index_text = row_num * screen_cols
                                line_text = full_screen_text[start_index_text:start_index_text + screen_cols]
                                captured_text_lines.append(line_text)
                            
                            from datetime import datetime
                            header = f"--- Step {step_index}.{sub_index}: Utility Text Capture at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ---"
                            full_capture_string = header + "\n" + "\n".join(captured_text_lines)
                            text_captures.append(full_capture_string)
                            print(f"Step {step_index}.{sub_index}: Utility screen text captured successfully")

                        elif utility_type == "random_input":
                            row = int(utility_step.get('row', 1))
                            col = int(utility_step.get('column', 1))
                            value = str(utility_step.get('value', '')).strip()
                            is_special_key = utility_step.get('is_special_key', False)
                            
                            if value:
                                autECLPS.SetCursorPos(row, col)
                                
                                if is_special_key:
                                    key_mapping = {
                                        "Enter Key": "[enter]", "Clear Key": "[clear]", "End Key": "[EraseEof]",
                                        "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                                        "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                                        "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                                    }
                                    pcomm_key = key_mapping.get(value, value)
                                    autECLPS.SendKeys(pcomm_key)
                                    print(f"Step {step_index}.{sub_index}: Sent special key '{value}' to position ({row}, {col})")
                                    time.sleep(1.0)
                                else:
                                    substituted_value = self.substitute_execution_variables(value, test_case_name)
                                    autECLPS.SendKeys(substituted_value)
                                    print(f"Step {step_index}.{sub_index}: Sent '{substituted_value}' to position ({row}, {col})")
                            
                        elif utility_type == "module_import":
                            # ✅ FIXED: Handle module import utility step (for both Input and Validation)
                            module_name = utility_step.get('module_name')
                            
                            if module_name in self.main_window.modules:
                                module_data = self.main_window.modules[module_name]
                                labels = module_data.get("labels", [])
                                
                                # ✅ FIXED: Process Input fields FIRST
                                for field in utility_step.get('fields', []):
                                    if self.stop_execution:
                                        test_was_stopped = True
                                        break
                                    
                                    action_type = field.get('action_type', 'Input')
                                    value = str(field.get('value', '')).strip()
                                    
                                    if action_type == 'Input' and value:
                                        field_name = field.get('field_name')
                                        
                                        # Find the label for this field
                                        for label in labels:
                                            label_name = label.get('name') or label.get('label') or label.get('text', '')
                                            if label_name == field_name:
                                                row = int(label.get('row', 1))
                                                col = int(label.get('column', 1))
                                                
                                                # Substitute variables before sending
                                                substituted_value = self.substitute_execution_variables(value, test_case_name)
                                                
                                                autECLPS.SetCursorPos(row, col)
                                                autECLPS.SendKeys(substituted_value)
                                                print(f"Step {step_index}.{sub_index}: Sent utility input '{substituted_value}' to {field_name}")
                                                time.sleep(0.1)
                                                break
                                
                                # ✅ FIXED: Process Validation fields AFTER inputs
                                for field in utility_step.get('fields', []):
                                    if self.stop_execution:
                                        test_was_stopped = True
                                        break
                                    
                                    action_type = field.get('action_type', 'Validate')
                                    value = str(field.get('value', '')).strip()
                                    
                                    if action_type == 'Validate' and value:
                                        field_name = field.get('field_name')
                                        
                                        # Find the label for this field
                                        for label in labels:
                                            label_name = label.get('name') or label.get('label') or label.get('text', '')
                                            if label_name == field_name:
                                                row = int(label.get('row', 1))
                                                col = int(label.get('column', 1))
                                                length = int(label.get('length', len(value)))

                                                
                                                # Read actual value from screen
                                                actual_value = autECLPS.GetText(row, col, length)  # ← Remove .strip() here
                                                
                                                # ✅ NEW: Use validation helper function
                                                validation_passed = self.validate_field_value(actual_value, value)
                                                
                                                if not validation_passed:
                                                    # ✅ ENHANCED: Better error message for {blank} validation
                                                    if value.lower() == '{blank}':
                                                        validation_failures.append({
                                                            "step": f"{step_index}.{sub_index}",
                                                            "field": field_name,
                                                            "expected": '<blank>',
                                                            "actual": f"'{actual_value.strip()}'" if actual_value.strip() else '<blank>'
                                                        })
                                                    else:
                                                        validation_failures.append({
                                                            "step": f"{step_index}.{sub_index}",
                                                            "field": field_name,
                                                            "expected": value,
                                                            "actual": actual_value.strip()
                                                        })
                                                    # ✅ Stop execution immediately on validation failure
                                                    print(f"❌ Utility validation failed at Step {step_index}.{sub_index} - Field '{field_name}': Expected '{value}', Got '{actual_value}'")
                                                    break
                                                
                                                print(f"Step {step_index}.{sub_index}: Validated {field_name} - Expected: '{value}', Actual: '{actual_value}'")
                                                time.sleep(0.1)
                                                break
                                            
                                        # ✅ Check if validation failed and stop utility steps
                                        if validation_failures:
                                            print(f"🛑 Stopping utility steps due to validation failure at Step {step_index}.{sub_index}")
                                            break
                                    
                                    # ✅ Check if validation failed and stop processing this step's utilities
                                    if validation_failures:
                                        break
                            else:
                                print(f"Warning: Module '{module_name}' not found for utility step")                                                
                      
                      
                        
                        time.sleep(0.1)

                # âœ… NEW: Check if validation failed in utility steps and stop main loop
                if validation_failures:
                    print(f"ðŸ›' Stopping test execution due to utility validation failure")
                    break  # Break out of main steps while loop
                
                time.sleep(0.1)                
                current_step_index += 1


            # Clean up COM
            pythoncom.CoUninitialize()

            # âœ… Check if execution was stopped
            if self.stop_execution:
                self.update_status(test_case_name, "Stopped")
                self.set_play_stop_button_state(test_case_name, False)
                
                # âœ… Create execution summary even for stopped execution
                from datetime import datetime
                execution_timestamp = datetime.now().strftime('%Y-%m-%d %H-%M-%S')
               
                execution_results = [{
                    'name': test_case_name,
                    'status': 'Stopped',
                    'error': 'Execution stopped by user',
                    'project': test_project
                }]
                docx_summary_path = self.create_execution_summary_docx(execution_results, execution_timestamp)
                
                if docx_summary_path:
                    QMessageBox.information(self, "Execution Stopped", 
                                          f"Test case '{test_case_name}' execution was stopped.\n\n"
                                          f"Summary saved to:\n{docx_summary_path}")
                else:
                    QMessageBox.information(self, "Execution Stopped", 
                                          f"Test case '{test_case_name}' execution was stopped.")
                return

            # Save captured text to file if any captures were made
            if text_captures:
                import os
                project_name = getattr(self.main_window, 'current_project_id', None)
                if project_name and project_name in self.main_window.projects:
                    project_name = self.main_window.projects[project_name]['name']
                
                output_dir = os.path.join(self.main_window.default_results_location, 'Results', project_name if project_name else 'Master')
                os.makedirs(output_dir, exist_ok=True)
                
                filename = os.path.join(output_dir, f"{test_case_name}.txt")
                
                with open(filename, "w", encoding="utf-8") as f:
                    f.write("\n\n\n".join(text_captures))
                
                print(f"All screen text for '{test_case_name}' saved to '{filename}'.")
            
            # Create single DOCX with all screenshots
            docx_path = None
            if docx_screenshots:
                try:
                    docx_path = self.main_window.create_test_case_docx(test_case_name, docx_screenshots)
                    print(f"DOCX document created: '{docx_path}' with {len(docx_screenshots)} screenshot(s)")
                except Exception as e:
                    print(f"Error creating DOCX: {e}")
                    validation_failures.append({
                        "step": "DOCX Generation",
                        "field": "Document Creation",
                        "expected": "Success",
                        "actual": f"Error: {str(e)}"
                    })

            # âœ… Create execution summary for single test case
            from datetime import datetime
            execution_timestamp = datetime.now().strftime('%Y-%m-%d %H-%M-%S')
            
            # Check if there were any validation failures
            if validation_failures:
                # âœ… Calculate execution time
                end_time = datetime.now()
                end_time_str = end_time.strftime('%Y-%m-%d %H:%M:%S')
                duration = end_time - start_time
                duration_str = str(duration).split('.')[0]
                
                # âœ… Store execution time
                self.execution_times[test_case_name] = {
                    'start_time': start_time_str,
                    'end_time': end_time_str,
                    'duration': duration_str
                }
                
                # âœ… Update UI
                for i in range(self.test_case_list.count()):
                    item = self.test_case_list.item(i)
                    item_data = item.data(Qt.ItemDataRole.UserRole)
                    
                    if item_data and item_data.get('type') == 'test_case' and item_data.get('name') == test_case_name:
                        widget = self.test_case_list.itemWidget(item)
                        time_label = widget.findChild(QLabel, "execution_time_label")
                        if time_label:
                            time_label.setText(f"{duration_str}")
                        break
                
                self.update_status(test_case_name, "Failed")
                
                # Build detailed failure message
                failure_msg = f"Test case '{test_case_name}' failed validation:\n\n"
                for failure in validation_failures:
                    failure_msg += (f"Step {failure['step']} - Field '{failure['field']}':\n"
                                  f"  Expected: '{failure['expected']}'\n"
                                  f"  Actual: '{failure['actual']}'\n\n")

                
                # âœ… Create execution summary
                execution_results = [{
                    'name': test_case_name,
                    'status': 'Failed',
                    'validation_failures': validation_failures,
                    'project': test_project,
                    'start_time': start_time_str,
                    'end_time': end_time_str,
                    'duration': duration_str
                }]
                docx_summary_path = self.create_execution_summary_docx(execution_results, execution_timestamp)
                
                if docx_summary_path:
                    failure_msg += f"\n\nExecution summary saved to:\n{docx_summary_path}"
                
                QMessageBox.critical(self, "Validation Failed", failure_msg)
            else:
                # âœ… Calculate execution time
                end_time = datetime.now()
                end_time_str = end_time.strftime('%Y-%m-%d %H:%M:%S')
                duration = end_time - start_time
                duration_str = str(duration).split('.')[0]  # Remove microseconds
                
                # âœ… Store execution time
                self.execution_times[test_case_name] = {
                    'start_time': start_time_str,
                    'end_time': end_time_str,
                    'duration': duration_str
                }
                
                # âœ… Update the execution time label in the UI
                for i in range(self.test_case_list.count()):
                    item = self.test_case_list.item(i)
                    item_data = item.data(Qt.ItemDataRole.UserRole)
                    
                    if item_data and item_data.get('type') == 'test_case' and item_data.get('name') == test_case_name:
                        widget = self.test_case_list.itemWidget(item)
                        time_label = widget.findChild(QLabel, "execution_time_label")
                        if time_label:
                            time_label.setText(f"{duration_str}")
                        break
                
                # Update status to Passed
                # Update status to Passed
                self.update_status(test_case_name, "Passed")
                success_msg = f"Test case '{test_case_name}' executed successfully!\nAll validations passed."

                # Add info about captured files
                file_info = []
                if text_captures:
                    file_info.append(f"{len(text_captures)} text screenshot(s)")
                if docx_path:
                    file_info.append(f"1 DOCX document with {len(docx_screenshots)} screenshot(s)")
                elif not self.main_window.document_config.get('generate_documentation', True):
                    file_info.append("Documentation generation disabled")

                if file_info:
                    success_msg += f"\n\nCaptured: {', '.join(file_info)}"
                
               
                # âœ… Create execution summary for passed test
                execution_results = [{
                    'name': test_case_name,
                    'status': 'Passed',
                    'project': test_project,
                    'start_time': start_time_str,
                    'end_time': end_time_str,
                    'duration': duration_str
                }]
                docx_summary_path = self.create_execution_summary_docx(execution_results, execution_timestamp)
                
                if docx_summary_path:
                    success_msg += f"\n\nExecution summary saved to:\n{docx_summary_path}"
                
                QMessageBox.information(self, "Success", success_msg)

        except Exception as e:
            try:
                pythoncom.CoUninitialize()
            except:
                pass
            self.update_status(test_case_name, "Failed")
            
            # âœ… Create execution summary even for errors
            from datetime import datetime
            execution_timestamp = datetime.now().strftime('%Y-%m-%d %H-%M-%S')
            execution_results = [{
                'name': test_case_name,
                'status': 'Failed',
                'error': str(e),
                'project': test_project
            }]
            docx_summary_path = self.create_execution_summary_docx(execution_results, execution_timestamp)
            
            error_msg = f"Failed to execute test case '{test_case_name}':\n\n{str(e)}"
            if docx_summary_path:
                error_msg += f"\n\nExecution summary saved to:\n{docx_summary_path}"
            
            QMessageBox.critical(self, "Error", error_msg)
        finally:
            # âœ… Always change button back to play icon
            self.set_play_stop_button_state(test_case_name, False)


    def execute_selected_tests(self):
        """Executes all selected test cases sequentially."""
        selected_tests = []
        selected_names = []
        
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            # Only process test case items
            if item_data and item_data.get('type') == 'test_case':
                widget = self.test_case_list.itemWidget(item)
                checkbox = widget.findChild(QCheckBox)
                
                if checkbox and checkbox.isChecked():
                    test_name = item_data['name']
                    selected_tests.append(test_name)
                    selected_names.append(test_name)

        if not selected_tests:
            QMessageBox.warning(self, "No Selection", "Please select at least one test case to execute.")
            return

        # Confirm execution
        reply = QMessageBox.question(
            self, 
            "Confirm Execution",
            f"Execute {len(selected_tests)} test case(s)?\n\n" + "\n".join(selected_names),
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply != QMessageBox.StandardButton.Yes:
            return

        self.stop_execution = False
        original_text = self.execute_button.text()
        original_icon = self.execute_button.icon()
        
        # Change button to Stop
        self.execute_button.setText("Stop Execution")
        self.execute_button.setIcon(self.main_window.style().standardIcon(QStyle.StandardPixmap.SP_MediaStop))
        
        # Disable the button temporarily to prevent multiple clicks
        #self.execute_button.setEnabled(False)
        #QApplication.processEvents()
        #self.execute_button.setEnabled(True)
        
        # Disconnect previous connection and connect to stop function
        try:
            self.execute_button.clicked.disconnect()
        except:
            pass
        self.execute_button.clicked.connect(self.request_stop_execution)
        
        # ✅ IMPORTANT: Process events so the button change is visible
        QApplication.processEvents()

        # Execute each test case sequentially
        passed_count = 0
        failed_count = 0
        stopped_count = 0
        results_summary = []
        execution_results = []
        from datetime import datetime
        execution_timestamp = datetime.now().strftime('%Y-%m-%d %H-%M-%S')
        
        for test_case_name in selected_tests:
            # ✅ IMPORTANT: Allow UI to process events (including stop button clicks)
            QApplication.processEvents()
            
            # Initialize test_project variable
            test_project = None
            for project_name, project_data in self.projects.items():
                if test_case_name in project_data['test_cases']:
                    test_project = project_name
                    break
            
            # ✅ NEW: Initialize test_was_stopped flag for this test case
            test_was_stopped = False
            
            # ✅ Check if stop was requested BEFORE starting this test case
            if self.stop_execution:
                print(f"\n⚠️ Execution stopped by user before test case '{test_case_name}'")
                # Mark remaining tests as stopped
                for remaining_test in selected_tests[selected_tests.index(test_case_name):]:
                    self.update_status(remaining_test, "Stopped")
                    stopped_count += 1
                    results_summary.append(f"⏸️ {remaining_test}: Stopped by user")
                   
                    execution_results.append({
                        'name': remaining_test,
                        'status': 'Stopped',
                        'project': test_project,
                        'error': 'Execution stopped by user before test started'
                    })
                break  # Exit the for loop completely
            
            print(f"\n{'='*60}")
            # âœ… Track test case execution time
            test_start_time = datetime.now()
            test_start_time_str = test_start_time.strftime('%Y-%m-%d %H:%M:%S')
            print(f"Executing Test Case: {test_case_name}")
            print(f"{'='*60}\n")
            
            # ✅ Check prerequisites first
            # In the execute_selected_tests method, find this section:

            # ✅ Check prerequisites first
            can_run, reason = self.check_prerequisites(test_case_name)
            if not can_run:
                # ✅ CHANGED: Custom message box with Override option
                msg_box = QMessageBox(self)
                msg_box.setIcon(QMessageBox.Icon.Warning)
                msg_box.setWindowTitle("Prerequisites Not Met")
                msg_box.setText(f"Cannot execute '{test_case_name}':\n\n{reason}\n\nWhat would you like to do?")
                
                # Add custom buttons
                skip_button = msg_box.addButton("Skip This Test", QMessageBox.ButtonRole.RejectRole)
                override_button = msg_box.addButton("Override && Execute", QMessageBox.ButtonRole.AcceptRole)
                stop_button = msg_box.addButton("Stop All", QMessageBox.ButtonRole.DestructiveRole)
                
                msg_box.setDefaultButton(skip_button)
                msg_box.exec()
                
                # Check which button was clicked
                if msg_box.clickedButton() == stop_button:
                    # Stop all remaining tests
                    self.stop_execution = True
                    for remaining_test in selected_tests[selected_tests.index(test_case_name):]:
                        self.update_status(remaining_test, "Stopped")
                        stopped_count += 1
                        results_summary.append(f"⏸️ {remaining_test}: Stopped by user")
                        execution_results.append({
                            'name': remaining_test,
                            'status': 'Stopped',
                            'project': test_project,
                            'error': 'Execution stopped by user'
                        })
                    break
                elif msg_box.clickedButton() == skip_button:
                    # Skip this test and mark as failed
                    self.update_status(test_case_name, "Failed")
                    failed_count += 1
                    results_summary.append(f"❌ {test_case_name}: Prerequisites not met - {reason}")
                    execution_results.append({
                        'name': test_case_name,
                        'status': 'Failed',
                        'project': test_project,
                        'error': f"Prerequisites not met: {reason}"
                    })
                    print(f"❌ Test '{test_case_name}' SKIPPED: {reason}")
                    continue
                # If override_button was clicked, continue execution normally
            
            # Update status to "Running"
            self.update_status(test_case_name, "Running")
            QApplication.processEvents()  # ✅ Update UI
            
            # Get test case data
            # ✅ FIXED: Get test case data from either projects or standalone
            test_case_data = self.get_test_case_data(test_case_name)
            if not test_case_data:
                self.update_status(test_case_name, "Failed")
                failed_count += 1
                results_summary.append(f"❌ {test_case_name}: Test case not found")
                
                execution_results.append({
                    'name': test_case_name,
                    'status': 'Failed',
                    'project': test_project,
                    'error': 'Test case not found'
                })
                continue

            # Track validation results and captured data for this test
            validation_failures = []
            text_captures = []
            docx_screenshots = []
            screen_rows = 24
            screen_cols = 80

            try:
                # Initialize COM for THIS test case
                # NEW CODE:
                pythoncom.CoInitialize()
                autECLSession = win32com.client.Dispatch("PCOMM.autECLSession")
                connection_name = self.main_window.get_connection_name_from_title(self.main_window.pcomm_window_title)
                autECLSession.SetConnectionByName(connection_name)
                autECLPS = autECLSession.autECLPS
                
                # Get the selected start step (0-based index)
                # Get the selected start step (0-based index)
                start_step_index = self.get_start_step_index(test_case_name)
                
                # Flag to track if this test was stopped
                test_was_stopped = False
                
                # ✅ CHANGED: Use while loop instead of for loop to allow dynamic step list updates
                all_steps = test_case_data.get("steps", [])
                current_step_index = start_step_index
                
                while current_step_index < len(all_steps):
                    step_index = current_step_index + 1  # Display as 1-based
                    step = all_steps[current_step_index]
                    
                    # ✅ CRITICAL: Allow UI to process events before each step
                    QApplication.processEvents()
                    
                    # ✅ Check if stop was requested during step execution
                    if self.stop_execution:
                        print(f"Execution stopped by user at step {step_index}")
                        test_was_stopped = True
                        break  # Break out of the step loop
                    
                    step_type = step.get("type")
                    
                  

                    if step_type == "module_import":
                        import os
                        
                        # Process module fields FIRST
                        module_name = step.get("module_name")
                        
                        # ✅ NEW: Capture "Before" screenshot if Screen Flow is enabled
                        capture_screen_flow = self.main_window.document_config.get('capture_screen_flow', False)
                        if capture_screen_flow:
                            try:
                                screen_flow_dir = os.path.join(self.main_window.default_results_location, 'Screen Flows')
                                os.makedirs(screen_flow_dir, exist_ok=True)
                                
                                before_filename = f"{test_case_name}_Step {step_index}_Before.jpg"
                                before_path = os.path.join(screen_flow_dir, before_filename)
                                
                                if self.main_window.capture_pcomm_screen_as_jpeg(before_path):
                                    print(f"Step {step_index}: Captured 'Before' screen flow: {before_filename}")
                            except Exception as e:
                                print(f"Error capturing before screenshot: {e}")
                        
                        # Process module fields for Input
                        for field in step.get("fields", []):
                            if self.stop_execution:
                                break
                                
                            action_type = field.get("action_type", "Input")
                            value = str(field.get("value", "")).strip()
                            
                            if action_type == "Input":
                                if not value:
                                    continue
                                
                                if module_name in self.main_window.modules:
                                    module_data = self.main_window.modules[module_name]
                                    labels = module_data.get("labels", [])
                                    
                                    field_name = field.get("field_name")
                                    for label in labels:
                                        label_name = label.get('name') or label.get('label') or label.get('text', '')
                                        if label_name == field_name:
                                            row = int(label.get("row", 1))
                                            col = int(label.get("column", 1))
                                            
                                            # Substitute variables in the value before sending
                                            substituted_value = self.substitute_execution_variables(value, test_case_name)
                                            
                                            autECLPS.SetCursorPos(row, col)
                                            autECLPS.SendKeys(substituted_value)
                                            time.sleep(0.1)
                                            break
                        
                        # ✅ NEW: Capture "After" screenshot with smart wait if Screen Flow is enabled
                        if capture_screen_flow:
                            try:
                                screen_flow_dir = os.path.join(self.main_window.default_results_location, 'Screen Flows')
                                after_filename = f"{test_case_name}_Step {step_index}_After.jpg"
                                after_path = os.path.join(screen_flow_dir, after_filename)
                                
                                # Small delay to ensure screen is updated
                                time.sleep(0.5)
                                
                                if self.main_window.capture_pcomm_screen_as_jpeg(after_path):
                                    print(f"Step {step_index}: Captured 'After' screen flow: {after_filename}")
                            except Exception as e:
                                print(f"Error capturing after screenshot: {e}")
                        
                        # Process validation actions for module fields
                        for field in step.get("fields", []):
                            # Check for stop during field processing
                            QApplication.processEvents()
                            if self.stop_execution:
                                test_was_stopped = True
                                break
                            
                            action_type = field.get("action_type", "Input")
                            value = str(field.get("value", "")).strip()
                            
                            if action_type == "Validate":
                                if not value:
                                    continue
                                
                                if module_name in self.main_window.modules:
                                    module_data = self.main_window.modules[module_name]
                                    labels = module_data.get("labels", [])
                                    
                                    field_name = field.get("field_name")
                                    for label in labels:
                                        label_name = label.get('name') or label.get('label') or label.get('text', '')
                                        if label_name == field_name:
                                            row = int(label.get("row", 1))
                                            col = int(label.get("column", 1))
                                            length = int(label.get("length", len(value)))

                                            actual_value = autECLPS.GetText(row, col, length)  # ← Remove .strip() here
                                            
                                            # ✅ NEW: Use validation helper function
                                            validation_passed = self.validate_field_value(actual_value, value)
                                            
                                            if not validation_passed:
                                                # ✅ ENHANCED: Better error message for {blank} validation
                                                if value.lower() == '{blank}':
                                                    validation_failures.append({
                                                        "step": step_index,
                                                        "field": field_name,
                                                        "expected": '<blank>',
                                                        "actual": f"'{actual_value.strip()}'" if actual_value.strip() else '<blank>'
                                                    })
                                                else:
                                                    validation_failures.append({
                                                        "step": step_index,
                                                        "field": field_name,
                                                        "expected": value,
                                                        "actual": actual_value.strip()
                                                    })
                                                # Stop execution immediately on validation failure
                                                print(f"❌ Validation failed at Step {step_index} - Field '{field_name}': Expected '{value}', Got '{actual_value}'")
                                                break
                                            
                                            time.sleep(0.1)
                                            break
                                
                                # Check if validation failed and stop test execution
                                if validation_failures:
                                    print(f"🛑 Stopping test execution due to validation failure at Step {step_index}")
                                    break  # Break out of fields loop

                        # Check if validation failed and stop processing steps
                        if validation_failures:
                            break  # Break out of steps loop
                    
                    elif step_type == "special_key":
                        key_value = step.get("key_value", "")
                        
                        key_mapping = {
                            "Enter Key": "[enter]",
                            "Clear Key": "[clear]",
                            "End Key": "[EraseEof]",
                            "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                            "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                            "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                            "F13": "[pf13]", "F14": "[pf14]", "F15": "[pf15]", "F16": "[pf16]",
                            "F17": "[pf17]", "F18": "[pf18]", "F19": "[pf19]", "F20": "[pf20]",
                            "F21": "[pf21]", "F22": "[pf22]", "F23": "[pf23]", "F24": "[pf24]",
                        }
                        
                        # Check if it's a combo key (contains + sign)
                        if "+" in key_value:
                            pcomm_key = self.main_window.convert_combo_key_to_pcomm(key_value)
                        else:
                            pcomm_key = key_mapping.get(key_value, key_value)
                        
                        # ✅ FIXED: Smart wait for special keys
                        # Keys that typically cause screen changes
                        action_keys = ["Enter Key"] + [f"F{i}" for i in range(1, 25)]
                        
                        if key_value in action_keys or "+" in key_value:
                            # Capture screen before action
                            before_screen = wait_for_pcomm_ready_smart(autECLPS, key_value)
                            
                            # Send the key
                            autECLPS.SendKeys(pcomm_key)
                            print(f"Step {step_index}: Sent {key_value}")
                            
                            # ✅ FIXED: Pass key_value as a string, not variable reference
                            success, elapsed = complete_pcomm_wait(
                                autECLPS, 
                                before_screen, 
                                action_description=f"Special Key: {key_value}",  # ✅ Fixed here
                                timeout=30
                            )
                            
                            if not success:
                                QMessageBox.warning(self, "Timeout Warning", 
                                    f"Step {step_index}: {key_value} did not complete within 30 seconds.\n\n"
                                    "The test will continue, but results may be unreliable.")
                        else:
                            # Non-action keys (just send with small delay)
                            autECLPS.SendKeys(pcomm_key)
                            time.sleep(0.5)
                    
                    elif step_type == "capture_screen_text":
                        print(f"Step {step_index}: Capturing screen text...")
                        screen_size = screen_rows * screen_cols
                        full_screen_text = autECLPS.GetText(1, screen_size)
                        
                        captured_text_lines = []
                        for row_num in range(screen_rows):
                            start_index_text = row_num * screen_cols
                            line_text = full_screen_text[start_index_text:start_index_text + screen_cols]
                            captured_text_lines.append(line_text)
                        
                        from datetime import datetime
                        header = f"--- Step {step_index}: Screen Text Capture at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ---"
                        full_capture_string = header + "\n" + "\n".join(captured_text_lines)
                        text_captures.append(full_capture_string)
                        print(f"Step {step_index}: Screen text captured successfully")
                    
                    elif step_type == "capture_screenshot":
                        if not self.main_window.document_config.get('generate_documentation', True):
                            print(f"Step {step_index}.{sub_index}: Utility screenshot skipped (documentation disabled)")
                            continue
                        
                        print(f"Step {step_index}.{sub_index}: Capturing utility screenshot for DOCX...")
                        screen_size = screen_rows * screen_cols
                        full_screen_text = autECLPS.GetText(1, screen_size)
                        
                        # ✅ FIXED: Get highlight information for utility screenshot
                        highlight_info = {}
                        reference_module = utility_step.get('reference_module')
                        
                        if reference_module and reference_module in self.main_window.modules:
                            module_data = self.main_window.modules[reference_module]
                            labels = module_data.get('labels', [])
                            
                            # Get highlight flags from utility step fields
                            for field in utility_step.get('fields', []):
                                if field.get('highlight', False):
                                    field_name = field.get('field_name')
                                    # Find the corresponding label
                                    for label in labels:
                                        label_name = label.get('name') or label.get('label') or label.get('text', '')
                                        if label_name == field_name:
                                            highlight_info[field_name] = {
                                                'row': label.get('row', 1),
                                                'column': label.get('column', 1),
                                                'length': label.get('length', 10)
                                            }
                                            break
                        
                        from datetime import datetime
                        docx_screenshots.append({
                            'step': f"{step_index}.{sub_index}",
                            'screen_text': full_screen_text,
                            'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                            'highlight_info': highlight_info  # ✅ FIXED: Add highlight info
                        })
                        print(f"Step {step_index}.{sub_index}: Utility screenshot captured for DOCX with {len(highlight_info)} highlighted field(s)")
                    
                    elif step_type == "random_input":
                        # Handle random input steps
                        row = int(step.get("row", 1))
                        col = int(step.get("column", 1))
                        value = str(step.get("value", "")).strip()
                        is_special_key = step.get("is_special_key", False)
                        
                        if value:
                            # Set cursor position
                            autECLPS.SetCursorPos(row, col)
                            
                            if is_special_key:
                                # Map to PCOMM key code
                                key_mapping = {
                                    "Enter Key": "[enter]",
                                    "Clear Key": "[clear]",
                                    "End Key": "[EraseEof]",
                                    "F1": "[pf1]",
                                    "F2": "[pf2]",
                                    "F3": "[pf3]",
                                    "F4": "[pf4]",
                                    "F5": "[pf5]",
                                    "F6": "[pf6]",
                                    "F7": "[pf7]",
                                    "F8": "[pf8]",
                                    "F9": "[pf9]",
                                    "F10": "[pf10]",
                                    "F11": "[pf11]",
                                    "F12": "[pf12]",
                                }
                                pcomm_key = key_mapping.get(value, value)
                                autECLPS.SendKeys(pcomm_key)
                                print(f"Step {step_index}: Sent special key '{value}' to position ({row}, {col})")
                                time.sleep(1.0)
                            else:
                                # Send text value
                                autECLPS.SendKeys(value)
                                print(f"Step {step_index}: Sent '{value}' to position ({row}, {col})")
                    
                    elif step_type == "wait":
                        seconds = float(step.get("seconds", 0))
                        if seconds > 0:
                            print(f"Step {step_index}: Waiting for {seconds} second(s)...")
                            # ✅ Break wait into smaller chunks to allow stop checking
                            wait_chunks = int(seconds * 10)  # Check every 0.1 seconds
                            for _ in range(wait_chunks):
                                QApplication.processEvents()
                                if self.stop_execution:
                                    test_was_stopped = True
                                    break
                                time.sleep(0.1)
                            
                            if not test_was_stopped:
                                print(f"Step {step_index}: Wait completed")
                    
                    elif step_type == "break":
                        message = step.get("message", "")
                        print(f"Step {step_index}: Break point reached")
                        
                        # Show break dialog
                        break_dialog = BreakExecutionDialog(message, self)
                        break_dialog.show()
                        
                        # Wait for user action in an event loop
                        while break_dialog.result_action is None:
                            QApplication.processEvents()
                            time.sleep(0.1)
                        
                        action = break_dialog.result_action
                        
                        if action == BreakExecutionDialog.STOP:
                            print("User chose to stop execution at break point")
                            self.stop_execution = True
                            break_dialog.close()
                            test_was_stopped = True
                            break
                        
                        elif action == BreakExecutionDialog.EDIT:
                            print("User chose to edit test case at break point")
                            break_dialog.result_action = None
                            
                            # Get current test case data
                            existing_steps = test_case_data.get('steps', [])
                            test_case_description = test_case_data.get('description', '')
                            test_case_assumptions = test_case_data.get('assumptions', '')
                            
                            # Open edit dialog
                            edit_dialog = EditTestCaseDialog(
                                existing_steps, 
                                self.main_window.modules, 
                                self.main_window,
                                test_case_name, 
                                test_case_description, 
                                test_case_assumptions
                            )
                            edit_dialog.setParent(break_dialog, edit_dialog.windowFlags())
                            edit_dialog.exec()
                            
                            # Get updated data
                            updated_steps = edit_dialog.get_updated_steps()
                            updated_description = edit_dialog.get_test_case_description()
                            updated_assumptions = edit_dialog.get_test_case_assumptions()
                            updated_prerequisites = edit_dialog.get_prerequisites()
                            
                            # Update in-memory test case data
                            test_case_data['steps'] = updated_steps
                            test_case_data['description'] = updated_description
                            test_case_data['assumptions'] = updated_assumptions
                            test_case_data['prerequisites'] = updated_prerequisites
                            
                            # Save to file
                            self.main_window.save_test_cases_to_file()
                            
                            # Refresh the all_steps list
                            all_steps = test_case_data.get("steps", [])
                            
                            print(f"Test case '{test_case_name}' updated during execution")
                            print(f"Total steps after update: {len(all_steps)}")
                            
                            QMessageBox.information(
                            break_dialog,
                            "Test Case Updated",
                            f"Test case '{test_case_name}' has been updated.\n"
                            f"Total steps: {len(all_steps)}\n"
                            f"Press 'Resume Execution' to continue from Step {step_index}."
                            )
                            
                            break_dialog.raise_()
                            break_dialog.activateWindow()
                            
                            # Continue waiting for the next action (Resume or Stop)
                            while break_dialog.result_action is None:
                                QApplication.processEvents()
                                time.sleep(0.1)
                            
                            action = break_dialog.result_action
                            
                            if action == BreakExecutionDialog.STOP:
                                self.stop_execution = True
                                break_dialog.close()
                                test_was_stopped = True
                                break
                            elif action == BreakExecutionDialog.RESUME:
                                break_dialog.close()
                        
                        elif action == BreakExecutionDialog.RESUME:
                            print("User chose to resume execution")
                            break_dialog.close()
                    
                    # ✅ NEW: Process utility steps for this main step
                    # Find this section in both execution methods:
                    if 'utility_steps' in step:
                        for sub_index, utility_step in enumerate(step['utility_steps'], 1):
                            if self.stop_execution:
                                print(f"Execution stopped at utility step {step_index}.{sub_index}")
                                break
                            
                            QApplication.processEvents()
                            
                            utility_type = utility_step.get("type")
                            print(f"Executing utility step {step_index}.{sub_index}: {utility_step.get('name', 'Unknown')}")
                            
                   
                            
                            if utility_type == "special_key":
                                key_value = utility_step.get("key_value", "")
                                
                                key_mapping = {
                                    "Enter Key": "[enter]",
                                    "Clear Key": "[clear]",
                                    "End Key": "[EraseEof]",
                                    "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                                    "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                                    "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                                    "F13": "[pf13]", "F14": "[pf14]", "F15": "[pf15]", "F16": "[pf16]",
                                    "F17": "[pf17]", "F18": "[pf18]", "F19": "[pf19]", "F20": "[pf20]",
                                    "F21": "[pf21]", "F22": "[pf22]", "F23": "[pf23]", "F24": "[pf24]",
                                }
                                
                                if "+" in key_value:
                                    pcomm_key = self.main_window.convert_combo_key_to_pcomm(key_value)
                                else:
                                    pcomm_key = key_mapping.get(key_value, key_value)
                                
                                # ✅ FIXED: Smart wait for utility special keys
                                action_keys = ["Enter Key"] + [f"F{i}" for i in range(1, 25)]
                                
                                if key_value in action_keys or "+" in key_value:
                                    before_screen = wait_for_pcomm_ready_smart(autECLPS, f"Utility {key_value}")
                                    autECLPS.SendKeys(pcomm_key)
                                    print(f"Step {step_index}.{sub_index}: Sent utility special key '{key_value}'")
                                    
                                    # ✅ FIXED: Pass action_description as a string
                                    success, elapsed = complete_pcomm_wait(
                                        autECLPS, 
                                        before_screen, 
                                        action_description=f"Utility Special Key: {key_value}",  # ✅ Fixed here
                                        timeout=30
                                    )
                                    
                                    if not success:
                                        print(f"⚠️ Utility {key_value} timeout at step {step_index}.{sub_index}")
                                else:
                                    autECLPS.SendKeys(pcomm_key)
                                    time.sleep(0.5)
                            
                            elif utility_type == "wait":
                                seconds = float(utility_step.get("seconds", 0))
                                if seconds > 0:
                                    print(f"Step {step_index}.{sub_index}: Utility wait for {seconds} second(s)...")
                                    time.sleep(seconds)
                                    print(f"Step {step_index}.{sub_index}: Utility wait completed")
                            
                            elif utility_type == "capture_screenshot":
                                if not self.main_window.document_config.get('generate_documentation', True):
                                    print(f"Step {step_index}.{sub_index}: Utility screenshot skipped (documentation disabled)")
                                    continue
                                
                                print(f"Step {step_index}.{sub_index}: Capturing utility screenshot for DOCX...")
                                screen_size = screen_rows * screen_cols
                                full_screen_text = autECLPS.GetText(1, screen_size)
                                
                                # Get highlight information for utility screenshot
                                highlight_info = {}
                                reference_module = utility_step.get('reference_module')
                                
                                if reference_module and reference_module in self.main_window.modules:
                                    module_data = self.main_window.modules[reference_module]
                                    labels = module_data.get('labels', [])
                                    
                                    for field in utility_step.get('fields', []):
                                        if field.get('highlight', False):
                                            field_name = field.get('field_name')
                                            for label in labels:
                                                label_name = label.get('name') or label.get('label') or label.get('text', '')
                                                if label_name == field_name:
                                                    highlight_info[field_name] = {
                                                        'row': label.get('row', 1),
                                                        'column': label.get('column', 1),
                                                        'length': label.get('length', 10)
                                                    }
                                                    break
                                
                                from datetime import datetime
                                docx_screenshots.append({
                                    'step': f"{step_index}.{sub_index}",
                                    'screen_text': full_screen_text,
                                    'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                                    'highlight_info': highlight_info
                                })
                                print(f"Step {step_index}.{sub_index}: Utility screenshot captured for DOCX with {len(highlight_info)} highlighted field(s)")
                            
                            elif utility_type == "capture_screen_text":
                                print(f"Step {step_index}.{sub_index}: Capturing utility screen text...")
                                screen_size = screen_rows * screen_cols
                                full_screen_text = autECLPS.GetText(1, screen_size)
                                
                                captured_text_lines = []
                                for row_num in range(screen_rows):
                                    start_index_text = row_num * screen_cols
                                    line_text = full_screen_text[start_index_text:start_index_text + screen_cols]
                                    captured_text_lines.append(line_text)
                                
                                from datetime import datetime
                                header = f"--- Step {step_index}.{sub_index}: Utility Text Capture at {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ---"
                                full_capture_string = header + "\n" + "\n".join(captured_text_lines)
                                text_captures.append(full_capture_string)
                                print(f"Step {step_index}.{sub_index}: Utility screen text captured successfully")
                            
                            elif utility_type == "random_input":
                                row = int(utility_step.get('row', 1))
                                col = int(utility_step.get('column', 1))
                                value = str(utility_step.get('value', '')).strip()
                                is_special_key = utility_step.get('is_special_key', False)
                                
                                if value:
                                    autECLPS.SetCursorPos(row, col)
                                    
                                    if is_special_key:
                                        key_mapping = {
                                            "Enter Key": "[enter]", "Clear Key": "[clear]", "End Key": "[EraseEof]",
                                            "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                                            "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                                            "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                                        }
                                        pcomm_key = key_mapping.get(value, value)
                                        autECLPS.SendKeys(pcomm_key)
                                        print(f"Step {step_index}.{sub_index}: Sent special key '{value}' to position ({row}, {col})")
                                        time.sleep(1.0)
                                    else:
                                        substituted_value = self.substitute_execution_variables(value, test_case_name)
                                        autECLPS.SendKeys(substituted_value)
                                        print(f"Step {step_index}.{sub_index}: Sent '{substituted_value}' to position ({row}, {col})")                            
                            
                            elif utility_type == "module_import":
                                # âœ… FIXED: Handle module import utility step (for both Input and Validation)
                                module_name = utility_step.get('module_name')
                                
                                if module_name in self.main_window.modules:
                                    module_data = self.main_window.modules[module_name]
                                    labels = module_data.get("labels", [])
                                    
                                    # âœ… FIXED: Process Input fields FIRST
                                    for field in utility_step.get('fields', []):
                                        if self.stop_execution:
                                            test_was_stopped = True
                                            break
                                        
                                        action_type = field.get('action_type', 'Input')
                                        value = str(field.get('value', '')).strip()
                                        
                                        if action_type == 'Input' and value:
                                            field_name = field.get('field_name')
                                            
                                            # Find the label for this field
                                            for label in labels:
                                                label_name = label.get('name') or label.get('label') or label.get('text', '')
                                                if label_name == field_name:
                                                    row = int(label.get('row', 1))
                                                    col = int(label.get('column', 1))
                                                    
                                                    # Substitute variables before sending
                                                    substituted_value = self.substitute_execution_variables(value, test_case_name)
                                                    
                                                    autECLPS.SetCursorPos(row, col)
                                                    autECLPS.SendKeys(substituted_value)
                                                    print(f"Step {step_index}.{sub_index}: Sent utility input '{substituted_value}' to {field_name}")
                                                    time.sleep(0.1)
                                                    break
                                    
                                    # âœ… FIXED: Process Validation fields AFTER inputs
                                    for field in utility_step.get('fields', []):
                                        if self.stop_execution:
                                            test_was_stopped = True
                                            break
                                        
                                        action_type = field.get('action_type', 'Validate')
                                        value = str(field.get('value', '')).strip()
                                        
                                        if action_type == 'Validate' and value:
                                            field_name = field.get('field_name')
                                            
                                            # Find the label for this field
                                            for label in labels:
                                                label_name = label.get('name') or label.get('label') or label.get('text', '')
                                                if label_name == field_name:
                                                    row = int(label.get('row', 1))
                                                    col = int(label.get('column', 1))
                                                    length = int(label.get('length', len(value)))

                                                    # Read actual value from screen
                                                    actual_value = autECLPS.GetText(row, col, length)  # ← Remove .strip() here
                                                    
                                                    # ✅ NEW: Use validation helper function
                                                    validation_passed = self.validate_field_value(actual_value, value)
                                                    
                                                    if not validation_passed:
                                                        # ✅ ENHANCED: Better error message for {blank} validation
                                                        if value.lower() == '{blank}':
                                                            validation_failures.append({
                                                                "step": f"{step_index}.{sub_index}",
                                                                "field": field_name,
                                                                "expected": '<blank>',
                                                                "actual": f"'{actual_value.strip()}'" if actual_value.strip() else '<blank>'
                                                            })
                                                        else:
                                                            validation_failures.append({
                                                                "step": f"{step_index}.{sub_index}",
                                                                "field": field_name,
                                                                "expected": value,
                                                                "actual": actual_value.strip()
                                                            })
                                                        # âœ… Stop execution immediately on validation failure
                                                        print(f"âŒ Utility validation failed at Step {step_index}.{sub_index} - Field '{field_name}': Expected '{value}', Got '{actual_value}'")
                                                        break
                                                    
                                                    print(f"Step {step_index}.{sub_index}: Validated {field_name} - Expected: '{value}', Actual: '{actual_value}'")
                                                    time.sleep(0.1)
                                                    break
                                            
                                            # âœ… Check if validation failed and stop utility steps
                                            if validation_failures:
                                                print(f"ðŸ›' Stopping utility steps due to validation failure at Step {step_index}.{sub_index}")
                                                break
                                        
                                        # âœ… Check if validation failed and stop processing this step's utilities
                                        if validation_failures:
                                            break
                                else:
                                    print(f"Warning: Module '{module_name}' not found for utility step")
                          
                            
                            time.sleep(0.1)

                    # âœ… NEW: Check if validation failed in utility steps and stop main loop
                    if validation_failures:
                        print(f"ðŸ›' Stopping test execution due to utility validation failure")
                        break  # Break out of main steps while loop
                    
                    time.sleep(0.1)                
                    current_step_index += 1


                
                # Handle if test was stopped during execution
                if test_was_stopped:
                    self.update_status(test_case_name, "Stopped")
                    stopped_count += 1
                    results_summary.append(f"⏸️ {test_case_name}: Stopped by user")
                    
                   
                    
                    execution_results.append({
                        'name': test_case_name,
                        'status': 'Stopped',
                        'project': test_project,
                        'error': 'Execution stopped by user during test'
                    })
                    print(f"⏸️ Test '{test_case_name}' STOPPED")
                    continue  # Skip to next test case

                # Everything below only runs if test was NOT stopped
                
                # Save captured text to file if any captures were made
                if text_captures:
                    import os
                    project_name = getattr(self.main_window, 'current_project_id', None)
                    if project_name and project_name in self.main_window.projects:
                        project_name = self.main_window.projects[project_name]['name']
                    
                    output_dir = os.path.join(self.main_window.default_results_location, 'Results', project_name if project_name else 'Master')
                    os.makedirs(output_dir, exist_ok=True)
                    
                    filename = os.path.join(output_dir, f"{test_case_name}.txt")
                    
                    with open(filename, "w", encoding="utf-8") as f:
                        f.write("\n\n\n".join(text_captures))
                    
                    print(f"All screen text for '{test_case_name}' saved to '{filename}'.")
                
                # Create single DOCX with all screenshots
                docx_path = None
                if docx_screenshots:
                    try:
                        docx_path = self.main_window.create_test_case_docx(test_case_name, docx_screenshots)
                        print(f"DOCX document created: '{docx_path}' with {len(docx_screenshots)} screenshot(s)")
                    except Exception as e:
                        print(f"Error creating DOCX: {e}")
                        validation_failures.append({
                            "step": "DOCX Generation",
                            "field": "Document Creation",
                            "expected": "Success",
                            "actual": f"Error: {str(e)}"
                        })

                # Check results
                if validation_failures:
                    self.update_status(test_case_name, "Failed")
                    failed_count += 1
                    
                    failure_summary = f"❌ {test_case_name}: {len(validation_failures)} validation(s) failed"
                    results_summary.append(failure_summary)
                   
                    execution_results.append({
                        'name': test_case_name,
                        'status': 'Failed',
                        'project': test_project,
                        'validation_failures': validation_failures
                    })
                    
                    print(f"\n⚠️ Test '{test_case_name}' FAILED with {len(validation_failures)} validation error(s)")
                    for failure in validation_failures:
                        print(f"  Step {failure['step']} - {failure['field']}: Expected '{failure['expected']}', Got '{failure['actual']}'")
                else:
                    # âœ… Calculate execution time
                    test_end_time = datetime.now()
                    test_end_time_str = test_end_time.strftime('%Y-%m-%d %H:%M:%S')
                    test_duration = test_end_time - test_start_time
                    test_duration_str = str(test_duration).split('.')[0]
                    
                    # âœ… Store execution time
                    self.execution_times[test_case_name] = {
                        'start_time': test_start_time_str,
                        'end_time': test_end_time_str,
                        'duration': test_duration_str
                    }
                    
                    # âœ… Update UI
                    # Update the UI
                    for i in range(self.test_case_list.count()):
                        item = self.test_case_list.item(i)
                        item_data = item.data(Qt.ItemDataRole.UserRole)
                        
                        if item_data and item_data.get('type') == 'test_case' and item_data.get('name') == test_case_name:
                            widget = self.test_case_list.itemWidget(item)
                            time_label = widget.findChild(QLabel, "execution_time_label")
                            if time_label:
                                time_label.setText(f"{test_duration_str}")
                            break

                    self.update_status(test_case_name, "Passed")
                    passed_count += 1

                    result_text = f"✅ {test_case_name}: Passed"

                    # Add info about captured files
                    file_info = []
                    if text_captures:
                        file_info.append(f"{len(text_captures)} text screenshot(s)")
                    if docx_path:
                        file_info.append(f"DOCX: {len(docx_screenshots)} screenshots")
                    elif not self.main_window.document_config.get('generate_documentation', True):
                        file_info.append("Documentation disabled")

                    if file_info:
                        result_text += f" ({', '.join(file_info)})"

                    results_summary.append(result_text)
                    


                    execution_results.append({
                        'name': test_case_name,
                        'status': 'Passed',
                        'project': test_project,
                        'start_time': test_start_time_str,
                        'end_time': test_end_time_str,
                        'duration': test_duration_str
                    })
                    
                    print(f"✅ Test '{test_case_name}' PASSED")

            except Exception as e:
                self.update_status(test_case_name, "Failed")
                failed_count += 1
                results_summary.append(f"❌ {test_case_name}: Error - {str(e)}")


                execution_results.append({
                    'name': test_case_name,
                    'status': 'Failed',
                    'project': test_project,
                    'error': str(e)
                })
                
                print(f"❌ Test '{test_case_name}' FAILED with error: {str(e)}")
            
            finally:
                # Clean up COM after EACH test case
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass
        
        # Restore button to original state after ALL tests complete or stopped
        try:
            self.execute_button.clicked.disconnect()
        except:
            pass
        self.execute_button.setText(original_text)
        self.execute_button.setIcon(original_icon)
        self.execute_button.clicked.connect(self.execute_selected_tests)
        self.execute_button.setEnabled(True)

        # Show final summary
        print(f"\n{'='*60}")
        print(f"EXECUTION SUMMARY")
        print(f"{'='*60}")
        print(f"Total Tests: {len(selected_tests)}")
        print(f"Passed: {passed_count}")
        print(f"Failed: {failed_count}")
        print(f"Stopped: {stopped_count}")
        print(f"{'='*60}\n")

        summary_message = f"Execution Complete!\n\n"
        summary_message += f"Total: {len(selected_tests)} | Passed: {passed_count} | Failed: {failed_count} | Stopped: {stopped_count}\n\n"
        summary_message += "Results:\n" + "\n".join(results_summary)

        # Create DOCX summary
        docx_summary_path = self.create_execution_summary_docx(execution_results, execution_timestamp)
        if docx_summary_path:
            summary_message += f"\n\nExecution summary saved to:\n{docx_summary_path}"
            print(f"Execution summary saved to: {docx_summary_path}")

        if stopped_count > 0:
            QMessageBox.warning(self, "Execution Stopped", summary_message)
        elif failed_count > 0:
            QMessageBox.warning(self, "Execution Complete", summary_message)
        else:
            QMessageBox.information(self, "Execution Complete", summary_message)
            
    def request_stop_execution(self):
        """Sets the stop flag to halt execution."""
        self.stop_execution = True
        self.execute_button.setEnabled(False)  # Disable button after stop is requested
        print("⏸️ Stop requested by user...")

    def create_execution_summary_docx(self, execution_results, timestamp):
        """
        Creates a DOCX file with test execution summary.
        
        Args:
            execution_results: List of dicts with test case results
            timestamp: Timestamp string for the filename
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn
            
            # Create a new document
            doc = Document()
            
            # Set document margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Add title
            title = doc.add_paragraph()
            title_run = title.add_run("Test Execution Summary")
            title_run.font.size = Pt(18)
            title_run.font.bold = True
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.paragraph_format.space_after = Pt(12)
            
            # Add timestamp
            time_para = doc.add_paragraph()
            time_run = time_para.add_run(f"Execution Time: {timestamp}")
            time_run.font.size = Pt(12)
            time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            time_para.paragraph_format.space_after = Pt(20)
            
            # Calculate summary statistics
            total_tests = len(execution_results)
            passed_tests = sum(1 for r in execution_results if r['status'] == 'Passed')
            failed_tests = sum(1 for r in execution_results if r['status'] == 'Failed')
            
            # Add summary statistics
            summary_para = doc.add_paragraph()
            summary_para.add_run("Summary:\n").font.bold = True
            summary_para.add_run(f"Total Test Cases: {total_tests}\n")
            
            passed_run = summary_para.add_run(f"Passed: {passed_tests}\n")
            passed_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            passed_run.font.bold = True
            
            failed_run = summary_para.add_run(f"Failed: {failed_tests}\n")
            failed_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            failed_run.font.bold = True
            
            summary_para.paragraph_format.space_after = Pt(20)
            
            # Add detailed results
            details_heading = doc.add_paragraph()
            details_heading.add_run("Detailed Results:").font.bold = True
            details_heading.paragraph_format.space_after = Pt(10)
            
            # Add each test case result
            for idx, result in enumerate(execution_results, 1):
                # Test case name and status
                result_para = doc.add_paragraph()
                
                # âœ… Add project information if available
                project_info = ""
                if result.get('project'):
                    project_info = f"[Project: {result['project']}] "
                
                result_para.add_run(f"{idx}. {project_info}{result['name']}: ").font.bold = True
                
                status_run = result_para.add_run(result['status'])
                status_run.font.bold = True
                if result['status'] == 'Passed':
                    status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
                else:
                    status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                
                # âœ… NEW: Add execution time details
                if 'start_time' in result and 'end_time' in result and 'duration' in result:
                    time_para = doc.add_paragraph()
                    time_para.paragraph_format.left_indent = Inches(0.5)
                    time_run = time_para.add_run(
                        f"Start Time: {result['start_time']}  |  "
                        f"End Time: {result['end_time']}  |  "
                        f"Duration: {result['duration']}"
                    )
                    time_run.font.size = Pt(9)
                    time_run.font.italic = True
                    time_run.font.color.rgb = RGBColor(75, 85, 99)  # Gray
                
                # Add error details if failed
                if result['status'] == 'Failed' and 'error' in result:
                    error_para = doc.add_paragraph()
                    error_para.paragraph_format.left_indent = Inches(0.5)
                    error_run = error_para.add_run(f"Error: {result['error']}")
                    error_run.font.size = Pt(10)
                    error_run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
                
                # Add validation failures if any
                if 'validation_failures' in result and result['validation_failures']:
                    failures_para = doc.add_paragraph()
                    failures_para.paragraph_format.left_indent = Inches(0.5)
                    failures_para.add_run("Validation Failures:\n").font.bold = True
                    
                    for failure in result['validation_failures']:
                        failure_detail = doc.add_paragraph()
                        failure_detail.paragraph_format.left_indent = Inches(0.75)
                        failure_detail.add_run(
                            f"Step {failure['step']} - {failure['field']}:\n"
                            f"  Expected: '{failure['expected']}'\n"
                            f"  Actual: '{failure['actual']}'\n"
                        ).font.size = Pt(9)
                
                # Add spacing between test cases
                result_para.paragraph_format.space_after = Pt(12)
            
            # Create output directory
            output_dir = os.path.join(self.main_window.default_results_location, 'Test Execution Summary')
            os.makedirs(output_dir, exist_ok=True)
            
            # Create filename with timestamp
            filename = os.path.join(output_dir, f"Test Execution - {timestamp}.docx")
            
            # Save the document
            doc.save(filename)
            
            return filename
            
        except Exception as e:
            print(f"Error creating execution summary DOCX: {e}")
            return None
        
    def check_prerequisites(self, test_case_name):
        """Checks if prerequisites are met."""
        # ✅ UPDATED: Use helper method
        test_case_data = self.get_test_case_data(test_case_name)
        if not test_case_data:
            return False, "Test case not found"
        
        prerequisites = test_case_data.get('prerequisites', [])
        
        if not prerequisites:
            return True, ""
        
        for prereq_name in prerequisites:
            prereq_status = None
            for i in range(self.test_case_list.count()):
                item = self.test_case_list.item(i)
                item_data = item.data(Qt.ItemDataRole.UserRole)
                
                if item_data and item_data.get('type') == 'test_case' and item_data.get('name') == prereq_name:
                    widget = self.test_case_list.itemWidget(item)
                    status_label = widget.findChild(QLabel, "status_label")
                    if status_label:
                        prereq_status = status_label.text()
                    break
            
            if prereq_status is None:
                return False, f"Prerequisite '{prereq_name}' is not in the execution list"
            
            if prereq_status == "Not Run":
                return False, f"Prerequisite '{prereq_name}' has not been executed yet"
            elif prereq_status == "Failed":
                return False, f"Prerequisite '{prereq_name}' has failed"
            elif prereq_status == "Running":
                return False, f"Prerequisite '{prereq_name}' is currently running"
        
        return True, ""
        
    def save_execution_data(self):
        """Saves execution data including projects."""
        # ✅ FIXED: Save complete test case data with projects
        for project_name, project_data in self.projects.items():
            # Ensure test_cases dictionary exists
            if 'test_cases' not in project_data:
                project_data['test_cases'] = {}
            
            # Status data is saved separately
            if 'status_data' not in project_data:
                project_data['status_data'] = {}
            
            for test_name in list(project_data['test_cases'].keys()):
                # Find and update status for this test case
                for i in range(self.test_case_list.count()):
                    item = self.test_case_list.item(i)
                    item_data = item.data(Qt.ItemDataRole.UserRole)
                    
                    if (item_data and item_data.get('type') == 'test_case' and 
                        item_data.get('name') == test_name and 
                        item_data.get('project') == project_name):
                        
                        widget = self.test_case_list.itemWidget(item)
                        status_label = widget.findChild(QLabel, "status_label")
                        step_combo = widget.findChild(QComboBox, "step_combo")
                        
                        # Store status in a separate dict (not in test_case_data itself)
                        project_data['status_data'][test_name] = {
                            'status': status_label.text() if status_label else "Not Run",
                            'selected_step': step_combo.currentIndex() if step_combo else 0
                        }
                        break
        
        execution_data = {
            'projects': self.projects,
            'standalone': {}
        }
        
        # Save standalone test cases with status
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case' and not item_data.get('project'):
                widget = self.test_case_list.itemWidget(item)
                status_label = widget.findChild(QLabel, "status_label")
                step_combo = widget.findChild(QComboBox, "step_combo")
                
                test_name = item_data['name']
                
                # ✅ CRITICAL FIX: Always get test case data from displayed_test_cases
                if test_name in self.displayed_test_cases:
                    execution_data['standalone'][test_name] = {
                        'status': status_label.text() if status_label else "Not Run",
                        'selected_step': step_combo.currentIndex() if step_combo else 0,
                        'test_case_data': self.displayed_test_cases[test_name]
                    }
        
        try:
            with open(self.execution_data_file, 'w') as f:
                json.dump(execution_data, f, indent=4)
        except Exception as e:
            print(f"Error saving: {e}")

    def load_execution_data(self):
        """Loads execution data including projects."""
        if not os.path.exists(self.execution_data_file):
            return
        
        try:
            with open(self.execution_data_file, 'r') as f:
                execution_data = json.load(f)
            
            # Load projects
            self.projects = execution_data.get('projects', {})
            
            # ✅ FIXED: Ensure each project has the required structure
            for project_name, project_data in self.projects.items():
                if 'test_cases' not in project_data:
                    project_data['test_cases'] = {}
                # ✅ IMPORTANT: Preserve the expanded state from saved data
                if 'expanded' not in project_data:
                    project_data['expanded'] = True  # Default to expanded if not specified
                if 'status_data' not in project_data:
                    project_data['status_data'] = {}
                
                # ✅ CRITICAL FIX: Ensure test_cases is a dict with proper structure
                if not isinstance(project_data['test_cases'], dict):
                    project_data['test_cases'] = {}
            
            # Load standalone test cases
            standalone_data = execution_data.get('standalone', {})
            for test_name, data in standalone_data.items():
                test_case_data = data.get('test_case_data')
                if test_case_data:
                    self.displayed_test_cases[test_name] = test_case_data
            
            self.refresh_test_case_list()
            
            # ✅ Restore status for all test cases after refresh
            self.restore_test_case_status(execution_data)
            
        except Exception as e:
            print(f"Error loading: {e}")
            
    def restore_test_case_status(self, execution_data):
        """Restores the status and selected step for all test cases."""
        # Restore project test case status
        for project_name, project_data in self.projects.items():
            status_data = project_data.get('status_data', {})
            for test_name, status_info in status_data.items():
                for i in range(self.test_case_list.count()):
                    item = self.test_case_list.item(i)
                    item_data = item.data(Qt.ItemDataRole.UserRole)
                    
                    if (item_data and item_data.get('type') == 'test_case' and 
                        item_data.get('name') == test_name and 
                        item_data.get('project') == project_name):
                        
                        widget = self.test_case_list.itemWidget(item)
                        status_label = widget.findChild(QLabel, "status_label")
                        step_combo = widget.findChild(QComboBox, "step_combo")
                        
                        if status_label:
                            saved_status = status_info.get('status', 'Not Run')
                            status_label.setText(saved_status)
                            if saved_status == "Passed":
                                status_label.setStyleSheet("color: green; font-weight: bold;")
                            elif saved_status == "Failed":
                                status_label.setStyleSheet("color: red; font-weight: bold;")
                            elif saved_status == "Stopped":
                                status_label.setStyleSheet("color: orange; font-weight: bold;")
                        
                        if step_combo:
                            saved_step = status_info.get('selected_step', 0)
                            if saved_step < step_combo.count():
                                step_combo.setCurrentIndex(saved_step)
                        break
        
        # Restore standalone test case status
        standalone_data = execution_data.get('standalone', {})
        for test_name, data in standalone_data.items():
            for i in range(self.test_case_list.count()):
                item = self.test_case_list.item(i)
                item_data = item.data(Qt.ItemDataRole.UserRole)
                
                if (item_data and item_data.get('type') == 'test_case' and 
                    item_data.get('name') == test_name and not item_data.get('project')):
                    
                    widget = self.test_case_list.itemWidget(item)
                    status_label = widget.findChild(QLabel, "status_label")
                    step_combo = widget.findChild(QComboBox, "step_combo")
                    
                    if status_label:
                        saved_status = data.get('status', 'Not Run')
                        status_label.setText(saved_status)
                        if saved_status == "Passed":
                            status_label.setStyleSheet("color: green; font-weight: bold;")
                        elif saved_status == "Failed":
                            status_label.setStyleSheet("color: red; font-weight: bold;")
                        elif saved_status == "Stopped":
                            status_label.setStyleSheet("color: orange; font-weight: bold;")
                    
                    if step_combo:
                        saved_step = data.get('selected_step', 0)
                        if saved_step < step_combo.count():
                            step_combo.setCurrentIndex(saved_step)
                    break
    
    def clear_all_test_cases(self):
        """Clears all projects and standalone test cases."""
        reply = QMessageBox.question(
            self, "Confirm Clear",
            "Clear all projects and test cases?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            self.projects.clear()
            self.displayed_test_cases.clear()
            self.test_case_list.clear()
            
            if os.path.exists(self.execution_data_file):
                os.remove(self.execution_data_file)
            
            QMessageBox.information(self, "Cleared", "All data cleared.")

    def toggle_play_stop(self, test_case_name):
        """Toggles between play and stop."""
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case' and item_data.get('name') == test_case_name:
                widget = self.test_case_list.itemWidget(item)
                button = widget.findChild(QPushButton, "play_stop_button")
                if button:
                    current_icon = button.icon()
                    play_icon = self.main_window.style().standardIcon(QStyle.StandardPixmap.SP_MediaPlay)
                    
                    if current_icon.pixmap(QSize(16, 16)).toImage() == play_icon.pixmap(QSize(16, 16)).toImage():
                        self.stop_execution = False
                        self.execute_single_test(test_case_name)
                    else:
                        self.stop_execution = True
                        self.update_status(test_case_name, "Stopped")
                        button.setIcon(play_icon)
                break
                
    def set_play_stop_button_state(self, test_case_name, is_playing):
        """Changes the play/stop button icon."""
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case' and item_data.get('name') == test_case_name:
                widget = self.test_case_list.itemWidget(item)
                button = widget.findChild(QPushButton, "play_stop_button")
                if button:
                    if is_playing:
                        button.setIcon(self.main_window.style().standardIcon(QStyle.StandardPixmap.SP_MediaStop))
                    else:
                        button.setIcon(self.main_window.style().standardIcon(QStyle.StandardPixmap.SP_MediaPlay))
                break

    def refresh_test_case(self, test_case_name):
        """Refreshes a test case from the library."""
        if test_case_name not in self.main_window.test_cases:
            QMessageBox.warning(self, "Test Case Not Found", 
                              f"Test case '{test_case_name}' no longer exists in the library.")
            return
        
        latest_test_case = self.main_window.test_cases[test_case_name]
        
        # ✅ FIXED: Find which project this test case belongs to (or if it's standalone)
        updated = False
        belongs_to_project = None
        
        # Check in projects first
        for project_name, project_data in self.projects.items():
            if test_case_name in project_data['test_cases']:
                project_data['test_cases'][test_case_name] = latest_test_case
                updated = True
                belongs_to_project = project_name
                break
        
        # Only check in standalone if it's NOT in a project
        if not updated and test_case_name in self.displayed_test_cases:
            self.displayed_test_cases[test_case_name] = latest_test_case
            updated = True
        
        if updated:
            # Update the step combo box for this specific test case
            for i in range(self.test_case_list.count()):
                item = self.test_case_list.item(i)
                item_data = item.data(Qt.ItemDataRole.UserRole)
                
                # ✅ FIXED: Match both name AND project to find the correct item
                if (item_data and item_data.get('type') == 'test_case' and 
                    item_data.get('name') == test_case_name and
                    item_data.get('project') == belongs_to_project):
                    
                    widget = self.test_case_list.itemWidget(item)
                    step_combo = widget.findChild(QComboBox, "step_combo")
                    if step_combo:
                        current_index = step_combo.currentIndex()
                        step_combo.clear()
                        num_steps = len(latest_test_case.get('steps', []))
                        for j in range(1, num_steps + 1):
                            step_combo.addItem(f"Step {j}")
                        
                        if current_index < num_steps:
                            step_combo.setCurrentIndex(current_index)
                        else:
                            step_combo.setCurrentIndex(0)
                    break
            
            self.save_execution_data()
            QMessageBox.information(self, "Refreshed", 
                                  f"Test case '{test_case_name}' has been refreshed.")
        else:
            QMessageBox.warning(self, "Not Found", 
                              f"Test case '{test_case_name}' not found in execution list.")
   
    def dropEvent(self, event):
        """
        Override dropEvent to prevent moving project headers and their test cases.
        Only allow reordering of test cases within their projects.
        """
        source_item = self.test_case_list.currentItem()
        
        if not source_item:
            event.ignore()
            return
        
        source_data = source_item.data(Qt.ItemDataRole.UserRole)
        
        # ✅ PREVENT: Don't allow dragging project headers
        if source_data and source_data.get('type') == 'project':
            event.ignore()
            QMessageBox.warning(
                self,
                "Cannot Move",
                "Project headers cannot be moved. Only test cases can be reordered."
            )
            return
        
        # Allow default drop behavior for test cases
        super().dropEvent(event)   
        
    def reorder_internal_data_after_drag(self, source_row, dest_row):
        """
        Reorders the internal data structures (projects and standalone test cases)
        to match the visual order after a drag-and-drop operation.
        """
        if source_row == dest_row or source_row < 0 or dest_row < 0:
            return
        
        # Collect all test case data in current visual order
        ordered_data = []
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case':
                test_name = item_data.get('name')
                project = item_data.get('project')
                
                # Get the actual test case data
                if project:
                    test_data = self.projects[project]['test_cases'].get(test_name)
                else:
                    test_data = self.displayed_test_cases.get(test_name)
                
                if test_data:
                    ordered_data.append({
                        'name': test_name,
                        'project': project,
                        'data': test_data
                    })
        
        # Rebuild data structures in new order
        # Clear existing data
        for project_data in self.projects.values():
            project_data['test_cases'].clear()
        self.displayed_test_cases.clear()
        
        # Re-add in new order
        for item in ordered_data:
            if item['project']:
                self.projects[item['project']]['test_cases'][item['name']] = item['data']
            else:
                self.displayed_test_cases[item['name']] = item['data']
        
        # Save the updated order
        self.save_execution_data()  

    def on_rows_moved(self, parent, start, end, destination, row):
        """
        Signal handler for when rows are moved via drag-and-drop.
        Updates internal data to match the new visual order.
        """
        # Small delay to ensure the UI has updated
        QTimer.singleShot(50, self.sync_data_with_visual_order)   

    def sync_data_with_visual_order(self):
        """
        Synchronizes internal data structures with the current visual order
        of items in the list widget.
        """
        # Collect all test case data in current visual order
        ordered_test_cases = []
        
        for i in range(self.test_case_list.count()):
            item = self.test_case_list.item(i)
            item_data = item.data(Qt.ItemDataRole.UserRole)
            
            if item_data and item_data.get('type') == 'test_case':
                test_name = item_data.get('name')
                project = item_data.get('project')
                
                # Get the actual test case data
                if project:
                    if project in self.projects and test_name in self.projects[project]['test_cases']:
                        test_data = copy.deepcopy(self.projects[project]['test_cases'][test_name])
                        ordered_test_cases.append({
                            'name': test_name,
                            'project': project,
                            'data': test_data
                        })
                else:
                    if test_name in self.displayed_test_cases:
                        test_data = copy.deepcopy(self.displayed_test_cases[test_name])
                        ordered_test_cases.append({
                            'name': test_name,
                            'project': None,
                            'data': test_data
                        })
        
        # Clear and rebuild data structures
        for project_data in self.projects.values():
            project_data['test_cases'].clear()
        self.displayed_test_cases.clear()
        
        # Re-add in new order
        for item in ordered_test_cases:
            if item['project']:
                self.projects[item['project']]['test_cases'][item['name']] = item['data']
            else:
                self.displayed_test_cases[item['name']] = item['data']
        
        # Save the updated order
        self.save_execution_data()

    def substitute_execution_variables(self, text, test_case_name):
        """
        Substitutes variables during test execution.
        Similar to main window's substitute_variables but focused on execution time.
        """
        from datetime import datetime
        
        now = datetime.now()
        
        # Get test case data
        test_case_data = self.get_test_case_data(test_case_name)
        test_description = test_case_data.get('description', '') if test_case_data else ''
        
        variables = {
            'date': now.strftime('%Y.%m.%d'),
            'time': now.strftime('%H:%M:%S'),
            'datetime': now.strftime('%Y-%m-%d %H:%M:%S'),
            'test_case_id': test_case_name,
            'test_description': test_description
        }
        
        result = text
        for var_name, var_value in variables.items():
            placeholder = '{' + var_name + '}'
            result = result.replace(placeholder, var_value)
        
        return result        

class CustomStepsListWidget(QListWidget):
    """Custom QListWidget that handles drag-and-drop for converting steps to utility steps."""
    
    def __init__(self, parent_dialog):
        super().__init__()
        self.parent_dialog = parent_dialog
        self.setDragEnabled(True)
        self.setAcceptDrops(True)
        self.setDragDropMode(QListWidget.DragDropMode.NoDragDrop)  # ✅ CHANGED: Disable default behavior
        self.setDefaultDropAction(Qt.DropAction.IgnoreAction)  # ✅ CHANGED: Ignore default action
        
        # ✅ NEW: Enable custom drag-drop
        self.setDragEnabled(True)
        self.setAcceptDrops(True)
    
    def dragEnterEvent(self, event):
        """Override to accept drag events."""
        if event.source() == self:
            event.accept()
        else:
            event.ignore()
    
    def dragMoveEvent(self, event):
        """Override to accept drag move events."""
        if event.source() == self:
            event.accept()
        else:
            event.ignore()
    
    def dropEvent(self, event):
        """Override dropEvent to use custom logic instead of default behavior."""
        # ✅ Call parent dialog's custom handler
        self.parent_dialog.handle_step_drop(event)
        
        # ✅ CRITICAL: Always ignore the event to prevent Qt's default behavior
        event.ignore()

class EditTestCaseDialog(QDialog):
    # Class variable to store copied steps across all instances
    copied_steps = []
    
    def __init__(self, existing_steps, modules, parent, test_case_name="", test_case_description="", test_case_assumptions=""):
        super().__init__(parent)  # Pass parent to QDialog
        
        self.main_window = parent  # ✅ Now this works correctly
        
        self.setWindowTitle(f"Edit Test Case: {test_case_name}" if test_case_name else "Edit Test Case")
        
        # ✅ Make it resizable and maximizable
        self.setMinimumSize(800, 900)
        self.setWindowFlags(
            self.windowFlags() | 
            Qt.WindowType.WindowMaximizeButtonHint |
            Qt.WindowType.WindowMinimizeButtonHint &
            ~Qt.WindowType.WindowContextHelpButtonHint
        )

        self.modules = modules
        self.added_steps = existing_steps
        self.selected_module_name = ""
        self.current_utility_step = None
        self.special_keys = ["Enter Key", "Clear Key", "End Key", "F1", "F2", "F3", "F4", "F5", "F6", "F7", "F8", "F9", "F10", "F11", "F12", "F13", "F14", "F15", "F16", "F17", "F18", "F19", "F20", "F21", "F22", "F23", "F24"]
        self.original_test_case_name = test_case_name

        # âœ… NEW: Add execution tracking variables
        self.is_executing = False
        self.execution_stop_flag = False
        self.main_window = parent
        
        main_layout = QVBoxLayout(self)

        # âœ… NEW: Test Case Name and Description Section - WITH PLAY/STOP BUTTON
        info_layout = QHBoxLayout()

        name_label = QLabel("Test Case Name:")
        name_label.setFixedWidth(105)
        self.test_case_name_input = QLineEdit()
        self.test_case_name_input.setText(test_case_name)
        self.test_case_name_input.setPlaceholderText("Enter test case name...")
        self.test_case_name_input.setFixedWidth(200)

        desc_label = QLabel("Description:")
        desc_label.setFixedWidth(85)
        self.test_case_description_input = QLineEdit()
        self.test_case_description_input.setText(test_case_description)
        self.test_case_description_input.setPlaceholderText("Enter test case description...")
        self.test_case_description_input.setFixedWidth(300)
        
        info_layout.addWidget(name_label)
        info_layout.addWidget(self.test_case_name_input)
        info_layout.addWidget(desc_label)
        info_layout.addWidget(self.test_case_description_input)
        
        # âœ… NEW: Add stretch space
        info_layout.addStretch()
        
        # âœ… NEW: Add Start Step option
        start_step_label = QLabel("Start from Step:")
        start_step_label.setStyleSheet("font-weight: bold;")
        self.start_step_combo = QComboBox()
        self.start_step_combo.setFixedWidth(120)
        info_layout.addWidget(start_step_label)
        info_layout.addWidget(self.start_step_combo)
        
        # âœ… NEW: Add End Step option (DROPDOWN with 'Last' option)
        end_step_label = QLabel("End at Step:")
        end_step_label.setStyleSheet("font-weight: bold;")
        self.end_step_combo = QComboBox()
        self.end_step_combo.setFixedWidth(120)
        info_layout.addWidget(end_step_label)
        info_layout.addWidget(self.end_step_combo)
        
        # âœ… NEW: Add Play/Stop button
        self.execute_button = QPushButton()
        self.execute_button.setIcon(self.main_window.style().standardIcon(QStyle.StandardPixmap.SP_MediaPlay) if self.main_window else QIcon())
        self.execute_button.setFixedSize(QSize(40, 40))
        self.execute_button.setToolTip("Execute test case steps")
        self.execute_button.clicked.connect(self.toggle_execution)
        self.execute_button.setStyleSheet("""
            QPushButton {
                border: 2px solid #6B2C91;
                border-radius: 6px;
                background-color: white;
                padding: 2px;
            }
            QPushButton:hover {
                background-color: #f3e8ff;
            }
            QPushButton:pressed {
                background-color: #e9d5f5;
            }
        """)
        info_layout.addWidget(self.execute_button)

        main_layout.addLayout(info_layout)

        # Assumptions Section - Second Row
        assumptions_layout = QHBoxLayout()

        assumptions_label = QLabel("Assumptions:")
        assumptions_label.setFixedWidth(105)
        assumptions_label.setAlignment(Qt.AlignmentFlag.AlignTop)
        
        self.test_case_assumptions_input = QTextEdit()
        if test_case_assumptions:
            self.test_case_assumptions_input.setHtml(test_case_assumptions)
        self.test_case_assumptions_input.setPlaceholderText("Enter test case assumptions...")
        self.test_case_assumptions_input.setMaximumHeight(80)
        self.test_case_assumptions_input.setMinimumHeight(60)
        self.test_case_assumptions_input.setFixedWidth(599)

        assumptions_layout.addWidget(assumptions_label)
        assumptions_layout.addWidget(self.test_case_assumptions_input)
        assumptions_layout.addStretch()

        main_layout.addLayout(assumptions_layout)

        # Prerequisites Section - Tag/Chip Style
        prereq_layout = QHBoxLayout()
        prereq_label = QLabel("Prerequisites:")
        prereq_label.setFixedWidth(105)

        # Container widget with underline effect for prerequisites
        self.prerequisites_container = QWidget()
        self.prerequisites_container.setFixedHeight(28)
        self.prerequisites_container.setFixedWidth(200 + 300)
        self.prerequisites_container.setStyleSheet("""
            QWidget {
                border-bottom: 1px solid #888;
                background-color: transparent;
            }
        """)

        # Horizontal layout for prerequisite chips/tags
        self.prerequisites_chips_layout = QHBoxLayout(self.prerequisites_container)
        self.prerequisites_chips_layout.setContentsMargins(2, 2, 2, 2)
        self.prerequisites_chips_layout.setSpacing(5)
        self.prerequisites_chips_layout.addStretch()

        # Initialize prerequisite chips dictionary
        self.prerequisite_chips = {}

        # Add button
        self.add_prereq_button = QPushButton("Add")
        self.add_prereq_button.clicked.connect(self.add_prerequisite)
        self.add_prereq_button.setFixedWidth(60)

        prereq_layout.addWidget(prereq_label)
        prereq_layout.addWidget(self.prerequisites_container)
        prereq_layout.addWidget(self.add_prereq_button)
        prereq_layout.addStretch()

        main_layout.addLayout(prereq_layout)
        main_layout.addSpacing(5)

        # --- Top section for adding new steps ---
        top_layout = QHBoxLayout()
        import_label = QLabel("Import:")
        self.import_type_combobox = QComboBox()
        self.import_type_combobox.addItem("Select Import Type...")
        self.import_type_combobox.addItems(["Import Module", "Special Keys", "Capture Screenshot", "Capture Text Screenshot", "Random Input", "Wait", "Break"])
        
        self.dynamic_list_combobox = QComboBox()
        self.dynamic_list_combobox.setFixedWidth(200)
        self.dynamic_list_combobox.hide()
        
        self.test_step_number_input = QLineEdit()
        self.test_step_number_input.setPlaceholderText("Test Step #")
        self.test_step_number_input.setFixedWidth(95)
        self.test_step_number_input.setValidator(QIntValidator(1, 999))
        self.test_step_number_input.hide()        
        
        # Random Input text fields
        self.random_input_row = QLineEdit()
        self.random_input_row.setPlaceholderText("Row")
        self.random_input_row.setFixedWidth(60)
        self.random_input_row.setValidator(QIntValidator(1, 999))
        self.random_input_row.hide()
        
        self.random_input_col = QLineEdit()
        self.random_input_col.setPlaceholderText("Column")
        self.random_input_col.setFixedWidth(80)
        self.random_input_col.setValidator(QIntValidator(1, 999))
        self.random_input_col.hide()
        
        self.random_input_value = QLineEdit()
        self.random_input_value.setPlaceholderText("Value")
        self.random_input_value.setFixedWidth(150)
        self.random_input_value.hide()
        
        self.random_input_special_key_combo = QComboBox()
        self.random_input_special_key_combo.addItem("(Text)")
        self.random_input_special_key_combo.addItems(self.special_keys)
        self.random_input_special_key_combo.setFixedWidth(120)
        self.random_input_special_key_combo.hide()
        
        self.wait_seconds_input = QLineEdit()
        self.wait_seconds_input.setPlaceholderText("Seconds")
        self.wait_seconds_input.setText("1")
        self.wait_seconds_input.setFixedWidth(100)
        self.wait_seconds_input.hide()

        self.add_step_button = QPushButton("Add as Test Step")
        self.add_step_button.setFixedWidth(140)

        # Add as Utility Step button
        self.add_utility_step_button = QPushButton("Add as Utility Step")
        self.add_utility_step_button.setFixedWidth(150)
        self.add_utility_step_button.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #34D399, stop:1 #10B981);
                color: white;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #4ADE80, stop:1 #34D399);
            }
            QPushButton:pressed {
                background: #059669;
            }
        """)
        self.add_utility_step_button.hide()  # Hidden by default

        top_layout.addWidget(import_label)
        top_layout.addWidget(self.import_type_combobox)
        top_layout.addWidget(self.dynamic_list_combobox)
        top_layout.addWidget(self.random_input_row)
        top_layout.addWidget(self.random_input_col)
        top_layout.addWidget(self.random_input_value)
        top_layout.addWidget(self.random_input_special_key_combo)
        top_layout.addWidget(self.wait_seconds_input)
        top_layout.addWidget(self.test_step_number_input)
        top_layout.addWidget(self.add_step_button)
        top_layout.addWidget(self.add_utility_step_button)

        # Action Type Radio Buttons (moved to same row as Add button)
        top_layout.addSpacing(20)
        action_type_label = QLabel("Action Type:")
        self.input_radio = QRadioButton("Input")
        self.validate_radio = QRadioButton("Validate")
        self.input_radio.setChecked(True)
        top_layout.addWidget(action_type_label)
        top_layout.addWidget(self.input_radio)
        top_layout.addWidget(self.validate_radio)

        top_layout.addStretch()

        main_layout.addLayout(top_layout)

        # ✅ Connect signals - THIS IS POINT 3
        self.import_type_combobox.currentIndexChanged.connect(self.update_dynamic_list_combobox)
        self.add_step_button.clicked.connect(self.add_step_from_selection)
        self.add_utility_step_button.clicked.connect(self.add_utility_step_from_selection)
        
        # ✅ NEW: Connect radio button changes to update utility button visibility
        self.input_radio.toggled.connect(lambda: self.on_action_type_changed())
        self.validate_radio.toggled.connect(lambda: self.on_action_type_changed())

        # Splitter for vertical division
        splitter = QSplitter(Qt.Orientation.Horizontal)
        main_layout.addWidget(splitter)

        # Left side: List of added steps
        left_widget = QWidget()
        left_main_layout = QVBoxLayout(left_widget)

        # Horizontal splitter for main steps and utility steps
        lists_splitter = QSplitter(Qt.Orientation.Horizontal)

        # Main steps widget
        main_steps_widget = QWidget()
        main_steps_layout = QVBoxLayout(main_steps_widget)
        main_steps_layout.setContentsMargins(0, 0, 0, 0)
        
        # ✅ CHANGE 1: Rename the label
        left_label = QLabel("Main Test Steps:")  # ✅ Changed from "Added Test Steps:"
        
        # ✅ CHANGE 2: Add styling to make it visible and bold
        left_label.setStyleSheet("font-weight: bold; font-size: 11pt; color: #6B2C91;")
        
        # ✅ CHANGE 3: ADD THE LABEL TO THE LAYOUT (this was missing!)
        main_steps_layout.addWidget(left_label)
        
        self.steps_list_widget = CustomStepsListWidget(self)  # ✅ Use custom widget
        self.steps_list_widget.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        self.steps_list_widget.setSelectionBehavior(QListWidget.SelectionBehavior.SelectRows)
        self.steps_list_widget.currentItemChanged.connect(self.display_module_details)
        self.steps_list_widget.itemClicked.connect(self.on_main_step_clicked)
        main_steps_layout.addWidget(self.steps_list_widget)
        
        lists_splitter.addWidget(main_steps_widget)

        # ✅ NEW: Utility steps widget (right side)
        utility_steps_widget = QWidget()
        utility_steps_layout = QVBoxLayout(utility_steps_widget)
        utility_steps_layout.setContentsMargins(0, 0, 0, 0)

        # Header with label and dynamic buttons
        utility_header_layout = QHBoxLayout()
        utility_label = QLabel("Utility Steps:")
        utility_label.setStyleSheet("font-weight: bold; font-size: 11pt; color: #6B2C91;")  # ✅ Added styling
        utility_header_layout.addWidget(utility_label)

        # Dynamic buttons for quick add (initially hidden)
        self.utility_input_button = QPushButton("➕ Input")
        self.utility_input_button.setStyleSheet("""
            QPushButton {
                background-color: #10b981;
                color: white;
                border-radius: 4px;
                padding: 4px 12px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #059669;
            }
        """)
        self.utility_input_button.setFixedHeight(26)
        self.utility_input_button.clicked.connect(lambda: self.add_quick_utility_module('Input'))
        self.utility_input_button.hide()  # Hidden by default
        utility_header_layout.addWidget(self.utility_input_button)

        self.utility_validate_button = QPushButton("➕ Validate")
        self.utility_validate_button.setStyleSheet("""
            QPushButton {
                background-color: #f59e0b;
                color: white;
                border-radius: 4px;
                padding: 4px 12px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d97706;
            }
        """)
        self.utility_validate_button.setFixedHeight(26)
        self.utility_validate_button.clicked.connect(lambda: self.add_quick_utility_module('Validate'))
        self.utility_validate_button.hide()  # Hidden by default
        utility_header_layout.addWidget(self.utility_validate_button)
        
        utility_header_layout.addSpacing(12)

        # ✅ NEW: Add Enter icon button
        self.utility_enter_icon = QPushButton("⏎")
        self.utility_enter_icon.setFixedSize(32, 26)
        self.utility_enter_icon.setStyleSheet("""
            QPushButton {
                background-color: #10b981;
                color: white;
                border: none;
                border-radius: 4px;
                font-size: 18px;
                font-weight: bold;
                padding: 0px;
            }
            QPushButton:hover {
                background-color: #059669;
            }
            QPushButton:pressed {
                background-color: #047857;
            }
        """)
        self.utility_enter_icon.setToolTip("Add Enter key with auto-wait")
        self.utility_enter_icon.clicked.connect(self.add_quick_enter_utility)
        self.utility_enter_icon.hide()  # Hidden by default
        utility_header_layout.addWidget(self.utility_enter_icon)

        # ✅ NEW: Add Clear icon button
        self.utility_clear_icon = QPushButton("⌫")
        self.utility_clear_icon.setFixedSize(32, 26)
        self.utility_clear_icon.setStyleSheet("""
            QPushButton {
                background-color: #ef4444;
                color: white;
                border: none;
                border-radius: 4px;
                font-size: 18px;
                font-weight: bold;
                padding: 0px;
            }
            QPushButton:hover {
                background-color: #dc2626;
            }
            QPushButton:pressed {
                background-color: #b91c1c;
            }
        """)
        self.utility_clear_icon.setToolTip("Add Clear key with auto-wait")
        self.utility_clear_icon.clicked.connect(self.add_quick_clear_utility)
        self.utility_clear_icon.hide()  # Hidden by default
        utility_header_layout.addWidget(self.utility_clear_icon)

        utility_header_layout.addStretch()
        utility_steps_layout.addLayout(utility_header_layout)

        utility_header_layout.addStretch()
        utility_steps_layout.addLayout(utility_header_layout)

        self.utility_steps_list_widget = QListWidget()
        self.utility_steps_list_widget.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)
        self.utility_steps_list_widget.setSelectionBehavior(QListWidget.SelectionBehavior.SelectRows)
        self.utility_steps_list_widget.currentItemChanged.connect(self.display_utility_step_details)
        utility_steps_layout.addWidget(self.utility_steps_list_widget)
        lists_splitter.addWidget(utility_steps_widget)

        # Set initial splitter sizes (50/50 split)
        #lists_splitter.setSizes([300, 300])
        
        lists_splitter.setStretchFactor(0, 50)  # Main Test Steps gets 60%
        lists_splitter.setStretchFactor(1, 50)  # Utility Steps gets 40%

        left_main_layout.addWidget(lists_splitter)
        
        # Add Copy, Paste, Delete, Move Up and Move Down buttons
        step_buttons_layout = QHBoxLayout()

        self.copy_steps_button = QPushButton("📋 Copy")
        self.copy_steps_button.clicked.connect(self.copy_selected_steps)
        self.copy_steps_button.setToolTip("Copy selected step(s) (Ctrl+C)")
        step_buttons_layout.addWidget(self.copy_steps_button)

        self.paste_steps_button = QPushButton("📄 Paste")
        self.paste_steps_button.clicked.connect(self.paste_copied_steps)
        self.paste_steps_button.setToolTip("Paste copied step(s) at the end (Ctrl+V)")
        step_buttons_layout.addWidget(self.paste_steps_button)

        self.delete_steps_button = QPushButton("🗑️ Delete")
        self.delete_steps_button.clicked.connect(self.delete_selected_steps)
        self.delete_steps_button.setToolTip("Delete selected step(s) (Delete)")
        step_buttons_layout.addWidget(self.delete_steps_button)

        step_buttons_layout.addStretch()

        self.move_step_up_button = QPushButton("↑ Move Up")
        self.move_step_up_button.clicked.connect(self.move_step_up)
        step_buttons_layout.addWidget(self.move_step_up_button)

        self.move_step_down_button = QPushButton("↓ Move Down")
        self.move_step_down_button.clicked.connect(self.move_step_down)
        step_buttons_layout.addWidget(self.move_step_down_button)

        left_main_layout.addLayout(step_buttons_layout)
        
        splitter.addWidget(left_widget)

        # Right side: Details of the selected step in a table
        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_label = QLabel("Module Details:")
        right_label.setStyleSheet("font-weight: bold; font-size: 11pt; color: #6B2C91;")  # ✅ Added styling
        self.details_table = QTableWidget()
        self.details_table.setColumnCount(2)
        self.details_table.setHorizontalHeaderLabels(["Field Name", "Value"])
        self.details_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.details_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        
        right_layout.addWidget(right_label)
        right_layout.addWidget(self.details_table)

        # Save button for the fields
        save_button = QPushButton("Save Step Fields")
        save_button.clicked.connect(self.save_table_data_to_step_with_message)
        right_layout.addWidget(save_button)
        
        splitter.addWidget(right_widget)
        #splitter.setSizes([200, 600])
        splitter.setStretchFactor(0, 45)  # Left side (Main + Utility) gets 55%
        splitter.setStretchFactor(1, 55)  # Module Details gets 45%

        # Dialog buttons
        close_button_layout = QHBoxLayout()
        close_button_layout.addStretch()
        close_button = QPushButton("Close")
        close_button.setFixedWidth(60)
        close_button.clicked.connect(self.accept)
        close_button_layout.addWidget(close_button)
        close_button_layout.addSpacing(10)
        main_layout.addLayout(close_button_layout)
        
        self.update_steps_list()
        
        # Load existing prerequisites if editing
        self.main_window = parent
        if test_case_name and self.main_window and test_case_name in self.main_window.test_cases:
            existing_prereqs = self.main_window.test_cases[test_case_name].get('prerequisites', [])
            for prereq in existing_prereqs:
                self.add_prerequisite_chip(prereq)

    def validate_field_value(self, actual_value, expected_value):
        """
        Validates a field value, with special handling for {blank} validation.
        """
        actual_stripped = actual_value.strip()
        expected_stripped = expected_value.strip()
        
        if expected_stripped.lower() == '{blank}':
            return actual_stripped == '' or actual_stripped.isspace() or len(actual_stripped) == 0
        
        return actual_stripped == expected_stripped

    def on_main_step_clicked(self, item):
        """When a main step is clicked, populate utility steps list with numbered sub-steps."""
        if not item:
            self.utility_steps_list_widget.clear()
            # Hide quick utility buttons
            self.utility_input_button.hide()
            self.utility_validate_button.hide()
            self.utility_enter_icon.hide()  # ✅ NEW
            self.utility_clear_icon.hide()  # ✅ NEW
            return
        
        current_step_index = self.steps_list_widget.row(item)
        if current_step_index < 0 or current_step_index >= len(self.added_steps):
            self.utility_steps_list_widget.clear()
            self.utility_input_button.hide()
            self.utility_validate_button.hide()
            self.utility_enter_icon.hide()  # ✅ NEW
            self.utility_clear_icon.hide()  # ✅ NEW
            return
        
        step_data = self.added_steps[current_step_index]
        
        # ✅ NEW: Show/hide quick utility buttons based on step type
        if step_data.get('type') == 'module_import':
            self.utility_input_button.show()
            self.utility_validate_button.show()
            self.utility_enter_icon.show()  # ✅ NEW
            self.utility_clear_icon.show()  # ✅ NEW
        else:
            self.utility_input_button.hide()
            self.utility_validate_button.hide()
            self.utility_enter_icon.hide()  # ✅ NEW
            self.utility_clear_icon.hide()  # ✅ NEW
        
        utility_steps = step_data.get('utility_steps', [])
        
        # Populate utility steps list with sub-step numbering
        self.utility_steps_list_widget.clear()
        main_step_num = current_step_index + 1
        
        for sub_idx, utility_step in enumerate(utility_steps, 1):
            step_name = utility_step.get('name', 'Utility')
            
            # Add sub-step number prefix (e.g., "Step 1.1: ...")
            numbered_name = f"Step {main_step_num}.{sub_idx}: {step_name}"
            
            # ✅ NEW: Check if this is a validate utility step
            is_validate_utility = False
            if utility_step.get('type') == 'module_import':
                fields = utility_step.get('fields', [])
                if fields and fields[0].get('action_type') == 'Validate':
                    is_validate_utility = True
            
            # Create custom widget with delete button
            list_item = QListWidgetItem()
            item_widget = QWidget()
            item_layout = QHBoxLayout(item_widget)
            item_layout.setContentsMargins(4, 2, 4, 2)
            item_layout.setSpacing(5)
            
            # Step name label
            name_label = QLabel(numbered_name)
            name_label.setMinimumWidth(200)
            name_label.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
            name_label.setWordWrap(False)
            item_layout.addWidget(name_label, 1)
            
            item_layout.addStretch()
            
            # Delete button
            delete_button = QPushButton("✕")
            delete_button.setFixedSize(18, 18)
            delete_font = QFont()
            delete_font.setBold(True)
            delete_font.setPointSize(12)
            delete_button.setFont(delete_font)
            delete_button.setStyleSheet("""
                QPushButton {
                    color: #dc2626;
                    border: none;
                    background-color: transparent;
                    padding: 0px;
                }
                QPushButton:hover {
                    color: #991b1b;
                    background-color: #fee2e2;
                    border-radius: 3px;
                }
            """)
            delete_button.setToolTip(f"Delete utility step {main_step_num}.{sub_idx}")
            delete_button.clicked.connect(
                lambda checked, ms=current_step_index, us=sub_idx-1: self.delete_utility_step(ms, us)
            )
            item_layout.addWidget(delete_button)
            
            # ✅ NEW: Apply yellow background for validate utility steps
            if is_validate_utility:
                item_widget.setStyleSheet("background-color: #fffacd;")
            
            list_item.setData(Qt.ItemDataRole.UserRole, {
                'main_step_index': current_step_index,
                'utility_step': utility_step,
                'sub_index': sub_idx
            })
            
            self.utility_steps_list_widget.addItem(list_item)
            self.utility_steps_list_widget.setItemWidget(list_item, item_widget)
            size_hint = list_item.sizeHint()
            size_hint.setHeight(max(32, size_hint.height()))
            list_item.setSizeHint(size_hint)

    def display_utility_step_details(self, current_item, previous_item):
        """Displays utility step details in the Module Details table."""
        if not current_item:
            return
        
        item_data = current_item.data(Qt.ItemDataRole.UserRole)
        if not item_data:
            return
        
        # ✅ CRITICAL FIX: Get the actual utility step from added_steps, not from item_data
        main_step_index = item_data['main_step_index']
        sub_index = item_data.get('sub_index', 1)  # 1-based sub_index
        
        if main_step_index >= len(self.added_steps):
            return
        
        main_step = self.added_steps[main_step_index]
        utility_steps = main_step.get('utility_steps', [])
        
        # Convert sub_index (1-based) to 0-based array index
        utility_step_index = sub_index - 1
        
        if utility_step_index < 0 or utility_step_index >= len(utility_steps):
            return
        
        # ✅ Get the actual utility step reference (not a copy)
        utility_step = utility_steps[utility_step_index]
        utility_type = utility_step.get('type')
        
        # Clear the table
        self.details_table.setRowCount(0)
        
        if utility_type == 'capture_screenshot':
            # Find reference module
            reference_module_name = None
            
            for i in range(main_step_index, -1, -1):
                check_step = self.added_steps[i]
                if check_step.get('type') == 'module_import':
                    reference_module_name = check_step.get('module_name')
                    break
            
            if reference_module_name:
                utility_step['reference_module'] = reference_module_name
                self.display_utility_screenshot_details(utility_step, reference_module_name)
        
        elif utility_type == 'module_import':
            module_name = utility_step.get('module_name')
            if module_name:
                self.display_utility_module_details(utility_step, module_name)

    def get_available_variables_info(self):
        """Returns a formatted string of available variables for tooltips."""
        return (
            "Available Variables:\n"
            "{date} - Current date (YYYY.MM.DD)\n"
            "{time} - Current time (HH:MM:SS)\n"
            "{datetime} - Date and time\n"
        )

    def substitute_field_variables(self, text):
        """
        Substitutes variables in field values during test execution.
        This is a preview - actual substitution happens during execution.
        """
        from datetime import datetime
        
        now = datetime.now()
        test_name = self.test_case_name_input.text().strip()
        test_desc = self.test_case_description_input.text().strip()
        
        variables = {
            'date': now.strftime('%Y.%m.%d'),
            'time': now.strftime('%H:%M:%S'),
            'datetime': now.strftime('%Y-%m-%d %H:%M:%S'),
            'test_case_id': test_name,
            'test_description': test_desc
        }
        
        result = text
        for var_name, var_value in variables.items():
            placeholder = '{' + var_name + '}'
            result = result.replace(placeholder, var_value)
        
        return result

    def keyPressEvent(self, event):
        """Handle keyboard shortcuts for copy/paste/delete operations."""
        # Check if Ctrl is pressed
        if event.modifiers() == Qt.KeyboardModifier.ControlModifier:
            if event.key() == Qt.Key.Key_C:
                # Ctrl+C: Copy
                self.copy_selected_steps()
                event.accept()
                return
            elif event.key() == Qt.Key.Key_V:
                # Ctrl+V: Paste
                self.paste_copied_steps()
                event.accept()
                return
        elif event.key() == Qt.Key.Key_Delete:
            # Delete key: Delete selected steps
            self.delete_selected_steps()
            event.accept()
            return
        
        # Call the parent implementation for other keys
        super().keyPressEvent(event)
        
    def on_action_type_changed(self, checked):
        """
        ✅ FIXED: Updates the utility step button visibility when action type changes.
        Called when Input/Validate radio buttons are toggled.
        
        Args:
            checked (bool): The checked state of the radio button
        """
        selected_type = self.import_type_combobox.currentText()
        
        # If Import Module is selected, update utility button visibility based on action type
        if selected_type == "Import Module":
            if self.validate_radio.isChecked():
                self.add_utility_step_button.show()
            else:
                self.add_utility_step_button.hide()       

    def update_dynamic_list_combobox(self, index):
        selected_type = self.import_type_combobox.currentText()
        self.dynamic_list_combobox.clear()
        
        # Hide all input fields by default
        self.random_input_row.hide()
        self.random_input_col.hide()
        self.random_input_value.hide()
        self.random_input_special_key_combo.hide()
        self.wait_seconds_input.hide()
        
        self.test_step_number_input.hide()
        
        # ✅ UPDATED: Show utility button for Import Module regardless of action type
        if selected_type == "Import Module":
            self.add_utility_step_button.show()
        else:
            # For other utility step types, always show the button
            utility_step_types = ["Special Keys", "Capture Screenshot", "Capture Text Screenshot", "Random Input","Wait"]
            if selected_type in utility_step_types:
                self.add_utility_step_button.show()
            else:
                self.add_utility_step_button.hide()
        
        if selected_type == "Import Module":
            # Make it editable
            self.dynamic_list_combobox.setEditable(True)
            self.dynamic_list_combobox.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
            
            # Add all modules
            self.dynamic_list_combobox.addItems(self.modules.keys())
            
            # Set placeholder text
            self.dynamic_list_combobox.lineEdit().setPlaceholderText("Select a module...")
            self.dynamic_list_combobox.clearEditText()
            
            # Configure completer for filtering
            self.dynamic_list_combobox.completer().setCompletionMode(
                self.dynamic_list_combobox.completer().CompletionMode.PopupCompletion
            )
            self.dynamic_list_combobox.completer().setFilterMode(Qt.MatchFlag.MatchContains)
            
            self.dynamic_list_combobox.show()
            self.dynamic_list_combobox.setFixedWidth(200)
            
            self.test_step_number_input.show()
        elif selected_type == "Special Keys":
            self.dynamic_list_combobox.setEditable(False)
            self.dynamic_list_combobox.addItems(self.special_keys)
            self.dynamic_list_combobox.setFixedWidth(120)
            self.dynamic_list_combobox.show()
            
            self.test_step_number_input.show()            
        elif selected_type == "Capture Text Screenshot":
            self.dynamic_list_combobox.hide()
            
            self.test_step_number_input.show()            
        elif selected_type == "Capture Screenshot":
            self.dynamic_list_combobox.hide()
            
            self.test_step_number_input.show()            
        elif selected_type == "Random Input":
            self.dynamic_list_combobox.hide()
            self.random_input_row.show()
            self.random_input_col.show()
            self.random_input_value.show()
            self.random_input_special_key_combo.show()
            
            self.test_step_number_input.show()            
        elif selected_type == "Wait":
            self.dynamic_list_combobox.hide()
            self.wait_seconds_input.setText("1")
            self.wait_seconds_input.show()
            
            self.test_step_number_input.show()   
        elif selected_type == "Break":
            self.dynamic_list_combobox.hide()
            self.test_step_number_input.show()            
        else:
            self.dynamic_list_combobox.hide()

    def add_step_from_selection(self):
        selected_type = self.import_type_combobox.currentText()
        
        if selected_type == "Import Module":
            # Get the current text from the line edit
            typed_text = self.dynamic_list_combobox.currentText().strip()
            
            # Find exact match in modules (case-sensitive)
            selected_item = None
            for module_name in self.modules.keys():
                if module_name == typed_text:
                    selected_item = module_name
                    break
            
            # If no exact match found, try the current index
            if not selected_item:
                current_index = self.dynamic_list_combobox.currentIndex()
                if current_index >= 0:
                    module_names = list(self.modules.keys())
                    if current_index < len(module_names):
                        selected_item = module_names[current_index]
            
            # If still no match, use the typed text as-is (will fail validation below)
            if not selected_item:
                selected_item = typed_text
        else:
            selected_item = self.dynamic_list_combobox.currentText().strip()
        
        # Handle "Capture Text Screenshot" separately
        if selected_type == "Capture Text Screenshot":
            self.add_capture_text_screenshot_step()
            return
        
        # Handle "Capture Screenshot" separately
        if selected_type == "Capture Screenshot":
            self.add_capture_screenshot_step()
            return
        
        # Handle "Random Input" separately
        if selected_type == "Random Input":
            self.add_random_input_step()
            return
            
        if selected_type == "Wait":
            self.add_wait_step()
            return
            
        if selected_type == "Break":
            self.add_break_step()
            return            
        
        # Check for empty or whitespace-only selection
        if selected_type == "Select Import Type..." or not selected_item:
            QMessageBox.warning(self, "No Selection", "Please select a specific item to add.")
            return

        # Get the selected action type from the radio buttons
        action_type = 'Input' if self.input_radio.isChecked() else 'Validate'

        # Get the target step number
        step_number_text = self.test_step_number_input.text().strip()
        insert_index = None

        if step_number_text:
            try:
                step_number = int(step_number_text)
                if step_number < 1:
                    QMessageBox.warning(self, "Invalid Step Number", "Step number must be at least 1.")
                    return
                insert_index = step_number - 1
                if insert_index > len(self.added_steps):
                    insert_index = len(self.added_steps)
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Please enter a valid step number.")
                return

        if selected_type == "Import Module":
            # Validate that the selected module exists
            if selected_item not in self.modules:
                QMessageBox.warning(self, "Invalid Module", f"Module '{selected_item}' does not exist.")
                return
                
            module_details = self.modules.get(selected_item, {})
            fields_from_module = copy.deepcopy(module_details.get('labels', []))
            
            # ✅ CRITICAL FIX: Calculate step number AFTER determining insert position
            # âœ… CRITICAL FIX: Calculate step number AFTER determining insert position
            if insert_index is not None:
                step_number = insert_index + 1
            else:
                step_number = len(self.added_steps) + 1

            fields = []
            for idx, label_data in enumerate(fields_from_module):
                field_name = label_data.get('label') or label_data.get('text') or label_data.get('name', 'N/A')
                
                # âœ… FIXED: Generate unique timestamp and random PER FIELD
                import time
                import random
                unique_timestamp = int(time.time() * 1000000)  # Microseconds
                random_suffix = random.randint(100000, 999999)  # 6-digit random
                
                # âœ… CRITICAL: Include module name in unique ID
                unique_field_id = f"{field_name}_{selected_item}_M{step_number}_{unique_timestamp}_{random_suffix}_{idx}"
                
                fields.append({
                    "field_name": field_name,  # Keep original for display
                    "internal_field_id": unique_field_id,  # ✅ UNIQUE identifier
                    "action_type": action_type,
                    "value": "",
                })

            new_step = {
                "name": f"Import Module: {selected_item}",
                "type": "module_import",
                "module_name": selected_item,
                "fields": fields,
                "utility_steps": []
            }
            
            # Check if this is a Validate step and add utility wait
            is_validate = action_type == 'Validate'
            if is_validate:
                # Add wait as a utility step instead of a separate test step
                wait_utility = {
                    "name": "Wait: 1 second(s)",
                    "type": "wait",
                    "seconds": 1
                }
                new_step['utility_steps'].append(wait_utility)
            
            # Insert at specific index or append
            if insert_index is not None:
                self.added_steps.insert(insert_index, new_step)
                if is_validate:
                    QMessageBox.information(self, "Success", 
                        f"'{selected_item}' inserted at step {step_number} with utility wait (Validate).")
                else:
                    QMessageBox.information(self, "Success", 
                        f"'{selected_item}' inserted at step {step_number}.")
            else:
                self.added_steps.append(new_step)
                if is_validate:
                    QMessageBox.information(self, "Success", 
                        f"'{selected_item}' added as a test step with utility wait (Validate).")
                else:
                    QMessageBox.information(self, "Success", 
                        f"'{selected_item}' added as a test step.")

        elif selected_type == "Special Keys":
            action_type = 'Input' if self.input_radio.isChecked() else 'Validate'
            
            new_step = {
                "name": f"Special Key: {selected_item}",
                "type": "special_key",
                "key_value": selected_item,
                "fields": [{
                    "field_name": "special_key",
                    "action_type": action_type,
                    "value": ""
                }],
                "utility_steps": []
            }
            
            # Always add wait as utility step for special keys
            wait_utility = {
                "name": "Wait: 1 second(s)",
                "type": "wait",
                "seconds": 1
            }
            new_step['utility_steps'].append(wait_utility)
            
            # Insert at specific index or append
            if insert_index is not None:
                self.added_steps.insert(insert_index, new_step)
                QMessageBox.information(self, "Success", 
                    f"'{selected_item}' inserted at step {step_number} with utility wait.")
            else:
                self.added_steps.append(new_step)
                QMessageBox.information(self, "Success", 
                    f"'{selected_item}' added as a special key step with utility wait.")
        
        # Clear the test step number input after adding
        self.test_step_number_input.clear()
        
        self.update_steps_list()
        
        self.input_radio.setChecked(True)
        self.update_step_combo_options()
    
                
    def add_random_input_step(self):
        """Adds a random input step to the test case."""
        row_text = self.random_input_row.text().strip()
        col_text = self.random_input_col.text().strip()
        value_text = self.random_input_value.text().strip()
        special_key = self.random_input_special_key_combo.currentText()
        
        # Validate that row and column are filled
        if not row_text or not col_text:
            QMessageBox.warning(self, "Missing Information", 
                              "Please provide Row and Column for the random input.")
            return
        
        # Check if either value or special key is selected
        if not value_text and special_key == "(Text)":
            QMessageBox.warning(self, "Missing Information", 
                              "Please provide either a Value or select a Special Key.")
            return
        
        try:
            row = int(row_text)
            col = int(col_text)
        except ValueError:
            QMessageBox.warning(self, "Invalid Input", "Row and Column must be valid numbers.")
            return
        
        # ✅ NEW: Get the target step number
        step_number_text = self.test_step_number_input.text().strip()
        insert_index = None

        if step_number_text:
            try:
                step_number = int(step_number_text)
                if step_number < 1:
                    QMessageBox.warning(self, "Invalid Step Number", "Step number must be at least 1.")
                    return
                insert_index = step_number - 1
                if insert_index > len(self.added_steps):
                    insert_index = len(self.added_steps)
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Please enter a valid step number.")
                return
        
        # Get the selected action type from the radio buttons
        action_type = 'Input' if self.input_radio.isChecked() else 'Validate'
        
        # Determine what to use - text value or special key
        if special_key != "(Text)":
            # Use special key
            display_value = f"[{special_key}]"
            actual_value = special_key
            is_special_key = True
        else:
            # Use text value
            display_value = value_text
            actual_value = value_text
            is_special_key = False
        
        # Create the new step
        new_step = {
            "name": f"Random Input: Row {row}, Col {col}, Value: {display_value}",
            "type": "random_input",
            "row": row,
            "column": col,
            "value": actual_value,
            "is_special_key": is_special_key,  # NEW: Flag to indicate if it's a special key
            "fields": [{
                "field_name": "random_input",
                "action_type": action_type,
                "value": actual_value
            }],
            "utility_steps": []
        }
        
        # ✅ NEW: Insert at specific index or append
        if insert_index is not None:
            self.added_steps.insert(insert_index, new_step)
            success_msg = f"Random input step inserted at step {step_number}."
        else:
            self.added_steps.append(new_step)
            success_msg = f"Random input step added at Row {row}, Column {col}."

        self.update_steps_list()

        # Clear the input fields
        self.random_input_row.clear()
        self.random_input_col.clear()
        self.random_input_value.clear()
        self.random_input_special_key_combo.setCurrentIndex(0)
        self.test_step_number_input.clear()  # ✅ Clear step number
        self.input_radio.setChecked(True)
        QMessageBox.information(self, "Success", 
                              f"Random input step added at Row {row}, Column {col}.")
                              
    def add_wait_step(self):
        """Adds a wait step to the test case."""
        seconds_text = self.wait_seconds_input.text().strip()
        
        # Validate that seconds field is filled
        if not seconds_text:
            QMessageBox.warning(self, "Missing Information", 
                              "Please provide the number of seconds to wait.")
            return
        
        try:
            seconds = float(seconds_text)
            if seconds <= 0:
                QMessageBox.warning(self, "Invalid Input", "Seconds must be a positive number.")
                return
        except ValueError:
            QMessageBox.warning(self, "Invalid Input", "Seconds must be a valid number.")
            return

        # âœ… NEW: Get the target step number
        step_number_text = self.test_step_number_input.text().strip()
        insert_index = None

        if step_number_text:
            try:
                step_number = int(step_number_text)
                if step_number < 1:
                    QMessageBox.warning(self, "Invalid Step Number", "Step number must be at least 1.")
                    return
                insert_index = step_number - 1
                if insert_index > len(self.added_steps):
                    insert_index = len(self.added_steps)
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Please enter a valid step number.")
                return

        # Create the new wait step
        new_step = {
            "name": f"Wait: {seconds} second(s)",
            "type": "wait",
            "seconds": seconds,
            "fields": [],
            "utility_steps": []
        }

        # âœ… NEW: Insert at specific index or append
        if insert_index is not None:
            self.added_steps.insert(insert_index, new_step)
            success_msg = f"Wait step inserted at step {step_number}."
        else:
            self.added_steps.append(new_step)
            success_msg = f"Wait step added for {seconds} second(s)."

        self.update_steps_list()

        # Clear the input fields
        self.wait_seconds_input.clear()
        self.test_step_number_input.clear()
        self.input_radio.setChecked(True)
        
        # ✅ ADD THIS LINE - Show the success message
        QMessageBox.information(self, "Success", success_msg)
                              
    def add_capture_screenshot_step(self):
        """Adds a screenshot capture step (DOCX) to the test case."""
        # âœ… NEW: Get the target step number
        step_number_text = self.test_step_number_input.text().strip()
        insert_index = None
        
        if step_number_text:
            try:
                step_number = int(step_number_text)
                if step_number < 1:
                    QMessageBox.warning(self, "Invalid Step Number", "Step number must be at least 1.")
                    return
                insert_index = step_number - 1
                if insert_index > len(self.added_steps):
                    insert_index = len(self.added_steps)
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Please enter a valid step number.")
                return
        
        new_step = {
            "name": "Capture Screenshot (DOCX)",
            "type": "capture_screenshot",
            "fields": [],
            "reference_module": None,
            "utility_steps": []
        }
        
        # âœ… NEW: Insert at specific index or append
        if insert_index is not None:
            self.added_steps.insert(insert_index, new_step)
            success_msg = f"Screenshot capture step inserted at step {step_number}."
        else:
            self.added_steps.append(new_step)
            success_msg = "Screenshot capture step added (will generate DOCX)."
        
        self.update_steps_list()
        self.test_step_number_input.clear()  # âœ… NEW
        self.input_radio.setChecked(True)
        QMessageBox.information(self, "Success", "Screenshot capture step added (will generate DOCX).")

        
    def on_action_type_toggled(self):
        """
        Handles the visibility of the expected value field based on radio button selection.
        """
        # No change to table visibility needed with the new layout.
        pass
        
    def update_steps_list(self):
        """
        Clears and repopulates the steps list widget with the current steps.
        ✅ CHANGED: Utility steps now numbered as sub-steps (1.1, 1.2, etc.)
        """
        
        current_row = self.steps_list_widget.currentRow()
        
        self.steps_list_widget.clear()
        self.utility_steps_list_widget.clear()
        self.update_step_combo_options()
        
        for i, step in enumerate(self.added_steps):
            step_type = step.get('type')
            step_name = f"Step {i + 1}: "
            
            is_validate_step = False
            if step.get('fields') and step['fields'][0].get('action_type') == 'Validate':
                is_validate_step = True

            if step_type == 'module_import':
                module_name = step.get('module_name', 'Unknown Module')
                step_name += module_name
            elif step_type == 'special_key':
                key_value = step.get('key_value', 'Unknown Key')
                step_name += f"Special Key: {key_value}"
            elif step_type == 'capture_screen_text':
                step_name += "Capture Text Screenshot"
            elif step_type == 'capture_screenshot':
                step_name += "Capture Screenshot (DOCX)"
            elif step_type == 'random_input':
                row = step.get('row', '?')
                col = step.get('column', '?')
                value = step.get('value', '?')
                step_name += f"Random Input (Row: {row}, Col: {col}, Value: {value})"
            elif step_type == 'wait':
                seconds = step.get('seconds', '?')
                step_name += f"Wait: {seconds} second(s)"
            elif step_type == 'break':
                step_name += "Break: Review & Decision Point"
            else:
                step_name += "Unknown Step"

            # Create custom widget with delete button
            list_item = QListWidgetItem()
            item_widget = QWidget()
            item_layout = QHBoxLayout(item_widget)
            item_layout.setContentsMargins(4, 2, 8, 2)
            item_layout.setSpacing(8)
            
            # Step name label
            # Step name label
            name_label = QLabel(step_name)
            name_label.setMinimumWidth(300)
            name_label.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)  # ✅ NEW
            name_label.setWordWrap(False)  # ✅ NEW: Prevent wrapping
            item_layout.addWidget(name_label, 1)  # ✅ CHANGED: Added stretch factor of 1
            
            item_layout.addStretch()
            
            # Delete button
            delete_button = QPushButton("✕")
            delete_button.setFixedSize(20, 20)
            delete_font = QFont()
            delete_font.setBold(True)
            delete_font.setPointSize(13)
            delete_button.setFont(delete_font)
            delete_button.setStyleSheet("""
                QPushButton {
                    color: #dc2626;
                    border: none;
                    background-color: transparent;
                    padding: 0px;
                }
                QPushButton:hover {
                    color: #991b1b;
                    background-color: #fee2e2;
                    border-radius: 3px;
                }
            """)
            delete_button.setToolTip(f"Delete '{step_name}'")
            delete_button.clicked.connect(lambda checked, idx=i: self.delete_step_by_index(idx))
            item_layout.addWidget(delete_button)
            
            if is_validate_step:
                item_widget.setStyleSheet("background-color: #fffacd;")
            
            self.steps_list_widget.addItem(list_item)
            self.steps_list_widget.setItemWidget(list_item, item_widget)
            # ✅ CHANGED: Set minimum height for better visibility
            size_hint = item_widget.sizeHint()
            size_hint.setHeight(max(32, size_hint.height()))  # Ensure minimum 40px height
            list_item.setSizeHint(size_hint)
            
        if 0 <= current_row < self.steps_list_widget.count():
            self.steps_list_widget.setCurrentRow(current_row)    

    def delete_single_test_step(self, step_name):
        """Deletes a single test step from the Added Test Steps list."""
        # ✅ FIXED: Find step by visual position, not by reconstructed name
        clicked_row = -1
        
        # Find which row was clicked by checking all items
        for i in range(self.steps_list_widget.count()):
            item = self.steps_list_widget.item(i)
            widget = self.steps_list_widget.itemWidget(item)
            if widget:
                # Check if this widget's label matches the step_name
                label = widget.findChild(QLabel)
                if label and label.text() == step_name:
                    clicked_row = i
                    break
        
        if clicked_row == -1:
            return
        
        # ✅ CRITICAL FIX: The visual row IS the correct index in added_steps
        # because update_steps_list() builds the list in the same order as added_steps
        step_index = clicked_row
        
        if step_index < 0 or step_index >= len(self.added_steps):
            return
        
        # Confirm deletion
        reply = QMessageBox.question(
            self,
            "Delete Step",
            f"Are you sure you want to delete '{step_name}'?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            del self.added_steps[step_index]
            self.update_steps_list()
            self.details_table.setRowCount(0)
            self.statusBar().showMessage(f"Step deleted.", 3000)


    def delete_step_by_index(self, step_index):
        """Deletes a test step by its index in the added_steps array."""
        if step_index < 0 or step_index >= len(self.added_steps):
            return
        
        step_data = self.added_steps[step_index]
        step_type = step_data.get('type')
        
        # Build step name for confirmation dialog
        step_name = f"Step {step_index + 1}: "
        if step_type == 'module_import':
            module_name = step_data.get('module_name', 'Unknown Module')
            step_name += module_name
        elif step_type == 'special_key':
            key_value = step_data.get('key_value', 'Unknown Key')
            step_name += f"Special Key: {key_value}"
        elif step_type == 'capture_screen_text':
            step_name += "Capture Text Screenshot"
        elif step_type == 'capture_screenshot':
            step_name += "Capture Screenshot (DOCX)"
        elif step_type == 'random_input':
            row = step_data.get('row', '?')
            col = step_data.get('column', '?')
            value = step_data.get('value', '?')
            step_name += f"Random Input (Row: {row}, Col: {col}, Value: {value})"
        elif step_type == 'wait':
            seconds = step_data.get('seconds', '?')
            step_name += f"Wait: {seconds} second(s)"
        elif step_type == 'break':
            step_name += "Break: Review & Decision Point"
        else:
            step_name += "Unknown Step"
        
        # Confirm deletion
        reply = QMessageBox.question(
            self,
            "Delete Step",
            f"Are you sure you want to delete '{step_name}'?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            del self.added_steps[step_index]
            self.update_steps_list()
            self.details_table.setRowCount(0)
            self.statusBar().showMessage(f"Step deleted.", 3000)

    def delete_utility_step(self, main_step_index, utility_step_index):
        """Deletes a utility step from a main step."""
        if main_step_index >= len(self.added_steps):
            return
        
        step = self.added_steps[main_step_index]
        utility_steps = step.get('utility_steps', [])
        
        if utility_step_index < 0 or utility_step_index >= len(utility_steps):
            return
        
        # Confirm deletion
        utility_step = utility_steps[utility_step_index]
        step_name = utility_step.get('name', 'Utility Step')
        
        reply = QMessageBox.question(
            self,
            "Delete Utility Step",
            f"Are you sure you want to delete utility step '{step_name}'?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            del utility_steps[utility_step_index]
            
            # ✅ FIXED: Use refresh_utility_steps_for_main_step instead of update_steps_list
            # This keeps the utility panel visible
            self.steps_list_widget.blockSignals(True)
            self.steps_list_widget.setCurrentRow(main_step_index)
            self.steps_list_widget.blockSignals(False)
            
            # Manually refresh utility steps list while keeping it visible
            self.refresh_utility_steps_for_main_step(main_step_index)
            
            self.statusBar().showMessage(f"Utility step deleted.", 3000)

    def delete_step(self, index):
        """
        Deletes a step from the list and the data model.
        """
        # Remove the step from the data list
        if 0 <= index < len(self.added_steps):
            self.added_steps.pop(index)
            # Clear the details table if the deleted item was selected
            if self.steps_list_widget.currentRow() == index:
                self.details_table.setRowCount(0)
            
            # Update the list widget and step numbers
            self.update_steps_list()
        self.update_step_combo_options()
        
    def move_step_up(self):
        """Moves the selected step up in the list (works for both main and utility steps)."""
        # Check which list has focus
        if self.utility_steps_list_widget.hasFocus() or len(self.utility_steps_list_widget.selectedItems()) > 0:
            # Moving a utility step
            self.move_utility_step_up()
        else:
            # Moving a main step
            self.move_main_step_up()

    def move_main_step_up(self):
        """Moves the selected main step up."""
        current_row = self.steps_list_widget.currentRow()
        
        # Can't move if nothing selected or already at top
        if current_row <= 0:
            return
        
        # Check if current item is a project header
        current_item = self.steps_list_widget.item(current_row)
        current_item_data = current_item.data(Qt.ItemDataRole.UserRole)
        
        if current_item_data and current_item_data.get('type') == 'project':
            QMessageBox.warning(self, "Cannot Move", "Project headers cannot be moved.")
            return
        
        # Swap the steps in the data
        self.added_steps[current_row], self.added_steps[current_row - 1] = \
            self.added_steps[current_row - 1], self.added_steps[current_row]
        
        # Update the UI
        self.update_steps_list()
        
        # Re-select the moved item at its new position
        self.steps_list_widget.setCurrentRow(current_row - 1)

    def move_utility_step_up(self):
        """Moves the selected utility step up within its parent main step."""
        current_utility_item = self.utility_steps_list_widget.currentItem()
        if not current_utility_item:
            return
        
        current_utility_row = self.utility_steps_list_widget.currentRow()
        
        # Can't move if already at top
        if current_utility_row <= 0:
            return
        
        # Get the item data
        item_data = current_utility_item.data(Qt.ItemDataRole.UserRole)
        if not item_data:
            return
        
        main_step_index = item_data['main_step_index']
        current_sub_index = item_data['sub_index']  # 1-based
        
        # Get the main step's utility steps array
        if main_step_index >= len(self.added_steps):
            return
        
        main_step = self.added_steps[main_step_index]
        utility_steps = main_step.get('utility_steps', [])
        
        # Convert to 0-based index for array manipulation
        current_idx = current_sub_index - 1
        target_idx = current_idx - 1
        
        if target_idx < 0 or current_idx >= len(utility_steps):
            return
        
        # Swap the utility steps in the array
        utility_steps[current_idx], utility_steps[target_idx] = \
            utility_steps[target_idx], utility_steps[current_idx]
        
        # ✅ FIXED: Don't call update_steps_list, just refresh utility steps
        # Update only the main steps list to keep it in sync
        self.steps_list_widget.blockSignals(True)
        self.steps_list_widget.setCurrentRow(main_step_index)
        self.steps_list_widget.blockSignals(False)
        
        # ✅ Manually refresh utility steps list while keeping it visible
        self.refresh_utility_steps_for_main_step(main_step_index)
        
        # Re-select the moved utility item at its new position
        self.utility_steps_list_widget.setCurrentRow(current_utility_row - 1)
        
    def move_step_down(self):
        """Moves the selected step down in the list (works for both main and utility steps)."""
        # Check which list has focus
        if self.utility_steps_list_widget.hasFocus() or len(self.utility_steps_list_widget.selectedItems()) > 0:
            # Moving a utility step
            self.move_utility_step_down()
        else:
            # Moving a main step
            self.move_main_step_down()

    def move_main_step_down(self):
        """Moves the selected main step down."""
        current_row = self.steps_list_widget.currentRow()
        
        # Can't move if nothing selected or already at bottom
        if current_row < 0 or current_row >= len(self.added_steps) - 1:
            return
        
        # Check if current item is a project header
        current_item = self.steps_list_widget.item(current_row)
        current_item_data = current_item.data(Qt.ItemDataRole.UserRole)
        
        if current_item_data and current_item_data.get('type') == 'project':
            QMessageBox.warning(self, "Cannot Move", "Project headers cannot be moved.")
            return
        
        # Swap the steps in the data
        self.added_steps[current_row], self.added_steps[current_row + 1] = \
            self.added_steps[current_row + 1], self.added_steps[current_row]
        
        # Update the UI
        self.update_steps_list()
        
        # Re-select the moved item at its new position
        self.steps_list_widget.setCurrentRow(current_row + 1)

    def move_utility_step_down(self):
        """Moves the selected utility step down within its parent main step."""
        current_utility_item = self.utility_steps_list_widget.currentItem()
        if not current_utility_item:
            return
        
        current_utility_row = self.utility_steps_list_widget.currentRow()
        
        # Get the item data
        item_data = current_utility_item.data(Qt.ItemDataRole.UserRole)
        if not item_data:
            return
        
        main_step_index = item_data['main_step_index']
        current_sub_index = item_data['sub_index']  # 1-based
        
        # Get the main step's utility steps array
        if main_step_index >= len(self.added_steps):
            return
        
        main_step = self.added_steps[main_step_index]
        utility_steps = main_step.get('utility_steps', [])
        
        # Convert to 0-based index for array manipulation
        current_idx = current_sub_index - 1
        target_idx = current_idx + 1
        
        # Can't move if already at bottom
        if current_idx < 0 or target_idx >= len(utility_steps):
            return
        
        # Swap the utility steps in the array
        utility_steps[current_idx], utility_steps[target_idx] = \
            utility_steps[target_idx], utility_steps[current_idx]
        
        # ✅ FIXED: Don't call update_steps_list, just refresh utility steps
        # Update only the main steps list to keep it in sync
        self.steps_list_widget.blockSignals(True)
        self.steps_list_widget.setCurrentRow(main_step_index)
        self.steps_list_widget.blockSignals(False)
        
        # ✅ Manually refresh utility steps list while keeping it visible
        self.refresh_utility_steps_for_main_step(main_step_index)
        
        # Re-select the moved utility item at its new position
        self.utility_steps_list_widget.setCurrentRow(current_utility_row + 1)
        
    def copy_selected_steps(self):
        """Copies the selected step(s) to the class clipboard."""
        selected_items = self.steps_list_widget.selectedItems()
        
        if not selected_items:
            QMessageBox.warning(self, "No Selection", "Please select at least one step to copy.")
            return
        
        # Get the indices of selected items
        selected_indices = [self.steps_list_widget.row(item) for item in selected_items]
        selected_indices.sort()  # Sort to maintain order
        
        # Deep copy the selected steps
        EditTestCaseDialog.copied_steps = [copy.deepcopy(self.added_steps[i]) for i in selected_indices]
        
        step_word = "step" if len(EditTestCaseDialog.copied_steps) == 1 else "steps"
        # Show status message (less intrusive than popup)
        self.statusBar().showMessage(f"Copied {len(EditTestCaseDialog.copied_steps)} {step_word} to clipboard.", 3000)
    
    def paste_copied_steps(self):
        """Pastes the copied step(s) at the end of the test steps list."""
        if not EditTestCaseDialog.copied_steps:
            QMessageBox.warning(self, "Nothing to Paste", "No steps have been copied yet.")
            return
        
        # Always paste at the end
        insert_index = len(self.added_steps)
        
        # Deep copy the steps to paste (to avoid reference issues)
        steps_to_paste = [copy.deepcopy(step) for step in EditTestCaseDialog.copied_steps]
        
        # Insert the copied steps at the end
        for step in steps_to_paste:
            self.added_steps.append(step)
        
        # Update the UI
        self.update_steps_list()
        
        # Select the first pasted step
        if insert_index < self.steps_list_widget.count():
            self.steps_list_widget.setCurrentRow(insert_index)
        
        step_word = "step" if len(steps_to_paste) == 1 else "steps"
        # Show status message instead of popup for smoother workflow
        self.statusBar().showMessage(f"Pasted {len(steps_to_paste)} {step_word} at the end.", 3000)
        
    def delete_selected_steps(self):
        """Deletes the selected step(s) from the test case."""
        selected_items = self.steps_list_widget.selectedItems()
        
        if not selected_items:
            QMessageBox.warning(self, "No Selection", "Please select at least one step to delete.")
            return
        
        # Get the indices of selected items
        selected_indices = [self.steps_list_widget.row(item) for item in selected_items]
        selected_indices.sort(reverse=True)  # Sort in reverse to delete from bottom to top
        
        # Confirm deletion
        step_word = "step" if len(selected_indices) == 1 else "steps"
        reply = QMessageBox.question(
            self,
            "Confirm Delete",
            f"Are you sure you want to delete {len(selected_indices)} {step_word}?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # Delete the steps (from bottom to top to maintain correct indices)
        for index in selected_indices:
            if 0 <= index < len(self.added_steps):
                del self.added_steps[index]
        
        # Update the UI
        self.update_steps_list()
        
        # Clear the details table
        self.details_table.setRowCount(0)
        
        self.statusBar().showMessage(f"Deleted {len(selected_indices)} {step_word}.", 3000)
    
    def statusBar(self):
        """Helper method to access the parent window's status bar."""
        parent = self.parent()
        if parent and hasattr(parent, 'statusBar'):
            return parent.statusBar()
        # Return a dummy object if no status bar is available
        class DummyStatusBar:
            def showMessage(self, msg, timeout=0):
                pass
        return DummyStatusBar()

    def add_module_as_step(self):
        selected_module_name = self.import_module_combobox.currentText()
        if selected_module_name and selected_module_name != "Select Module...":
            new_step = {
                "name": f"Import Module: {selected_module_name}",
                "type": "module_import",
                "module_name": selected_module_name,
                "fields": [] 
            }
            self.added_steps.append(new_step)
            self.update_steps_list()
            QMessageBox.information(self, "Success", f"'{selected_module_name}' added as a test step.")

    def toggle_execution(self):
        """Toggles between play and stop execution."""
        if self.is_executing:
            # Stop execution
            self.execution_stop_flag = True
            self.is_executing = False
            self.execute_button.setIcon(self.main_window.style().standardIcon(QStyle.StandardPixmap.SP_MediaPlay))
            self.statusBar().showMessage("Test execution stopped by user.", 3000)
        else:
            # Start execution
            self.execute_test_case_preview()
    
    def execute_test_case_preview(self):
        """Executes the test case from the specified start step to end step."""
        if not self.main_window or not hasattr(self.main_window, 'pcomm_window_title'):
            QMessageBox.warning(self, "Error", "Cannot execute: Main window not available.")
            return
        
        # Get selected steps from dropdowns
        start_step_data = self.start_step_combo.currentData()
        end_step_data = self.end_step_combo.currentData()
        
        if start_step_data is None or end_step_data is None:
            QMessageBox.warning(self, "Error", "Please select valid steps to execute.")
            return
        
        # ✅ NEW: Parse start step (can be int or tuple)
        if isinstance(start_step_data, tuple):
            start_main_step = start_step_data[0] - 1  # 0-based main step
            start_utility_step = start_step_data[1] - 1  # 0-based utility step
        else:
            start_main_step = start_step_data - 1
            start_utility_step = 0  # Start from beginning of main step
        
        # ✅ NEW: Parse end step (can be int, tuple, or -1 for "Last")
        if end_step_data == -1:
            # "Last" option - execute all steps including utilities
            end_main_step = len(self.added_steps)
            end_utility_step = -1  # Execute all utilities
        elif isinstance(end_step_data, tuple):
            end_main_step = end_step_data[0]
            end_utility_step = end_step_data[1] - 1  # 0-based utility step
        else:
            end_main_step = end_step_data
            end_utility_step = -1  # Execute all utilities in the last main step
        
        # Validate ranges
        if start_main_step < 0 or start_main_step >= len(self.added_steps):
            QMessageBox.warning(self, "Invalid Step", "Start step is out of range.")
            return
        
        if end_main_step < start_main_step + 1 or end_main_step > len(self.added_steps):
            QMessageBox.warning(self, "Invalid Step", "End step is out of range.")
            return
        
        self.is_executing = True
        self.execution_stop_flag = False
        self.execute_button.setIcon(self.main_window.style().standardIcon(QStyle.StandardPixmap.SP_MediaStop))
        
        try:
            import pythoncom
            import win32com.client
            import time
            
            pythoncom.CoInitialize()
            autECLSession = win32com.client.Dispatch("PCOMM.autECLSession")
            connection_name = self.main_window.get_connection_name_from_title(self.main_window.pcomm_window_title)
            autECLSession.SetConnectionByName(connection_name)
            autECLPS = autECLSession.autECLPS
            
            # ✅ NEW: Execute main steps from start to end
            # ✅ CHANGED: Use while loop to handle dynamic step list updates
            step_idx = start_main_step
            while step_idx < len(self.added_steps):  # ✅ Check against current length
                if self.execution_stop_flag:
                    break
                
                # ✅ Check if we've reached the end step
                if step_idx >= end_main_step:
                    break
                
                QApplication.processEvents()
                
                step = self.added_steps[step_idx]
                step_type = step.get('type')
                step_num = step_idx + 1
                
                # ✅ FIXED: Determine which utility steps to execute and whether to skip main step
                execute_utilities = True
                utility_start = 0
                utility_end = -1
                skip_main_step = False  # ✅ NEW flag
                
                if step_idx == start_main_step and start_utility_step > 0:
                    # Starting from a utility step - skip main step execution
                    skip_main_step = True
                    utility_start = start_utility_step
                
                if step_idx == end_main_step - 1:
                    # Last step - end at specified utility
                    utility_end = end_utility_step
                
                print(f"Executing Step {step_num}: {step.get('name', 'Unknown')}")
                
                # ✅ FIXED: Only execute main step if not skipping
                if not skip_main_step:
                    # Execute main step logic
                    if step_type == 'module_import':
                        module_name = step.get('module_name')
                        
                        # Process input fields
                        for field in step.get('fields', []):
                            if self.execution_stop_flag:
                                break
                            
                            action_type = field.get('action_type', 'Input')
                            value = str(field.get('value', '')).strip()
                            
                            if action_type == 'Input' and value:
                                if module_name in self.modules:
                                    module_data = self.modules[module_name]
                                    labels = module_data.get('labels', [])
                                    
                                    field_name = field.get('field_name')
                                    for label in labels:
                                        label_name = label.get('name') or label.get('label') or label.get('text', '')
                                        if label_name == field_name:
                                            row = int(label.get('row', 1))
                                            col = int(label.get('column', 1))
                                            
                                            autECLPS.SetCursorPos(row, col)
                                            autECLPS.SendKeys(value)
                                            time.sleep(0.2)
                                            break
                        
                        # Process validation fields
                        for field in step.get('fields', []):
                            if self.execution_stop_flag:
                                break
                            
                            action_type = field.get('action_type', 'Input')
                            value = str(field.get('value', '')).strip()
                            
                            if action_type == 'Validate' and value:
                                if module_name in self.modules:
                                    module_data = self.modules[module_name]
                                    labels = module_data.get('labels', [])
                                    
                                    field_name = field.get('field_name')
                                    for label in labels:
                                        label_name = label.get('name') or label.get('label') or label.get('text', '')
                                        if label_name == field_name:
                                            row = int(label.get('row', 1))
                                            col = int(label.get('column', 1))
                                            length = int(label.get('length', len(value)))

                                            actual_value = autECLPS.GetText(row, col, length)
                                            
                                            # ✅ NEW: Use validation helper function
                                            validation_passed = self.validate_field_value(actual_value, value)
                                            
                                            if not validation_passed:
                                                # ✅ ENHANCED: Better error message for {blank} validation
                                                if value.lower() == '{blank}':
                                                    QMessageBox.warning(self, "Validation Failed",
                                                        f"Step {step_num}: Field '{field_name}'\n"
                                                        f"Expected: <blank>\n"
                                                        f"Actual: '{actual_value.strip()}' (not blank)")
                                                else:
                                                    QMessageBox.warning(self, "Validation Failed",
                                                        f"Step {step_num}: Field '{field_name}'\n"
                                                        f"Expected: '{value}'\n"
                                                        f"Actual: '{actual_value.strip()}'")
                                                self.execution_stop_flag = True
                                            
                                            time.sleep(0.1)
                                            break
                    
                    elif step_type == 'special_key':
                        key_value = step.get('key_value', '')
                        key_mapping = {
                            "Enter Key": "[enter]", "Clear Key": "[clear]", "End Key": "[EraseEof]",
                            "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                            "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                            "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                        }
                        pcomm_key = key_mapping.get(key_value, key_value)
                        autECLPS.SendKeys(pcomm_key)
                        time.sleep(2.0)
                    
                    elif step_type == 'wait':
                        seconds = float(step.get('seconds', 1))
                        print(f"Waiting {seconds} second(s)...")
                        for _ in range(int(seconds * 10)):
                            if self.execution_stop_flag:
                                break
                            time.sleep(0.1)
                            QApplication.processEvents()
                                
                    elif step_type == 'random_input':
                        row = int(step.get('row', 1))
                        col = int(step.get('column', 1))
                        value = str(step.get('value', '')).strip()
                        is_special_key = step.get('is_special_key', False)
                        
                        if value:
                            print(f"Step {step_num}: Random Input at ({row}, {col})")
                            
                            # Capture screen before action
                            before_screen = wait_for_pcomm_ready_smart(autECLPS, "Random Input")
                            
                            # Set cursor position
                            autECLPS.SetCursorPos(row, col)
                            
                            # ✅ Check if it's a special key
                            if is_special_key:
                                key_mapping = {
                                    "Enter Key": "[enter]", "Clear Key": "[clear]", "End Key": "[EraseEof]",
                                    "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                                    "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                                    "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                                }
                                pcomm_key = key_mapping.get(value, value)
                                autECLPS.SendKeys(pcomm_key)
                                print(f"Step {step_num}: Sent special key '{value}' -> '{pcomm_key}'")
                            else:
                                # Regular text input
                                autECLPS.SendKeys(value)
                                print(f"Step {step_num}: Sent text '{value}'")
                            
                            # Wait for screen change
                            complete_pcomm_wait(autECLPS, before_screen, action_description="Random Input")
        
                    elif step_type == 'break':
                        message = step.get('message', '')
                        print(f"Step {step_num}: Break point reached")
                        
                        # Show break dialog
                        break_dialog = BreakExecutionDialog(message, self)
                        break_dialog.show()
                        
                        # Wait for user action
                        while break_dialog.result_action is None:
                            QApplication.processEvents()
                            time.sleep(0.1)
                        
                        action = break_dialog.result_action
                        
                        if action == BreakExecutionDialog.STOP:
                            print("User chose to stop execution at break point")
                            self.execution_stop_flag = True
                            break_dialog.close()
                            break
                        
                        elif action == BreakExecutionDialog.EDIT:
                            print("User chose to edit test case at break point")
                            break_dialog.result_action = None
                            
                            # Get current test case data
                            existing_steps = self.added_steps
                            test_case_description = self.test_case_description_input.text()
                            test_case_assumptions = self.test_case_assumptions_input.toHtml()
                            
                            # Open edit dialog
                            edit_dialog = EditTestCaseDialog(
                                existing_steps,
                                self.modules,
                                self.main_window,
                                self.test_case_name_input.text(),
                                test_case_description,
                                test_case_assumptions
                            )
                            
                            edit_dialog.setParent(break_dialog, edit_dialog.windowFlags())
                            edit_dialog.exec()
                            
                            # Get updated data
                            updated_steps = edit_dialog.get_updated_steps()
                            updated_description = edit_dialog.get_test_case_description()
                            updated_assumptions = edit_dialog.get_test_case_assumptions()
                            
                            # ✅ Update the current test case data
                            self.added_steps = updated_steps
                            self.test_case_description_input.setText(updated_description)
                            self.test_case_assumptions_input.setHtml(updated_assumptions)
                            
                            # ✅ CRITICAL: Update end_main_step to reflect new step count
                            if end_step_data == -1:
                                end_main_step = len(self.added_steps)
                            
                            print(f"Test case updated during execution. Total steps: {len(self.added_steps)}")
                            
                            QMessageBox.information(
                                break_dialog,
                                "Test Case Updated",
                                f"Test case has been updated.\nTotal steps: {len(self.added_steps)}\n"
                                f"Press 'Resume Execution' to continue from Step {step_num}."
                            )
                            
                            break_dialog.raise_()
                            break_dialog.activateWindow()
                            
                            # Wait for next action
                            while break_dialog.result_action is None:
                                QApplication.processEvents()
                                time.sleep(0.1)
                            
                            action = break_dialog.result_action
                            
                            if action == BreakExecutionDialog.STOP:
                                self.execution_stop_flag = True
                                break_dialog.close()
                                break
                            elif action == BreakExecutionDialog.RESUME:
                                break_dialog.close()
                        
                        elif action == BreakExecutionDialog.RESUME:
                            print("User chose to resume execution")
                            break_dialog.close()
                else:
                    print(f"  Skipping main step {step_num}, starting from utility step {utility_start + 1}")
                
                time.sleep(0.1)
        
                # ✅ FIXED: Execute utility steps with range control
                if execute_utilities:
                    utility_steps = step.get('utility_steps', [])
                    
                    # ✅ FIXED: Determine actual utility range
                    u_start = utility_start  # Use the calculated utility_start
                    u_end = len(utility_steps) if utility_end == -1 else min(utility_end + 1, len(utility_steps))
                    
                    if u_start < len(utility_steps):  # ✅ Only log if we have utilities to execute
                        print(f"  Executing utility steps {u_start + 1} to {u_end} for Step {step_num}")
                    
                    for utility_idx in range(u_start, u_end):
                        if self.execution_stop_flag:
                            break
                        
                        QApplication.processEvents()
                        
                        utility_step = utility_steps[utility_idx]
                        utility_type = utility_step.get('type')
                        print(f"  Utility Step {step_num}.{utility_idx + 1}: {utility_step.get('name', 'Unknown')}")
                        
                        if utility_type == 'special_key':
                            key_value = utility_step.get('key_value', '')
                            key_mapping = {
                                "Enter Key": "[enter]", "Clear Key": "[clear]", "End Key": "[EraseEof]",
                                "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                                "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                                "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                            }
                            pcomm_key = key_mapping.get(key_value, key_value)
                            autECLPS.SendKeys(pcomm_key)
                            time.sleep(2.0)
                        
                        elif utility_type == 'wait':
                            seconds = float(utility_step.get('seconds', 1))
                            print(f"    Waiting {seconds} second(s)...")
                            for _ in range(int(seconds * 10)):
                                if self.execution_stop_flag:
                                    break
                                time.sleep(0.1)
                                QApplication.processEvents()
                        
                        elif utility_type == 'module_import':
                            module_name = utility_step.get('module_name')
                            if module_name in self.modules:
                                module_data = self.modules[module_name]
                                labels = module_data.get('labels', [])
                                
                                # ✅ FIXED: Process Input fields first
                                for field in utility_step.get('fields', []):
                                    if self.execution_stop_flag:
                                        break
                                    
                                    action_type = field.get('action_type', 'Input')
                                    value = str(field.get('value', '')).strip()
                                    
                                    if action_type == 'Input' and value:
                                        field_name = field.get('field_name')
                                        for label in labels:
                                            label_name = label.get('name') or label.get('label') or label.get('text', '')
                                            if label_name == field_name:
                                                row = int(label.get('row', 1))
                                                col = int(label.get('column', 1))
                                                
                                                autECLPS.SetCursorPos(row, col)
                                                autECLPS.SendKeys(value)
                                                print(f"Utility Step {step_num}.{utility_idx + 1}: Sent input '{value}' to {field_name}")
                                                time.sleep(0.2)
                                                break
                                
                                # ✅ FIXED: Then process Validation fields
                                for field in utility_step.get('fields', []):
                                    if self.execution_stop_flag:
                                        break
                                    
                                    action_type = field.get('action_type', 'Validate')
                                    value = str(field.get('value', '')).strip()
                                    
                                    if action_type == 'Validate' and value:
                                        field_name = field.get('field_name')
                                        for label in labels:
                                            label_name = label.get('name') or label.get('label') or label.get('text', '')
                                            if label_name == field_name:
                                                row = int(label.get('row', 1))
                                                col = int(label.get('column', 1))
                                                length = int(label.get('length', len(value)))

                                                actual_value = autECLPS.GetText(row, col, length)
                                                
                                                # ✅ NEW: Use validation helper function
                                                validation_passed = self.validate_field_value(actual_value, value)
                                                
                                                if not validation_passed:
                                                    # ✅ ENHANCED: Better error message for {blank} validation
                                                    if value.lower() == '{blank}':
                                                        QMessageBox.warning(self, "Validation Failed",
                                                            f"Utility Step {step_num}.{utility_idx + 1}: Field '{field_name}'\n"
                                                            f"Expected: <blank>\n"
                                                            f"Actual: '{actual_value.strip()}' (not blank)")
                                                    else:
                                                        QMessageBox.warning(self, "Validation Failed",
                                                            f"Utility Step {step_num}.{utility_idx + 1}: Field '{field_name}'\n"
                                                            f"Expected: '{value}'\n"
                                                            f"Actual: '{actual_value.strip()}'")
                                                    self.execution_stop_flag = True
                                                
                                                time.sleep(0.1)
                                                break
                        
                        elif utility_type == 'random_input':
                            row = int(utility_step.get('row', 1))
                            col = int(utility_step.get('column', 1))
                            value = str(utility_step.get('value', '')).strip()
                            is_special_key = utility_step.get('is_special_key', False)
                            
                            if value:
                                print(f"  Utility Step {step_num}.{utility_idx + 1}: Random Input at ({row}, {col})")
                                
                                # Capture screen before action
                                before_screen = wait_for_pcomm_ready_smart(autECLPS, "Random Input")
                                
                                # Set cursor position
                                autECLPS.SetCursorPos(row, col)
                                
                                # ✅ Check if it's a special key
                                if is_special_key:
                                    key_mapping = {
                                        "Enter Key": "[enter]", "Clear Key": "[clear]", "End Key": "[EraseEof]",
                                        "F1": "[pf1]", "F2": "[pf2]", "F3": "[pf3]", "F4": "[pf4]",
                                        "F5": "[pf5]", "F6": "[pf6]", "F7": "[pf7]", "F8": "[pf8]",
                                        "F9": "[pf9]", "F10": "[pf10]", "F11": "[pf11]", "F12": "[pf12]",
                                    }
                                    pcomm_key = key_mapping.get(value, value)
                                    autECLPS.SendKeys(pcomm_key)
                                    print(f"  Utility Step {step_num}.{utility_idx + 1}: Sent special key '{value}' -> '{pcomm_key}'")
                                else:
                                    # Regular text input
                                    autECLPS.SendKeys(value)
                                    print(f"  Utility Step {step_num}.{utility_idx + 1}: Sent text '{value}'")
                                
                                # Wait for screen change
                                complete_pcomm_wait(autECLPS, before_screen, action_description="Random Input")
                        
                        time.sleep(0.1)
                
                time.sleep(0.1)
                step_idx += 1
            
            pythoncom.CoUninitialize()
            
            if not self.execution_stop_flag:
                start_desc = self.start_step_combo.currentText()
                end_desc = self.end_step_combo.currentText()
                QMessageBox.information(self, "Success", 
                    f"Test execution completed successfully!\nExecuted from {start_desc} to {end_desc}.")
            
        except Exception as e:
            QMessageBox.critical(self, "Execution Error", f"An error occurred during execution:\n\n{str(e)}")
        
        finally:
            try:
                import pythoncom
                pythoncom.CoUninitialize()
            except:
                pass
            
            self.is_executing = False
            self.execution_stop_flag = False
            self.execute_button.setIcon(self.main_window.style().standardIcon(QStyle.StandardPixmap.SP_MediaPlay))
        
    def update_step_combo_options(self):
        """Updates the dropdown options for start and end step based on number of steps."""
        self.start_step_combo.blockSignals(True)
        self.end_step_combo.blockSignals(True)
        
        self.start_step_combo.clear()
        self.end_step_combo.clear()
        
        num_steps = len(self.added_steps)
        
        if num_steps == 0:
            self.start_step_combo.addItem("No steps")
            self.end_step_combo.addItem("No steps")
        else:
            # Populate start step dropdown - Include main steps AND utility sub-steps
            for i in range(1, num_steps + 1):
                self.start_step_combo.addItem(f"Step {i}", i)
                
                # ✅ NEW: Add utility sub-steps to start dropdown
                step = self.added_steps[i-1]
                utility_steps = step.get('utility_steps', [])
                for j in range(1, len(utility_steps) + 1):
                    self.start_step_combo.addItem(f"Step {i}.{j}", (i, j))
            
            # Populate end step dropdown - Include main steps AND utility sub-steps
            for i in range(1, num_steps + 1):
                self.end_step_combo.addItem(f"Step {i}", i)
                
                # ✅ NEW: Add utility sub-steps to end dropdown
                step = self.added_steps[i-1]
                utility_steps = step.get('utility_steps', [])
                for j in range(1, len(utility_steps) + 1):
                    self.end_step_combo.addItem(f"Step {i}.{j}", (i, j))
            
            # Add "Last (with utilities)" option to end step
            self.end_step_combo.addItem("Last (with utilities)", -1)
            
            # Set defaults
            self.start_step_combo.setCurrentIndex(0)
            self.end_step_combo.setCurrentIndex(self.end_step_combo.count() - 1)  # Select "Last" by default
        
        self.start_step_combo.blockSignals(False)
        self.end_step_combo.blockSignals(False)
    
    def statusBar(self):
        """Helper method to access status bar."""
        if hasattr(self, 'main_window') and hasattr(self.main_window, 'statusBar'):
            return self.main_window.statusBar()
        class DummyStatusBar:
            def showMessage(self, msg, timeout=0):
                pass
        return DummyStatusBar()

    def get_test_case_name(self):
        """Returns the updated test case name."""
        return self.test_case_name_input.text().strip()

    def get_test_case_description(self):  # ✅ ADD THIS METHOD
        """Returns the updated test case description."""
        return self.test_case_description_input.text().strip()

    def get_test_case_assumptions(self):
        """Returns the updated test case assumptions."""
        return self.test_case_assumptions_input.toHtml()
        
    def get_updated_steps(self):
        """Returns the updated list of test case steps."""
        return self.added_steps         
        
    def display_module_details(self, current_item, previous_item):
        """
        Displays the details of the selected module on the right side in a table.
        UPDATED: Now handles 'capture_screenshot' steps by showing the previous module's labels with highlight checkboxes.
        """
        # Auto-save previous item if it exists
        if previous_item:
            # Remove bold from previous item
            widget = self.steps_list_widget.itemWidget(previous_item)
            if widget:
                label = widget.findChild(QLabel)
                if label:
                    font = label.font()
                    font.setBold(False)
                    label.setFont(font)

        # Clear table if multiple items are selected
        if len(self.steps_list_widget.selectedItems()) > 1:
            self.details_table.setRowCount(0)
            return

        self.details_table.setRowCount(0)
        
        if current_item:
            # Make current item bold
            widget = self.steps_list_widget.itemWidget(current_item)
            if widget:
                label = widget.findChild(QLabel)
                if label:
                    font = label.font()
                    font.setBold(True)
                    label.setFont(font)
            
            current_step_index = self.steps_list_widget.row(current_item)
            step_data = self.added_steps[current_step_index]
            step_type = step_data.get('type')
            
            if step_type == 'break':
                self.details_table.setColumnCount(2)
                self.details_table.setHorizontalHeaderLabels(["Property", "Value"])
                self.details_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
                self.details_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
                
                self.details_table.setRowCount(1)
                
                # Property name
                property_item = QTableWidgetItem("Display Message")
                property_item.setFlags(property_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                self.details_table.setItem(0, 0, property_item)
                
                # Message text edit
                message_text_edit = QTextEdit()
                message_text_edit.setPlaceholderText("Enter break message here...")
                message_text_edit.setMinimumHeight(200)
                message_text_edit.setMaximumHeight(400)
                message_text_edit.setAcceptRichText(True)
                
                # Add keyboard shortcuts for text formatting
                from PyQt6.QtGui import QKeySequence, QAction
                
                # Bold (Ctrl+B)
                bold_action = QAction(message_text_edit)
                bold_action.setShortcut(QKeySequence.StandardKey.Bold)
                bold_action.triggered.connect(lambda: message_text_edit.setFontWeight(
                    QFont.Weight.Normal if message_text_edit.fontWeight() == QFont.Weight.Bold else QFont.Weight.Bold
                ))
                message_text_edit.addAction(bold_action)
                
                # Italic (Ctrl+I)
                italic_action = QAction(message_text_edit)
                italic_action.setShortcut(QKeySequence.StandardKey.Italic)
                italic_action.triggered.connect(lambda: message_text_edit.setFontItalic(
                    not message_text_edit.fontItalic()
                ))
                message_text_edit.addAction(italic_action)
                
                # Underline (Ctrl+U)
                underline_action = QAction(message_text_edit)
                underline_action.setShortcut(QKeySequence.StandardKey.Underline)
                underline_action.triggered.connect(lambda: message_text_edit.setFontUnderline(
                    not message_text_edit.fontUnderline()
                ))
                message_text_edit.addAction(underline_action)
                
                saved_message = step_data.get('message', '')
                if saved_message:
                    message_text_edit.setHtml(saved_message)
                
                self.details_table.setCellWidget(0, 1, message_text_edit)
                self.details_table.setRowHeight(0, 250)
                
                return            
            
            # ✅ NEW: Handle capture_screenshot step
            if step_type == 'capture_screenshot':
                # Find the most recent module_import step before this screenshot
                reference_module_name = None
                for i in range(current_step_index - 1, -1, -1):
                    prev_step = self.added_steps[i]
                    if prev_step.get('type') == 'module_import':
                        reference_module_name = prev_step.get('module_name')
                        break
                
                if not reference_module_name:
                    # No module found before this screenshot
                    self.details_table.setRowCount(1)
                    info_item = QTableWidgetItem("No module found before this screenshot step")
                    info_item.setFlags(info_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                    self.details_table.setItem(0, 0, info_item)
                    return
                
                # Store the reference module in step data
                step_data['reference_module'] = reference_module_name
                
                # Get the module details
                module_details = self.modules.get(reference_module_name, {})
                labels = module_details.get('labels', [])
                
                # Set table headers for screenshot step (3 columns: Field Name, Position, Highlight)
                self.details_table.setColumnCount(3)
                self.details_table.setHorizontalHeaderLabels(["Field Name", "Position (Row, Col)", "Highlight"])
                self.details_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
                self.details_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
                self.details_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
                
                self.details_table.setRowCount(len(labels))
                
                # Get existing highlight flags from step data
                saved_highlights = {field['field_name']: field.get('highlight', False) 
                                  for field in step_data.get('fields', [])}
                
                self.details_table.blockSignals(True)
                
                for i, label_data in enumerate(labels):
                    field_name = label_data.get('label') or label_data.get('text') or label_data.get('name', 'N/A')
                    row = label_data.get('row', '')
                    col = label_data.get('column', '')
                    
                    # Field name (read-only)
                    field_name_item = QTableWidgetItem(field_name)
                    field_name_item.setFlags(field_name_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                    self.details_table.setItem(i, 0, field_name_item)
                    
                    # Position (read-only)
                    position_item = QTableWidgetItem(f"({row}, {col})")
                    position_item.setFlags(position_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                    self.details_table.setItem(i, 1, position_item)
                    
                    # Highlight checkbox
                    highlight_checkbox = QCheckBox()
                    highlight_checkbox.setChecked(saved_highlights.get(field_name, False))
                    
                    # Center the checkbox in the cell
                    checkbox_widget = QWidget()
                    checkbox_layout = QHBoxLayout(checkbox_widget)
                    checkbox_layout.addWidget(highlight_checkbox)
                    checkbox_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    checkbox_layout.setContentsMargins(0, 0, 0, 0)
                    
                    self.details_table.setCellWidget(i, 2, checkbox_widget)
                
                self.details_table.blockSignals(False)
                return
            
            # ✅ EXISTING: Handle module_import step
            if step_type != 'module_import':
                return
                
            module_name = step_data.get('module_name')

            # Get action type from the saved fields and set radio buttons
            saved_fields_list = step_data.get('fields', [])

            if saved_fields_list:
                first_field = saved_fields_list[0]
                action_type = first_field.get('action_type', 'Input')
                if action_type == 'Input':
                    self.input_radio.setChecked(True)
                else:
                    self.validate_radio.setChecked(True)
            else:
                self.input_radio.setChecked(True)

            # Reset table headers for module import (2 columns)
            self.details_table.setColumnCount(2)
            self.details_table.setHorizontalHeaderLabels(["Field Name", "Value"])
            self.details_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
            self.details_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)

            # ✅ CRITICAL FIX: Display EXACTLY what's in the step's fields array
            # Do NOT look up module definition - just show what this step has
            self.details_table.setRowCount(len(saved_fields_list))

            self.details_table.blockSignals(True)

            for i, field_data in enumerate(saved_fields_list):
                field_name = field_data.get('field_name', 'N/A')
                field_value = field_data.get('value', '')
                
                # Display the field name
                field_name_item = QTableWidgetItem(field_name)
                field_name_item.setFlags(field_name_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                self.details_table.setItem(i, 0, field_name_item)
                
                # Display the field value
                value_item = QLineEdit()
                value_item.setText(field_value)
                value_item.setToolTip(self.get_available_variables_info())
                
                self.details_table.setCellWidget(i, 1, value_item)

            self.details_table.blockSignals(False)
        
    def save_table_data_to_step_with_message(self):
        """
        Saves the current state with a success message.
        Called when the 'Save Step Fields' button is clicked.
        """
        current_item = self.steps_list_widget.currentItem()
        if not current_item:
            QMessageBox.warning(self, "No Step Selected", "Please select a test step first.")
            return
        
        step_index = self.steps_list_widget.row(current_item)
        self.save_table_data_to_step(step_index, show_message=True)

    def on_action_type_changed(self):
        """
        ✅ REMOVED: No longer controls utility button visibility for Import Module.
        The utility button is now always shown for Import Module regardless of action type.
        This method is kept for potential future use but currently does nothing.
        """
        # The utility button visibility is now controlled entirely by update_dynamic_list_combobox
        pass

    def save_table_data_to_step(self, step_index=None, show_message=False):
        """
        Saves the current state of the details table back to the step data.
        FIXED: Uses internal_field_id for precise field matching instead of field_name.
        ✅ FIXED: Now handles QTextEdit for Break steps correctly.
        """
        # ✅ NEW: Get the currently selected main step to determine what we're editing
        if step_index is None:
            current_item = self.steps_list_widget.currentItem()
            if not current_item:
                return
            step_index = self.steps_list_widget.row(current_item)
        
        if step_index < 0 or step_index >= len(self.added_steps):
            return
        
        step_data = self.added_steps[step_index]
        step_type = step_data.get('type')
        
        # ✅ CRITICAL CHECK: Only process utility steps if we're actually viewing one
        # Check if the utility steps list has focus/selection
        has_utility_selection = (
            hasattr(self, 'utility_steps_list_widget') and 
            self.utility_steps_list_widget.currentItem() is not None
        )
        
        if has_utility_selection and hasattr(self, 'current_utility_step') and self.current_utility_step:
            # We're editing a UTILITY step
            utility_type = self.current_utility_step.get('type')
            
            if utility_type == 'capture_screenshot':
                # Existing screenshot handling
                fields = []
                for row in range(self.details_table.rowCount()):
                    field_name_item = self.details_table.item(row, 0)
                    checkbox_widget = self.details_table.cellWidget(row, 2)
                    
                    if field_name_item and checkbox_widget:
                        field_name = field_name_item.text()
                        checkbox = checkbox_widget.findChild(QCheckBox)
                        if checkbox:
                            is_highlighted = checkbox.isChecked()
                            fields.append({
                                "field_name": field_name,
                                "highlight": is_highlighted
                            })
                
                self.current_utility_step['fields'] = fields
                
                if show_message:
                    QMessageBox.information(self, "Saved", "Utility screenshot highlight settings saved successfully.")
            
            elif utility_type == 'module_import':
                action_type = self.current_utility_step.get("action_type", "Input")
                
                # Build fields list using CURRENT table values
                fields = []
                for row in range(self.details_table.rowCount()):
                    field_name_item = self.details_table.item(row, 0)
                    field_value_widget = self.details_table.cellWidget(row, 1)
                    
                    if field_name_item and field_value_widget:
                        field_name = field_name_item.text()
                        
                        if isinstance(field_value_widget, QLineEdit):
                            field_value = field_value_widget.text()
                        elif isinstance(field_value_widget, QTextEdit):
                            field_value = field_value_widget.toHtml()
                        else:
                            field_value = ""
                        
                        # Generate NEW unique ID per save
                        import time
                        import random
                        unique_timestamp = int(time.time() * 1000000)
                        random_suffix = random.randint(100000, 999999)
                        
                        field_dict = {
                            "field_name": field_name,
                            "internal_field_id": f"{field_name}_UTIL_{unique_timestamp}_{random_suffix}_{row}",
                            "action_type": action_type,
                            "value": field_value,
                        }
                        
                        fields.append(field_dict)
                
                self.current_utility_step['fields'] = fields
                
                if show_message:
                    QMessageBox.information(self, "Saved", "Utility module validation fields saved successfully.")

            self.current_utility_step = None
            return
        
        # ✅ Below here is for MAIN STEPS only - clear utility reference
        self.current_utility_step = None
        
        # Handle break step with QTextEdit
        if step_type == 'break':
            message_widget = self.details_table.cellWidget(0, 1)
            if message_widget and isinstance(message_widget, QTextEdit):
                import time
                import random
                unique_timestamp = int(time.time() * 1000000)
                random_suffix = random.randint(100000, 999999)
                
                step_data['message'] = message_widget.toHtml()
                step_data['message_field_id'] = f"break_message_{step_index + 1}_{unique_timestamp}_{random_suffix}"
                
                if show_message:
                    QMessageBox.information(self, "Saved", "Break message saved successfully.")
            return           
        
        # Handle capture_screenshot step
        if step_type == 'capture_screenshot':
            fields = []
            for row in range(self.details_table.rowCount()):
                field_name_item = self.details_table.item(row, 0)
                checkbox_widget = self.details_table.cellWidget(row, 2)
                
                if field_name_item and checkbox_widget:
                    field_name = field_name_item.text()
                    checkbox = checkbox_widget.findChild(QCheckBox)
                    if checkbox:
                        is_highlighted = checkbox.isChecked()
                        
                        fields.append({
                            "field_name": field_name,
                            "highlight": is_highlighted
                        })
            
            self.added_steps[step_index]['fields'] = fields
            
            if show_message:
                QMessageBox.information(self, "Saved", "Screenshot highlight settings saved successfully.")
            return
        
        # Handle module_import step
        if step_type != 'module_import':
            return
        
        action_type = 'Input' if self.input_radio.isChecked() else 'Validate'

        # Rebuild fields from table
        fields = []
        for row in range(self.details_table.rowCount()):
            field_name_item = self.details_table.item(row, 0)
            field_value_widget = self.details_table.cellWidget(row, 1)
            
            if field_name_item and field_value_widget:
                field_name = field_name_item.text()
                
                if isinstance(field_value_widget, QLineEdit):
                    field_value = field_value_widget.text()
                elif isinstance(field_value_widget, QTextEdit):
                    field_value = field_value_widget.toHtml()
                else:
                    field_value = ""
                
                # Generate NEW unique ID per save
                import time
                import random
                unique_timestamp = int(time.time() * 1000000)
                random_suffix = random.randint(100000, 999999)
                
                field_dict = {
                    "field_name": field_name,
                    "internal_field_id": f"{field_name}_MAIN_M{step_index + 1}_{unique_timestamp}_{random_suffix}_{row}",
                    "action_type": action_type,
                    "value": field_value,
                }
                
                fields.append(field_dict)

        self.added_steps[step_index]['fields'] = fields

        if show_message:
            QMessageBox.information(self, "Saved", "Module fields saved successfully.")
        else:
            print(f"IndexError avoided: Invalid step index {step_index}")
            
    def add_capture_text_screenshot_step(self):
        """Adds a text screenshot capture step to the test case."""
        # ✅ NEW: Get the target step number
        step_number_text = self.test_step_number_input.text().strip()
        insert_index = None
        
        if step_number_text:
            try:
                step_number = int(step_number_text)
                if step_number < 1:
                    QMessageBox.warning(self, "Invalid Step Number", "Step number must be at least 1.")
                    return
                insert_index = step_number - 1
                if insert_index > len(self.added_steps):
                    insert_index = len(self.added_steps)
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Please enter a valid step number.")
                return
        
        new_step = {
            "name": "Capture Text Screenshot",
            "type": "capture_screen_text",
            "fields": [],
            "utility_steps": []
        }
        
        # ✅ NEW: Insert at specific index or append
        if insert_index is not None:
            self.added_steps.insert(insert_index, new_step)
            success_msg = f"Text screenshot capture step inserted at step {step_number}."
        else:
            self.added_steps.append(new_step)
            success_msg = "Text screenshot capture step added."
        self.input_radio.setChecked(True)
        
        self.update_steps_list()
        self.test_step_number_input.clear()  # ✅ NEW: Clear step number

        QMessageBox.information(self, "Success", "Text screenshot capture step added.")
        
    def add_utility_step_from_selection(self):
        """Adds a utility step to the specified test step (or last step if not specified)."""
        selected_type = self.import_type_combobox.currentText()
        
        if not self.added_steps:
            QMessageBox.warning(self, "No Test Steps", "Please add at least one main test step before adding utility steps.")
            return
        
        # Get the target step number from input field
        step_number_text = self.test_step_number_input.text().strip()
        
        if step_number_text:
            try:
                target_step_number = int(step_number_text)
                if target_step_number < 1:
                    QMessageBox.warning(self, "Invalid Step Number", "Step number must be at least 1.")
                    return
                if target_step_number > len(self.added_steps):
                    QMessageBox.warning(self, "Invalid Step Number", 
                        f"Step number cannot exceed {len(self.added_steps)}. Please enter a valid step number.")
                    return
                target_step_index = target_step_number - 1  # Convert to 0-based index
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Please enter a valid step number.")
                return
        else:
            # If no step number specified, use the last step
            target_step_index = len(self.added_steps) - 1
        
        # Get the target step
        target_step = self.added_steps[target_step_index]
        
        # Ensure utility_steps field exists
        if 'utility_steps' not in target_step:
            target_step['utility_steps'] = []
        
        utility_step = None
        auto_wait_step = None  # Variable to hold auto-wait step
        
        if selected_type == "Special Keys":
            selected_item = self.dynamic_list_combobox.currentText().strip()
            if not selected_item:
                QMessageBox.warning(self, "No Selection", "Please select a special key.")
                return
            
            utility_step = {
                "name": f"Special Key: {selected_item}",
                "type": "special_key",
                "key_value": selected_item
            }
            
            # Auto-add wait step after special key
            auto_wait_step = {
                "name": "Wait: 1 second(s)",
                "type": "wait",
                "seconds": 1
            }
            
            # Clear inputs
            self.dynamic_list_combobox.setCurrentIndex(0)
        
        elif selected_type == "Wait":
            seconds_text = self.wait_seconds_input.text().strip()
            if not seconds_text:
                QMessageBox.warning(self, "Missing Information", "Please provide the number of seconds to wait.")
                return
            
            try:
                seconds = float(seconds_text)
                if seconds <= 0:
                    QMessageBox.warning(self, "Invalid Input", "Seconds must be a positive number.")
                    return
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Seconds must be a valid number.")
                return
            
            utility_step = {
                "name": f"Wait: {seconds} second(s)",
                "type": "wait",
                "seconds": seconds
            }
            
            # Clear inputs
            self.wait_seconds_input.clear()
        
        elif selected_type == "Capture Screenshot":
            utility_step = {
                "name": "Capture Screenshot (DOCX)",
                "type": "capture_screenshot",
                "fields": [],
                "reference_module": None
            }
        
        elif selected_type == "Capture Text Screenshot":
            utility_step = {
                "name": "Capture Text Screenshot",
                "type": "capture_screen_text",
                "fields": []
            }
        
        elif selected_type == "Random Input":
            row_text = self.random_input_row.text().strip()
            col_text = self.random_input_col.text().strip()
            value_text = self.random_input_value.text().strip()
            special_key = self.random_input_special_key_combo.currentText()
            
            if not row_text or not col_text:
                QMessageBox.warning(self, "Missing Information", "Please provide Row and Column.")
                return
            
            if not value_text and special_key == "(Text)":
                QMessageBox.warning(self, "Missing Information", "Please provide either a Value or select a Special Key.")
                return
            
            try:
                row = int(row_text)
                col = int(col_text)
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Row and Column must be valid numbers.")
                return
            
            if special_key != "(Text)":
                display_value = f"[{special_key}]"
                actual_value = special_key
                is_special_key = True
            else:
                display_value = value_text
                actual_value = value_text
                is_special_key = False
            
            utility_step = {
                "name": f"Random Input: Row {row}, Col {col}, Value: {display_value}",
                "type": "random_input",
                "row": row,
                "column": col,
                "value": actual_value,
                "is_special_key": is_special_key
            }
            
            # Clear inputs
            self.random_input_row.clear()
            self.random_input_col.clear()
            self.random_input_value.clear()
            self.random_input_special_key_combo.setCurrentIndex(0)        
        
        elif selected_type == "Import Module":
            # ✅ CRITICAL FIX: Capture the action type BEFORE any other operations
            # Store in a local variable to prevent it from being modified
            captured_action_type = 'Input' if self.input_radio.isChecked() else 'Validate'
            
            selected_item = self.dynamic_list_combobox.currentText().strip()
            
            # Get typed text for editable combo box
            if not selected_item:
                typed_text = self.dynamic_list_combobox.currentText().strip()
                selected_item = None
                for module_name in self.modules.keys():
                    if module_name == typed_text:
                        selected_item = module_name
                        break
                
                if not selected_item:
                    current_index = self.dynamic_list_combobox.currentIndex()
                    if current_index >= 0:
                        module_names = list(self.modules.keys())
                        if current_index < len(module_names):
                            selected_item = module_names[current_index]
                
                if not selected_item:
                    selected_item = typed_text
            
            if not selected_item or selected_item not in self.modules:
                QMessageBox.warning(self, "Invalid Module", "Please select a valid module.")
                return
            
            # Get the module details to create fields
            module_details = self.modules.get(selected_item, {})
            fields_from_module = json.loads(json.dumps(module_details.get('labels', [])))

            
            # Generate unique field identifiers for utility steps
            # Generate unique field identifiers for utility steps
            step_number = len(self.added_steps) if target_step_index is None else target_step_index + 1
            utility_number = len(target_step['utility_steps']) + 1

            fields = []
            for idx, label_data in enumerate(fields_from_module):
                field_name = label_data.get('label') or label_data.get('text') or label_data.get('name', 'N/A')
                
                # âœ… FIXED: Generate unique timestamp and random PER FIELD
                import time
                import random
                unique_timestamp = int(time.time() * 1000000)  # Microseconds for more precision
                random_suffix = random.randint(100000, 999999)
                
                unique_field_name = f"{field_name}_U{step_number}_{utility_number}_{unique_timestamp}_{random_suffix}_{idx}"
                
                fields.append({
                    "field_name": field_name,
                    "internal_field_id": unique_field_name,
                    "action_type": captured_action_type,  # ✅ FIXED: Use the captured local variable
                    "value": "",
                })
            
            utility_step = {
                "name": f"Import Module: {selected_item}",
                "type": "module_import",
                "module_name": selected_item,
                "action_type": captured_action_type,
                "fields": fields
            }
            
            # Clear inputs
            self.dynamic_list_combobox.clearEditText()
            self.dynamic_list_combobox.setCurrentIndex(0)
        
        if utility_step:
            target_step['utility_steps'].append(utility_step)
            
            # Add auto-wait step if it exists
            if auto_wait_step:
                target_step['utility_steps'].append(auto_wait_step)
            
            # ✅ FIXED: Don't call update_steps_list, just refresh utility steps
            self.steps_list_widget.blockSignals(True)
            self.steps_list_widget.setCurrentRow(target_step_index)
            self.steps_list_widget.blockSignals(False)
            
            # ✅ Manually refresh utility steps list while keeping it visible
            self.refresh_utility_steps_for_main_step(target_step_index)
            
            self.input_radio.setChecked(True)  # Reset radio button AFTER appending
            
            # Clear the test step number input after adding
            self.test_step_number_input.clear()
            
            # Show success message with step number
            if auto_wait_step:
                QMessageBox.information(self, "Success", 
                    f"Utility step added to Step {target_step_index + 1} with auto-wait.")
            else:
                QMessageBox.information(self, "Success", 
                    f"Utility step added to Step {target_step_index + 1}.")
        else:
            QMessageBox.warning(self, "Invalid Selection", 
                "Please select a valid utility step type (Wait, Special Keys, Screenshot, or Module).")

    def remove_utility_step(self, step_index, utility_step):
        """Removes a utility step from a main step."""
        if step_index >= len(self.added_steps):
            return
        
        step = self.added_steps[step_index]
        if 'utility_steps' in step and utility_step in step['utility_steps']:
            step['utility_steps'].remove(utility_step)
            self.update_steps_list()                

    def add_prerequisite(self):
        """Opens a dialog to select a prerequisite test case."""
        available_test_cases = [tc for tc in self.main_window.test_cases.keys() 
                               if tc != self.original_test_case_name]
        
        if not available_test_cases:
            QMessageBox.information(self, "No Test Cases", "No other test cases available to add as prerequisites.")
            return
        
        # Create selection dialog
        selection_dialog = QDialog(self)
        selection_dialog.setWindowTitle("Select Prerequisite Test Case(s)")  # ✅ CHANGED: Added (s)
        selection_dialog.setMinimumSize(400, 300)
        
        layout = QVBoxLayout(selection_dialog)
        
        # ✅ NEW: Add search bar
        search_bar = QLineEdit()
        search_bar.setPlaceholderText("Search test cases...")
        layout.addWidget(search_bar)
        
        # ✅ NEW: Add instruction label
        instruction_label = QLabel("Hold Ctrl to select multiple test cases")
        instruction_label.setStyleSheet("color: #6b7280; font-size: 9pt; font-style: italic;")
        layout.addWidget(instruction_label)
        
        list_widget = QListWidget()
        list_widget.setSelectionMode(QListWidget.SelectionMode.ExtendedSelection)  # ✅ CHANGED: Enable multi-select
        list_widget.addItems(available_test_cases)
        layout.addWidget(list_widget)
        
        # ✅ NEW: Connect search bar to filter function
        def filter_prerequisites(query):
            query = query.strip().lower()
            for i in range(list_widget.count()):
                item = list_widget.item(i)
                item_text = item.text().lower()
                if query in item_text:
                    item.setHidden(False)
                else:
                    item.setHidden(True)
        
        search_bar.textChanged.connect(filter_prerequisites)
        
        button_box = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(selection_dialog.accept)
        button_box.rejected.connect(selection_dialog.reject)
        layout.addWidget(button_box)
        
        if selection_dialog.exec() == QDialog.DialogCode.Accepted:
            selected_items = list_widget.selectedItems()  # ✅ CHANGED: Now returns multiple items
            if selected_items:
                added_count = 0
                skipped_count = 0
                skipped_names = []
                
                # ✅ NEW: Process all selected items
                for item in selected_items:
                    prereq_name = item.text()
                    # Check if already added
                    if prereq_name not in self.prerequisite_chips:
                        self.add_prerequisite_chip(prereq_name)
                        added_count += 1
                    else:
                        skipped_count += 1
                        skipped_names.append(prereq_name)
                
                # ✅ NEW: Show summary message
                if added_count > 0 and skipped_count == 0:
                    QMessageBox.information(self, "Success", 
                                          f"Added {added_count} prerequisite(s) successfully.")
                elif added_count > 0 and skipped_count > 0:
                    QMessageBox.information(self, "Partially Added", 
                                          f"Added {added_count} prerequisite(s).\n\n"
                                          f"Skipped {skipped_count} (already in list): {', '.join(skipped_names)}")
                elif skipped_count > 0:
                    QMessageBox.information(self, "Already Added", 
                                          f"All selected test cases are already in the prerequisites list:\n{', '.join(skipped_names)}")

    def add_prerequisite_chip(self, prereq_name):
        """Adds a chip/tag button for a prerequisite."""
        # Create a container widget for the chip
        chip_widget = QWidget()
        chip_layout = QHBoxLayout(chip_widget)
        chip_layout.setContentsMargins(6, 1, 1, 1)  # ✅ Reduced margins
        chip_layout.setSpacing(3)  # ✅ Reduced spacing
        
        # Style the chip to look like a tag
        chip_widget.setStyleSheet("""
            QWidget {
                background-color: #f3e8ff;
                border: 1px solid #d4b3e6;
                border-radius: 10px;
            }
        """)
        chip_widget.setFixedHeight(20)  # ✅ Reduced from 24 to 20
        
        # Label with the test case name
        name_label = QLabel(prereq_name)
        name_label.setStyleSheet("background-color: transparent; border: none; color: #6B2C91; font-weight: bold; font-size: 9pt;")
        chip_layout.addWidget(name_label)
        
        # Close button (X)
        close_button = QPushButton("×")
        close_button.setFixedSize(18, 18)
        close_button.setStyleSheet("""
            QPushButton {
                background-color: transparent;
                border: none;
                color: #6B2C91;
                font-size: 16px;
                font-weight: bold;
                padding: 0px;
            }
            QPushButton:hover {
                color: #d32f2f;
                background-color: #ffebee;
                border-radius: 9px;
            }
        """)
        close_button.setCursor(Qt.CursorShape.PointingHandCursor)
        close_button.clicked.connect(lambda: self.remove_prerequisite_chip(prereq_name))
        chip_layout.addWidget(close_button)
        
        # Insert the chip before the stretch
        # Remove the stretch temporarily
        stretch_item = self.prerequisites_chips_layout.takeAt(self.prerequisites_chips_layout.count() - 1)
        
        # Add the chip
        self.prerequisites_chips_layout.addWidget(chip_widget)
        
        # Add the stretch back
        self.prerequisites_chips_layout.addItem(stretch_item)
        
        # Store reference
        self.prerequisite_chips[prereq_name] = chip_widget

    def remove_prerequisite_chip(self, prereq_name):
        """Removes a prerequisite chip."""
        if prereq_name in self.prerequisite_chips:
            chip_widget = self.prerequisite_chips[prereq_name]
            self.prerequisites_chips_layout.removeWidget(chip_widget)
            chip_widget.deleteLater()
            del self.prerequisite_chips[prereq_name]

    def get_prerequisites(self):
        """Returns the list of prerequisite test case names."""
        return list(self.prerequisite_chips.keys())
        
    def display_utility_screenshot_details(self, utility_step, reference_module_name):
        """
        Displays the fields of a utility screenshot step in the Module Details table.
        Similar to display_module_details but for utility steps.
        """
        self.details_table.setRowCount(0)
        
        # Get the module details
        module_details = self.modules.get(reference_module_name, {})
        labels = module_details.get('labels', [])
        
        if not labels:
            QMessageBox.information(
                self,
                "No Fields",
                f"Module '{reference_module_name}' has no fields to display."
            )
            return
        
        # Set table headers for screenshot step (3 columns)
        self.details_table.setColumnCount(3)
        self.details_table.setHorizontalHeaderLabels(["Field Name", "Position (Row, Col)", "Highlight"])
        self.details_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.details_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.details_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        
        self.details_table.setRowCount(len(labels))
        
        # Get existing highlight flags from utility step
        saved_highlights = {field['field_name']: field.get('highlight', False) 
                          for field in utility_step.get('fields', [])}
        
        self.details_table.blockSignals(True)
        
        for i, label_data in enumerate(labels):
            field_name = label_data.get('label') or label_data.get('text') or label_data.get('name', 'N/A')
            row = label_data.get('row', '')
            col = label_data.get('column', '')
            
            # Field name (read-only)
            field_name_item = QTableWidgetItem(field_name)
            field_name_item.setFlags(field_name_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.details_table.setItem(i, 0, field_name_item)
            
            # Position (read-only)
            position_item = QTableWidgetItem(f"({row}, {col})")
            position_item.setFlags(position_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.details_table.setItem(i, 1, position_item)
            
            # Highlight checkbox
            highlight_checkbox = QCheckBox()
            highlight_checkbox.setChecked(saved_highlights.get(field_name, False))
            
            # Center the checkbox in the cell
            checkbox_widget = QWidget()
            checkbox_layout = QHBoxLayout(checkbox_widget)
            checkbox_layout.addWidget(highlight_checkbox)
            checkbox_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)
            checkbox_layout.setContentsMargins(0, 0, 0, 0)
            
            self.details_table.setCellWidget(i, 2, checkbox_widget)
        
        self.details_table.blockSignals(False)
        
        # Store reference to the utility step being edited
        self.current_utility_step = utility_step
        
        # Show a message in status bar or add a save button hint
        QMessageBox.information(
            self,
            "Utility Screenshot Preview",
            f"Showing fields from module '{reference_module_name}'.\n\n"
            "Check the boxes to highlight fields in the screenshot.\n"
            "Click 'Save Step Fields' to save your changes."
        )     
        
    def add_break_step(self):
        """Adds a break/pause step to the test case."""
        step_number_text = self.test_step_number_input.text().strip()
        insert_index = None
        
        if step_number_text:
            try:
                step_number = int(step_number_text)
                if step_number < 1:
                    QMessageBox.warning(self, "Invalid Step Number", "Step number must be at least 1.")
                    return
                insert_index = step_number - 1
                if insert_index > len(self.added_steps):
                    insert_index = len(self.added_steps)
            except ValueError:
                QMessageBox.warning(self, "Invalid Input", "Please enter a valid step number.")
                return
        
        new_step = {
            "name": "Break: Review & Decision Point",
            "type": "break",
            "message": "",  # Will be filled by user in Module Details
            "fields": [],
            "utility_steps": []
        }
        
        if insert_index is not None:
            self.added_steps.insert(insert_index, new_step)
            success_msg = f"Break step inserted at step {step_number}."
        else:
            self.added_steps.append(new_step)
            success_msg = "Break step added."
        
        self.update_steps_list()
        self.test_step_number_input.clear()
        self.input_radio.setChecked(True)
        QMessageBox.information(self, "Success", success_msg)    

    def display_utility_module_details(self, utility_step, module_name):
        """
        Displays the fields of a utility module step with independent data.
        """
        self.details_table.setRowCount(0)
        
        # Reset table headers for module import (2 columns)
        self.details_table.setColumnCount(2)
        self.details_table.setHorizontalHeaderLabels(["Field Name", "Value"])
        self.details_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        self.details_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        
        # Display EXACTLY what's in the utility step's fields array
        saved_fields_list = utility_step.get('fields', [])
        
        if not saved_fields_list:
            QMessageBox.information(
                self,
                "No Fields",
                f"No fields found in this utility step."
            )
            return
        
        self.details_table.setRowCount(len(saved_fields_list))
        
        self.details_table.blockSignals(True)
        
        for i, field_data in enumerate(saved_fields_list):
            field_name = field_data.get('field_name', 'N/A')
            field_value = field_data.get('value', '')
            
            # Display the field name
            field_name_item = QTableWidgetItem(field_name)
            field_name_item.setFlags(field_name_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.details_table.setItem(i, 0, field_name_item)
            
            # Display the field value
            value_item = QLineEdit()
            value_item.setText(field_value)
            value_item.setToolTip(self.get_available_variables_info())
            
            self.details_table.setCellWidget(i, 1, value_item)
        
        self.details_table.blockSignals(False)
        
        # Store reference to the utility step being edited
        self.current_utility_step = utility_step

        # -----------------------------------------------------
        # NEW LOGIC TO SHOW CORRECT MESSAGE BASED ON ACTION TYPE
        # -----------------------------------------------------
        # Determine action type of this utility step
        action_type = None
        if saved_fields_list:
            action_type = saved_fields_list[0].get('action_type', 'Input')

        # Build dynamic title & instruction text
        title = (
            "Utility Module - INPUT Mode"
            if action_type == "Input"
            else "Utility Module - VALIDATION Mode"
        )

        action_line = (
            "Enter values to INPUT into this module."
            if action_type == "Input"
            else "Enter expected values for VALIDATION."
        )

        # Dynamic popup message
        QMessageBox.information(
            self,
            title,
            f"✅ Independent Fields\n\n"
            f"Showing fields from module '{module_name}'.\n\n"
            f"These fields are INDEPENDENT from the main step.\n"
            f"Changes here will NOT affect the main step data.\n\n"
            f"{action_line}\n"
            f"Click 'Save Step Fields' to save your changes."
        )

    def refresh_utility_steps_for_main_step(self, main_step_index):
        """Refreshes the utility steps list for a specific main step without clearing it unnecessarily."""
        if main_step_index < 0 or main_step_index >= len(self.added_steps):
            self.utility_steps_list_widget.clear()
            return
        
        step_data = self.added_steps[main_step_index]
        utility_steps = step_data.get('utility_steps', [])
        
        # Clear and repopulate utility steps list
        self.utility_steps_list_widget.clear()
        
        main_step_num = main_step_index + 1  # 1-based step number
        
        for sub_idx, utility_step in enumerate(utility_steps, 1):
            step_name = utility_step.get('name', 'Utility')
            
            # Add sub-step number prefix (e.g., "Step 1.1: ...")
            numbered_name = f"Step {main_step_num}.{sub_idx}: {step_name}"
            
            # Check if this is a validate utility step
            is_validate_utility = False
            if utility_step.get('type') == 'module_import':
                fields = utility_step.get('fields', [])
                if fields and fields[0].get('action_type') == 'Validate':
                    is_validate_utility = True
            
            # Create custom widget with delete button
            list_item = QListWidgetItem()
            item_widget = QWidget()
            item_layout = QHBoxLayout(item_widget)
            item_layout.setContentsMargins(4, 2, 4, 2)
            item_layout.setSpacing(5)
            
            # Step name label
            name_label = QLabel(numbered_name)
            name_label.setMinimumWidth(200)
            name_label.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
            name_label.setWordWrap(False)
            item_layout.addWidget(name_label, 1)
            
            item_layout.addStretch()
            
            # Delete button
            delete_button = QPushButton("✕")
            delete_button.setFixedSize(18, 18)
            delete_font = QFont()
            delete_font.setBold(True)
            delete_font.setPointSize(12)
            delete_button.setFont(delete_font)
            delete_button.setStyleSheet("""
                QPushButton {
                    color: #dc2626;
                    border: none;
                    background-color: transparent;
                    padding: 0px;
                }
                QPushButton:hover {
                    color: #991b1b;
                    background-color: #fee2e2;
                    border-radius: 3px;
                }
            """)
            delete_button.setToolTip(f"Delete utility step {main_step_num}.{sub_idx}")
            delete_button.clicked.connect(
                lambda checked, ms=main_step_index, us=sub_idx-1: self.delete_utility_step(ms, us)
            )
            item_layout.addWidget(delete_button)
            
            # Apply yellow background for validate utility steps
            if is_validate_utility:
                item_widget.setStyleSheet("background-color: #fffacd;")
            
            list_item.setData(Qt.ItemDataRole.UserRole, {
                'main_step_index': main_step_index,
                'utility_step': utility_step,
                'sub_index': sub_idx
            })
            
            self.utility_steps_list_widget.addItem(list_item)
            self.utility_steps_list_widget.setItemWidget(list_item, item_widget)
            size_hint = list_item.sizeHint()
            size_hint.setHeight(max(32, size_hint.height()))
            list_item.setSizeHint(size_hint)
            
            
    # Add this new method to the EditTestCaseDialog class

    def add_quick_utility_module(self, action_type):
        """
        Quickly adds a utility step with the same module as the selected main step.
        
        Args:
            action_type: Either 'Input' or 'Validate'
        """
        # Get the currently selected main step
        current_item = self.steps_list_widget.currentItem()
        if not current_item:
            QMessageBox.warning(self, "No Selection", "Please select a main step first.")
            return
        
        current_step_index = self.steps_list_widget.row(current_item)
        if current_step_index < 0 or current_step_index >= len(self.added_steps):
            return
        
        step_data = self.added_steps[current_step_index]
        
        # Only works for module_import steps
        if step_data.get('type') != 'module_import':
            QMessageBox.warning(self, "Invalid Step", 
                              "Quick utility buttons only work for module import steps.")
            return
        
        module_name = step_data.get('module_name')
        if not module_name or module_name not in self.modules:
            QMessageBox.warning(self, "Module Not Found", 
                              f"Module '{module_name}' not found.")
            return
        
        # Ensure utility_steps field exists
        if 'utility_steps' not in step_data:
            step_data['utility_steps'] = []
        
        # Get module fields
        module_data = self.modules[module_name]
        fields_from_module = json.loads(json.dumps(module_data.get('labels', [])))
        
        # Generate unique field identifiers
        step_number = current_step_index + 1
        utility_number = len(step_data['utility_steps']) + 1
        
        fields = []
        for idx, label_data in enumerate(fields_from_module):
            field_name = label_data.get('label') or label_data.get('text') or label_data.get('name', 'N/A')
            
            import time
            import random
            unique_timestamp = int(time.time() * 1000000)
            random_suffix = random.randint(100000, 999999)
            
            unique_field_name = f"{field_name}_U{step_number}_{utility_number}_{unique_timestamp}_{random_suffix}_{idx}"
            
            fields.append({
                "field_name": field_name,
                "internal_field_id": unique_field_name,
                "action_type": action_type,
                "value": "",
            })
        
        # Create utility step
        utility_step = {
            "name": f"Import Module: {module_name}",
            "type": "module_import",
            "module_name": module_name,
            "action_type": action_type,
            "fields": fields
        }
        
        # Add to utility steps
        step_data['utility_steps'].append(utility_step)
        
        # Refresh utility steps list while keeping main step selected
        self.steps_list_widget.blockSignals(True)
        self.steps_list_widget.setCurrentRow(current_step_index)
        self.steps_list_widget.blockSignals(False)
        
        # Manually refresh utility steps list
        self.refresh_utility_steps_for_main_step(current_step_index)
        
        # Show success message
        self.statusBar().showMessage(
            f"Added {action_type} utility step for module '{module_name}'", 
            3000
        )            
    
    def handle_step_drop(self, event):
        """
        Handles dropping a main step onto another main step to convert it to a utility step.
        ✅ FIXED: Now also moves all utility steps from the source step.
        """
        source_item = self.steps_list_widget.currentItem()
        
        # Get the drop position
        target_item = self.steps_list_widget.itemAt(event.position().toPoint())
        
        if not source_item or not target_item:
            return
        
        source_row = self.steps_list_widget.row(source_item)
        target_row = self.steps_list_widget.row(target_item)
        
        # Validate that both items exist in added_steps
        if source_row < 0 or source_row >= len(self.added_steps):
            return
        
        if target_row < 0 or target_row >= len(self.added_steps):
            return
        
        source_step = self.added_steps[source_row]
        target_step = self.added_steps[target_row]
        
        # Don't allow dropping on itself
        if source_row == target_row:
            return
        
        # Ask user for confirmation
        reply = QMessageBox.question(
            self,
            "Convert to Utility Step",
            f"Convert Step {source_row + 1} into a utility step under Step {target_row + 1}?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.Yes
        )
        
        if reply != QMessageBox.StandardButton.Yes:
            # User cancelled - just refresh to ensure clean state
            self.update_steps_list()
            return
        
        # Ensure target step has utility_steps array
        if 'utility_steps' not in target_step:
            target_step['utility_steps'] = []
        
        # ✅ NEW: Create a deep copy to avoid reference issues
        import copy
        utility_step_copy = copy.deepcopy(source_step)
        
        # ✅ NEW: Extract utility steps from the source step BEFORE adding it as a utility
        source_utility_steps = source_step.get('utility_steps', [])
        
        # ✅ NEW: Remove utility_steps from the copy since it will become a utility step itself
        # (utility steps shouldn't have nested utility steps)
        if 'utility_steps' in utility_step_copy:
            del utility_step_copy['utility_steps']
        
        # Add source step as a utility step to target (without its utility_steps)
        target_step['utility_steps'].append(utility_step_copy)
        
        # ✅ NEW: Now add all of the source step's utility steps to the target
        if source_utility_steps:
            for util_step in source_utility_steps:
                util_step_copy = copy.deepcopy(util_step)
                target_step['utility_steps'].append(util_step_copy)
        
        # Remove source step from main steps
        del self.added_steps[source_row]
        
        # Calculate the correct target index after removal
        adjusted_target_row = target_row if source_row > target_row else target_row - 1
        
        # Block signals during refresh
        self.steps_list_widget.blockSignals(True)
        self.utility_steps_list_widget.blockSignals(True)
        
        try:
            # Refresh the UI completely
            self.update_steps_list()
            
            # Select the target step and show its utility steps
            if 0 <= adjusted_target_row < self.steps_list_widget.count():
                self.steps_list_widget.setCurrentRow(adjusted_target_row)
                
                # Get the item at the adjusted row
                target_item_after_refresh = self.steps_list_widget.item(adjusted_target_row)
                if target_item_after_refresh:
                    # Manually trigger the utility steps display
                    self.on_main_step_clicked(target_item_after_refresh)
        
        finally:
            # Unblock signals
            self.steps_list_widget.blockSignals(False)
            self.utility_steps_list_widget.blockSignals(False)
        
        # Process pending events
        QApplication.processEvents()
        
        # ✅ UPDATED: Show success message with count of utility steps moved
        utility_count = len(source_utility_steps)
        if utility_count > 0:
            QMessageBox.information(
                self,
                "Success",
                f"Step {source_row + 1} and its {utility_count} utility step(s) converted to utility steps under Step {adjusted_target_row + 1}!"
            )
        else:
            QMessageBox.information(
                self,
                "Success",
                f"Step {source_row + 1} converted to utility step under Step {adjusted_target_row + 1}!"
            )
    
    def add_quick_enter_utility(self):
        """Quickly adds Enter key with auto-wait as utility steps."""
        current_item = self.steps_list_widget.currentItem()
        if not current_item:
            QMessageBox.warning(self, "No Selection", "Please select a main step first.")
            return
        
        current_step_index = self.steps_list_widget.row(current_item)
        if current_step_index < 0 or current_step_index >= len(self.added_steps):
            return
        
        step_data = self.added_steps[current_step_index]
        
        # Only works for module_import steps
        if step_data.get('type') != 'module_import':
            QMessageBox.warning(self, "Invalid Step", 
                              "Quick utility buttons only work for module import steps.")
            return
        
        # Ensure utility_steps field exists
        if 'utility_steps' not in step_data:
            step_data['utility_steps'] = []
        
        # Add Enter key utility step
        enter_utility = {
            "name": "Special Key: Enter Key",
            "type": "special_key",
            "key_value": "Enter Key"
        }
        step_data['utility_steps'].append(enter_utility)
        
        # Add Wait utility step
        wait_utility = {
            "name": "Wait: 1 second(s)",
            "type": "wait",
            "seconds": 1
        }
        step_data['utility_steps'].append(wait_utility)
        
        # Refresh utility steps list while keeping main step selected
        self.steps_list_widget.blockSignals(True)
        self.steps_list_widget.setCurrentRow(current_step_index)
        self.steps_list_widget.blockSignals(False)
        
        # Manually refresh utility steps list
        self.refresh_utility_steps_for_main_step(current_step_index)
        
        # Show success message
        self.statusBar().showMessage(
            f"Added Enter key with auto-wait utility steps", 
            3000
        )

    def add_quick_clear_utility(self):
        """Quickly adds Clear key with auto-wait as utility steps."""
        current_item = self.steps_list_widget.currentItem()
        if not current_item:
            QMessageBox.warning(self, "No Selection", "Please select a main step first.")
            return
        
        current_step_index = self.steps_list_widget.row(current_item)
        if current_step_index < 0 or current_step_index >= len(self.added_steps):
            return
        
        step_data = self.added_steps[current_step_index]
        
        # Only works for module_import steps
        if step_data.get('type') != 'module_import':
            QMessageBox.warning(self, "Invalid Step", 
                              "Quick utility buttons only work for module import steps.")
            return
        
        # Ensure utility_steps field exists
        if 'utility_steps' not in step_data:
            step_data['utility_steps'] = []
        
        # Add Clear key utility step
        clear_utility = {
            "name": "Special Key: Clear Key",
            "type": "special_key",
            "key_value": "Clear Key"
        }
        step_data['utility_steps'].append(clear_utility)
        
        # Add Wait utility step
        wait_utility = {
            "name": "Wait: 1 second(s)",
            "type": "wait",
            "seconds": 1
        }
        step_data['utility_steps'].append(wait_utility)
        
        # Refresh utility steps list while keeping main step selected
        self.steps_list_widget.blockSignals(True)
        self.steps_list_widget.setCurrentRow(current_step_index)
        self.steps_list_widget.blockSignals(False)
        
        # Manually refresh utility steps list
        self.refresh_utility_steps_for_main_step(current_step_index)
        
        # Show success message
        self.statusBar().showMessage(
            f"Added Clear key with auto-wait utility steps", 
            3000
        )    
    
class StepListItemWidget(QFrame):
    def __init__(self, step_name, is_validate_step=False, with_utility_steps=False, parent=None):
        super().__init__(parent)
        self.step_name = step_name
        self.parent_dialog = parent
        
        # Set fixed height to prevent text cutting
        self.setFixedHeight(28)
        
        self.main_layout = QHBoxLayout(self)
        self.main_layout.setContentsMargins(4, 0, 4, 4)
        self.main_layout.setSpacing(8)
        
        self.label = QLabel(step_name)
        self.label.setWordWrap(False)
        self.label.setMinimumWidth(200)
        self.label.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.main_layout.addWidget(self.label, 1)
        
        # Add stretch to push utility steps and delete button to the right
        self.main_layout.addStretch()
        
        # ✅ Utility steps container (only if with_utility_steps is True)
        if with_utility_steps:
            self.utility_container = QWidget()
            self.utility_container.setObjectName("utility_steps_container")
            self.utility_layout = QHBoxLayout(self.utility_container)
            self.utility_layout.setContentsMargins(0, 0, 8, 0)
            self.utility_layout.setSpacing(4)
            self.utility_container.setStyleSheet("QWidget { margin: 0px; padding: 0px; }")
            self.main_layout.addWidget(self.utility_container, 0, Qt.AlignmentFlag.AlignVCenter)
        
        self.delete_button = QPushButton("✕")
        self.delete_button.setFixedSize(20, 20)

        delete_font = QFont()
        delete_font.setBold(True)
        delete_font.setPointSize(14)
        self.delete_button.setFont(delete_font)

        self.delete_button.setStyleSheet("""
            QPushButton {
                color: #dc2626;
                border: none;
                background-color: transparent;
                margin-bottom: 0px;
                padding: 0px;
                font-weight: bold;
            }
            QPushButton:hover {
                color: #991b1b;
                background-color: #fee2e2;
                border-radius: 3px;
            }
            QPushButton:pressed {
                color: #7f1d1d;
                background-color: #fecaca;
            }
        """)
        
        self.main_layout.addWidget(self.delete_button, 0, Qt.AlignmentFlag.AlignVCenter)
        
        if is_validate_step:
            self.setStyleSheet("""
                QFrame { 
                    background-color: #fffacd; 
                    margin-bottom: 2px;
                    border-radius: 2px;
                }
            """)
        else:
            self.setStyleSheet("""
                QFrame { 
                    background-color: transparent; 
                    margin-bottom: 2px;
                }
            """)
        
    def add_utility_chip(self, utility_step, on_delete_callback):
        """Adds a utility step chip to the container with preview button for screenshots and modules."""
        if not hasattr(self, 'utility_container'):
            return
        
        # ✅ NEW: Check if this is a validate utility step
        is_validate_utility = False
        if utility_step.get('type') == 'module_import':
            fields = utility_step.get('fields', [])
            if fields and fields[0].get('action_type') == 'Validate':
                is_validate_utility = True
        
        # Create chip widget
        chip_widget = QWidget()
        chip_layout = QHBoxLayout(chip_widget)
        chip_layout.setContentsMargins(8, 0, 4, 0)
        chip_layout.setSpacing(4)
        chip_layout.setAlignment(Qt.AlignmentFlag.AlignVCenter)
        
        # ✅ CHANGED: Set background color based on action type
        if is_validate_utility:
            chip_widget.setStyleSheet("""
                QWidget {
                    background-color: #fffacd;
                    border: 1px solid #f0e68c;
                    border-radius: 9px;
                    margin: 0px;
                    padding: 0px;
                }
            """)
        else:
            chip_widget.setStyleSheet("""
                QWidget {
                    background-color: #dcfce7;
                    border: 1px solid #86efac;
                    border-radius: 9px;
                    margin: 0px;
                    padding: 0px;
                }
            """)
        
        chip_widget.setFixedHeight(22)
        chip_widget.setMinimumWidth(120)
        
        # Determine short name for the chip - WITHOUT icons
        step_name = utility_step.get('name', 'Utility')
        import re
        step_match = re.match(r'Step (\d+\.\d+):', step_name)
        step_prefix = f"{step_match.group(1)}: " if step_match else ""
        
        if step_name.startswith("Wait:"):
            seconds = utility_step.get('seconds', '?')
            short_name = f"{step_prefix}Wait {seconds}s"
        elif step_name.startswith("Special Key:"):
            key = utility_step.get('key_value', '?')
            short_name = f"{step_prefix}{key}"
        elif "Screenshot" in step_name and "DOCX" in step_name:
            short_name = f"{step_prefix}Screenshot Doc"
        elif "Screenshot" in step_name:
            short_name = f"{step_prefix}Screenshot Txt"
        elif step_name.startswith("Import Module:"):
            module_name = utility_step.get('module_name', 'Module')
            short_name = f"{step_prefix}{module_name}"
        else:
            short_name = (step_prefix + step_name[:15]) if not step_prefix else step_name[:18]
        
        # ✅ CHANGED: Adjust label color based on validate status
        label_color = "#806000" if is_validate_utility else "#15803d"
        
        name_label = QLabel(short_name)
        name_label.setStyleSheet(f"""
            QLabel {{
                background-color: transparent; 
                border: none; 
                color: {label_color}; 
                font-weight: bold; 
                font-size: 8pt;
                padding: 0px;
                margin: 0px;
            }}
        """)
        name_label.setAlignment(Qt.AlignmentFlag.AlignVCenter)
        chip_layout.addWidget(name_label, 0, Qt.AlignmentFlag.AlignVCenter)
        
        # Add eye button for module import AND screenshot utility steps
        if utility_step.get('type') in ['capture_screenshot', 'module_import']:
            eye_button = QPushButton("👁")
            eye_button.setFixedSize(16, 16)
            
            eye_font = QFont()
            eye_font.setPointSize(10)
            eye_button.setFont(eye_font)
            
            # ✅ CHANGED: Adjust button style based on validate status
            eye_bg_color = "#f0e68c" if is_validate_utility else "#bef264"
            eye_button.setStyleSheet(f"""
                QPushButton {{
                    background-color: transparent;
                    border: none;
                    color: {label_color};
                    padding: 0px;
                    margin: 0px;
                }}
                QPushButton:hover {{
                    background-color: {eye_bg_color};
                    border-radius: 8px;
                }}
            """)
            eye_button.setToolTip("Preview fields")
            eye_button.setCursor(Qt.CursorShape.PointingHandCursor)
            
            # Connect to preview handler
            eye_button.clicked.connect(lambda: self.preview_utility_step(utility_step))
            chip_layout.addWidget(eye_button, 0, Qt.AlignmentFlag.AlignVCenter)
        
        # Close button
        close_button = QPushButton("×")
        close_button.setFixedSize(16, 16)
        
        close_font = QFont()
        close_font.setPointSize(14)
        close_font.setBold(True)
        close_button.setFont(close_font)
        
        # ✅ CHANGED: Adjust close button style based on validate status
        close_hover_bg = "#fff9e6" if is_validate_utility else "#ffebee"
        close_button.setStyleSheet(f"""
            QPushButton {{
                background-color: transparent;
                border: none;
                color: {label_color};
                font-weight: bold;
                padding: 0px;
                margin: 0px;
                margin-bottom: 2px;
            }}
            QPushButton:hover {{
                color: #d32f2f;
                background-color: {close_hover_bg};
                border-radius: 8px;
            }}
        """)
        close_button.setCursor(Qt.CursorShape.PointingHandCursor)
        close_button.clicked.connect(on_delete_callback)
        chip_layout.addWidget(close_button, 0, Qt.AlignmentFlag.AlignVCenter)
        
        self.utility_layout.addWidget(chip_widget)

    def preview_utility_step(self, utility_step):
        """✅ NEW: Handles preview button click for utility steps (both screenshot and module import)."""
        if not self.parent_dialog:
            return
        
        # Get the step index of this utility step's parent
        step_index = None
        for i in range(self.parent_dialog.steps_list_widget.count()):
            item = self.parent_dialog.steps_list_widget.item(i)
            widget = self.parent_dialog.steps_list_widget.itemWidget(item)
            if widget == self:
                step_index = i
                break
        
        if step_index is None:
            return
        
        utility_type = utility_step.get('type')
        
        if utility_type == 'capture_screenshot':
            # Existing screenshot preview logic
            self.preview_utility_screenshot(utility_step, step_index)
        elif utility_type == 'module_import':
            # ✅ NEW: Module validation preview logic
            self.preview_utility_module(utility_step, step_index)

    def preview_utility_screenshot(self, utility_step, step_index):
        """Preview screenshot utility step (existing code)."""
        # Find the reference module
        reference_module_name = None
        for i in range(step_index, -1, -1):
            check_step = self.parent_dialog.added_steps[i]
            if check_step.get('type') == 'module_import':
                reference_module_name = check_step.get('module_name')
                break
        
        if not reference_module_name:
            QMessageBox.information(
                self.parent_dialog,
                "No Module Found",
                "No module found before this screenshot utility step."
            )
            return
        
        # Store reference in utility step
        utility_step['reference_module'] = reference_module_name
        
        # Display the fields in the details table
        self.parent_dialog.display_utility_screenshot_details(utility_step, reference_module_name)

    def preview_utility_module(self, utility_step, step_index):
        """✅ NEW: Preview module import utility step for validation."""
        module_name = utility_step.get('module_name')
        
        if not module_name or module_name not in self.parent_dialog.modules:
            QMessageBox.information(
                self.parent_dialog,
                "Module Not Found",
                f"Module '{module_name}' not found."
            )
            return
        
        # Display the fields in the details table for validation
        self.parent_dialog.display_utility_module_details(utility_step, module_name)        

# --- UPDATED CLASS: Dialog to get the test case name ---
class TestCaseNameDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("New Test Case")
        self.setMinimumSize(500, 400)
        self.resize(600, 450)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        self.test_case_name = ""
        self.test_case_description = ""
        self.test_case_assumptions = ""

        layout = QVBoxLayout(self)
        
        form_layout = QFormLayout()
        self.name_input = QLineEdit()
        self.name_input.setPlaceholderText("Enter test case name...")
        form_layout.addRow("Test Case Name:", self.name_input)

        # Add description field
        self.description_input = QTextEdit()
        self.description_input.setPlaceholderText("Enter test case description...")
        self.description_input.setMaximumHeight(80)
        form_layout.addRow("Description:", self.description_input)

        # ✅ UPDATED: Assumptions field with paste event handling
        self.assumptions_input = QTextEdit()
        self.assumptions_input.setPlaceholderText("Enter test case assumptions (supports tables)...")
        self.assumptions_input.setMinimumHeight(150)
        self.assumptions_input.setAcceptRichText(True)
        
        # ✅ NEW: Install event filter to intercept paste events
        self.assumptions_input.installEventFilter(self)

        # Toolbar for table creation
        assumptions_toolbar = QHBoxLayout()

        add_table_button = QPushButton("Insert Table")
        add_table_button.setFixedWidth(110)
        add_table_button.clicked.connect(self.insert_table)
        add_table_button.setToolTip("Insert a table into assumptions")
        assumptions_toolbar.addWidget(add_table_button)

        remove_table_button = QPushButton("Clear Format")
        remove_table_button.setFixedWidth(110)
        remove_table_button.clicked.connect(self.clear_formatting)
        remove_table_button.setToolTip("Remove all formatting and convert to plain text")
        assumptions_toolbar.addWidget(remove_table_button)

        assumptions_toolbar.addStretch()

        form_layout.addRow("Assumptions:", self.assumptions_input)
        form_layout.addRow("", assumptions_toolbar)
        
        layout.addLayout(form_layout)

        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept_dialog)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
    
    # ✅ NEW: Event filter to handle paste operations
# ✅ FIXED: Event filter to handle paste operations
    def eventFilter(self, obj, event):
        """Intercept paste events to handle table formatting."""
        from PyQt6.QtGui import QKeySequence
        
        if obj == self.assumptions_input and event.type() == event.Type.KeyPress:
            # Check for Ctrl+V paste
            if event.matches(QKeySequence.StandardKey.Paste):
                self.paste_with_formatting()
                return True  # Event handled, don't propagate
        return super().eventFilter(obj, event)
        
    # ✅ NEW: Custom paste function that preserves table borders
    def paste_with_formatting(self):
        """Paste content from clipboard while preserving table formatting."""
        from PyQt6.QtGui import QTextDocument, QTextCursor
        from PyQt6.QtCore import QMimeData
        
        clipboard = QApplication.clipboard()
        mime_data = clipboard.mimeData()
        
        if mime_data.hasHtml():
            # Get HTML content from clipboard
            html_content = mime_data.html()
            
            # Ensure tables have borders by modifying the HTML
            html_content = self.ensure_table_borders(html_content)
            
            # Insert the modified HTML
            cursor = self.assumptions_input.textCursor()
            cursor.insertHtml(html_content)
        elif mime_data.hasText():
            # Fallback to plain text if no HTML
            cursor = self.assumptions_input.textCursor()
            cursor.insertText(mime_data.text())
    
    # ✅ NEW: Helper function to add borders to pasted tables
    def ensure_table_borders(self, html):
        """Ensures that tables in HTML have visible borders."""
        import re
        
        # Add border styling to <table> tags if not present
        def add_table_border(match):
            table_tag = match.group(0)
            
            # Check if border is already defined
            if 'border' not in table_tag.lower():
                # Insert border style before the closing >
                table_tag = table_tag[:-1] + ' border="1" style="border-collapse: collapse; border: 1px solid black;">'
            elif 'style=' in table_tag.lower():
                # Add border to existing style
                table_tag = re.sub(
                    r'style="([^"]*)"',
                    r'style="\1; border-collapse: collapse; border: 1px solid black;"',
                    table_tag,
                    flags=re.IGNORECASE
                )
            
            return table_tag
        
        # Apply border to all <table> tags
        html = re.sub(r'<table[^>]*>', add_table_border, html, flags=re.IGNORECASE)
        
        # Also ensure <td> and <th> cells have borders
        html = re.sub(
            r'<(td|th)([^>]*)>',
            r'<\1\2 style="border: 1px solid black; padding: 4px;">',
            html,
            flags=re.IGNORECASE
        )
        
        return html
    
    def insert_table(self):
        """Inserts a table into the assumptions field."""
        from PyQt6.QtWidgets import QInputDialog
        
        # Ask user for table dimensions
        rows, ok1 = QInputDialog.getInt(self, "Table Rows", "Enter number of rows:", 3, 1, 20)
        if not ok1:
            return
        
        cols, ok2 = QInputDialog.getInt(self, "Table Columns", "Enter number of columns:", 3, 1, 10)
        if not ok2:
            return
        
        # Get the text cursor
        cursor = self.assumptions_input.textCursor()
        
        # Create table format with borders
        from PyQt6.QtGui import QTextTableFormat, QTextFrameFormat
        table_format = QTextTableFormat()
        table_format.setBorderStyle(QTextFrameFormat.BorderStyle.BorderStyle_Solid)
        table_format.setBorder(1)
        table_format.setCellPadding(4)
        table_format.setCellSpacing(0)
        
        # Insert the table
        cursor.insertTable(rows, cols, table_format)

    def clear_formatting(self):
        """Removes all formatting from the assumptions field."""
        reply = QMessageBox.question(
            self,
            "Clear Formatting",
            "Remove all tables and formatting? This will convert to plain text.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            # Get plain text and reset
            plain_text = self.assumptions_input.toPlainText()
            self.assumptions_input.clear()
            self.assumptions_input.setPlainText(plain_text)

    def accept_dialog(self):
        self.test_case_name = self.name_input.text().strip()
        self.test_case_description = self.description_input.toPlainText().strip()
        self.test_case_assumptions = self.assumptions_input.toHtml()
        if not self.test_case_name:
            QMessageBox.warning(self, "Invalid Name", "Test case name cannot be empty.")
        else:
            self.accept()

class SplashScreen(QWidget):
    """
    Custom splash screen with fade-in effect.
    """
    def __init__(self):
        super().__init__()
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.WindowStaysOnTopHint)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        
        # Set up the layout
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Load and display the image
        # Load and display the image
        self.label = QLabel()
        pixmap = QPixmap("splash_screen.png")
        
        # Scale the image to a smaller size
        if not pixmap.isNull():
            pixmap = pixmap.scaled(600, 400, Qt.AspectRatioMode.KeepAspectRatio, 
                                   Qt.TransformationMode.SmoothTransformation)
        
        self.label.setPixmap(pixmap)
        self.label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.label)
        
        # Adjust window size to fit the image
        self.adjustSize()
        
        # Center the splash screen (must be called after adjustSize)
        self.center_on_screen()
        
        # Set initial opacity to 0 for fade-in effect
        self.setWindowOpacity(0.0)
        
        # Create fade-in animation
        self.fade_animation = None
        self.start_fade_in()
        
    def center_on_screen(self):
        """Centers the splash screen on the screen."""
        screen = QApplication.primaryScreen().geometry()
        size = self.geometry()
        self.move(
            (screen.width() - size.width()) // 2,
            (screen.height() - size.height()) // 2
        )
    
    def start_fade_in(self):
        """Starts the fade-in animation."""
        from PyQt6.QtCore import QPropertyAnimation, QEasingCurve
        
        self.fade_animation = QPropertyAnimation(self, b"windowOpacity")
        self.fade_animation.setDuration(800)  # 800ms fade-in duration
        self.fade_animation.setStartValue(0.0)
        self.fade_animation.setEndValue(1.0)
        self.fade_animation.setEasingCurve(QEasingCurve.Type.InOutQuad)
        self.fade_animation.start()

class CustomPCOMMTextEdit(QTextEdit):
    """Custom QTextEdit that only shows label dialog on mouse release."""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.main_window = None
        self.is_selecting = False
    
    def mousePressEvent(self, event):
        """Track when mouse button is pressed."""
        if event.button() == Qt.MouseButton.LeftButton:
            self.is_selecting = True
        super().mousePressEvent(event)
    
    def mouseReleaseEvent(self, event):
        """Override mouse release to trigger label dialog only after selection is complete."""
        # Call parent implementation first
        super().mouseReleaseEvent(event)
        
        # Only proceed if left mouse button was released after selecting
        if event.button() == Qt.MouseButton.LeftButton and self.is_selecting:
            self.is_selecting = False
            
            # Check if there's actually a selection
            cursor = self.textCursor()
            if self.main_window and cursor.hasSelection():
                start_pos = cursor.selectionStart()
                start_cursor = self.textCursor()
                start_cursor.setPosition(start_pos, QTextCursor.MoveMode.MoveAnchor)
                
                start_row = start_cursor.blockNumber() + 1
                start_col = start_cursor.columnNumber() + 1
                selected_text = cursor.selectedText()
                selection_length = len(selected_text.replace('\n', ''))
                
                # Update status bar first
                self.main_window.update_selection_info_status_only(start_row, start_col, selected_text, selection_length)
                
                # Use QTimer to delay the dialog so mouse release completes first
                QTimer.singleShot(150, lambda: self.main_window.prompt_define_label(start_row, start_col, selection_length))

class BreakExecutionDialog(QDialog):
    """Dialog shown when execution reaches a Break step."""
    
    STOP = 0
    RESUME = 1
    EDIT = 2
    
    def __init__(self, message, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Break Point - Execution Paused")
        self.setMinimumSize(500, 300)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        
        self.result_action = None
        
        layout = QVBoxLayout(self)
        
        # Title
        title_label = QLabel("⏸️ Execution Paused at Break Point")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(title_label)
        
        layout.addSpacing(10)
        
        # Message display
        message_label = QLabel("Message:")
        message_label.setStyleSheet("font-weight: bold;")
        layout.addWidget(message_label)
        
        message_display = QTextEdit()
        message_display.setReadOnly(True)
        message_display.setAcceptRichText(True)  # ✅ NEW: Enable rich text
        
        # ✅ CHANGED: Use setHtml to preserve formatting (bold, italic, etc.)
        if message:
            message_display.setHtml(message)
        else:
            message_display.setPlainText("No message provided.")
        
        message_display.setMinimumHeight(150)  # ✅ CHANGED: Use minimum instead of maximum
        message_display.setMaximumHeight(300)  # ✅ NEW: Allow larger display
        layout.addWidget(message_display)
        
        layout.addSpacing(20)
        
        # Buttons
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        stop_button = QPushButton("🛑 Stop Execution")
        stop_button.setStyleSheet("""
            QPushButton {
                background-color: #dc2626;
                color: white;
                padding: 8px 16px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #991b1b;
            }
        """)
        stop_button.clicked.connect(lambda: self.set_result(self.STOP))
        button_layout.addWidget(stop_button)
        
        edit_button = QPushButton("✏️ Edit Test Case")
        edit_button.setStyleSheet("""
            QPushButton {
                background-color: #f59e0b;
                color: white;
                padding: 8px 16px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #d97706;
            }
        """)
        edit_button.clicked.connect(lambda: self.set_result(self.EDIT))
        button_layout.addWidget(edit_button)
        
        resume_button = QPushButton("▶️ Resume Execution")
        resume_button.setStyleSheet("""
            QPushButton {
                background-color: #16a34a;
                color: white;
                padding: 8px 16px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #15803d;
            }
        """)
        resume_button.clicked.connect(lambda: self.set_result(self.RESUME))
        button_layout.addWidget(resume_button)
        
        button_layout.addStretch()
        layout.addLayout(button_layout)
    
    def set_result(self, action):
        """Sets the result action and closes the dialog."""
        self.result_action = action
        
        # ✅ CHANGED: Only close for STOP and RESUME, not for EDIT
        if action in (self.STOP, self.RESUME):
            self.accept()
    
    def get_action(self):
        """Returns the selected action."""
        return self.result_action

class PCOMMMainFrame(QMainWindow):
    """
    The main window for the PCOMM desktop application,
    mimicking the TOSCA UI layout.
    """
    copied_labels = []
    def __init__(self):
        super().__init__()
        palette = self.palette()
        palette.setColor(QPalette.ColorRole.Highlight, QColor("#6B2C91"))
        self.setPalette(palette)
        self.setGeometry(100, 100, 1000, 700)

        self.setWindowTitle("InstaRun v1.0")
        self.set_application_icon()


        # --- Style Sheet remains the same ---
        self.setStyleSheet("""
            /* Main Window */
            QMainWindow { 
                background-color: #f5f7fa; 
            }
            
            /* Dock Widgets - Modern card style */
            QDockWidget { 
                border: none;
                titlebar-close-icon: url(close.png);
                titlebar-normal-icon: url(float.png);
            }
            QDockWidget::title {
                text-align: left; 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #ffffff, stop:1 #f0f2f5);
                padding: 8px;
                border-bottom: 1px solid #d1d5db;
                font-weight: 600;
                color: #1f2937;
            }
            QDockWidget::close-button, QDockWidget::float-button {
                border: none;
                background: transparent;
                padding: 3px;
            }
            QDockWidget::close-button:hover, QDockWidget::float-button:hover {
                background: #e5e7eb;
                border-radius: 4px;
            }
            
            /* Menu Bar - Clean modern look */
            QMenuBar { 
                background-color: #ffffff; 
                border-bottom: 1px solid #e5e7eb;
                padding: 2px;
            }
            QMenuBar::item {
                padding: 3px 12px;
                border-radius: 4px;
                color: #374151;
            }
            QMenuBar::item:selected {
                background-color: #f3f4f6;
            }
            QMenuBar::item:pressed {
                background-color: #e5e7eb;
            }
            
            /* Toolbar - Modern with icons */
            QToolBar { 
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #ffffff, stop:1 #f9fafb);
                border-bottom: 1px solid #e5e7eb;
                spacing: 3px;
                padding: 2px 6px;
            }
            QToolBar::separator {
                background: #d1d5db;
                width: 1px;
                margin: 2px 6px;
            }

            /* Tool Buttons - Elevated card style */
            QToolButton {
                color: #1f2937;
                border: none;
                border-radius: 6px;
                padding: 1px 2px;
                margin: 1px;
                background-color: transparent;
                font-weight: 500;
            }
            QToolButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #f3f4f6, stop:1 #e5e7eb);
            }
            QToolButton:pressed {
                background: #d1d5db;
            }
            QToolButton:checked {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #e9d5f5, stop:1 #d4b3e6);
                color: #6B2C91;
                border: 1px solid #b388cc;
            }
            
            /* Status Bar */
            QStatusBar { 
                background-color: #f9fafb;
                border-top: 1px solid #e5e7eb;
                color: #6b7280;
                padding: 4px;
            }
                        
            /* Tree Widget - Clean with hover effects */
            QTreeWidget { 
                background-color: #ffffff; 
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 4px;
                selection-background-color: #e9d5f5;
                selection-color: #6B2C91;
                outline: none;
            }
            QTreeWidget::item {
                padding: 3px 6px;
                border-radius: 4px;
            }
            QTreeWidget::item:hover {
                background-color: #f3f4f6;
            }
            QTreeWidget::item:selected {
                background-color: #e9d5f5;
                color: #6B2C91;
            }
            
            /* List Widget */
            QListWidget {
                background-color: #ffffff;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                padding: 4px;
                selection-background-color: #e9d5f5;
                selection-color: #6B2C91;
            }
            QListWidget::item {
                padding: 4px 8px;
                border-radius: 4px;
                margin: 1px;
            }
            QListWidget::item:hover {
                background-color: #f3f4f6;
            }
            QListWidget::item:selected {
                background-color: #e9d5f5;
                color: #6B2C91;
            }
            
            /* Table Widget - Modern grid */
            QTableWidget {
                background-color: #ffffff;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                gridline-color: #f3f4f6;
                selection-background-color: #e9d5f5;
                selection-color: #6B2C91;
            }
            QTableWidget::item {
                padding: 3px 6px;
            }
            QTableWidget::item:hover {
                background-color: #f9fafb;
            }
            QHeaderView::section {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #f9fafb, stop:1 #f3f4f6);
                color: #374151;
                padding: 1px 2px;
                border: none;
                border-bottom: 2px solid #e5e7eb;
                border-right: 1px solid #e5e7eb;
                font-weight: 600;
            }
            QHeaderView::section:hover {
                background: #f3f4f6;
            }

            /* Explicitly target TreeWidget headers */
            QTreeWidget QHeaderView::section {
                padding: 4px 8px;
                min-height: 20px;
            }
            
            /* Central Frame */
            QFrame#centralFrame {
                background-color: #ffffff;
                border: 1px solid #e5e7eb;
                border-radius: 8px;
            }
            
            /* Tab Widget - Modern tabs */
            QTabWidget::pane {
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                background-color: #ffffff;
                top: -1px;
            }
            QTabWidget::tab-bar {
                left: 8px;
            }
            QTabBar::tab {
                background: transparent;
                border: none;
                border-bottom: 3px solid transparent;
                padding: 3px 12px;
                margin-right: 4px;
                color: #6b7280;
                font-weight: 500;
                font-size: 9pt;
            }
            QTabBar::tab:selected {
                color: #6B2C91;
                border-bottom-color: #6B2C91;
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 transparent, stop:1 #f3e8ff);
            }
            QTabBar::tab:hover:!selected {
                color: #374151;
                background-color: #f9fafb;
                border-radius: 4px 4px 0 0;
            }
            QTabBar::close-button {
                subcontrol-position: right;
                background: transparent;
                border: none;
                padding: 0px;
                margin: 4px;
                width: 16px;
                height: 16px;
            }
            QTabBar::close-button:hover {
                background: #fee2e2;
                border-radius: 8px;
            }
            
            /* Buttons - Modern with elevation */
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #7C3AA3, stop:1 #6B2C91);
                color: white;
                border: none;
                border-radius: 4px;
                padding: 4px 12px;
                font-weight: 500;
                min-height: 18px;
            }
            QPushButton:hover {
                background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                    stop:0 #8B44B2, stop:1 #7C3AA3);
            }
            QPushButton:pressed {
                background: #5A2380;
            }
            QPushButton:disabled {
                background: #e5e7eb;
                color: #9ca3af;
            }

            /* Dialog Buttons - Smaller */
            QDialog QPushButton {
                padding: 3px 10px;
                min-height: 16px;
                font-size: 9pt;
            }

            /* Table Cell Widgets - Compact */
            QTableWidget QLineEdit {
                padding: 2px 4px;
                min-height: 16px;
            }

            QTableWidget QComboBox {
                padding: 2px 4px;
                min-height: 16px;
            }

            /* Line Edit - Clean input fields */
            QLineEdit {
                background-color: #ffffff;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 4px 8px;
                color: #1f2937;
                selection-background-color: #e9d5f5;
                min-height: 18px;
            }
            QLineEdit:focus {
                border: 2px solid #6B2C91;
                padding: 3px 7px;
            }
            QLineEdit:hover {
                border-color: #9ca3af;
            }
            
            /* Text Edit */
            QTextEdit {
                background-color: #ffffff;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 6px;
                color: #1f2937;
                selection-background-color: #e9d5f5;
            }
            QTextEdit:focus {
                border: 2px solid #6B2C91;
            }
            
            /* Combo Box - Modern dropdown */
            QComboBox {
                background-color: #ffffff;
                border: 1px solid #d1d5db;
                border-radius: 4px;
                padding: 4px 8px;
                color: #1f2937;
                min-height: 18px;
            }
            QComboBox:hover {
                border-color: #9ca3af;
            }
            QComboBox:focus {
                border: 2px solid #6B2C91;
            }
            QComboBox::drop-down {
                border: none;
                padding-right: 8px;
            }
            QComboBox::down-arrow {
                width: 12px;
                height: 12px;
            }
            QComboBox QAbstractItemView {
                background-color: #ffffff;
                border: 1px solid #e5e7eb;
                border-radius: 6px;
                selection-background-color: #e9d5f5;
                selection-color: #6B2C91;
                padding: 4px;
            }
            
            /* Scrollbar - Minimal modern style */
            QScrollBar:vertical {
                background: #f9fafb;
                width: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:vertical {
                background: #d1d5db;
                border-radius: 6px;
                min-height: 30px;
            }
            QScrollBar::handle:vertical:hover {
                background: #9ca3af;
            }
            QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                height: 0px;
            }
            QScrollBar:horizontal {
                background: #f9fafb;
                height: 12px;
                border-radius: 6px;
            }
            QScrollBar::handle:horizontal {
                background: #d1d5db;
                border-radius: 6px;
                min-width: 30px;
            }
            QScrollBar::handle:horizontal:hover {
                background: #9ca3af;
            }
            QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
                width: 0px;
            }
            
            /* Splitter */
            QSplitter::handle {
                background-color: #e5e7eb;
            }
            QSplitter::handle:hover {
                background-color: #d1d5db;
            }
            
            /* Message Box */
            QMessageBox {
                background-color: #ffffff;
            }
            QMessageBox QPushButton {
                min-width: 80px;
            }

        """)

        self.pcomm_window_title = 'SessionA'
        self.modules = {}
        self.module_counter = 0
        self.module_file = 'captured_modules.json'
        
        self.test_cases = {}
        self.test_case_file = 'captured_test_cases.json'
        self.test_case_counter = 0
        self.templates = {}  # {template_name: {excel_path, sheet_name, test_cases: []}}
        self.template_file = 'templates.json'
        self.template_tree_root = None
        
        self.document_config = {'text_elements': [], 'highlight_color': 'Yellow', 'generate_documentation': True}  # ✅ Updated default structure
        self.document_config_file = 'document_config.json'
        self.load_document_config()
        
        self.masking_enabled = False
        self.masking_patterns = []
        self.masking_config_file = 'masking_config.json'
        self.load_masking_config()
        
        # ADD THESE THREE LINES HERE:
        self.default_results_location = os.path.join(os.path.expanduser("~"), "Desktop")
        self.default_results_location_file = 'default_location_config.json'
        self.load_default_location_config()
        
        self.load_pcomm_window_config()

        self.num_rows = 24
        self.num_cols = 80
        self.module_tree_root = None
        self.test_case_tree_root = None
        self.current_selected_module = None
        self.libraries_dock = None
        self.bottom_dock = None
        self.toggle_modules_action = None
        self.toggle_modules_button_action = None
        self.toggle_properties_action = None
        self.toggle_test_cases_action = None
        self.toggle_test_cases_button_action = None
        
        self.pcomm_canvas_text_edit = None
        self.pcomm_preview_title = None
        self.central_frame = None

        self.modules_widget = None
        self.test_cases_widget = None

        # Correct placement of test_cases_deck
        self.test_cases_deck = QDockWidget("Test Cases", self)
        self.test_cases_deck.setAllowedAreas(Qt.DockWidgetArea.AllDockWidgetAreas)
        self.test_case_tree = QTreeWidget()
        self.test_case_tree.setHeaderLabels(['Test Case ID', 'Description', 'Action', 'Expected Result'])
        self.test_cases_deck.setWidget(self.test_case_tree)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.test_cases_deck)
        self.test_cases_deck.setVisible(False)  # This line hides the window on startup
        
        self.tab_states = {
            'modules': True,
            'test_cases': True,
            'templates': True  # ADD THIS
        }
        
        self.create_menus_and_toolbar()
        self.create_dock_widgets()
        self.create_central_widget()
        self.setStatusBar(QStatusBar(self))
        self.load_modules_from_file()
        self.load_test_cases_from_file()
        self.load_templates_from_file()
        
        self.toggle_modules_button_action.toggled.connect(
            lambda checked: self.toggle_libraries_dock_and_tabs(checked, 'modules')
        )
        self.toggle_test_cases_button_action.toggled.connect(
            lambda checked: self.toggle_libraries_dock_and_tabs(checked, 'test_cases')
        )
        
        # Connect the Libraries menu action to toggle the dock visibility
        self.toggle_libraries_action.toggled.connect(self.toggle_libraries_dock)
        self.libraries_dock.visibilityChanged.connect(self.toggle_libraries_action.setChecked)
        
        self.toggle_properties_action.toggled.connect(self.bottom_dock.setVisible)
        self.bottom_dock.visibilityChanged.connect(self.toggle_properties_action.setChecked)
        self.libraries_tabs.tabCloseRequested.connect(self.close_library_tab)
    
    def update_selection_info_status_only(self, start_row, start_col, selected_text, selection_length):
        """Updates only the status bar without triggering any popups."""
        message = (
            f"Selection: '{selected_text.replace('\n', '')}' | "
            f"Start: ({start_row}, {start_col}) | "
            f"Length: {selection_length}"
        )
        self.statusBar().showMessage(message, 0)       
       
    def execute_testcase_on_pcomm(self, testcase_data):
        """
        Send defined test case values to PCOMM Session A screen.
        testcase_data: list of dicts with {label, row, col, value}
        Example:
            [
              {"label": "CARD NUMBER", "row": 5, "col": 12, "value": "4567123412341234"},
              {"label": "EXPIRY", "row": 5, "col": 40, "value": "1226"},
              {"label": "AMOUNT", "row": 7, "col": 15, "value": "5000"},
            ]
        """
        try:
            pythoncom.CoInitialize()

            # Connect to PCOMM Session A
            autECLSession = win32com.client.Dispatch("PCOMM.autECLSession")
            autECLSession.SetConnectionByName("A")
            autECLPS = autECLSession.autECLPS

            # 1️⃣ Create empty screen (24x80 = 1920 spaces)
            screen = [[" " for _ in range(80)] for _ in range(24)]

            # 2️⃣ Fill in each field value at given row/col
            for field in testcase_data:
                val = str(field.get("value", "")).strip()
                if not val:
                    continue  # skip empty
                row = int(field["row"]) - 1
                col = int(field["col"]) - 1
                for i, ch in enumerate(val):
                    if col + i < 80:
                        screen[row][col + i] = ch

            # 3️⃣ Join into one full-screen string
            screen_text = "\n".join("".join(r) for r in screen)

            # 4️⃣ Move to home and send all text
            autECLPS.SetCursorPos(1, 1)
            autECLPS.SendKeys(screen_text)

            QMessageBox.information(self, "Success", "Test case values pushed to PCOMM screen successfully.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to send test data to PCOMM.\n\n{str(e)}")
        finally:
            pythoncom.CoUninitialize()

        
    def create_menus_and_toolbar(self):
        """Creates the main menu and a simple toolbar."""
        self.toolbar_visible = True
        menu_bar = self.menuBar()
        
      
        # === FILE MENU ===
        file_menu = menu_bar.addMenu("File")
        
        # Import actions
        import_modules_action = QAction("Import Modules...", self)
        import_modules_action.triggered.connect(self.import_modules)
        file_menu.addAction(import_modules_action)

        import_test_cases_action = QAction("Import Test Cases...", self)
        import_test_cases_action.triggered.connect(self.import_test_cases)
        file_menu.addAction(import_test_cases_action)
        
        file_menu.addSeparator()
        
        # Configure action
        configure_action = QAction("Configure...", self)
        configure_action.triggered.connect(self.open_settings_dialog)
        file_menu.addAction(configure_action)
        
        file_menu.addSeparator()
        file_menu.addAction("Exit").triggered.connect(self.close)

        # Create a new 'Window' menu
        window_menu = menu_bar.addMenu("Window")

        # Libraries action
        self.toggle_libraries_action = QAction("Libraries", self, checkable=True)
        window_menu.addAction(self.toggle_libraries_action)

        # Module properties action
        self.toggle_properties_action = QAction("Module & Properties", self, checkable=True)
        window_menu.addAction(self.toggle_properties_action)

        # Show Toolbar option
        window_menu.addSeparator()
        self.toggle_toolbar_action = QAction("Show Toolbar", self, checkable=True)
        self.toggle_toolbar_action.setChecked(True)
        self.toggle_toolbar_action.triggered.connect(self.toggle_toolbar_visibility)
        window_menu.addAction(self.toggle_toolbar_action)

        # Create main toolbar
        toolbar = QToolBar("Main Toolbar")
        toolbar.setObjectName("main_toolbar")
        self.addToolBar(toolbar)
        toolbar.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextUnderIcon)

        # Create a new icon from an SVG string for Scan

        # Create a new icon from an SVG string for Scan
        svg_scan_icon = """
        <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M15.5 14H14.71L14.43 13.73C15.41 12.59 16 11.11 16 9.5C16 5.91 13.09 3 9.5 3C5.91 3 3 5.91 3 9.5C3 13.09 5.91 16 9.5 16C11.11 16 12.59 15.41 13.73 14.43L14 14.71V15.5L19 20.49L20.49 19L15.5 14ZM9.5 14C7.01 14 5 11.99 5 9.5C5 7.01 7.01 5 9.5 5C11.99 5 14 7.01 14 9.5C14 11.99 11.99 14 9.5 14Z" fill="#6B2C91"/>
        </svg>
        """
        scan_icon = QIcon()
        pixmap = QPixmap()
        pixmap.loadFromData(QByteArray(svg_scan_icon.encode('utf-8')))
        scan_icon.addPixmap(pixmap, QIcon.Mode.Normal, QIcon.State.Off)

        scan_action = QAction(scan_icon, "SCAN\nPCOMM", self)
        scan_action.triggered.connect(self.scan_pcomm_screen)
        toolbar.addAction(scan_action)
        
        # --- NEW: Import Modules button ---
        svg_import_icon = """
        <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M14 2H6C5.44772 2 5 2.44772 5 3V19C5 19.5523 5.44772 20 6 20H18C18.5523 20 19 19.5523 19 19V8L14 2ZM13 3.5L17.5 8H13V3.5ZM17 22H7C6.44772 22 6 21.5523 6 21V19H18V21C18 21.5523 17.5523 22 17 22ZM12 11V18" stroke="#6B2C91" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M9 14L12 11L15 14" stroke="#6B2C91" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>
        """
        import_icon = QIcon()
        pixmap_import = QPixmap()
        pixmap_import.loadFromData(QByteArray(svg_import_icon.encode('utf-8')))
        import_icon.addPixmap(pixmap_import, QIcon.Mode.Normal, QIcon.State.Off)
        
        import_action = QAction(import_icon, "Import\nModules", self)
        import_action.triggered.connect(self.import_modules)
        toolbar.addAction(import_action)

        # --- NEW: Import Test Cases button ---
        svg_import_test_cases_icon = """
        <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M14 2H6C5.44772 2 5 2.44772 5 3V19C5 19.5523 5.44772 20 6 20H18C18.5523 20 19 19.5523 19 19V8L14 2ZM13 3.5L17.5 8H13V3.5ZM17 22H7C6.44772 22 6 21.5523 6 21V19H18V21C18 21.5523 17.5523 22 17 22ZM12 11V18" stroke="#6B2C91" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/>
            <path d="M9 14L12 11L15 14" stroke="#6B2C91" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"/>
        </svg>
        """
        import_test_cases_icon = QIcon()
        pixmap_import_test_cases = QPixmap()
        pixmap_import_test_cases.loadFromData(QByteArray(svg_import_test_cases_icon.encode('utf-8')))
        import_test_cases_icon.addPixmap(pixmap_import_test_cases, QIcon.Mode.Normal, QIcon.State.Off)
        
        import_test_cases_action = QAction(import_test_cases_icon, "Import\nTest Cases", self)
        import_test_cases_action.triggered.connect(self.import_test_cases)
        toolbar.addAction(import_test_cases_action)

        # --- UPDATED: Rename 'Modules' to 'Modules Library' ---
        svg_modules_icon = """
        <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M11 17H7C6.46957 17 5.96086 16.7893 5.58579 16.4142C5.21071 16.0391 5 15.5304 5 15V9C5 8.46957 5.21071 7.96086 5.58579 7.58579C5.96086 7.21071 6.46957 7 7 7H11C11.5304 7 12.0391 7.21071 12.4142 7.58579C12.7893 7.96086 13 8.46957 13 9V15C13 15.5304 12.7893 16.0391 12.4142 16.4142C12.0391 16.7893 11.5304 16.7893 11 17Z" fill="#6B2C91"/>
            <path d="M17 17H13C12.4696 17 11.9609 16.7893 11.5858 16.4142C11.2107 16.0391 11 15.5304 11 15V9C11 8.46957 11.2107 7.96086 11.5858 7.58579C11.9609 7.21071 12.4696 7 13 7H17C17.5304 7 18.0391 7.21071 18.4142 7.58579C18.7893 7.96086 19 8.46957 19 9V15C19 15.5304 18.7893 16.0391 18.4142 16.4142C18.0391 16.7893 17.5304 17 17 17Z" fill="#6B2C91"/>
            <path d="M13 15C13 15.5523 12.5523 16 12 16C11.4477 16 11 15.5523 11 15V9C11 8.44772 11.4477 8 12 8C12.5523 8 13 8.44772 13 9V15Z" fill="#6B2C91"/>
            <circle cx="9" cy="9" r="2" fill="#6B2C91"/>
            <circle cx="15" cy="15" r="2" fill="#6B2C91"/>
        </svg>
        """
        modules_icon = QIcon()
        pixmap_modules = QPixmap()
        pixmap_modules.loadFromData(QByteArray(svg_modules_icon.encode('utf-8')))
        modules_icon.addPixmap(pixmap_modules, QIcon.Mode.Normal, QIcon.State.Off)

        self.toggle_modules_button_action = QAction(modules_icon, "Modules\nLibrary", self, checkable=True)
        toolbar.addAction(self.toggle_modules_button_action)

        # --- NEW: Add 'Test cases Library' button ---
        svg_test_cases_icon = """
        <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M21 7.28571V17C21 17.5304 20.7893 18.0391 20.4142 18.4142C20.0391 18.7893 19.5304 19 19 19H5C4.46957 19 3.96086 18.7893 3.58579 18.4142C3.21071 18.0391 3 17.5304 3 17V7.28571C3 6.7553 3.21071 6.24659 3.58579 5.87152C3.96086 5.49645 4.46957 5.28571 5 5.28571H19C19.5304 5.28571 20.0391 5.49645 20.4142 5.87152C20.7893 6.24659 21 6.7553 21 7.28571ZM19 7.28571H5V17H19V7.28571ZM8 9.28571C8 9.81614 7.78929 10.3248 7.41421 10.6999C7.03914 11.075 6.53043 11.2857 6 11.2857C5.46957 11.2857 4.96086 11.075 4.58579 10.6999C4.21071 10.3248 4 9.81614 4 9.28571C4 8.7553 4.21071 8.24659 4.58579 7.87152C4.96086 7.49645 5.46957 7.28571 6 7.28571C6.53043 7.28571 7.03914 7.49645 7.41421 7.87152C7.78929 8.24659 8 8.7553 8 9.28571Z" fill="#6B2C91"/>
        </svg>
        """
        test_cases_icon = QIcon()
        pixmap_test_cases = QPixmap()
        pixmap_test_cases.loadFromData(QByteArray(svg_test_cases_icon.encode('utf-8')))
        test_cases_icon.addPixmap(pixmap_test_cases, QIcon.Mode.Normal, QIcon.State.Off)
        
        self.toggle_test_cases_button_action = QAction(test_cases_icon, "Test cases\nLibrary", self, checkable=True)
        toolbar.addAction(self.toggle_test_cases_button_action)

        # Connect the checkable states
        # --- REMOVED: Connections moved to __init__ ---

        # --- NEW: Create a test case button with an icon ---
        svg_create_test_case_icon = """
        <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M12 2C6.48 2 2 6.48 2 12C2 17.52 6.48 22 12 22C17.52 22 22 17.52 22 12C22 6.48 17.52 2 12 2ZM17 13H13V17H11V13H7V11H11V7H13V11H17V13Z" fill="#6B2C91"/>
        </svg>
        """
        create_test_case_icon = QIcon()
        pixmap_create_test_case = QPixmap()
        pixmap_create_test_case.loadFromData(QByteArray(svg_create_test_case_icon.encode('utf-8')))
        create_test_case_icon.addPixmap(pixmap_create_test_case, QIcon.Mode.Normal, QIcon.State.Off)

        create_test_case_action = QAction(create_test_case_icon, "Create a\nTest Case", self)
        create_test_case_action.triggered.connect(self.create_test_case)
        toolbar.addAction(create_test_case_action)

        # --- NEW: Test Execution button with a play icon ---
# --- NEW: Test Execution button with a play icon ---
        svg_test_execution_icon = """
        <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M8 5V19L19 12L8 5Z" fill="#6B2C91"/>
        </svg>
        """
        test_execution_icon = QIcon()
        pixmap_test_execution = QPixmap()
        pixmap_test_execution.loadFromData(QByteArray(svg_test_execution_icon.encode('utf-8')))
        test_execution_icon.addPixmap(pixmap_test_execution, QIcon.Mode.Normal, QIcon.State.Off)

        test_execution_action = QAction(test_execution_icon, "Test\nExecution", self)
        test_execution_action.triggered.connect(self.test_execution)
        toolbar.addAction(test_execution_action)
        
        # --- NEW: Convert to Template button ---
        svg_template_icon = """
        <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M19 3H5C3.89 3 3 3.89 3 5V19C3 20.11 3.89 21 5 21H19C20.11 21 21 20.11 21 19V5C21 3.89 20.11 3 19 3ZM19 19H5V5H19V19Z" fill="#6B2C91"/>
            <path d="M7 7H17V9H7V7ZM7 11H17V13H7V11ZM7 15H13V17H7V15Z" fill="#6B2C91"/>
        </svg>
        """
        template_icon = QIcon()
        pixmap_template = QPixmap()
        pixmap_template.loadFromData(QByteArray(svg_template_icon.encode('utf-8')))
        template_icon.addPixmap(pixmap_template, QIcon.Mode.Normal, QIcon.State.Off)

        convert_template_action = QAction(template_icon, "Convert to\nTemplate", self)
        convert_template_action.triggered.connect(self.convert_to_template)
        toolbar.addAction(convert_template_action)
        
        # --- NEW: Link Data Source button ---
        svg_link_icon = """
        <svg width="24" height="24" viewBox="0 0 24 24" fill="none" xmlns="http://www.w3.org/2000/svg">
            <path d="M3.9 12C3.9 10.29 5.29 8.9 7 8.9H11V7H7C4.24 7 2 9.24 2 12C2 14.76 4.24 17 7 17H11V15.1H7C5.29 15.1 3.9 13.71 3.9 12ZM8 13H16V11H8V13ZM17 7H13V8.9H17C18.71 8.9 20.1 10.29 20.1 12C20.1 13.71 18.71 15.1 17 15.1H13V17H17C19.76 17 22 14.76 22 12C22 9.24 19.76 7 17 7Z" fill="#6B2C91"/>
        </svg>
        """
        link_icon = QIcon()
        pixmap_link = QPixmap()
        pixmap_link.loadFromData(QByteArray(svg_link_icon.encode('utf-8')))
        link_icon.addPixmap(pixmap_link, QIcon.Mode.Normal, QIcon.State.Off)

        link_datasource_action = QAction(link_icon, "Link Data\nSource", self)
        link_datasource_action.triggered.connect(self.link_data_source)
        toolbar.addAction(link_datasource_action)

    def toggle_masking(self, checked):
        """Toggles text masking on/off."""
        self.masking_enabled = checked
        self.save_masking_config()
        
        status = "enabled" if checked else "disabled"
        self.statusBar().showMessage(f"Text masking {status}.", 3000)

    def configure_pcomm_window(self):
        """Opens a dialog to configure the PCOMM window title."""
        dialog = QDialog(self)
        dialog.setWindowTitle("PCOMM Window Configuration")
        dialog.setFixedSize(450, 150)
        dialog.setWindowFlags(dialog.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title_label = QLabel("PCOMM Window Title")
        title_font = QFont()
        title_font.setPointSize(11)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        # Description
        desc_label = QLabel("Enter the exact window title as it appears in your PCOMM application:")
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)
        layout.addSpacing(10)
        
        # Input field
        input_layout = QHBoxLayout()
        input_label = QLabel("Window Title:")
        input_label.setFixedWidth(100)
        input_field = QLineEdit()
        input_field.setText(self.pcomm_window_title)
        input_field.setPlaceholderText("e.g., SessionA or Session A")
        input_layout.addWidget(input_label)
        input_layout.addWidget(input_field)
        layout.addLayout(input_layout)
        
        layout.addSpacing(10)
        
        # Buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_title = input_field.text().strip()
            if new_title:
                self.pcomm_window_title = new_title
                self.save_pcomm_window_config()
                QMessageBox.information(self, "Success", f"PCOMM window title set to: '{new_title}'")
            else:
                QMessageBox.warning(self, "Invalid Input", "Window title cannot be empty.")

    def create_central_widget(self):
        """
        Creates the main PCOMM screenshot view in the center of the window,
        styled to match the 24x80 terminal grid.
        
        FIXED: Made fully flexible with scrollbars, title bar has fixed height.
        """
        self.central_frame = QFrame(self)
        self.central_frame.setObjectName("centralFrame")
        self.setCentralWidget(self.central_frame)

        layout = QVBoxLayout(self.central_frame)
        layout.setContentsMargins(0, 0, 0, 0)

        # --- FIXED: Title bar with fixed height ---
        title_bar_layout = QHBoxLayout()
        title_bar_layout.setContentsMargins(5, 2, 5, 2)  # Reduced vertical margins
        
        self.pcomm_preview_title = QLabel("PCOMM Screen Preview")
        self.pcomm_preview_title.setStyleSheet("""
            font-weight: bold;
        """)
                
        close_button = QPushButton("✕")
        close_button.setFixedSize(20, 20)
        close_button.setStyleSheet("""
            QPushButton {
                border: none;
                background-color: transparent;
                font-weight: bold;
                font-size: 16px;
                color: black;
                border-radius: 3px;
                padding: 0px;
            }
            QPushButton:hover {
                background-color: #fee2e2;
                color: #dc2626;
            }
            QPushButton:pressed {
                background-color: #fecaca;
            }
        """)
        close_button.setCursor(Qt.CursorShape.PointingHandCursor)  # ✅ NEW: Add cursor pointer
        close_button.clicked.connect(self.hide_pcomm_preview)

        title_bar_layout.addWidget(self.pcomm_preview_title)
        title_bar_layout.addStretch()
        title_bar_layout.addWidget(close_button)

        # In create_central_widget method, update the title_bar_frame fixed height:

        title_bar_frame = QFrame()
        title_bar_frame.setLayout(title_bar_layout)
        # Changed from setFixedHeight(40) to setFixedHeight(28)
        title_bar_frame.setFixedHeight(28)
        title_bar_frame.setStyleSheet("""
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #ffffff, stop:1 #f3f4f6);
            border-bottom: 2px solid #e5e7eb;
            border-radius: 8px 8px 0 0;
        """)
        layout.addWidget(title_bar_frame)
        
        # --- FIXED: QTextEdit is now flexible with scrollbars ---
        self.pcomm_canvas_text_edit = CustomPCOMMTextEdit(self)
        self.pcomm_canvas_text_edit.main_window = self
        self.pcomm_canvas_text_edit.setObjectName("pcomm_preview")
        self.pcomm_canvas_text_edit.setReadOnly(True)
        
        # Connect the selectionChanged signal
        #self.pcomm_canvas_text_edit.selectionChanged.connect(self.update_selection_info)
        
        # Use QFontMetrics to calculate the font size
        font = QFont("Courier New")
        font.setStyleHint(QFont.StyleHint.Monospace)
        font.setPointSize(10)  # Use a fixed reasonable size
        
        self.pcomm_canvas_text_edit.setFont(font)
        
        # FIXED: Remove setFixedSize() - let it be flexible
        # Set size policy to allow expansion and contraction
        self.pcomm_canvas_text_edit.setSizePolicy(
            QSizePolicy.Policy.Expanding, 
            QSizePolicy.Policy.Expanding
        )
        
        # Set minimum size so it doesn't disappear completely
        self.pcomm_canvas_text_edit.setMinimumHeight(150)
        self.pcomm_canvas_text_edit.setMinimumWidth(400)
        
        # Enable word wrap for better viewing
        self.pcomm_canvas_text_edit.setLineWrapMode(QTextEdit.LineWrapMode.NoWrap)
        
        # Enable scrollbars when content is larger than viewport
        self.pcomm_canvas_text_edit.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.pcomm_canvas_text_edit.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        
        # Set the preview screen colors and font styling
        self.pcomm_canvas_text_edit.setStyleSheet("""
            #pcomm_preview {
                background-color: #1e1e1e;
                color: #00ffff;
                font-family: 'Courier New', monospace;
                padding: 12px;
                border: 2px solid #6B2C91;
                border-radius: 8px;
            }
            #pcomm_preview:focus {
                border: 2px solid #8B44B2;
            }
        """)

        self.pcomm_canvas_text_edit.setText("PCOMM Screenshot Preview")
        
        # FIXED: Add to layout with stretch factor to make it flexible
        layout.addWidget(self.pcomm_canvas_text_edit, 1)  # stretch factor = 1

# --- REVISED: Function to create and configure the tabbed interface ---
    def create_libraries_tab_widget(self):
        """
        Creates and returns a QTabWidget containing the Modules and Test Cases tabs.
        This replaces the two separate dock widgets.
        """
        tab_widget = QTabWidget()
        tab_widget.setTabsClosable(True)
        tab_widget.setMovable(True)
        
        # Set stylesheet for tab close buttons to make them visible like dock close buttons
        tab_widget.setStyleSheet("""
            QTabWidget::pane {
                border-top: 2px solid #C2C7CB;
            }
            QTabBar::tab {
                background: #e0e0e0;
                border: 1px solid #C4C4C3;
                border-bottom-color: #C2C7CB;
                border-top-left-radius: 4px;
                border-top-right-radius: 4px;
                padding: 5px;
                min-width: 80px;
            }
            QTabBar::tab:selected {
                background: #f9f9f9;
                border-color: #9B9B9B;
                border-bottom-color: transparent;
            }
            QTabBar::close-button {
                subcontrol-position: right;
                background: transparent;
                border: none;
                padding: 0px;
                margin: 0px 4px 0px 4px;
                width: 16px;
                height: 16px;
                image: none;
            }
            QTabBar::close-button:hover {
                background: #ddd;
                border-radius: 8px;
            }
        """)
        
        # Modules Tab
        self.modules_widget = QWidget()
        modules_layout = QVBoxLayout(self.modules_widget)
        self.modules_search_bar = QLineEdit()
        self.modules_search_bar.setPlaceholderText("Search modules...")
        self.modules_search_bar.textChanged.connect(self.filter_modules)
        modules_layout.addWidget(self.modules_search_bar)

        self.module_tree = QTreeWidget()
        self.module_tree.setColumnCount(2)
        self.module_tree.setRootIsDecorated(False)
        self.module_tree.setSelectionMode(QTreeWidget.SelectionMode.ExtendedSelection)  # ✅ ADD THIS LINE
        self.module_tree.header().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.module_tree.header().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.module_tree.header().setStretchLastSection(False)
        self.module_tree.header().setVisible(False)
        self.module_tree.itemClicked.connect(self.display_module_details_and_screenshot)
        self.module_tree.setEditTriggers(QTreeWidget.EditTrigger.DoubleClicked)
        self.module_tree.itemChanged.connect(self.handle_rename_finish)
        self.module_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.module_tree.customContextMenuRequested.connect(self.show_context_menu)
        self.module_tree.installEventFilter(self)  # ✅ ADD THIS LINE
        modules_layout.addWidget(self.module_tree)
        
        # Test Cases Tab
        self.test_cases_widget = QWidget()
        test_cases_layout = QVBoxLayout(self.test_cases_widget)
        self.test_cases_search_bar = QLineEdit()
        self.test_cases_search_bar.setPlaceholderText("Search test cases...")
        self.test_cases_search_bar.textChanged.connect(self.filter_test_cases)
        test_cases_layout.addWidget(self.test_cases_search_bar)

        self.test_case_tree = QTreeWidget()
        self.test_case_tree.setHeaderLabels(["Test Case", ""])
        self.test_case_tree.setRootIsDecorated(False)
        self.test_case_tree.setSelectionMode(QTreeWidget.SelectionMode.ExtendedSelection)  # ✅ ADD THIS LINE
        self.test_case_tree.header().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.test_case_tree.header().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.test_case_tree.header().setStretchLastSection(False)
        self.test_case_tree.header().setVisible(False)
        self.test_case_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.test_case_tree.customContextMenuRequested.connect(self.show_test_case_context_menu)
        self.test_case_tree.itemDoubleClicked.connect(self.on_test_case_item_double_clicked)
        self.test_case_tree.installEventFilter(self)  # ✅ ADD THIS LINE
        test_cases_layout.addWidget(self.test_case_tree)

        self.templates_widget = QWidget()
        templates_layout = QVBoxLayout(self.templates_widget)
        self.templates_search_bar = QLineEdit()
        self.templates_search_bar.setPlaceholderText("Search templates...")
        self.templates_search_bar.textChanged.connect(self.filter_templates)
        templates_layout.addWidget(self.templates_search_bar)

        self.template_tree = QTreeWidget()
        self.template_tree.setHeaderLabels(["Template", ""])
        self.template_tree.setRootIsDecorated(True)
        self.template_tree.setSelectionMode(QTreeWidget.SelectionMode.ExtendedSelection)
        self.template_tree.header().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.template_tree.header().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.template_tree.header().setStretchLastSection(False)
        self.template_tree.header().setVisible(False)
        self.template_tree.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.template_tree.customContextMenuRequested.connect(self.show_template_context_menu)
        self.template_tree.itemDoubleClicked.connect(self.on_template_item_double_clicked)
        self.template_tree.installEventFilter(self)
        templates_layout.addWidget(self.template_tree)        
        
        
        # Add the tabs initially
        tab_widget.addTab(self.modules_widget, "Modules Library")
        tab_widget.addTab(self.test_cases_widget, "Test Cases Library")
        tab_widget.addTab(self.templates_widget, "Templates Library")
        
        # Replace default close buttons with custom styled buttons
        tab_bar = tab_widget.tabBar()
        for i in range(tab_bar.count()):
            # Create a custom close button
            close_btn = QPushButton("×")
            close_btn.setFixedSize(16, 16)
            close_btn.setStyleSheet("""
                QPushButton {
                    border: none;
                    background-color: transparent;
                    font-weight: bold;
                    font-size: 18px;
                    color: #555;
                    padding: 0px;
                }
                QPushButton:hover {
                    background-color: #e0e0e0;
                    border-radius: 3px;
                    color: #000;
                }
            """)
            close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
            
            # Connect to close tab functionality
            close_btn.clicked.connect(lambda checked, idx=i: self.close_library_tab(idx))
            
            # Set the custom button as the tab button
            tab_bar.setTabButton(i, QTabBar.ButtonPosition.RightSide, close_btn)
        
        return tab_widget
        
    def update_template_tree(self):
        """Updates the template tree widget with the latest template data."""
        self.template_tree.clear()
        
        # ✅ ALWAYS create root if there are templates
        if self.templates:
            self.template_tree_root = QTreeWidgetItem(self.template_tree, ["Templates"])
            self.template_tree_root.setExpanded(True)

            for template_name, template_data in self.templates.items():
                template_item = QTreeWidgetItem(self.template_tree_root, [template_name])
                template_item.setData(0, Qt.ItemDataRole.UserRole, template_name)
                
                # Button widget
                button_widget = QWidget()
                button_layout = QHBoxLayout(button_widget)
                button_layout.setContentsMargins(0, 0, 0, 0)
                button_layout.setSpacing(5)

                delete_button = QPushButton("✕")
                delete_button.setFixedSize(18, 18)
                delete_font = QFont()
                delete_font.setBold(True)
                delete_font.setPointSize(12)
                delete_button.setFont(delete_font)
                delete_button.setStyleSheet("""
                    QPushButton {
                        color: #dc2626;
                        border: none;
                        border-radius: 4px;
                        background-color: transparent;
                        padding: 0px;
                    }
                    QPushButton:hover {
                        background-color: #fee2e2;
                    }
                    QPushButton:pressed {
                        background-color: #fecaca;
                    }
                """)
                delete_button.setCursor(Qt.CursorShape.PointingHandCursor)
                delete_button.setToolTip(f"Delete '{template_name}'.")
                delete_button.clicked.connect(lambda _, item=template_item: self.delete_template(item))
                button_layout.addWidget(delete_button)
                
                spacer = QSpacerItem(10, 0, QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Minimum)
                button_layout.addItem(spacer)
                
                self.template_tree.setItemWidget(template_item, 1, button_widget)
                
                # Add child test cases if expanded
                test_cases = template_data.get('test_cases', [])
                if test_cases:
                    for tc in test_cases:
                        tc_item = QTreeWidgetItem(template_item, [tc['test_case_id']])
                        tc_item.setData(0, Qt.ItemDataRole.UserRole, {'type': 'template_test_case', 'data': tc})
        else:
            # ✅ Set root to None when no templates exist
            self.template_tree_root = None      
        
    def show_template_context_menu(self, position: QPoint):
        """Shows context menu for template tree."""
        item = self.template_tree.itemAt(position)
        if not item or item.parent() is None:
            return
        
        # Check if it's a parent template item
        if item.parent() == self.template_tree_root:
            menu = QMenu()
            
            template_name = item.text(0)
            
            # Break into test cases action
            break_action = QAction("🔓 Break into Test Cases", self)
            break_action.triggered.connect(lambda: self.break_template_into_test_cases(template_name))
            menu.addAction(break_action)
            
            menu.addSeparator()
            
            # View info action
            info_action = QAction("ℹ️ View Info", self)
            info_action.triggered.connect(lambda: self.view_template_info(template_name))
            menu.addAction(info_action)
            
            menu.exec(self.template_tree.viewport().mapToGlobal(position))        
    
    def break_template_into_test_cases(self, template_name):
        """Breaks template into individual test cases and adds them to Test Cases Library."""
        template_data = self.templates.get(template_name)
        if not template_data:
            return
        
        # ✅ NEW: Load data from Excel at break time
        excel_path = template_data.get('excel_path')
        sheet_name = template_data.get('sheet_name')
        
        if not os.path.exists(excel_path):
            QMessageBox.warning(self, "File Not Found", 
                              f"Excel file not found: {excel_path}")
            return
        
        try:
            import openpyxl
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            
            if sheet_name not in wb.sheetnames:
                QMessageBox.warning(self, "Sheet Not Found", 
                                  f"Sheet '{sheet_name}' not found in Excel file.")
                return
            
            ws = wb[sheet_name]
            
            # Read field labels from column A (starting from row 7)
            field_labels = []
            row_idx = 7
            while True:
                cell_value = ws.cell(row=row_idx, column=1).value
                if not cell_value:
                    break
                field_labels.append(str(cell_value))
                row_idx += 1
            
            if not field_labels:
                QMessageBox.warning(self, "No Fields", "No field labels found in the Excel file.")
                return
            
            # Read test case data from columns (B onwards)
            test_cases_list = []
            col_idx = 2  # Start from column B
            
            while True:
                # Check if column has header
                test_case_id = ws.cell(row=6, column=col_idx).value
                if not test_case_id:
                    break
                
                # Read data for this test case column
                row_data = {}
                for field_row_idx, field_label in enumerate(field_labels, 7):
                    cell_value = ws.cell(row=field_row_idx, column=col_idx).value
                    row_data[field_label] = str(cell_value) if cell_value is not None else ""
                
                test_cases_list.append({
                    'test_case_id': str(test_case_id),
                    'data': row_data
                })
                
                col_idx += 1
            
        except Exception as e:
            QMessageBox.critical(self, "Error Reading Excel", 
                               f"Failed to read Excel file:\n\n{str(e)}")
            return
        
        if not test_cases_list:
            QMessageBox.warning(self, "No Data", "No test cases found in this template.")
            return
        
        # Get the base test case
        base_test_case_name = template_data.get('base_test_case')
        if base_test_case_name not in self.test_cases:
            QMessageBox.warning(self, "Base Test Case Not Found", 
                              f"Base test case '{base_test_case_name}' not found in library.")
            return
        
        base_test_case = self.test_cases[base_test_case_name]
        
        # Create individual test cases
        created_count = 0
        for tc_data in test_cases_list:
            new_test_case_id = tc_data['test_case_id']
            
            # Skip if already exists
            if new_test_case_id in self.test_cases:
                continue
            
            # Deep copy the base test case
            new_test_case = copy.deepcopy(base_test_case)
            
            # Update the description
            new_test_case['description'] = f"Generated from template: {template_name}"
            
            # Fill in the field values from the data
            field_data = tc_data['data']
            
            for step_idx, step in enumerate(new_test_case.get('steps', []), 1):
                if step.get('type') == 'module_import':
                    module_name = step.get('module_name', '')
                    for field in step.get('fields', []):
                        if field.get('action_type') == 'Input':
                            field_name = field.get('field_name', '')
                            field_label = f"Step {step_idx} - {module_name}.{field_name}"
                            
                            if field_label in field_data:
                                field['value'] = field_data[field_label]
                
                elif step.get('type') == 'random_input':
                    row = step.get('row', '?')
                    col = step.get('column', '?')
                    field_label = f"Step {step_idx} - RandomInput (R{row},C{col})"
                    
                    if field_label in field_data:
                        step['value'] = field_data[field_label]
            
            # Add to test cases library
            self.test_cases[new_test_case_id] = new_test_case
            created_count += 1
        
        # Save and update
        self.save_test_cases_to_file()
        self.update_test_case_tree()
        
        # Ensure Test Cases tab is visible
        test_cases_tab_index = -1
        for i in range(self.libraries_tabs.count()):
            if self.libraries_tabs.tabText(i) == "Test Cases Library":
                test_cases_tab_index = i
                break
        
        if test_cases_tab_index == -1:
            test_cases_tab_index = self.libraries_tabs.addTab(self.test_cases_widget, "Test Cases Library")
            self.setup_tab_close_button(test_cases_tab_index)
            self.tab_states['test_cases'] = True
        
        # Show the Libraries dock if hidden
        if not self.libraries_dock.isVisible():
            self.libraries_dock.setVisible(True)
            self.toggle_libraries_action.setChecked(True)
        
        # Switch to Test Cases tab
        self.libraries_tabs.setCurrentIndex(test_cases_tab_index)
        
        QMessageBox.information(
            self,
            "Success",
            f"Template '{template_name}' broken into test cases.\n\n"
            f"Created {created_count} test case(s) in Test Cases Library."
        )
    
    def setup_tab_close_button(self, tab_index):
        """
        Sets up a custom close button for a specific tab.
        """
        tab_bar = self.libraries_tabs.tabBar()
        
        # Create a custom close button
        close_btn = QPushButton("×")
        close_btn.setFixedSize(16, 16)
        close_btn.setStyleSheet("""
            QPushButton {
                border: none;
                background-color: transparent;
                font-weight: bold;
                font-size: 18px;
                color: #555;
                padding: 0px;
            }
            QPushButton:hover {
                background-color: #e0e0e0;
                border-radius: 3px;
                color: #000;
            }
        """)
        close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        
        # Connect to close tab functionality
        close_btn.clicked.connect(lambda: self.close_library_tab(tab_index))
        
        # Set the custom button as the tab button
        tab_bar.setTabButton(tab_index, QTabBar.ButtonPosition.RightSide, close_btn)
        
    # In MainWindow class
    def on_test_case_item_double_clicked(self, item, column):
        if item.parent() == self.test_case_tree_root:
            test_case_id = item.text(0)
            
            existing_steps = self.test_cases.get(test_case_id, {}).get('steps', [])
            test_case_description = self.test_cases.get(test_case_id, {}).get('description', '')
            test_case_assumptions = self.test_cases.get(test_case_id, {}).get('assumptions', '')  # ✅ NEW
            
            dialog = EditTestCaseDialog(existing_steps, self.modules, self, test_case_id, test_case_description, test_case_assumptions)  # ✅ CHANGED
            dialog.exec()
            
            updated_steps = dialog.get_updated_steps()
            updated_name = dialog.get_test_case_name()
            updated_description = dialog.get_test_case_description()
            updated_assumptions = dialog.get_test_case_assumptions()  # ✅ NEW
            
            if not updated_name:
                QMessageBox.warning(self, "Invalid Name", "Test case name cannot be empty.")
                return
            
            if updated_name != test_case_id and updated_name in self.test_cases:
                QMessageBox.warning(self, "Duplicate Name", f"A test case with the name '{updated_name}' already exists.")
                return
            
            if updated_name != test_case_id:
                self.test_cases[updated_name] = self.test_cases.pop(test_case_id)
                test_case_id = updated_name
            
            self.test_cases[test_case_id]['steps'] = updated_steps
            self.test_cases[test_case_id]['description'] = updated_description
            self.test_cases[test_case_id]['assumptions'] = updated_assumptions  # ✅ NEW
            self.test_cases[test_case_id]['prerequisites'] = dialog.get_prerequisites()
            self.save_test_cases_to_file()
            self.update_test_case_tree()
            QMessageBox.information(self, "Success", f"Test case '{test_case_id}' updated successfully.")
            
           

    def add_step_to_test_case(self, test_case_id, module_name):
        """
        Adds a new step (module) to the specified test case.
        """
        if test_case_id in self.test_cases:
            new_step = {
                "name": f"Import Module: {module_name}",
                "type": "module_import",
                "module_name": module_name
            }
            self.test_cases[test_case_id]['steps'].append(new_step)
            
            # Refresh the UI to show the new step
            self.update_test_case_tree()
            self.save_test_cases_to_file()
            QMessageBox.information(self, "Updated", f"Added '{module_name}' to test case '{test_case_id}'.")
        else:
            QMessageBox.warning(self, "Error", f"Test case '{test_case_id}' not found.")

    # --- REPLACED: The create_dock_widgets function to use a single QDockWidget and a QTabWidget ---
    def create_dock_widgets(self):
        """
        Creates the dockable widgets, combining the modules and test cases
        into a single tabbed dock.
        
        FIXED: Module Properties dock is now fully flexible and can shrink/expand.
        """
        # --- Left Dock Widget (Combined Libraries) ---
        # --- Left Dock Widget (Combined Libraries) ---
        self.libraries_dock = QDockWidget("Libraries", self)
        self.libraries_dock.setObjectName("libraries_dock")
        self.libraries_dock.setFeatures(QDockWidget.DockWidgetFeature.DockWidgetClosable)
        self.libraries_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea)

        # ✅ ADD: Custom title bar with close button
        libraries_title_widget = QWidget()
        libraries_title_layout = QHBoxLayout(libraries_title_widget)
        libraries_title_layout.setContentsMargins(8, 4, 4, 4)
        libraries_title_layout.setSpacing(0)

        libraries_title_label = QLabel("Libraries")
        libraries_title_label.setStyleSheet("font-weight: bold; color: #1f2937;")
        libraries_title_layout.addWidget(libraries_title_label)
        libraries_title_layout.addStretch()

        libraries_close_btn = QPushButton("✕")
        libraries_close_btn.setFixedSize(20, 20)
        libraries_close_btn.setStyleSheet("""
            QPushButton {
                color: black;
                font-weight: bold;
                font-size: 16px;
                border: none;
                border-radius: 3px;
                background-color: transparent;
                padding: 0px;
            }
            QPushButton:hover {
                background-color: #fee2e2;
                color: #dc2626;
            }
            QPushButton:pressed {
                background-color: #fecaca;
            }
        """)
        libraries_close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        libraries_close_btn.clicked.connect(self.libraries_dock.close)
        libraries_title_layout.addWidget(libraries_close_btn)

        libraries_title_widget.setStyleSheet("""
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #ffffff, stop:1 #f0f2f5);
            border-bottom: 1px solid #d1d5db;
        """)

        self.libraries_dock.setTitleBarWidget(libraries_title_widget)
        
        self.libraries_tabs = self.create_libraries_tab_widget()
        self.libraries_dock.setWidget(self.libraries_tabs)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.libraries_dock)

                # --- Bottom Dock Widget (Module & Properties) ---
        # --- Bottom Dock Widget (Module & Properties) ---
        self.bottom_dock = QDockWidget("Module & Properties", self)
        self.bottom_dock.setFeatures(QDockWidget.DockWidgetFeature.DockWidgetClosable)
        self.bottom_dock.setAllowedAreas(Qt.DockWidgetArea.BottomDockWidgetArea)

        # ✅ ADD: Custom title bar with close button
        bottom_title_widget = QWidget()
        bottom_title_layout = QHBoxLayout(bottom_title_widget)
        bottom_title_layout.setContentsMargins(8, 4, 4, 4)
        bottom_title_layout.setSpacing(0)

        bottom_title_label = QLabel("Module & Properties")
        bottom_title_label.setStyleSheet("font-weight: bold; color: #1f2937;")
        bottom_title_layout.addWidget(bottom_title_label)
        bottom_title_layout.addStretch()

        bottom_close_btn = QPushButton("✕")
        bottom_close_btn.setFixedSize(18, 18)
        bottom_close_btn.setStyleSheet("""
            QPushButton {
                color: black;
                font-weight: bold;
                font-size: 16px;
                border: none;
                border-radius: 3px;
                background-color: transparent;
                padding: 0px;
            }
            QPushButton:hover {
                background-color: #fee2e2;
                color: #dc2626;
            }
            QPushButton:pressed {
                background-color: #fecaca;
            }
        """)
        bottom_close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        bottom_close_btn.clicked.connect(self.bottom_dock.close)
        bottom_title_layout.addWidget(bottom_close_btn)

        bottom_title_widget.setStyleSheet("""
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #ffffff, stop:1 #f0f2f5);
            border-bottom: 1px solid #d1d5db;
        """)

        self.bottom_dock.setTitleBarWidget(bottom_title_widget)
        
        # FIXED: Set size constraints for the dock
        self.bottom_dock.setMinimumHeight(80)  # Minimum height when shrunk
        self.bottom_dock.setMaximumHeight(16777215)  # No maximum limit (Qt default max)

        self.main_tabs = QTabWidget()
        
        # FIXED: Set size policy for tab widget
        self.main_tabs.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Expanding
        )
        
        self.module_details_tab_widget = QWidget()
        module_details_tab_layout = QVBoxLayout(self.module_details_tab_widget)
        
        label_button_layout = QHBoxLayout()
        
        # All buttons in one row on the left
        self.add_label_button = QPushButton("Add New Label")
        self.add_label_button.clicked.connect(self.add_new_label)
        label_button_layout.addWidget(self.add_label_button)

        self.copy_label_button = QPushButton("📋 Copy")
        self.copy_label_button.clicked.connect(self.copy_selected_labels)
        self.copy_label_button.setToolTip("Copy selected label(s) (Ctrl+C)")
        label_button_layout.addWidget(self.copy_label_button)

        self.paste_label_button = QPushButton("📄 Paste")
        self.paste_label_button.clicked.connect(self.paste_copied_labels)
        self.paste_label_button.setToolTip("Paste copied label(s) at the end (Ctrl+V)")
        label_button_layout.addWidget(self.paste_label_button)

        self.delete_labels_button = QPushButton("🗑️ Delete")
        self.delete_labels_button.clicked.connect(self.delete_selected_labels)
        self.delete_labels_button.setToolTip("Delete selected label(s) (Delete)")
        label_button_layout.addWidget(self.delete_labels_button)


        self.move_up_button = QPushButton("↑ Move Up")
        self.move_up_button.clicked.connect(self.move_label_up)
        label_button_layout.addWidget(self.move_up_button)
        
        self.move_down_button = QPushButton("↓ Move Down")
        self.move_down_button.clicked.connect(self.move_label_down)
        label_button_layout.addWidget(self.move_down_button)

        # Add stretch to push all buttons to the left
        label_button_layout.addStretch()

        module_details_tab_layout.addLayout(label_button_layout)
        
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # FIXED: Set size policy for splitter to allow flexibility
        splitter.setSizePolicy(
            QSizePolicy.Policy.Expanding, 
            QSizePolicy.Policy.Expanding
        )

        self.labels_table = QTableWidget()
        self.labels_table.setColumnCount(5)
        self.labels_table.setHorizontalHeaderLabels(["Field name", "Row", "Column", "Length", ""])
        self.labels_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.labels_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
        self.labels_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
        self.labels_table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
        self.labels_table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeMode.Fixed)
        self.labels_table.horizontalHeader().setStretchLastSection(False)
        self.labels_table.setColumnWidth(4, 30)
        self.labels_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.labels_table.setSelectionMode(QTableWidget.SelectionMode.ExtendedSelection)
        self.labels_table.cellClicked.connect(self.display_selected_label_properties)
        self.labels_table.setEditTriggers(QTableWidget.EditTrigger.DoubleClicked)
        self.labels_table.cellChanged.connect(self.handle_label_edit_finish)
        
        # FIXED: Set size policies to allow flexible shrinking/expanding
        self.labels_table.setSizePolicy(
            QSizePolicy.Policy.Expanding, 
            QSizePolicy.Policy.Expanding
        )
        self.labels_table.setMinimumHeight(40)  # Very small minimum
        self.labels_table.setMinimumWidth(200)  # Reasonable minimum width
        
        self.labels_table.installEventFilter(self)

        self.properties_tree = QTreeWidget()
        self.properties_tree.setHeaderLabels(["Properties", "Value"])
        self.properties_tree.header().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.properties_tree.setColumnCount(2)
        
        # FIXED: Set size policies to allow flexible shrinking/expanding
        self.properties_tree.setSizePolicy(
            QSizePolicy.Policy.Expanding, 
            QSizePolicy.Policy.Expanding
        )
        self.properties_tree.setMinimumHeight(40)  # Very small minimum
        self.properties_tree.setMinimumWidth(200)  # Reasonable minimum width

        splitter.addWidget(self.labels_table)
        splitter.addWidget(self.properties_tree)

        module_details_tab_layout.addWidget(splitter)
        self.main_tabs.addTab(self.module_details_tab_widget, "Module Details")

        self.bottom_dock.setWidget(self.main_tabs)
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, self.bottom_dock)
        self.bottom_dock.setVisible(True)
        
        
        
    def eventFilter(self, obj, event):
        """Handle keyboard shortcuts for the labels table, module tree, and test case tree."""
        # Handle labels table shortcuts
        if hasattr(self, 'labels_table') and obj == self.labels_table and event.type() == event.Type.KeyPress:
            # Check if Ctrl is pressed
            if event.modifiers() == Qt.KeyboardModifier.ControlModifier:
                if event.key() == Qt.Key.Key_C:
                    # Ctrl+C: Copy
                    self.copy_selected_labels()
                    return True
                elif event.key() == Qt.Key.Key_V:
                    # Ctrl+V: Paste
                    self.paste_copied_labels()
                    return True
            elif event.key() == Qt.Key.Key_Delete:
                # Delete key: Delete selected labels
                self.delete_selected_labels()
                return True
        
        # Handle module tree shortcuts
        elif hasattr(self, 'module_tree') and obj == self.module_tree and event.type() == event.Type.KeyPress:
            if event.key() == Qt.Key.Key_Delete:
                self.delete_selected_modules()
                return True
        
        # Handle test case tree shortcuts
        elif hasattr(self, 'test_case_tree') and obj == self.test_case_tree and event.type() == event.Type.KeyPress:
            if event.key() == Qt.Key.Key_Delete:
                self.delete_selected_test_cases()
                return True
        
        # Call the parent implementation for other events
        return super().eventFilter(obj, event)

    def toggle_libraries_dock(self, checked):
        """
        Toggles the visibility of the entire Libraries dock window.
        This is connected to the Window menu "Libraries" option.
        """
        self.libraries_dock.setVisible(checked)

        # --- REVISED: toggle_libraries_dock_and_tabs to re-add tabs ---
    def toggle_libraries_dock_and_tabs(self, checked, library_name):
        """
        Toggles the visibility of the libraries dock and sets the current tab.
        'library_name' can be 'modules' or 'test_cases'.
        """
        if not self.libraries_dock:
            return

        # Check the current state of the main dock and the button
        dock_visible = self.libraries_dock.isVisible()
        
        if checked:
            # If the button is checked, ensure the dock is visible and the tab is open
            if not dock_visible:
                self.libraries_dock.setVisible(True)

            if library_name == 'modules':
                if not self.tab_states['modules']:
                    # The tab is "closed", re-add it
                    index = self.libraries_tabs.addTab(self.modules_widget, "Modules Library")
                    self.setup_tab_close_button(index)
                    self.tab_states['modules'] = True
                self.libraries_tabs.setCurrentWidget(self.modules_widget)
                
                # Uncheck the other button if it's checked
                if self.toggle_test_cases_button_action.isChecked():
                    self.toggle_test_cases_button_action.setChecked(False)
            
            elif library_name == 'test_cases':
                if not self.tab_states['test_cases']:
                    # The tab is "closed", re-add it
                    index = self.libraries_tabs.addTab(self.test_cases_widget, "Test Cases Library")
                    self.setup_tab_close_button(index)
                    self.tab_states['test_cases'] = True
                self.libraries_tabs.setCurrentWidget(self.test_cases_widget)
                
                # Uncheck the other button if it's checked
                if self.toggle_modules_button_action.isChecked():
                    self.toggle_modules_button_action.setChecked(False)

        else:
            # If a button is un-checked, it means the tab should be closed.
            # This logic is handled by the close_library_tab function,
            # which is connected to the tab's 'x' button.
            # So we just need to ensure the dock is hidden if all tabs are closed.
            if self.libraries_tabs.count() == 0:
                self.libraries_dock.setVisible(False)

    # --- NEW: Function to handle tab closing ---
    # --- REVISED: Function to handle tab closing (no destruction) ---
    def close_library_tab(self, index):
        """
        Handles the closing of a tab by removing it from the tab widget
        and unchecking the corresponding toolbar button. The widget is not destroyed.
        """
        tab_title = self.libraries_tabs.tabText(index)
        
        # Uncheck the corresponding button
        if tab_title == "Modules Library":
            self.tab_states['modules'] = False
            self.toggle_modules_button_action.setChecked(False)
        elif tab_title == "Test Cases Library":
            self.tab_states['test_cases'] = False
            self.toggle_test_cases_button_action.setChecked(False)

        self.libraries_tabs.removeTab(index)
        
        # If no tabs are left, hide the dock
        if self.libraries_tabs.count() == 0:
            self.libraries_dock.setVisible(False)
            # Also uncheck the Libraries menu action
            self.toggle_libraries_action.setChecked(False)

        
    def filter_modules(self, query: str):
        """
        Filters the modules tree based on the search query.
        """
        if not self.module_tree_root:
            return
            
        query = query.strip().lower()
        for i in range(self.module_tree_root.childCount()):
            item = self.module_tree_root.child(i)
            item_text = item.text(0).lower()
            if query in item_text:
                item.setHidden(False)
            else:
                item.setHidden(True)

    def filter_test_cases(self, query: str):
        """
        Filters the test case tree based on the search query.
        This also filters the steps within each test case.
        """
        if not self.test_case_tree_root:
            return

        query = query.strip().lower()
        for i in range(self.test_case_tree_root.childCount()):
            test_case_item = self.test_case_tree_root.child(i)
            found_match = False
            
            # Check the test case name itself
            if query in test_case_item.text(0).lower():
                found_match = True
            
            # Check the steps within the test case
            for j in range(test_case_item.childCount()):
                step_item = test_case_item.child(j)
                step_text = step_item.text(0).lower()
                if query in step_text:
                    found_match = True
                
            test_case_item.setHidden(not found_match)

    # --- NEW: Method to hide the PCOMM preview
    def hide_pcomm_preview(self):
        """Hides the central PCOMM preview frame."""
        self.central_frame.setVisible(False)
        
    def create_test_case(self):
        dialog = TestCaseNameDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            test_case_name = dialog.test_case_name
            test_case_description = dialog.test_case_description
            test_case_assumptions = dialog.test_case_assumptions  # ✅ NEW
            
            if test_case_name in self.test_cases:
                QMessageBox.warning(self, "Duplicate Name", f"A test case with the name '{test_case_name}' already exists.")
                return

            self.test_cases[test_case_name] = {
                "steps": [],
                "description": test_case_description,
                "assumptions": test_case_assumptions,  # ✅ NEW
                "prerequisites": []
            }
            self.save_test_cases_to_file()
            self.update_test_case_tree()
            self.statusBar().showMessage(f"Test case '{test_case_name}' created.", 5000)

    def show_test_case_context_menu(self, position: QPoint):
        """
        Shows the context menu for the test case tree with template and data source options.
        """
        item = self.test_case_tree.itemAt(position)
        print(f"DEBUG CONTEXT MENU: Item clicked: {item}")
        
        if item and item.parent() is self.test_case_tree_root:
            menu = QMenu()
            
            test_case_name = item.text(0)
            print(f"DEBUG CONTEXT MENU: Test case name: {test_case_name}")
            
            # ✅ NEW: Add to Test Execution option
            add_to_execution_action = QAction("▶️ Add to Test Execution", self)
            add_to_execution_action.triggered.connect(lambda: self.add_test_case_to_execution(test_case_name))
            menu.addAction(add_to_execution_action)
            
            menu.addSeparator()
            
            # Add template conversion option
            convert_action = QAction("📋 Convert to Template", self)
            convert_action.triggered.connect(lambda: self.convert_specific_test_case_to_template(test_case_name))
            menu.addAction(convert_action)
            
            # Add link data source option
            link_action = QAction("🔗 Link Data Source", self)
            link_action.triggered.connect(lambda: (print(f"DEBUG: Link action triggered for {test_case_name}"), self.link_data_source_to_test_case(test_case_name)))
            menu.addAction(link_action)
            
            # Check if data source is already linked
            test_case_data = self.test_cases.get(test_case_name, {})
            if 'data_source' in test_case_data:
                menu.addSeparator()
                
                # View data source info
                view_action = QAction("ℹ️ View Data Source Info", self)
                view_action.triggered.connect(lambda: self.view_data_source_info(test_case_name))
                menu.addAction(view_action)
                
                # Unlink data source
                unlink_action = QAction("❌ Unlink Data Source", self)
                unlink_action.triggered.connect(lambda: self.unlink_data_source(test_case_name))
                menu.addAction(unlink_action)
            
            print("DEBUG CONTEXT MENU: Showing menu...")
            menu.exec(self.test_case_tree.viewport().mapToGlobal(position))
            print("DEBUG CONTEXT MENU: Menu closed")
    
    def add_test_case_to_execution(self, test_case_name):
        """
        Adds the selected test case to Test Execution as a standalone test (without opening the dialog).
        """
        if test_case_name not in self.test_cases:
            QMessageBox.warning(self, "Test Case Not Found", 
                              f"Test case '{test_case_name}' not found in library.")
            return
        
        # Load existing execution data
        execution_data_file = 'test_execution_data.json'
        
        try:
            if os.path.exists(execution_data_file):
                with open(execution_data_file, 'r') as f:
                    execution_data = json.load(f)
            else:
                execution_data = {
                    'projects': {},
                    'standalone': {}
                }
            
            # Get test case data and add it as standalone
            test_case_data = copy.deepcopy(self.test_cases[test_case_name])
            
            # Add to standalone section
            execution_data['standalone'][test_case_name] = {
                'status': 'Not Run',
                'selected_step': 0,
                'test_case_data': test_case_data
            }
            
            # Save back to file
            with open(execution_data_file, 'w') as f:
                json.dump(execution_data, f, indent=4)
            
            # Show success message
            QMessageBox.information(
                self, 
                "Success", 
                f"Test case '{test_case_name}' has been added to Test Execution.\n\n"
                "Click 'Test Execution' button to run it."
            )
            
            self.statusBar().showMessage(f"Test case '{test_case_name}' added to Test Execution.", 5000)
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add test case to execution:\n\n{str(e)}")    
    
    def add_new_step(self, test_case_item: QTreeWidgetItem):
        """
        Adds a new step under the selected test case with a dropdown
        of available modules.
        """
        test_case_name = test_case_item.text(0)
        step_item = QTreeWidgetItem(test_case_item, ["New Step"])
        
        # Create a widget to hold the QComboBox and a line edit for the name
        widget = QWidget()
        layout = QHBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)
        
        # Dropdown to select modules
        module_combo = QComboBox()
        module_combo.addItem("Select Module...")
        for module_name in self.modules.keys():
            module_combo.addItem(module_name)
        
        module_combo.currentIndexChanged.connect(lambda index: self.save_step_module(test_case_item, step_item, index))
        
        layout.addWidget(module_combo)
        
        step_item.setSizeHint(0, QSize(150, 25))
        self.test_case_tree.setItemWidget(step_item, 0, widget)

        # Update the data structure and save
        step_data = {"module_name": None}
        self.test_cases[test_case_name]["steps"].append(step_data)
        self.save_test_cases_to_file()
        self.statusBar().showMessage(f"New step added to test case '{test_case_name}'.", 5000)

    def save_step_module(self, test_case_item: QTreeWidgetItem, step_item: QTreeWidgetItem, index: int):
        """
        Saves the selected module name for a step.
        """
        test_case_name = test_case_item.text(0)
        step_index = test_case_item.indexOfChild(step_item)
        
        if test_case_name in self.test_cases and step_index >= 0 and index > 0:
            module_name = self.test_case_tree.itemWidget(step_item, 0).findChild(QComboBox).currentText()
            self.test_cases[test_case_name]["steps"][step_index]["module_name"] = module_name
            step_item.setText(0, f"Step {step_index + 1}: {module_name}")
            self.save_test_cases_to_file()

    def test_cases_library(self):
        """
        Handles the 'Test cases Library' button click.
        Toggles the visibility of the Test Cases dock and the PCOMM preview.
        This method is simplified as there is no step addition process now.
        """
        is_visible = not self.test_cases_deck.isVisible()
        self.test_cases_deck.setVisible(is_visible)
        self.central_frame.setVisible(not is_visible)
        self.toggle_test_cases_action.setChecked(is_visible)


    def test_execution(self):
        """Handles the Test Execution workflow via a pop-up dialog."""
        dialog = TestExecutionDialog(self, test_cases_data=[])
        dialog.exec()

    # --- NEW: Method to show a dialog and add a new label with user input ---
    def add_new_label(self):
        """
        Shows a dialog to get label details and adds it to the selected module.
        âœ… UPDATED: New label is added at the top of the list (index 0).
        """
        selected_module_item = self.module_tree.currentItem()
        if not selected_module_item or selected_module_item.parent() is None:
            QMessageBox.warning(self, "No Module Selected", "Please select a module first.")
            return

        dialog = AddLabelDialog(self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            new_label = dialog.get_data()
            
            module_name = selected_module_item.data(0, Qt.ItemDataRole.UserRole)
            module_data = self.modules.get(module_name)
            
            if module_data:
                if 'labels' not in module_data:
                    module_data['labels'] = []
                
                # âœ… CHANGED: Insert at index 0 instead of appending
                module_data['labels'].insert(0, new_label)
                
                self.save_modules_to_file()
                self.display_module_details_and_screenshot(selected_module_item, 0)
                self.statusBar().showMessage("New label added at the top.", 5000)
                
    def copy_selected_labels(self):
        """Copies the selected label(s) to the class clipboard."""
        selected_rows = self.labels_table.selectionModel().selectedRows()
        
        if not selected_rows:
            QMessageBox.warning(self, "No Selection", "Please select at least one label to copy.")
            return
        
        selected_module_item = self.module_tree.currentItem()
        if not selected_module_item or selected_module_item.parent() is None:
            QMessageBox.warning(self, "No Module Selected", "Please select a module first.")
            return
        
        module_name = selected_module_item.data(0, Qt.ItemDataRole.UserRole)
        module_data = self.modules.get(module_name)
        
        if not module_data or 'labels' not in module_data:
            return
        
        # Get the indices of selected rows
        selected_indices = [row.row() for row in selected_rows]
        selected_indices.sort()  # Sort to maintain order
        
        # Deep copy the selected labels
        PCOMMMainFrame.copied_labels = [copy.deepcopy(module_data['labels'][i]) for i in selected_indices]
        
        label_word = "label" if len(PCOMMMainFrame.copied_labels) == 1 else "labels"
        self.statusBar().showMessage(f"Copied {len(PCOMMMainFrame.copied_labels)} {label_word} to clipboard.", 3000)
        
    def paste_copied_labels(self):
        """Pastes the copied label(s) at the end of the labels list."""
        if not PCOMMMainFrame.copied_labels:
            QMessageBox.warning(self, "Nothing to Paste", "No labels have been copied yet.")
            return
        
        selected_module_item = self.module_tree.currentItem()
        if not selected_module_item or selected_module_item.parent() is None:
            QMessageBox.warning(self, "No Module Selected", "Please select a module first.")
            return
        
        module_name = selected_module_item.data(0, Qt.ItemDataRole.UserRole)
        module_data = self.modules.get(module_name)
        
        if not module_data:
            return
        
        if 'labels' not in module_data:
            module_data['labels'] = []
        
        # Deep copy the labels to paste (to avoid reference issues)
        labels_to_paste = [copy.deepcopy(label) for label in PCOMMMainFrame.copied_labels]
        
        # Insert the copied labels at the end
        for label in labels_to_paste:
            module_data['labels'].append(label)
        
        # Save and refresh
        self.save_modules_to_file()
        self.display_module_details_and_screenshot(selected_module_item, 0)
        
        label_word = "label" if len(labels_to_paste) == 1 else "labels"
        self.statusBar().showMessage(f"Pasted {len(labels_to_paste)} {label_word} at the end.", 3000)
        
    def delete_selected_labels(self):
        """Deletes the selected label(s) from the module."""
        selected_rows = self.labels_table.selectionModel().selectedRows()
        
        if not selected_rows:
            QMessageBox.warning(self, "No Selection", "Please select at least one label to delete.")
            return
        
        selected_module_item = self.module_tree.currentItem()
        if not selected_module_item or selected_module_item.parent() is None:
            return
        
        module_name = selected_module_item.data(0, Qt.ItemDataRole.UserRole)
        module_data = self.modules.get(module_name)
        
        if not module_data or 'labels' not in module_data:
            return
        
        # Get the indices of selected rows
        selected_indices = [row.row() for row in selected_rows]
        selected_indices.sort(reverse=True)  # Sort in reverse to delete from bottom to top
        
        # Get label names for confirmation message
        label_names = [module_data['labels'][i].get('name', 'Unnamed') for i in selected_indices]
        
        # Confirm deletion
        label_word = "label" if len(selected_indices) == 1 else "labels"
        reply = QMessageBox.question(
            self,
            "Confirm Delete",
            f"Are you sure you want to delete {len(selected_indices)} {label_word}?\n\n{', '.join(label_names)}",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply != QMessageBox.StandardButton.Yes:
            return
        
        # Delete the labels (from bottom to top to maintain correct indices)
        for index in selected_indices:
            if 0 <= index < len(module_data['labels']):
                del module_data['labels'][index]
        
        # Save and refresh
        self.save_modules_to_file()
        self.display_module_details_and_screenshot(selected_module_item, 0)
        self.properties_tree.clear()
        
        self.statusBar().showMessage(f"Deleted {len(selected_indices)} {label_word}.", 3000)

    # --- NEW: Methods to move a label up or down ---
    def move_label_up(self):
        """
        Moves the currently selected label up in the list.
        """
        self.move_label(-1)

    def move_label_down(self):
        """
        Moves the currently selected label down in the list.
        """
        self.move_label(1)

    def move_label(self, direction):
        """
        Generic method to move a label up or down.
        Only works with single selection.
        ✅ FIXED: Removed duplicate properties tree population.
        """
        selected_rows = self.labels_table.selectionModel().selectedRows()
        
        if not selected_rows:
            QMessageBox.warning(self, "No Label Selected", "Please select a label to move.")
            return
        
        if len(selected_rows) > 1:
            QMessageBox.warning(self, "Multiple Selection", "Please select only one label to move.")
            return
        
        current_row = selected_rows[0].row()
        new_row = current_row + direction
        
        if not (0 <= new_row < self.labels_table.rowCount()):
            return  # Cannot move beyond the bounds

        selected_module_item = self.module_tree.currentItem()
        if not selected_module_item or selected_module_item.parent() is None:
            return

        module_name = selected_module_item.data(0, Qt.ItemDataRole.UserRole)
        module_data = self.modules.get(module_name)
        
        if module_data and 'labels' in module_data:
            # Swap the items in the underlying data structure
            labels_list = module_data['labels']
            labels_list[current_row], labels_list[new_row] = labels_list[new_row], labels_list[current_row]
            
            self.save_modules_to_file()
            
            # ✅ Block signals to prevent triggering cellClicked
            self.labels_table.blockSignals(True)
            
            # Refresh the table and re-select the moved item
            self.display_module_details_and_screenshot(selected_module_item, 0)
            self.labels_table.setCurrentCell(new_row, 0)
            
            # ✅ Make the moved row bold
            for j in range(self.labels_table.columnCount() - 1):  # Exclude delete button column
                item = self.labels_table.item(new_row, j)
                if item:
                    font = item.font()
                    font.setBold(True)
                    item.setFont(font)
            
            # ✅ FIXED: Just call display_selected_label_properties - it already updates the tree!
            self.display_selected_label_properties(new_row, 0)
            
            # ✅ Unblock signals
            self.labels_table.blockSignals(False)
            
            self.statusBar().showMessage(f"Label moved {'up' if direction == -1 else 'down'}.", 5000)

            
    def toggle_toolbar_visibility(self):
        """Toggles the visibility of the main toolbar."""
        toolbar = self.findChild(QToolBar, "main_toolbar")
        if toolbar:
            self.toolbar_visible = not self.toolbar_visible
            toolbar.setVisible(self.toolbar_visible)
            self.toggle_toolbar_action.setChecked(self.toolbar_visible)
            
            # Update the status message
            status_msg = "Toolbar shown" if self.toolbar_visible else "Toolbar hidden"
            self.statusBar().showMessage(status_msg, 3000)

    def show_context_menu(self, position: QPoint):
        """
        Shows the context menu for the module tree.
        The 'Rename' action has been removed as it is now triggered by double-clicking.
        """
        # The parent() check ensures we only show the menu for the module items, not the root
        item = self.module_tree.itemAt(position)
        if item and item.parent() == self.module_tree_root:
            menu = QMenu()
            # Only add other actions here if needed in the future
            menu.exec(self.module_tree.viewport().mapToGlobal(position))

    def scan_pcomm_screen(self):
        """
        Scans the PCOMM screen, extracts text, creates a new module, and
        displays the captured text and screenshot in the preview window.
        """
        try:
            # --- NEW: Hide the test cases dock before scanning ---
          
            pcomm_window = gw.getWindowsWithTitle(self.pcomm_window_title)[0]
            
            # Use Win32 API's PrintWindow for a more reliable capture
            hwnd = pcomm_window._hWnd
            left, top, right, bottom = win32gui.GetWindowRect(hwnd)
            w = right - left
            h = bottom - top

            # Get the device context and create a compatible one
            hwnd_dc = win32gui.GetWindowDC(hwnd)
            mfc_dc = win32ui.CreateDCFromHandle(hwnd_dc)
            save_dc = mfc_dc.CreateCompatibleDC()

            # Create a bitmap to hold the image
            save_bitmap = win32ui.CreateBitmap()
            save_bitmap.CreateCompatibleBitmap(mfc_dc, w, h)
            save_dc.SelectObject(save_bitmap)
            
            # --- The key step: Use PrintWindow to get a reliable screenshot ---
            # Try to use PW_CLIENTONLY, but fall back to 0 if it's not defined
            try:
                flag = win32con.PW_CLIENTONLY
            except AttributeError:
                # Fallback to 0 if PW_CLIENTONLY is not found
                print("win32con.PW_CLIENTONLY not found, falling back to 0.")
                flag = 0
                
            result = windll.user32.PrintWindow(hwnd, save_dc.GetSafeHdc(), flag)

            if result != 1:
                raise Exception("Failed to capture window with PrintWindow.")
                
            # Convert the bitmap to a Pillow Image
            bmpinfo = save_bitmap.GetInfo()
            bmpstr = save_bitmap.GetBitmapBits(True)
            im = Image.frombuffer(
                'RGB',
                (bmpinfo['bmWidth'], bmpinfo['bmHeight']),
                bmpstr, 'raw', 'BGRX', 0, 1
            )
            
            # Clean up the device contexts and bitmap
            win32gui.DeleteObject(save_bitmap.GetHandle())
            save_dc.DeleteDC()
            mfc_dc.DeleteDC()
            win32gui.ReleaseDC(hwnd, hwnd_dc)

            # --- Capture Text for basic overview ---
            pcomm_window.activate()
            time.sleep(0.5)
            pyautogui.hotkey('ctrl', 'c')
            time.sleep(0.5)
            copied_text = pyperclip.paste()

            # Process the text and screenshot to identify fields and create a new module
            module_name = f"Module_{self.module_counter}"
            
            # Use the new visual processing method
            module_data = self.process_pcomm_screen_from_text(copied_text)
            
            # Add other necessary data
            # Add other necessary data
            screenshot_dir = 'Modules Screenshots'
            module_data["screenshot"] = os.path.join(screenshot_dir, f"{module_name}.png")
            module_data["captured_text"] = copied_text # Store the raw text for later display

            # Create the directory if it doesn't exist
            os.makedirs(screenshot_dir, exist_ok=True)

            # Save the image
            im.save(module_data["screenshot"])
            
            self.modules[module_name] = module_data
            self.module_counter += 1
            self.save_modules_to_file()
            self.update_module_tree()
            self.statusBar().showMessage(f"Scanned screen and created new module: {module_name}", 5000)

            # Update the central PCOMM preview canvas to show the text
            self.display_text_in_preview(copied_text)
            
            # --- NEW: Ensure the PCOMM preview is visible when a scan is performed ---
            self.central_frame.setVisible(True)


        except IndexError:
            QMessageBox.warning(self, "Error", f"PCOMM window with title '{self.pcomm_window_title}' not found.")
        except Exception as e:
            QMessageBox.critical(self, "Error", str(e))
        
    def process_pcomm_screen_from_text(self, text):
        """
        Parses the PCOMM screen text line by line to extract labels and text boxes.
        This version relies on text content, not pixels, and specifically looks for 'a' characters.
        UPDATED: Finds 'a+' blocks with spaces before and after, treats preceding text as label.
        Example: "Company: aaaaa Product: aaa" -> Label "Company:" for first field, "Product:" for second.
        """
        lines = text.split('\n')
        labels = []
        
        for row, line in enumerate(lines):
            # Find all occurrences of 'a+' that are surrounded by spaces
            # Use regex to find all matches with their positions
            pattern = r'(?:^|\s)(a+)(?=\s|$)'
            matches = re.finditer(pattern, line)
            
            last_field_end = 0  # Track where the last field ended
            
            for match in matches:
                # Get the 'a+' field
                field_text = match.group(1)  # Group 1 is the 'a+' part
                field_start = match.start(1)  # Start position of 'a+' (not including leading space)
                field_end = match.end(1)
                field_length = len(field_text)
                
                # Extract label text: everything between last field end and current field start
                label_text = line[last_field_end:field_start].strip()
                
                if label_text:
                    labels.append({
                        "name": label_text,
                        "row": row + 1,
                        "column": field_start + 1,  # Convert to 1-based
                        "length": field_length
                    })
                
                # Update last_field_end to track where this field ended
                last_field_end = field_end

        return {"labels": labels}

    def get_whitespace_length(self, line, start_index):
        """
        Calculates the length of the whitespace block to determine field size.
        """
        length = 0
        for char in line[start_index:]:
            if char.isspace():
                length += 1
            else:
                break
        return length

    def display_text_in_preview(self, text):
        """
        Loads and displays the captured text in the main preview window.
        
        UPDATED: Now uses QTextEdit and highlights consecutive 'a' characters in red,
        but only if they are surrounded by spaces on both sides.
        """
        # Ensure the central widget is the QTextEdit before trying to set text
        if isinstance(self.centralWidget(), QFrame):
            import re
            
            # Updated regex pattern to match 'a+' only if preceded and followed by a space
            # Positive lookbehind (?<=\s) ensures there's a space before
            # Positive lookahead (?=\s) ensures there's a space after
            html_text = re.sub(
                r'(?<=\s)(a+)(?=\s)', 
                r'<span style="color: red;">\1</span>', 
                text
            )
            
            # Wrap in HTML with pre-formatted text to preserve spacing
            html_content = f'<pre style="color: #00ffff; font-family: Courier New, monospace;">{html_text}</pre>'
            
            self.pcomm_canvas_text_edit.setHtml(html_content)
            
            # Get the current text cursor
            cursor = self.pcomm_canvas_text_edit.textCursor()
            
            # Move the cursor to the top-left position (start of the document)
            cursor.movePosition(QTextCursor.MoveOperation.Start)
            
            # Set the new cursor position and ensure it's visible
            self.pcomm_canvas_text_edit.setTextCursor(cursor)
    
    def display_image_in_preview(self, image_path):
        """
        NOTE: This method is now obsolete as the central widget is a QTextEdit.
        It is being kept as a placeholder in case the functionality is
        needed in the future.
        """
        pass

    def format_pcomm_text_for_display(self, text):
        """
        Formats the copied text to fit the 24x80 grid for display.
        This is not used currently since we are displaying the image,
        but it's good to keep in case we want to revert back.
        """
        lines = text.split('\n')
        formatted_lines = [line.ljust(self.num_cols) for line in lines[:self.num_rows]]
        return "\n".join(formatted_lines)
        
    
    def update_selection_info_only(self):
        """
        Only updates the status bar with selection info (no popup).
        """
        cursor = self.pcomm_canvas_text_edit.textCursor()
        
        if cursor.hasSelection():
            start_pos = cursor.selectionStart()
            end_pos = cursor.selectionEnd()

            start_cursor = self.pcomm_canvas_text_edit.textCursor()
            end_cursor = self.pcomm_canvas_text_edit.textCursor()
            
            start_cursor.setPosition(start_pos, QTextCursor.MoveMode.MoveAnchor)
            end_cursor.setPosition(end_pos, QTextCursor.MoveMode.MoveAnchor)
            
            start_row = start_cursor.blockNumber() + 1
            start_col = start_cursor.columnNumber() + 1
            
            selected_text = cursor.selectedText()
            
            if selected_text:
                end_cursor.movePosition(QTextCursor.MoveOperation.PreviousCharacter)

            end_row = end_cursor.blockNumber() + 1
            end_col = end_cursor.columnNumber() + 1

            selection_length = len(selected_text.replace('\n', ''))
            
            message = (
                f"Selection: '{selected_text.replace('\n', '')}' | "
                f"Start: ({start_row}, {start_col}) | "
                f"End: ({end_row}, {end_col})"
            )
            self.statusBar().showMessage(message, 0)
        else:
            self.statusBar().clearMessage()        
        
    def update_selection_info(self):
        """
        Displays the selected text's start and end coordinates in the status bar.
        Also prompts the user to define a label if text is selected.
        """
        cursor = self.pcomm_canvas_text_edit.textCursor()
        
        if cursor.hasSelection():
            # Get the starting and ending positions of the selection
            start_pos = cursor.selectionStart()
            end_pos = cursor.selectionEnd()

            # Create cursors for the start and end positions
            start_cursor = self.pcomm_canvas_text_edit.textCursor()
            end_cursor = self.pcomm_canvas_text_edit.textCursor()
            
            start_cursor.setPosition(start_pos, QTextCursor.MoveMode.MoveAnchor)
            end_cursor.setPosition(end_pos, QTextCursor.MoveMode.MoveAnchor)
            
            # Get the row and column numbers (Qt is 0-indexed, so add 1)
            start_row = start_cursor.blockNumber() + 1
            start_col = start_cursor.columnNumber() + 1
            
            # Get the selected text to determine if it's empty
            selected_text = cursor.selectedText()
            
            # The selectionEnd() position is after the last character.
            # We need to get the column of the last character itself.
            if selected_text:
                end_cursor.movePosition(QTextCursor.MoveOperation.PreviousCharacter)

            end_row = end_cursor.blockNumber() + 1
            end_col = end_cursor.columnNumber() + 1

            # Calculate selection length (single line selection)
            selection_length = len(selected_text.replace('\n', ''))
            
            message = (
                f"Selection: '{selected_text.replace('\n', '')}' | "
                f"Start: ({start_row}, {start_col}) | "
                f"End: ({end_row}, {end_col})"
            )
            self.statusBar().showMessage(message, 0) # 0 means the message stays until a new one is set
            
            # ✅ NEW: Prompt user to define a label
            self.prompt_define_label(start_row, start_col, selection_length)
        else:
            self.statusBar().clearMessage() # Clear the status bar if no text is selected
    
            
    def prompt_define_label(self, row, column, length):
        """
        Prompts the user to define a label for the selected text.
        If yes, opens the AddLabelDialog with pre-filled values.
        âœ… UPDATED: New label is added at the top of the list (index 0).
        """
        reply = QMessageBox.question(
            self,
            "Define Label",
            "Want to define a Label?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.Yes
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            # Check if a module is currently selected
            selected_module_item = self.module_tree.currentItem()
            if not selected_module_item or selected_module_item.parent() is None:
                QMessageBox.warning(self, "No Module Selected", 
                                  "Please select a module first before defining a label.")
                return
            
            # Open the AddLabelDialog with pre-filled values
            dialog = AddLabelDialog(self)
            dialog.row_input.setText(str(row))
            dialog.column_input.setText(str(column))
            dialog.length_input.setText(str(length))
            
            if dialog.exec() == QDialog.DialogCode.Accepted:
                new_label = dialog.get_data()
                
                module_name = selected_module_item.data(0, Qt.ItemDataRole.UserRole)
                module_data = self.modules.get(module_name)
                
                if module_data:
                    if 'labels' not in module_data:
                        module_data['labels'] = []
                    
                    # âœ… CHANGED: Insert at index 0 instead of appending
                    module_data['labels'].insert(0, new_label)
                    
                    self.save_modules_to_file()
                    self.display_module_details_and_screenshot(selected_module_item, 0)
                    self.statusBar().showMessage("New label added from selection at the top.", 5000)

    def update_module_tree(self):
        """
        Updates the module tree widget with the latest module data.
        Ensures a clean slate before populating.
        """
        # Clear the entire tree to remove all items.
        self.module_tree.clear()

        if self.modules:
            # Create the 'Modules' root item.
            self.module_tree_root = QTreeWidgetItem(self.module_tree, ["Modules"])
            self.module_tree_root.setExpanded(True)
            
            for module_name, module_data in self.modules.items():
                module_item = QTreeWidgetItem(self.module_tree_root, [module_name])
                module_item.setData(0, Qt.ItemDataRole.UserRole, module_name)
                # Explicitly set the item to be editable
                module_item.setFlags(module_item.flags() | Qt.ItemFlag.ItemIsEditable)
                
                # Use a container widget to hold both buttons
                button_widget = QWidget()
                button_layout = QHBoxLayout(button_widget)
                button_layout.setContentsMargins(0, 0, 0, 0)
                button_layout.setSpacing(5)

                save_button = QPushButton("💾")
                save_button.setFixedSize(17, 17)
                save_button.setStyleSheet("""
                    QPushButton {
                        border: none;
                        background-color: transparent;
                        font-size: 16px;
                        padding: 2px;
                    }
                    QPushButton:hover {
                        background-color: #e0e0e0;
                        border-radius: 4px;
                    }
                """)
                save_button.setCursor(Qt.CursorShape.PointingHandCursor)
                save_button.setToolTip(f"Save '{module_name}' to a file.")
                # Pass the entire item to the lambda for a more robust connection
                save_button.clicked.connect(lambda _, item=module_item: self.save_single_module(item.data(0, Qt.ItemDataRole.UserRole)))
                button_layout.addWidget(save_button)
                
                delete_button = QPushButton("✕")
                delete_button.setFixedSize(20, 20)
                
                # ✅ Set font explicitly before stylesheet
                delete_font = QFont()
                delete_font.setBold(True)
                delete_font.setPointSize(14)
                delete_button.setFont(delete_font)
                
                delete_button.setStyleSheet("""
                    QPushButton {
                        color: #dc2626;
                        border: none;
                        border-radius: 4px;
                        background-color: transparent;
                        padding: 0px;
                        margin-left: 4px;
                        margin-right: 4px;
                    }
                    QPushButton:hover {
                        background-color: #fee2e2;
                    }
                    QPushButton:pressed {
                        background-color: #fecaca;
                    }
                """)
                delete_button.setCursor(Qt.CursorShape.PointingHandCursor)
                delete_button.setToolTip(f"Delete '{module_name}'.")
                delete_button.clicked.connect(lambda _, item=module_item: self.delete_module(item))
                button_layout.addWidget(delete_button)
                
                # MINIMAL CHANGE: Use the correct enum members for QSizePolicy
                spacer = QSpacerItem(10, 0, QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Minimum)
                button_layout.addItem(spacer)
                
                self.module_tree.setItemWidget(module_item, 1, button_widget)
        else:
            # If no modules exist, reset the root reference.
            self.module_tree_root = None
    
    # NEW: Method to update the test case tree
    def update_test_case_tree(self):
        """
        Updates the test case tree widget with the latest test case data.
        Shows all step types with descriptive names matching the Edit Test Case dialog.
        """
        self.test_case_tree.clear()
        
        if self.test_cases:
            self.test_case_tree_root = QTreeWidgetItem(self.test_case_tree, ["Test Cases"])
            self.test_case_tree_root.setExpanded(True)

            for test_case_name, test_case_data in self.test_cases.items():
                test_case_item = QTreeWidgetItem(self.test_case_tree_root, [test_case_name])
                test_case_item.setData(0, Qt.ItemDataRole.UserRole, test_case_name)
                
                # Use a container widget to hold both buttons
                button_widget = QWidget()
                button_layout = QHBoxLayout(button_widget)
                button_layout.setContentsMargins(0, 0, 0, 0)
                button_layout.setSpacing(5)

                save_button = QPushButton("💾")
                save_button.setFixedSize(17, 17)
                save_button.setStyleSheet("""
                    QPushButton {
                        border: none;
                        background-color: transparent;
                        font-size: 16px;
                        padding: 2px;
                    }
                    QPushButton:hover {
                        background-color: #e0e0e0;
                        border-radius: 4px;
                    }
                """)
                save_button.setCursor(Qt.CursorShape.PointingHandCursor)
                save_button.setToolTip(f"Save '{test_case_name}' to a file.")
                save_button.clicked.connect(lambda _, name=test_case_name: self.save_single_test_case(name))
                button_layout.addWidget(save_button)

                delete_button = QPushButton("✕")
                delete_button.setFixedSize(18, 18)  # ✅ Slightly larger for consistency
                
                # ✅ Set font explicitly before stylesheet
                delete_font = QFont()
                delete_font.setBold(True)
                delete_font.setPointSize(12)
                delete_button.setFont(delete_font)
                
                delete_button.setStyleSheet("""
                    QPushButton {
                        color: #dc2626;
                        border: none;
                        border-radius: 4px;
                        background-color: transparent;
                        padding: 0px;
                    }
                    QPushButton:hover {
                        background-color: #fee2e2;
                    }
                    QPushButton:pressed {
                        background-color: #fecaca;
                    }
                """)
                delete_button.setCursor(Qt.CursorShape.PointingHandCursor)
                delete_button.setToolTip(f"Delete '{test_case_name}'.")
                delete_button.clicked.connect(lambda _, item=test_case_item: self.delete_test_case(item))
                button_layout.addWidget(delete_button)
                
                # MINIMAL CHANGE: Use the correct enum members for QSizePolicy
                spacer = QSpacerItem(10, 0, QSizePolicy.Policy.Fixed, QSizePolicy.Policy.Minimum)
                button_layout.addItem(spacer)

                self.test_case_tree.setItemWidget(test_case_item, 1, button_widget)

                # Add existing steps as children items - using same logic as update_steps_list in EditTestCaseDialog
                for i, step_data in enumerate(test_case_data.get("steps", []), 1):
                    step_type = step_data.get("type", "")
                    step_text = f"Step {i}: "
                    
                    if step_type == "module_import":
                        module_name = step_data.get("module_name", "Unknown Module")
                        step_text += module_name
                    elif step_type == "special_key":
                        key_value = step_data.get("key_value", "Unknown Key")
                        step_text += f"Special Key: {key_value}"
                    elif step_type == "capture_screen_text":
                        step_text += "Capture Text Screenshot"
                    elif step_type == "capture_screenshot":
                        step_text += "Capture Screenshot (DOCX)"
                    elif step_type == "random_input":
                        row = step_data.get('row', '?')
                        col = step_data.get('column', '?')
                        value = step_data.get('value', '?')
                        step_text += f"Random Input (Row: {row}, Col: {col}, Value: {value})"
                    elif step_type == "wait":
                        seconds = step_data.get('seconds', '?')
                        step_text += f"Wait: {seconds} second(s)"
                    elif step_type == "break":  # ✅ NEW
                        step_text += "Break: Review & Decision Point"                        
                    else:
                        step_text += "Unknown Step"
                    
                    step_item = QTreeWidgetItem(test_case_item, [step_text])
                    
                    # Add field details only for module_import steps
                    if step_type == "module_import" and 'fields' in step_data:
                        for field in step_data['fields']:
                            field_name = field.get('field_name', 'N/A')
                            field_value = field.get('value', 'N/A')
                            field_item = QTreeWidgetItem(step_item)
                            field_item.setText(0, f"{field_name}: {field_value}")
        
    # NEW helper method to delete a test case
    def delete_test_case(self, item):
        test_case_name = item.text(0)
        reply = QMessageBox.question(
            self,
            'Delete Test Case',
            f"Are you sure you want to delete the test case '{test_case_name}'?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        if reply == QMessageBox.StandardButton.Yes:
            if test_case_name in self.test_cases:
                del self.test_cases[test_case_name]
                self.save_test_cases_to_file()
                self.update_test_case_tree()
                self.statusBar().showMessage(f"Test case '{test_case_name}' deleted.", 5000)

    # NEW helper method to handle saving from the tree item
    def save_single_module(self, module_name):
        """
        Saves a single module's data to a new JSON file.
        """
        # --- NEW: Add a check to ensure module_name exists in the data ---
        if module_name in self.modules:
            file_dialog = QFileDialog()
            # Suggest a default filename based on the module name
            suggested_filename = f"{module_name}.json"
            file_path, _ = file_dialog.getSaveFileName(
                self,
                "Save Module",
                suggested_filename,
                "JSON Files (*.json);;All Files (*)"
            )
            
            if file_path:
                try:
                    with open(file_path, 'w') as f:
                        json.dump({module_name: self.modules[module_name]}, f, indent=4)
                    self.statusBar().showMessage(f"Module '{module_name}' saved to '{os.path.basename(file_path)}'", 5000)
                except Exception as e:
                    QMessageBox.critical(self, "Save Error", f"An error occurred while saving the file: {e}")
        else:
            self.statusBar().showMessage(f"Error: Module '{module_name}' not found. Cannot save.", 5000)

    def import_modules(self):
        """
        Prompts the user to select one or more JSON files and imports the modules
        into the current project.
        """
        file_dialog = QFileDialog()
        file_paths, _ = file_dialog.getOpenFileNames(
            self,
            "Import Modules",
            "",
            "JSON Files (*.json);;All Files (*)"
        )
        
        if file_paths:
            skipped_modules = []
            
            for file_path in file_paths:
                try:
                    with open(file_path, 'r') as f:
                        imported_modules = json.load(f)
                    
                    # Merge the imported modules with the existing ones
                    for name, data in imported_modules.items():
                        if name in self.modules:
                            skipped_modules.append(name)
                        else:
                            self.modules[name] = data
                
                except Exception as e:
                    QMessageBox.critical(self, "Import Error", f"An error occurred while importing the file '{file_path}': {e}")
            
            # Update the module counter to be the largest number plus one
            numeric_modules = [int(re.search(r'\d+', key).group()) for key in self.modules.keys() if re.search(r'\d+', key)]
            self.module_counter = max(numeric_modules) + 1 if numeric_modules else 0
                    
            self.save_modules_to_file()
            self.update_module_tree()
            
            if skipped_modules:
                message = (f"Successfully imported modules. The following modules were skipped due to name conflicts: "
                           f"{', '.join(skipped_modules)}")
                QMessageBox.warning(self, "Import Warning", message)
            else:
                self.statusBar().showMessage("Successfully imported modules.", 5000)

    def delete_module(self, item):
        """
        Deletes the selected module from the tree, the internal data, and the screenshot file.
        """
        module_name_to_delete = item.data(0, Qt.ItemDataRole.UserRole)
        
        # Confirm deletion with the user
        reply = QMessageBox.question(
            self, 
            'Delete Module',
            f"Are you sure you want to delete the module '{module_name_to_delete}'?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, 
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            if module_name_to_delete in self.modules:
                # Delete the associated screenshot file
                screenshot_path = self.modules[module_name_to_delete].get("screenshot")
                if screenshot_path and os.path.exists(screenshot_path):
                    try:
                        os.remove(screenshot_path)
                    except Exception as e:
                        print(f"Error deleting screenshot file: {e}")

                del self.modules[module_name_to_delete]
                self.save_modules_to_file()
                self.update_module_tree()
                self.labels_table.setRowCount(0)
                self.properties_tree.clear()
                self.pcomm_canvas_text_edit.setText("PCOMM Screenshot Preview")
                self.statusBar().showMessage(f"Module '{module_name_to_delete}' deleted.", 5000)
                
    def delete_selected_modules(self):
        """Deletes all selected modules from the tree."""
        selected_items = self.module_tree.selectedItems()
        
        # Filter out the root item and get only module items
        module_items = [item for item in selected_items if item.parent() == self.module_tree_root]
        
        if not module_items:
            QMessageBox.warning(self, "No Selection", "Please select at least one module to delete.")
            return
        
        module_names = [item.data(0, Qt.ItemDataRole.UserRole) for item in module_items]
        
        # Confirm deletion
        module_word = "module" if len(module_names) == 1 else "modules"
        reply = QMessageBox.question(
            self,
            'Delete Modules',
            f"Are you sure you want to delete {len(module_names)} {module_word}?\n\n{', '.join(module_names)}",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            for module_name in module_names:
                if module_name in self.modules:
                    # Delete the associated screenshot file
                    screenshot_path = self.modules[module_name].get("screenshot")
                    if screenshot_path and os.path.exists(screenshot_path):
                        try:
                            os.remove(screenshot_path)
                        except Exception as e:
                            print(f"Error deleting screenshot file: {e}")
                    
                    del self.modules[module_name]
            
            self.save_modules_to_file()
            self.update_module_tree()
            self.labels_table.setRowCount(0)
            self.properties_tree.clear()
            self.pcomm_canvas_text_edit.setText("PCOMM Screenshot Preview")
            self.statusBar().showMessage(f"{len(module_names)} {module_word} deleted.", 5000)

    def delete_selected_test_cases(self):
        """Deletes all selected test cases from the tree."""
        selected_items = self.test_case_tree.selectedItems()
        
        # Filter out the root item and get only test case items
        test_case_items = [item for item in selected_items if item.parent() == self.test_case_tree_root]
        
        if not test_case_items:
            QMessageBox.warning(self, "No Selection", "Please select at least one test case to delete.")
            return
        
        test_case_names = [item.text(0) for item in test_case_items]
        
        # Confirm deletion
        test_word = "test case" if len(test_case_names) == 1 else "test cases"
        reply = QMessageBox.question(
            self,
            'Delete Test Cases',
            f"Are you sure you want to delete {len(test_case_names)} {test_word}?\n\n{', '.join(test_case_names)}",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            for test_case_name in test_case_names:
                if test_case_name in self.test_cases:
                    del self.test_cases[test_case_name]
            
            self.save_test_cases_to_file()
            self.update_test_case_tree()
            self.statusBar().showMessage(f"{len(test_case_names)} {test_word} deleted.", 5000)                

    def display_module_details_and_screenshot(self, item, column):
        """
        Displays the details of a selected module in the labels table
        and shows its corresponding captured text.
        
        UPDATED: Now uses captured text instead of screenshot image.
        """
        self.labels_table.setRowCount(0)
        self.properties_tree.clear()
        
        # Check if the clicked item is a top-level module (not the root item)
        if item.parent() is None:
            # If it's the root, clear the preview
            self.pcomm_canvas_text_edit.setText("PCOMM Screenshot Preview (24x80 Grid)")
            self.central_frame.setVisible(True) # Make sure the preview is visible
            return

        module_name = item.data(0, Qt.ItemDataRole.UserRole)
        module_data = self.modules.get(module_name)
        
        if module_data:
            # Display the saved text
            captured_text = module_data.get('captured_text', "No captured text found.")
            # Ensure the central widget is the QTextEdit before trying to set text
            if isinstance(self.centralWidget(), QFrame):
                self.display_text_in_preview(captured_text)
                self.central_frame.setVisible(True) # Ensure the preview is visible
            
            # Display labels and a delete button for each one
            if 'labels' in module_data:
                self.labels_table.setRowCount(len(module_data['labels']))
                for i, label in enumerate(module_data['labels']):
                    # ✅ NEW: Create items without bold (will be made bold on selection)
                    name_item = QTableWidgetItem(label.get('name', ''))
                    row_item = QTableWidgetItem(str(label.get('row', '')))
                    col_item = QTableWidgetItem(str(label.get('column', '')))
                    length_item = QTableWidgetItem(str(label.get('length', '')))
                    
                    self.labels_table.setItem(i, 0, name_item)
                    self.labels_table.setItem(i, 1, row_item)
                    self.labels_table.setItem(i, 2, col_item)
                    self.labels_table.setItem(i, 3, length_item)
                    
                    delete_label_button = QPushButton("✕")
                    delete_label_button.setFixedSize(18, 18)
                    
                    # ✅ Set font explicitly before stylesheet
                    delete_font = QFont()
                    delete_font.setBold(True)
                    delete_font.setPointSize(12)
                    delete_label_button.setFont(delete_font)
                    
                    delete_label_button.setStyleSheet("""
                        QPushButton {
                            color: #dc2626;
                            border: none;
                            border-radius: 3px;
                            background-color: transparent;
                            padding: 0px;
                        }
                        QPushButton:hover {
                            background-color: #fee2e2;
                        }
                        QPushButton:pressed {
                            background-color: #fecaca;
                        }
                    """)
                    delete_label_button.setCursor(Qt.CursorShape.PointingHandCursor)
                    
                    # Connect the button to a new method that receives the button object itself
                    delete_label_button.clicked.connect(lambda _, btn=delete_label_button: self.delete_label(btn))
                    
                    self.labels_table.setCellWidget(i, 4, delete_label_button)

    def delete_label(self, button):
        """
        Deletes a specific label from a module based on the button clicked.
        FIXED: Now correctly identifies the row of the clicked button.
        """
        # Find the row index by iterating through all rows and checking which one contains this button
        row_index = -1
        for i in range(self.labels_table.rowCount()):
            cell_widget = self.labels_table.cellWidget(i, 4)
            if cell_widget == button:
                row_index = i
                break
        
        # If we couldn't find the button, return
        if row_index == -1:
            return
            
        selected_module_item = self.module_tree.currentItem()
        if not selected_module_item or selected_module_item.parent() is None:
            return
            
        module_name = selected_module_item.data(0, Qt.ItemDataRole.UserRole)
        module_data = self.modules.get(module_name)

        if module_data and 'labels' in module_data and 0 <= row_index < len(module_data['labels']):
            # Get the label name for the confirmation message
            label_name = module_data['labels'][row_index].get('name', 'this label')
            
            # Confirm deletion with the user
            reply = QMessageBox.question(
                self, 
                'Delete Label',
                f"Are you sure you want to delete the label '{label_name}'?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No, 
                QMessageBox.StandardButton.No
            )
            
            if reply == QMessageBox.StandardButton.Yes:
                del module_data['labels'][row_index]
                self.save_modules_to_file()
                
                # Refresh the labels table to reflect the change
                self.labels_table.removeRow(row_index)
                
                # Clear the properties tree as the selection is now invalid
                self.properties_tree.clear()
                self.statusBar().showMessage(f"Label '{label_name}' deleted from module '{module_name}'.", 5000)

    def display_selected_label_properties(self, row, column):
        """Displays properties of the selected label in the properties tree and selects text in preview."""
        if getattr(self, "_updating_properties", False):
            return
        self._updating_properties = True

        try:
            self.properties_tree.clear()
            selected_module_item = self.module_tree.currentItem()
            if not selected_module_item or selected_module_item.parent() is None:
                return

            module_name = selected_module_item.data(0, Qt.ItemDataRole.UserRole)
            module_data = self.modules.get(module_name)

            if module_data and 'labels' in module_data:
                # remove bolding from all rows
                for i in range(self.labels_table.rowCount()):
                    for j in range(self.labels_table.columnCount() - 1):
                        item = self.labels_table.item(i, j)
                        if item:
                            font = item.font()
                            font.setBold(False)
                            item.setFont(font)

                # make selected row bold
                for j in range(self.labels_table.columnCount() - 1):
                    item = self.labels_table.item(row, j)
                    if item:
                        font = item.font()
                        font.setBold(True)
                        item.setFont(font)

                label_data = module_data['labels'][row]
                properties = {
                    "Text": label_data.get('name', ''),
                    "Row": str(label_data.get('row', '')),
                    "Column": str(label_data.get('column', '')),
                    "Length": str(label_data.get('length', 'N/A'))
                }

                for key, value in properties.items():
                    QTreeWidgetItem(self.properties_tree, [key, value])
                
                # ✅ NEW: Select the corresponding text in PCOMM preview
                self.select_text_in_preview(
                    label_data.get('row', 1),
                    label_data.get('column', 1),
                    label_data.get('length', 0)
                )

        finally:
            self._updating_properties = False


    def handle_rename_finish(self, item, column):
        """Handles the completion of a module rename operation."""
        # Check if the item is a top-level module to prevent the TypeError
        if item.parent() != self.module_tree_root:
            return

        new_name = item.text(0)
        old_name = item.data(0, Qt.ItemDataRole.UserRole)
        
        if old_name and new_name != old_name and new_name not in self.modules:
            # Rename the associated screenshot file
            old_path = self.modules[old_name].get("screenshot")
            new_path = f"Modules Screenshots/{new_name}.png"
            if old_path and os.path.exists(old_path):
                try:
                    os.rename(old_path, new_path)
                except Exception as e:
                    print(f"Error renaming screenshot file: {e}")
                    new_path = old_path # Revert if renaming fails

            self.modules[new_name] = self.modules.pop(old_name)
            self.modules[new_name]["screenshot"] = new_path # Update the path
            item.setData(0, Qt.ItemDataRole.UserRole, new_name)
            self.save_modules_to_file()
            self.statusBar().showMessage(f"Renamed module from '{old_name}' to '{new_name}'", 5000)
        elif new_name == old_name:
            # Name did not change, do nothing
            pass
        else:
            QMessageBox.warning(self, "Invalid Name", f"The name '{new_name}' is already in use or invalid.")
            self.update_module_tree() # Revert the item back to the old name
            
    def handle_label_edit_finish(self, row, column):
        """
        Handles the completion of a cell edit in the labels table,
        validates input, updates the internal data, and refreshes the UI.
        """
        # Block signals to prevent infinite loops when we programmatically change a cell
        self.labels_table.blockSignals(True)
        
        selected_module_item = self.module_tree.currentItem()
        if not selected_module_item or selected_module_item.parent() is None:
            self.labels_table.blockSignals(False)
            return

        module_name = selected_module_item.data(0, Qt.ItemDataRole.UserRole)
        module_data = self.modules.get(module_name)
        
        if module_data and 'labels' in module_data and row < len(module_data['labels']):
            label = module_data['labels'][row]
            new_value = self.labels_table.item(row, column).text()

            try:
                if column == 0:  # Field name
                    label['name'] = new_value
                    message = f"Updated Field name to '{new_value}'."
                elif column == 1:  # Row
                    label['row'] = int(new_value)
                    message = f"Updated Row to '{new_value}'."
                elif column == 2:  # Column
                    label['column'] = int(new_value)
                    message = f"Updated Column to '{new_value}'."
                elif column == 3:  # NEW: Handle editing of the Length column
                    label['length'] = int(new_value)
                    message = f"Updated Length to '{new_value}'."
                else:
                    message = "No change."
                    
                # Update the Properties Tree with the latest data
                self.display_selected_label_properties(row, column)
                
                # Save the updated modules data to the file
                self.save_modules_to_file()
                
                self.statusBar().showMessage(message, 5000)

            except ValueError:
                # Handle case where non-numeric input is provided for row/column/length
                QMessageBox.warning(self, "Invalid Input", "Row, Column, and Length values must be numbers.")
                # Revert the cell to its original value from the data
                if column == 1:
                    self.labels_table.setItem(row, column, QTableWidgetItem(str(label.get('row', ''))))
                elif column == 2:
                    self.labels_table.setItem(row, column, QTableWidgetItem(str(label.get('column', ''))))
                elif column == 3:
                    self.labels_table.setItem(row, column, QTableWidgetItem(str(label.get('length', ''))))
        
        self.labels_table.blockSignals(False)

    def save_pcomm_window_config(self):
        """Saves the PCOMM window title configuration."""
        config_file = 'pcomm_config.json'
        try:
            with open(config_file, 'w') as f:
                json.dump({'window_title': self.pcomm_window_title}, f, indent=4)
        except Exception as e:
            QMessageBox.warning(self, "Save Error", f"Failed to save PCOMM configuration: {e}")

    def load_pcomm_window_config(self):
        """Loads the PCOMM window title configuration."""
        config_file = 'pcomm_config.json'
        if os.path.exists(config_file):
            try:
                with open(config_file, 'r') as f:
                    config = json.load(f)
                    self.pcomm_window_title = config.get('window_title', 'SessionA')
            except Exception as e:
                print(f"Error loading PCOMM config: {e}")
                self.pcomm_window_title = 'SessionA'
        else:
            self.pcomm_window_title = 'SessionA'

    def save_modules_to_file(self):
        """Saves the captured module data to a JSON file."""
        with open(self.module_file, 'w') as f:
            json.dump(self.modules, f, indent=4)
        
    def load_modules_from_file(self):
        """Loads captured module data from a JSON file on startup."""
        if os.path.exists(self.module_file):
            with open(self.module_file, 'r') as f:
                self.modules = json.load(f)
            self.module_counter = len(self.modules)
            self.update_module_tree()

    # --- NEW: Methods for Test Cases data persistence ---
    def save_test_cases_to_file(self):
        """Saves the test cases data to a JSON file."""
        with open(self.test_case_file, 'w') as f:
            json.dump(self.test_cases, f, indent=4)

    def load_test_cases_from_file(self):
        """Loads test cases from a JSON file on startup."""
        if os.path.exists(self.test_case_file):
            with open(self.test_case_file, 'r') as f:
                self.test_cases = json.load(f)
            self.test_case_counter = len(self.test_cases)
            self.update_test_case_tree()
            
    def save_single_test_case(self, test_case_name):
        """Saves a single test case to a file chosen by the user."""
        test_case_data = self.test_cases.get(test_case_name)
        if not test_case_data:
            QMessageBox.warning(self, "Error", "Test case not found.")
            return

        # Ensure the test case name is saved with the data for consistent imports
        test_case_to_save = test_case_data.copy()
        test_case_to_save['name'] = test_case_name

        options = QFileDialog.Option.DontUseNativeDialog
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Test Case", f"{test_case_name}.json",
                                                   "JSON Files (*.json);;All Files (*)", options=options)

        if file_path:
            try:
                with open(file_path, 'w') as f:
                    json.dump(test_case_to_save, f, indent=4)
                QMessageBox.information(self, "Success", f"Test Case '{test_case_name}' saved successfully.")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save test case: {e}")
# --- NEW: Methods for Test Cases data persistence ---
    def save_test_cases_to_file(self):
        """Saves the test cases data to a JSON file."""
        with open(self.test_case_file, 'w') as f:
            json.dump(self.test_cases, f, indent=4)

    def load_test_cases_from_file(self):
        """Loads test cases from a JSON file on startup."""
        if os.path.exists(self.test_case_file):
            with open(self.test_case_file, 'r') as f:
                self.test_cases = json.load(f)
            self.update_test_case_tree()
            
    def import_test_cases(self):
        """
        Allows the user to import test cases from a JSON file.
        The imported test cases are added to the existing ones.
        """
        options = QFileDialog.Option.ReadOnly
        file_path, _ = QFileDialog.getOpenFileName(self, "Import Test Cases", "",
                                                   "JSON Files (*.json);;All Files (*)", options=options)
        
        if not file_path:
            return

        try:
            with open(file_path, 'r') as f:
                imported_data = json.load(f)

            imported_count = 0
            
            # Case 1: The file contains a single test case (saved from 'Save Test Case' button)
            if isinstance(imported_data, dict) and 'name' in imported_data and 'steps' in imported_data:
                test_case_name = imported_data['name']
                # Add a unique suffix if a test case with the same name already exists
                if test_case_name in self.test_cases:
                    suffix = 1
                    while f"{test_case_name} ({suffix})" in self.test_cases:
                        suffix += 1
                    new_name = f"{test_case_name} ({suffix})"
                    imported_data['name'] = new_name
                    test_case_name = new_name
                
                self.test_cases[test_case_name] = imported_data
                imported_count = 1
            
            # Case 2: The file contains a dictionary of multiple test cases (saved as a library)
            elif isinstance(imported_data, dict) and all(isinstance(v, dict) and 'name' in v and 'steps' in v for v in imported_data.values()):
                for test_case_name, test_case_data in imported_data.items():
                    if test_case_name in self.test_cases:
                        suffix = 1
                        while f"{test_case_name} ({suffix})" in self.test_cases:
                            suffix += 1
                        new_name = f"{test_case_name} ({suffix})"
                        self.test_cases[new_name] = test_case_data
                    else:
                        self.test_cases[test_case_name] = test_case_data
                    imported_count += 1
            
            if imported_count > 0:
                self.save_test_cases_to_file()
                self.update_test_case_tree()
                QMessageBox.information(self, "Import Successful", 
                                        f"Successfully imported {imported_count} test case(s).")
            else:
                QMessageBox.warning(self, "Import Failed", "The selected file does not contain valid test case data.")

        except json.JSONDecodeError:
            QMessageBox.critical(self, "Import Error", "Failed to decode JSON from the selected file. Please ensure the file is a valid JSON format.")
        except Exception as e:
            QMessageBox.critical(self, "Import Error", f"An unexpected error occurred: {e}")
            
    # Insert this function into your PCOMMMainFrame class or as a standalone helper
    def get_pcomm_session(window_title="Session A"):
        """
        Connects to an active PCOMM session by its window title.
        Returns the PCOMM Session object or None if not found.
        """
        try:
            # PCOMM's COM Automation object
            autObj = win32com.client.Dispatch("PCOMM.autobj")

            # Iterate through all available sessions
            for i in range(1, autObj.Count + 1):
                session = autObj.GetObject(i)
                # Use the window title (or session ID) to find the correct session
                if session.Name.strip() == window_title.strip() or \
                   session.OIA.HostName.strip().startswith(window_title.strip()):
                    return session

            # If the session is not found
            QMessageBox.warning(None, "PCOMM Error",
                                f"PCOMM session with title '{window_title}' not found.")
            return None

        except Exception as e:
            QMessageBox.critical(None, "PCOMM COM Error",
                                 f"Failed to connect to PCOMM COM object. Error: {e}")
            return None

    def send_data_to_pcomm_field(session, row, col, value):
        """
        Sends data to a specific screen position (row, column) in PCOMM.

        Args:
            session: The active PCOMM session COM object.
            row (int): The screen row number (1-based).
            col (int): The screen column number (1-based).
            value (str): The string value to write.
        """
        if not session:
            return False
        
        try:
            # SetCursorPos places the cursor at the desired 1-based coordinates
            session.SetCursorPos(row, col)
            
            # SendKeys sends the string, overwriting what is currently at the position
            session.SendKeys(value)
            
            return True
        except Exception as e:
            QMessageBox.critical(None, "PCOMM Write Error",
                                 f"Failed to write value at ({row}, {col}). Error: {e}")
            return False

    def send_special_key_to_pcomm(session, key):
        """
        Sends a special PCOMM key (e.g., [enter], [tab])
        """
        if not session:
            return False
        try:
            session.SendKeys(key)
            return True
        except Exception as e:
            QMessageBox.critical(None, "PCOMM Key Error",
                                 f"Failed to send key '{key}'. Error: {e}")
            return False
            
    def open_document_config(self):
        """Opens the document configuration dialog."""
        existing_text_elements = self.document_config.get('text_elements', [])
        existing_color = self.document_config.get('highlight_color', 'Yellow')
        
        # ✅ CHANGED: Pass highlight color to dialog constructor
        dialog = DocumentConfigDialog(
            self, 
            existing_config=existing_text_elements,
            existing_highlight_color=existing_color
        )
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            config_result = dialog.get_configuration()
            self.document_config = config_result
            self.save_document_config()
            QMessageBox.information(self, "Success", "Document configuration saved successfully.")
    
    def save_document_config(self):
        """Saves the document configuration to a JSON file."""
        try:
            with open(self.document_config_file, 'w') as f:
                json.dump(self.document_config, f, indent=4)
        except Exception as e:
            QMessageBox.warning(self, "Save Error", f"Failed to save configuration: {e}")
    
    def load_document_config(self):
        """Loads the document configuration from a JSON file."""
        if os.path.exists(self.document_config_file):
            try:
                with open(self.document_config_file, 'r') as f:
                    loaded_config = json.load(f)
                    
                # Handle old format (just a list) vs new format (dict with text_elements and highlight_color)
                if isinstance(loaded_config, list):
                    # Old format - convert to new format
                    self.document_config = {
                        'text_elements': loaded_config,
                        'highlight_color': 'Yellow',
                        'generate_documentation': True
                    }
                else:
                    # New format
                    self.document_config = loaded_config
                    # Ensure all keys exist
                    if 'text_elements' not in self.document_config:
                        self.document_config['text_elements'] = []
                    if 'highlight_color' not in self.document_config:
                        self.document_config['highlight_color'] = 'Yellow'
                    if 'generate_documentation' not in self.document_config:
                        self.document_config['generate_documentation'] = True
                    if 'capture_screen_flow' not in self.document_config:  # ✅ NEW
                        self.document_config['capture_screen_flow'] = False  # ✅ NEW
            except Exception as e:
                print(f"Error loading document config: {e}")
                self.document_config = {'text_elements': [], 'highlight_color': 'Yellow', 'generate_documentation': True, 'capture_screen_flow': False}  # ✅ UPDATED
        else:
            self.document_config = {'text_elements': [], 'highlight_color': 'Yellow', 'generate_documentation': True, 'capture_screen_flow': False}  # ✅ UPDATED
                
    def set_application_icon(self):
        """
        Sets the application icon for the main window and taskbar.
        Creates a custom icon with a play/run symbol.
        """
        # Create an SVG icon with a play/run symbol
        svg_app_icon = """
        <svg width="64" height="64" viewBox="0 0 64 64" fill="none" xmlns="http://www.w3.org/2000/svg">
            <!-- Background circle -->
            <circle cx="32" cy="32" r="30" fill="#6B2C91"/>
            <!-- Play/Run triangle -->
            <path d="M24 18 L24 46 L46 32 Z" fill="white"/>
            <!-- Small accent circle -->
            <circle cx="48" cy="16" r="6" fill="#FFC107"/>
        </svg>
        """
        
        # Convert SVG to QIcon
        app_icon = QIcon()
        pixmap = QPixmap()
        pixmap.loadFromData(QByteArray(svg_app_icon.encode('utf-8')))
        app_icon.addPixmap(pixmap, QIcon.Mode.Normal, QIcon.State.Off)
        
        # Set the window icon
        self.setWindowIcon(app_icon)
        
        # Also set as the application icon (for taskbar)
        QApplication.setWindowIcon(app_icon)

    def create_test_case_docx(self, test_case_name, screenshots_data):
        """
        Creates a single DOCX document for a test case with all screenshots.
        UPDATED: Now highlights fields marked for highlighting in yellow.
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            
            # Create a new document
            doc = Document()
            
            # Set document margins (narrower for better space usage)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Add configured text elements at the top
            text_elements = self.document_config.get('text_elements', [])
            if text_elements:
                for config_item in text_elements:
                    if config_item.get('type') == 'blank_line':
                        doc.add_paragraph()
                        continue
                    
                    paragraph = doc.add_paragraph()
                    original_text = config_item.get('text', '')
                    substituted_text = self.substitute_variables(
                        original_text, 
                        test_case_name, 
                        len(screenshots_data)
                    )
                    
                    run = paragraph.add_run(substituted_text)
                    font = run.font
                    font.name = config_item.get('font_name', 'Arial')
                    font.size = Pt(config_item.get('font_size', 12))
                    font.bold = config_item.get('bold', False)
                    font.italic = config_item.get('italic', False)
                    
                    paragraph_format = paragraph.paragraph_format
                    paragraph.style.font.name = config_item.get('font_name', 'Arial')
                    
                    alignment_map = {
                        'Left': WD_ALIGN_PARAGRAPH.LEFT,
                        'Center': WD_ALIGN_PARAGRAPH.CENTER,
                        'Right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'Justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    alignment = config_item.get('alignment', 'Left')
                    paragraph.alignment = alignment_map.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
                    paragraph.paragraph_format.space_after = Pt(6)

            # ✅ NEW: Get the configured highlight color
            highlight_color_name = self.document_config.get('highlight_color', 'Yellow')
            color_index_map = {
                'Yellow': 7,
                'Bright Green': 4,
                'Turquoise': 3,
                'Pink': 5,
                'Blue': 2,
                'Red': 6,
                'Dark Blue': 9,
                'Dark Cyan': 10,
                'Dark Green': 11,
                'Dark Magenta': 12,
                'Dark Red': 13,
                'Dark Yellow': 14,
                'Gray 25%': 16,
                'Gray 50%': 15
            }
            highlight_color_index = color_index_map.get(highlight_color_name, 7)
            
            # Add screenshots with highlighting
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn
            
            for idx, screenshot in enumerate(screenshots_data):
                screen_text = screenshot['screen_text']
                
                # ✅ APPLY MASKING IF ENABLED
                if self.masking_enabled:
                    screen_text = self.apply_masking_to_text(screen_text)
                
                highlight_info = screenshot.get('highlight_info', {})
                
                if idx > 0 and idx % 2 == 0:
                    doc.add_page_break()
                
                if idx > 0 and idx % 2 != 0:
                    doc.add_paragraph()
                    doc.add_paragraph()
                
                # Format the screen text into lines (24 rows x 80 columns)
                screen_rows = 24
                screen_cols = 80
                
                screen_lines = []
                for row_num in range(screen_rows):
                    start_idx = row_num * screen_cols
                    end_idx = start_idx + screen_cols
                    if start_idx < len(screen_text):
                        line = screen_text[start_idx:end_idx]
                        screen_lines.append(line)
                    else:
                        screen_lines.append(' ' * screen_cols)
                
                # ✅ FIXED: Create paragraph with character-level highlighting
                screen_para = doc.add_paragraph()

                # Convert highlight info to line-based positions (accounting for newlines)
                highlight_ranges = []
                for field_name, field_info in highlight_info.items():
                    row = field_info['row'] - 1  # Convert to 0-based
                    col = field_info['column'] - 1
                    length = field_info['length']
                    
                    # Calculate the start position in the text WITH newlines
                    # Each line is screen_cols chars + 1 newline char
                    start_pos = (row * (screen_cols + 1)) + col
                    end_pos = start_pos + length
                    
                    highlight_ranges.append((start_pos, end_pos))

                # Merge overlapping ranges and sort them
                highlight_ranges.sort()
                merged_ranges = []
                for start, end in highlight_ranges:
                    if merged_ranges and start <= merged_ranges[-1][1]:
                        # Overlapping or adjacent, merge them
                        merged_ranges[-1] = (merged_ranges[-1][0], max(merged_ranges[-1][1], end))
                    else:
                        merged_ranges.append((start, end))

                # Build the paragraph with highlighting
                full_text = '\n'.join(screen_lines)

                if not merged_ranges:
                    # No highlighting needed, add entire text as one run
                    run = screen_para.add_run(full_text)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(8)
                else:
                    # Add text in segments with appropriate highlighting
                    current_pos = 0
                    
                    for start, end in merged_ranges:
                        # Add non-highlighted text before this range
                        if current_pos < start:
                            run = screen_para.add_run(full_text[current_pos:start])
                            run.font.name = 'Courier New'
                            run.font.size = Pt(8)
                        
                        # Add highlighted text
                        run = screen_para.add_run(full_text[start:end])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8)
                        run.font.highlight_color = highlight_color_index  # ✅ Use configured color instead of hardcoded 7
                        
                        current_pos = end
                    
                    # Add any remaining non-highlighted text
                    if current_pos < len(full_text):
                        run = screen_para.add_run(full_text[current_pos:])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8)

                # Add border to the paragraph
                pPr = screen_para._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')

                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '12')
                    border.set(qn('w:space'), '4')
                    border.set(qn('w:color'), '808080')
                    pBdr.append(border)

                pPr.append(pBdr)

                # Add shading
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), 'F0F0F0')
                pPr.append(shading_elm)

                screen_para.paragraph_format.space_before = Pt(0)
                screen_para.paragraph_format.space_after = Pt(0)
                screen_para.paragraph_format.left_indent = Inches(0.2)
                screen_para.paragraph_format.right_indent = Inches(0.2)
            
            # Determine output directory

            # Determine output directory
            project_name = getattr(self, 'current_project_id', None)
            if project_name and hasattr(self, 'projects') and project_name in self.projects:
                project_name = self.projects[project_name]['name']

            output_dir = os.path.join(self.default_results_location, 'Results', project_name if project_name else 'Master')
            os.makedirs(output_dir, exist_ok=True)

            # ✅ NEW: Add timestamp to filename
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = os.path.join(output_dir, f"{test_case_name}_{timestamp}.docx")
            doc.save(filename)
            
            return filename
            
        except Exception as e:
            print(f"Error creating DOCX: {e}")
            raise

    def substitute_variables(self, text, test_case_name, total_screenshots):
        """
        Substitutes variables in text with actual values.
        
        Args:
            text: Text containing variables like {test_case_id}
            test_case_name: Name of the test case
            total_screenshots: Number of screenshots in the test case
        
        Returns:
            str: Text with variables replaced
        """
        from datetime import datetime
        
        # Get current date and time
        now = datetime.now()
        
        # Get test case description
        test_description = ""
        if test_case_name in self.test_cases:
            test_description = self.test_cases[test_case_name].get('description', '')
        
        # Define variable mappings
        variables = {
            'test_case_id': test_case_name,
            'test_description': test_description,
            'date': now.strftime('%Y-%m-%d'),
            'time': now.strftime('%H:%M:%S'),
            'datetime': now.strftime('%Y-%m-%d %H:%M:%S'),
            'total_screenshots': str(total_screenshots),
            'space': ' '  # Single space character
        }
        
        # Replace each variable in the text
        result = text
        for var_name, var_value in variables.items():
            placeholder = '{' + var_name + '}'
            result = result.replace(placeholder, var_value)
        
        return result
           
    def convert_combo_key_to_pcomm(self, combo_key):
        """
        Converts a combo key string (e.g., "Shift+F11", "Ctrl+P") to PCOMM format.
        
        PCOMM Combo Key Mapping:
        - Shift+F1 through F12: @a through @l
        - Ctrl+Character: ^character (e.g., ^p for Ctrl+P)
        - Alt+Character: %character (e.g., %x for Alt+X)
        - Ctrl+Function: @C[pf#]
        - Alt+Function: @A[pf#]
        
        Examples:
        - Shift+F11 → "@k"
        - Ctrl+P → "^p"
        - Alt+X → "%x"
        - Ctrl+F3 → "@C[pf3]"
        """
        parts = [p.strip() for p in combo_key.split("+")]
        
        if len(parts) != 2:
            # Invalid format, return as-is
            return combo_key
        
        modifier = parts[0]
        key = parts[1]
        
        # Shift+Function key mapping (F1=@a, F2=@b, ... F12=@l)
        if modifier == "Shift" and key.startswith("F"):
            shift_fn_map = {
                "F1": "@a", "F2": "@b", "F3": "@c", "F4": "@d",
                "F5": "@e", "F6": "@f", "F7": "@g", "F8": "@h",
                "F9": "@i", "F10": "@j", "F11": "@k", "F12": "@l"
            }
            return shift_fn_map.get(key, combo_key)
        
        # Ctrl+Character (e.g., Ctrl+P → ^p)
        if modifier == "Ctrl" and len(key) == 1 and key.isalpha():
            return f"^{key.lower()}"
        
        # Alt+Character (e.g., Alt+X → %x)
        if modifier == "Alt" and len(key) == 1 and key.isalpha():
            return f"%{key.lower()}"
        
        # Ctrl+Function key (e.g., Ctrl+F3 → @C[pf3])
        if modifier == "Ctrl" and key.startswith("F"):
            try:
                fn_num = int(key[1:])
                if 1 <= fn_num <= 12:
                    return f"@C[pf{fn_num}]"
            except:
                pass
        
        # Alt+Function key (e.g., Alt+F3 → @A[pf3])
        if modifier == "Alt" and key.startswith("F"):
            try:
                fn_num = int(key[1:])
                if 1 <= fn_num <= 12:
                    return f"@A[pf{fn_num}]"
            except:
                pass
        
        # If no match found, return original
        return combo_key

    def get_connection_name_from_title(self, window_title):
        """
        Extracts the connection name (A, B, C, etc.) from the PCOMM window title.
        
        Examples:
        - "SessionA" -> "A"
        - "Session A" -> "A"
        - "PCOMM - SessionB" -> "B"
        - "SessionC - [24x80]" -> "C"
        
        Returns the last letter found in the title, or "A" as default.
        """
        import re
        
        # Try to find "Session" followed by a letter
        match = re.search(r'Session\s*([A-E])', window_title, re.IGNORECASE)
        if match:
            return match.group(1).upper()
        
        # If not found, try to find any single uppercase letter A-E
        match = re.search(r'\b([A-E])\b', window_title)
        if match:
            return match.group(1).upper()
        
        # Default to "A" if nothing found
        return "A"
        
    def load_masking_config(self):
        """Loads masking configuration from file."""
        if os.path.exists(self.masking_config_file):
            try:
                with open(self.masking_config_file, 'r') as f:
                    config = json.load(f)
                    self.masking_enabled = config.get('enabled', False)
                    self.masking_patterns = config.get('patterns', [])
            except Exception as e:
                print(f"Error loading masking config: {e}")
                self.masking_enabled = False
                self.masking_patterns = []
        else:
            self.masking_enabled = False
            self.masking_patterns = []

    def save_masking_config(self):
        """Saves masking configuration to file."""
        try:
            config = {
                'enabled': self.masking_enabled,
                'patterns': self.masking_patterns
            }
            with open(self.masking_config_file, 'w') as f:
                json.dump(config, f, indent=4)
        except Exception as e:
            QMessageBox.warning(self, "Save Error", f"Failed to save masking configuration: {e}")

    def apply_masking_to_text(self, text):
        """Applies masking patterns to text if masking is enabled."""
        if not self.masking_enabled or not self.masking_patterns:
            return text
        
        masked_text = text
        
        for pattern_obj in self.masking_patterns:
            regex = pattern_obj.get('regex')
            mask_indices = pattern_obj.get('mask_indices')
            
            if not regex or not mask_indices:
                continue
            
            try:
                def mask_match(match):
                    original = match.group(0)
                    masked = list(original)
                    for idx in mask_indices:
                        if 0 <= idx < len(masked):
                            masked[idx] = 'x'
                    return ''.join(masked)
                
                masked_text = re.sub(regex, mask_match, masked_text)
            except re.error:
                continue
        
        return masked_text

    def open_settings_dialog(self):
        """Opens the unified settings dialog."""
        dialog = SettingsDialog(self)
        dialog.exec()        

    def select_text_in_preview(self, row, column, length):
        """
        Selects text in the PCOMM preview based on row, column, and length.
        
        Args:
            row (int): 1-based row number
            column (int): 1-based column number
            length (int): Length of text to select
        """
        if not self.pcomm_canvas_text_edit or length <= 0:
            return
        
        # Get the text cursor
        cursor = self.pcomm_canvas_text_edit.textCursor()
        
        # Move to the start of the document
        cursor.movePosition(QTextCursor.MoveOperation.Start)
        
        # Move down to the correct row (row is 1-based, so row-1 moves)
        for _ in range(row - 1):
            cursor.movePosition(QTextCursor.MoveOperation.Down)
        
        # Move right to the correct column (column is 1-based, so column-1 moves)
        for _ in range(column - 1):
            cursor.movePosition(QTextCursor.MoveOperation.Right)
        
        # Select the text for the given length
        for _ in range(length):
            cursor.movePosition(QTextCursor.MoveOperation.Right, QTextCursor.MoveMode.KeepAnchor)
        
        # Set the cursor with selection
        self.pcomm_canvas_text_edit.setTextCursor(cursor)
        
        # Ensure the selection is visible by scrolling to it
        self.pcomm_canvas_text_edit.ensureCursorVisible()

    def load_default_location_config(self):
        """Loads the default results location from file."""
        if os.path.exists(self.default_results_location_file):
            try:
                with open(self.default_results_location_file, 'r') as f:
                    config = json.load(f)
                    self.default_results_location = config.get('location', os.path.join(os.path.expanduser("~"), "Desktop"))
            except Exception as e:
                print(f"Error loading default location config: {e}")
                self.default_results_location = os.path.join(os.path.expanduser("~"), "Desktop")
        else:
            self.default_results_location = os.path.join(os.path.expanduser("~"), "Desktop")

    def save_default_location_config(self):
        """Saves the default results location to file."""
        try:
            config = {'location': self.default_results_location}
            with open(self.default_results_location_file, 'w') as f:
                json.dump(config, f, indent=4)
        except Exception as e:
            QMessageBox.warning(self, "Save Error", f"Failed to save default location configuration: {e}")
            
    def set_initial_dock_sizes(self):
        """Sets the initial sizes of dock widgets to be equal on startup."""
        # Get the total width of the main window
        total_width = self.width()
        
        # Set the libraries dock to 50% of the window width
        left_width = int(total_width * 0.5)
        
        # Resize the left dock
        self.resizeDocks([self.libraries_dock], [left_width], Qt.Orientation.Horizontal)  

    def capture_pcomm_screen_as_jpeg(self, output_path):
        """
        Captures the PCOMM screen and saves it as JPEG.
        
        Args:
            output_path: Full path where the JPEG should be saved
        
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            pcomm_window = gw.getWindowsWithTitle(self.pcomm_window_title)[0]
            
            # Use Win32 API's PrintWindow for a more reliable capture
            hwnd = pcomm_window._hWnd
            left, top, right, bottom = win32gui.GetWindowRect(hwnd)
            w = right - left
            h = bottom - top

            # Get the device context and create a compatible one
            hwnd_dc = win32gui.GetWindowDC(hwnd)
            mfc_dc = win32ui.CreateDCFromHandle(hwnd_dc)
            save_dc = mfc_dc.CreateCompatibleDC()

            # Create a bitmap to hold the image
            save_bitmap = win32ui.CreateBitmap()
            save_bitmap.CreateCompatibleBitmap(mfc_dc, w, h)
            save_dc.SelectObject(save_bitmap)
            
            # Try to use PW_CLIENTONLY, but fall back to 0 if it's not defined
            try:
                flag = win32con.PW_CLIENTONLY
            except AttributeError:
                flag = 0
                
            result = windll.user32.PrintWindow(hwnd, save_dc.GetSafeHdc(), flag)

            if result != 1:
                raise Exception("Failed to capture window with PrintWindow.")
                
            # Convert the bitmap to a Pillow Image
            bmpinfo = save_bitmap.GetInfo()
            bmpstr = save_bitmap.GetBitmapBits(True)
            im = Image.frombuffer(
                'RGB',
                (bmpinfo['bmWidth'], bmpinfo['bmHeight']),
                bmpstr, 'raw', 'BGRX', 0, 1
            )
            
            # Clean up the device contexts and bitmap
            win32gui.DeleteObject(save_bitmap.GetHandle())
            save_dc.DeleteDC()
            mfc_dc.DeleteDC()
            win32gui.ReleaseDC(hwnd, hwnd_dc)

            # Save as JPEG
            im.save(output_path, 'JPEG', quality=95)
            
            return True
            
        except Exception as e:
            print(f"Error capturing PCOMM screen: {e}")
            return False        


    def convert_to_template(self):
        """Opens dialog to select a test case and convert it to an Excel template."""
        if not self.test_cases:
            QMessageBox.warning(self, "No Test Cases", "No test cases available to convert.")
            return
        
        # Create a custom dialog with search
        dialog = QDialog(self)
        dialog.setWindowTitle("Convert to Template")
        dialog.setMinimumSize(400, 500)
        dialog.setWindowFlags(dialog.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title_label = QLabel("Select Test Case to Convert")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        layout.addSpacing(10)
        
        # Search bar
        search_bar = QLineEdit()
        search_bar.setPlaceholderText("Search test cases...")
        layout.addWidget(search_bar)
        
        # List widget
        list_widget = QListWidget()
        list_widget.addItems(sorted(self.test_cases.keys()))
        layout.addWidget(list_widget)
        
        # Connect search functionality
        def filter_list(query):
            query = query.strip().lower()
            for i in range(list_widget.count()):
                item = list_widget.item(i)
                item_text = item.text().lower()
                item.setHidden(query not in item_text if query else False)
        
        search_bar.textChanged.connect(filter_list)
        
        # Double-click to select
        list_widget.itemDoubleClicked.connect(dialog.accept)
        
        # Buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)
        
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
        
        selected_items = list_widget.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "No Selection", "Please select a test case.")
            return
        
        test_case_name = selected_items[0].text()
        
        # Get save location
        default_filename = f"{test_case_name}_Template.xlsx"
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Template As",
            default_filename,
            "Excel Files (*.xlsx);;All Files (*)"
        )
        
        if not file_path:
            return
        
        try:
            self._create_excel_template(test_case_name, file_path)
            QMessageBox.information(
                self,
                "Success",
                f"Template created successfully!\n\nFile: {file_path}\n\n"
                "You can now fill in the template with test data and link it to the test case."
            )
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create template:\n\n{str(e)}")
    
    def _create_excel_template(self, test_case_name, file_path):
        """Creates an Excel template from a test case with column-wise data layout."""
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        
        test_case_data = self.test_cases.get(test_case_name)
        if not test_case_data:
            raise ValueError(f"Test case '{test_case_name}' not found")
        
        # Create workbook
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Test Data"
        
        # Styles
        header_fill = PatternFill(start_color="6B2C91", end_color="6B2C91", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=11)
        field_label_fill = PatternFill(start_color="E9D5F5", end_color="E9D5F5", fill_type="solid")
        field_label_font = Font(bold=True, size=10)
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Add metadata section
        ws['A1'] = 'Test Case Template'
        ws['A1'].font = Font(bold=True, size=14)
        ws['A2'] = f'Test Case: {test_case_name}'
        ws['A3'] = f'Description: {test_case_data.get("description", "")}'
        ws['A4'] = f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
        
        # Extract all input fields from test case steps
        field_rows = []  # Store (field_label, step_index, module_name)
        
        for step_index, step in enumerate(test_case_data.get('steps', []), 1):
            step_type = step.get('type')
            
            if step_type == 'module_import':
                module_name = step.get('module_name', 'Unknown')
                for field in step.get('fields', []):
                    if field.get('action_type') == 'Input':
                        field_name = field.get('field_name', '')
                        if field_name:
                            field_label = f"Step {step_index} - {module_name}.{field_name}"
                            field_rows.append((field_label, step_index, module_name, field_name))
            
            elif step_type == 'random_input':
                row = step.get('row', '?')
                col = step.get('column', '?')
                field_label = f"Step {step_index} - RandomInput (R{row},C{col})"
                field_name = f"RandomInput_R{row}C{col}"
                field_rows.append((field_label, step_index, 'Random', field_name))
        
        # COLUMN-WISE LAYOUT
        # Row 6: Headers (Field Name, Test Case 1, Test Case 2, ...)
        start_row = 6
        ws.cell(row=start_row, column=1, value="Field Name")
        ws.cell(row=start_row, column=1).fill = header_fill
        ws.cell(row=start_row, column=1).font = header_font
        ws.cell(row=start_row, column=1).alignment = Alignment(horizontal='center', vertical='center')
        ws.cell(row=start_row, column=1).border = border
        
        # Add headers for 5 test cases
        for col_idx in range(2, 7):  # Columns B to F (5 test cases)
            cell = ws.cell(row=start_row, column=col_idx, value=f"Test Case {col_idx-1}")
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = border
        
        # Add field labels in rows (starting from row 7)
        for row_idx, (field_label, step_idx, module_name, field_name) in enumerate(field_rows, start_row + 1):
            # Column A: Field label
            cell = ws.cell(row=row_idx, column=1, value=field_label)
            cell.fill = field_label_fill
            cell.font = field_label_font
            cell.alignment = Alignment(horizontal='left', vertical='center')
            cell.border = border
            
            # Columns B-F: Empty cells (no placeholder)
            for col_idx in range(2, 7):
                cell = ws.cell(row=row_idx, column=col_idx, value="")
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # Auto-adjust column widths
        ws.column_dimensions['A'].width = 40  # Field name column
        for col_idx in range(2, 7):
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = 20
        
        # Add instructions sheet with field mapping
        instructions_ws = wb.create_sheet("Instructions")
        instructions_ws['A1'] = "How to Use This Template"
        instructions_ws['A1'].font = Font(bold=True, size=14)
        
        instructions = [
            "",
            "1. Fill in the 'Test Data' sheet with your test data",
            "2. Each COLUMN (starting from column B) represents one test execution",
            "3. Each ROW represents a field to be filled",
            "4. Field Name: Shows which step and field the data is for",
            "5. Enter your test data in the corresponding cells",
            "",
            "Field Mapping:",
            ""
        ]
        
        for idx, instr in enumerate(instructions, 2):
            instructions_ws[f'A{idx}'] = instr
        
        # Add field mapping information
        mapping_row = len(instructions) + 2
        instructions_ws[f'A{mapping_row}'] = "Field Label"
        instructions_ws[f'B{mapping_row}'] = "Step"
        instructions_ws[f'C{mapping_row}'] = "Module"
        instructions_ws[f'D{mapping_row}'] = "Field Name"
        
        for idx, (field_label, step_idx, module_name, field_name) in enumerate(field_rows, mapping_row + 1):
            instructions_ws[f'A{idx}'] = field_label
            instructions_ws[f'B{idx}'] = f"Step {step_idx}"
            instructions_ws[f'C{idx}'] = module_name
            instructions_ws[f'D{idx}'] = field_name
        
        # Save workbook
        wb.save(file_path)
    
    def link_data_source(self):
        """Opens dialog to link an Excel data source to a test case."""
        if not self.test_cases:
            QMessageBox.warning(self, "No Test Cases", "No test cases available.")
            return
        
        # Create a custom dialog with search
        dialog = QDialog(self)
        dialog.setWindowTitle("Link Data Source")
        dialog.setMinimumSize(400, 500)
        dialog.setWindowFlags(dialog.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        
        layout = QVBoxLayout(dialog)
        
        # Title
        title_label = QLabel("Select Test Case to Link Data Source")
        title_font = QFont()
        title_font.setPointSize(12)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        layout.addSpacing(10)
        
        # Search bar
        search_bar = QLineEdit()
        search_bar.setPlaceholderText("Search test cases...")
        layout.addWidget(search_bar)
        
        # List widget
        list_widget = QListWidget()
        list_widget.addItems(sorted(self.test_cases.keys()))
        layout.addWidget(list_widget)
        
        # Connect search functionality
        def filter_list(query):
            query = query.strip().lower()
            for i in range(list_widget.count()):
                item = list_widget.item(i)
                item_text = item.text().lower()
                item.setHidden(query not in item_text if query else False)
        
        search_bar.textChanged.connect(filter_list)
        
        # Double-click to select
        list_widget.itemDoubleClicked.connect(dialog.accept)
        
        # Buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)
        
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return
        
        selected_items = list_widget.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "No Selection", "Please select a test case.")
            return
        
        test_case_name = selected_items[0].text()
        
        # Now call the actual linking method
        self.link_data_source_to_test_case(test_case_name)
    
    def convert_specific_test_case_to_template(self, test_case_name):
        """Converts a specific test case to template (called from context menu)."""
        default_filename = f"{test_case_name}_Template.xlsx"
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Template As",
            default_filename,
            "Excel Files (*.xlsx);;All Files (*)"
        )
        
        if not file_path:
            return
        
        try:
            self._create_excel_template(test_case_name, file_path)
            QMessageBox.information(
                self,
                "Success",
                f"Template created successfully!\n\nFile: {file_path}\n\n"
                "Fill in the template with test data and use 'Link Data Source' to connect it."
            )
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create template:\n\n{str(e)}")
        
    def link_data_source_to_test_case(self, test_case_name):
        """Links a data source to a specific test case (called from context menu)."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Excel Data Source",
            "",
            "Excel Files (*.xlsx *.xls);;All Files (*)"
        )
        
        if not file_path:
            return
        
        # Ask for sheet name
        sheet_name, ok = QInputDialog.getText(
            self,
            "Sheet Name",
            "Enter the sheet name:",
            text="Test Data"
        )
        
        if not ok:
            return
        
        sheet_name = sheet_name.strip() or "Test Data"
        
        # Validate the Excel file and sheet (but don't read data yet)
        try:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            # Check if sheet exists
            if sheet_name not in wb.sheetnames:
                QMessageBox.warning(self, "Invalid Sheet", f"Sheet '{sheet_name}' not found in the Excel file.")
                return
            
            # Create template with just file reference (no data loading)
            template_name = os.path.splitext(os.path.basename(file_path))[0]
            
            self.templates[template_name] = {
                'excel_path': file_path,
                'sheet_name': sheet_name,
                'base_test_case': test_case_name,
                'test_cases': [],  # Empty - will be loaded when breaking into test cases
                'expanded': False,
                'linked_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            }
            
            # Save templates
            self.save_templates_to_file()
            
            # Ensure Templates tab is visible
            templates_tab_index = -1
            for i in range(self.libraries_tabs.count()):
                if self.libraries_tabs.tabText(i) == "Templates Library":
                    templates_tab_index = i
                    break
            
            # If tab doesn't exist, add it
            if templates_tab_index == -1:
                templates_tab_index = self.libraries_tabs.addTab(self.templates_widget, "Templates Library")
                self.setup_tab_close_button(templates_tab_index)
                self.tab_states['templates'] = True
            
            # Show the Libraries dock if it's hidden
            if not self.libraries_dock.isVisible():
                self.libraries_dock.setVisible(True)
                self.toggle_libraries_action.setChecked(True)
            
            # Update tree
            self.update_template_tree()
            
            # Clear search bar
            self.templates_search_bar.clear()
            
            # Switch to Templates tab
            self.libraries_tabs.setCurrentIndex(templates_tab_index)
            
            # Store the link in test case
            if test_case_name in self.test_cases:
                self.test_cases[test_case_name]['data_source'] = {
                    'excel_path': file_path,
                    'sheet_name': sheet_name,
                    'template_name': template_name,  # ✅ NEW: Store template name
                    'linked_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                }
                self.save_test_cases_to_file()
            
            QMessageBox.information(
                self,
                "Success",
                f"Template created successfully!\n\n"
                f"Template Name: {template_name}\n"
                f"Excel File: {os.path.basename(file_path)}\n"
                f"Sheet: {sheet_name}\n\n"
                f"The template is now visible in 'Templates Library'.\n"
                f"Use 'Break into Test Cases' to generate test cases from the Excel data."
            )
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to create template:\n\n{str(e)}")
    
    def save_templates_to_file(self):
        """Saves templates to JSON file."""
        try:
            with open(self.template_file, 'w') as f:
                json.dump(self.templates, f, indent=4)
        except Exception as e:
            print(f"Error saving templates: {e}")

    def load_templates_from_file(self):
        """Loads templates from JSON file."""
        if os.path.exists(self.template_file):
            try:
                with open(self.template_file, 'r') as f:
                    self.templates = json.load(f)
                self.update_template_tree()
            except Exception as e:
                print(f"Error loading templates: {e}")
                self.templates = {}    
    
    def filter_templates(self, query: str):
        """Filters templates based on search query."""
        if not self.template_tree_root:
            return
        
        query = query.strip().lower()
        for i in range(self.template_tree_root.childCount()):
            item = self.template_tree_root.child(i)
            item_text = item.text(0).lower()
            if query:
                item.setHidden(query not in item_text)
            else:
                item.setHidden(False)

    def delete_template(self, item):
        """Deletes a template and unlinks associated test cases."""
        template_name = item.text(0)
        reply = QMessageBox.question(
            self,
            'Delete Template',
            f"Are you sure you want to delete template '{template_name}'?\n\n"
            "This will also unlink any test cases using this template.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            if template_name in self.templates:
                # ✅ NEW: Find and unlink associated test cases
                for test_case_name, test_case_data in self.test_cases.items():
                    data_source = test_case_data.get('data_source', {})
                    if data_source.get('template_name') == template_name:
                        # Remove the data source link
                        del test_case_data['data_source']
                
                # Save test cases with updated links
                self.save_test_cases_to_file()
                
                # Delete the template
                del self.templates[template_name]
                self.save_templates_to_file()
                self.update_template_tree()
                
                self.statusBar().showMessage(
                    f"Template '{template_name}' deleted and associated test cases unlinked.", 
                    5000
                )

    def view_template_info(self, template_name):
        """Shows template information."""
        template_data = self.templates.get(template_name)
        if not template_data:
            return
        
        info = (
            f"Template: {template_name}\n"
            f"{'='*40}\n\n"
            f"Excel File: {os.path.basename(template_data['excel_path'])}\n"
            f"Sheet: {template_data['sheet_name']}\n"
            f"Base Test Case: {template_data['base_test_case']}\n"
            f"Test Cases: {len(template_data['test_cases'])}\n"
            f"Created: {template_data['linked_date']}\n"
            f"Status: {'Expanded' if template_data.get('expanded') else 'Collapsed'}"
        )
        
        QMessageBox.information(self, "Template Info", info)

    def on_template_item_double_clicked(self, item, column):
        """Handles double-click on template items."""
        # Expand/collapse on double-click
        if item.parent() == self.template_tree_root:
            item.setExpanded(not item.isExpanded())    
    
    def view_data_source_info(self, test_case_name):
        """Displays information about the linked data source."""
        test_case_data = self.test_cases.get(test_case_name, {})
        data_source = test_case_data.get('data_source', {})
        
        if not data_source:
            QMessageBox.information(self, "No Data Source", "No data source linked to this test case.")
            return
        
        excel_path = data_source.get('excel_path', 'Unknown')
        sheet_name = data_source.get('sheet_name', 'Unknown')
        linked_date = data_source.get('linked_date', 'Unknown')
        
        # Count data rows
        row_count = 0
        try:
            import openpyxl
            if os.path.exists(excel_path):
                wb = openpyxl.load_workbook(excel_path, data_only=True)
                if sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    # Count non-empty rows starting from row 7 (after header at row 6)
                    for row_idx in range(7, ws.max_row + 1):
                        if ws.cell(row=row_idx, column=1).value:
                            row_count += 1
        except:
            pass
        
        info_text = (
            f"Data Source Information\n"
            f"{'=' * 40}\n\n"
            f"Test Case: {test_case_name}\n"
            f"Excel File: {os.path.basename(excel_path)}\n"
            f"Full Path: {excel_path}\n"
            f"Sheet Name: {sheet_name}\n"
            f"Linked Date: {linked_date}\n"
            f"Data Rows: {row_count}\n\n"
            f"Status: {'✓ File exists' if os.path.exists(excel_path) else '✗ File not found'}"
        )
        
        QMessageBox.information(self, "Data Source Info", info_text)
    
    def unlink_data_source(self, test_case_name):
        """Unlinks the data source from a test case."""
        reply = QMessageBox.question(
            self,
            "Confirm Unlink",
            f"Are you sure you want to unlink the data source from '{test_case_name}'?\n\n"
            "The Excel file will not be deleted, only the link will be removed.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No
        )
        
        if reply == QMessageBox.StandardButton.Yes:
            if test_case_name in self.test_cases:
                if 'data_source' in self.test_cases[test_case_name]:
                    del self.test_cases[test_case_name]['data_source']
                    self.save_test_cases_to_file()
                    QMessageBox.information(self, "Success", f"Data source unlinked from '{test_case_name}'.")


class LinkDataSourceDialog(QDialog):
    """Dialog for linking an Excel data source to a test case."""
    
    def __init__(self, parent=None, test_case_names=None):
        super().__init__(parent)
        self.test_case_names = test_case_names or []
        self.setup_ui()
    
    def setup_ui(self):
        self.setWindowTitle("Link Data Source")
        self.setMinimumSize(600, 400)
        self.setWindowFlags(self.windowFlags() & ~Qt.WindowType.WindowContextHelpButtonHint)
        
        layout = QVBoxLayout(self)
        
        # Title
        title_label = QLabel("Link Excel Data Source")
        title_font = QFont()
        title_font.setPointSize(14)
        title_font.setBold(True)
        title_label.setFont(title_font)
        layout.addWidget(title_label)
        
        layout.addSpacing(10)
        
        # Description
        desc_label = QLabel(
            "Link an Excel spreadsheet to a test case for data-driven testing.\n"
            "The test case will be executed once for each data row in the linked sheet."
        )
        desc_label.setWordWrap(True)
        layout.addWidget(desc_label)
        
        layout.addSpacing(20)
        
        # Form layout
        form_layout = QFormLayout()
        
        # Test case selection
        self.test_case_combo = QComboBox()
        self.test_case_combo.addItems(self.test_case_names)
        self.test_case_combo.setEditable(True)
        self.test_case_combo.completer().setCompletionMode(
            self.test_case_combo.completer().CompletionMode.PopupCompletion
        )
        self.test_case_combo.completer().setFilterMode(Qt.MatchFlag.MatchContains)
        form_layout.addRow("Test Case:", self.test_case_combo)
        
        # Excel file selection
        excel_layout = QHBoxLayout()
        self.excel_path_input = QLineEdit()
        self.excel_path_input.setPlaceholderText("Select Excel file...")
        self.excel_path_input.setReadOnly(True)
        excel_layout.addWidget(self.excel_path_input)
        
        browse_button = QPushButton("Browse...")
        browse_button.setFixedWidth(100)
        browse_button.clicked.connect(self.browse_excel)
        excel_layout.addWidget(browse_button)
        
        form_layout.addRow("Excel File:", excel_layout)
        
        # Sheet name input
        self.sheet_name_input = QLineEdit()
        self.sheet_name_input.setPlaceholderText("Enter sheet name (default: Test Data)")
        self.sheet_name_input.setText("Test Data")
        form_layout.addRow("Sheet Name:", self.sheet_name_input)
        
        layout.addLayout(form_layout)
        
        layout.addSpacing(20)
        
        # Preview section
        preview_group = QWidget()
        preview_layout = QVBoxLayout(preview_group)
        preview_layout.setContentsMargins(0, 0, 0, 0)
        
        preview_label = QLabel("Preview:")
        preview_label.setStyleSheet("font-weight: bold;")
        preview_layout.addWidget(preview_label)
        
        self.preview_table = QTableWidget()
        self.preview_table.setMaximumHeight(150)
        preview_layout.addWidget(self.preview_table)
        
        layout.addWidget(preview_group)
        
        # Connect excel path change to preview update
        self.excel_path_input.textChanged.connect(self.update_preview)
        self.sheet_name_input.textChanged.connect(self.update_preview)
        
        layout.addStretch()
        
        # Buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(self.accept)
        button_box.rejected.connect(self.reject)
        layout.addWidget(button_box)
    
    def browse_excel(self):
        """Opens file dialog to select Excel file."""
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Select Excel File",
            "",
            "Excel Files (*.xlsx *.xls);;All Files (*)"
        )
        
        if file_path:
            self.excel_path_input.setText(file_path)
    
    def update_preview(self):
        """Updates the preview table with data from the selected Excel file."""
        excel_path = self.excel_path_input.text()
        sheet_name = self.sheet_name_input.text() or "Test Data"
        
        if not excel_path or not os.path.exists(excel_path):
            self.preview_table.setRowCount(0)
            self.preview_table.setColumnCount(0)
            return
        
        try:
            import openpyxl
            
            wb = openpyxl.load_workbook(excel_path, data_only=True)
            
            if sheet_name not in wb.sheetnames:
                self.preview_table.setRowCount(0)
                self.preview_table.setColumnCount(0)
                return
            
            ws = wb[sheet_name]
            
            # Find header row (row 6 in our template)
            header_row = 6
            headers = []
            for cell in ws[header_row]:
                if cell.value:
                    headers.append(str(cell.value))
                else:
                    break
            
            if not headers:
                self.preview_table.setRowCount(0)
                self.preview_table.setColumnCount(0)
                return
            
            # Setup table
            self.preview_table.setColumnCount(len(headers))
            self.preview_table.setHorizontalHeaderLabels(headers)
            
            # Load first 5 data rows
            data_rows = []
            for row_idx in range(header_row + 1, header_row + 6):
                row_data = []
                for col_idx in range(1, len(headers) + 1):
                    cell_value = ws.cell(row=row_idx, column=col_idx).value
                    row_data.append(str(cell_value) if cell_value is not None else "")
                
                # Only add non-empty rows
                if any(row_data):
                    data_rows.append(row_data)
            
            self.preview_table.setRowCount(len(data_rows))
            
            for row_idx, row_data in enumerate(data_rows):
                for col_idx, value in enumerate(row_data):
                    self.preview_table.setItem(row_idx, col_idx, QTableWidgetItem(value))
            
            # Auto-resize columns
            self.preview_table.resizeColumnsToContents()
            
        except Exception as e:
            print(f"Error loading preview: {e}")
            self.preview_table.setRowCount(0)
            self.preview_table.setColumnCount(0)
    
    def get_test_case_name(self):
        """Returns the selected test case name."""
        return self.test_case_combo.currentText().strip()
    
    def get_excel_path(self):
        """Returns the selected Excel file path."""
        return self.excel_path_input.text().strip()
    
    def get_sheet_name(self):
        """Returns the sheet name."""
        return self.sheet_name_input.text().strip() or "Test Data"

if __name__ == "__main__":
    import os
    os.environ["QT_ENABLE_HIGHDPI_SCALING"] = "0"
    app = QApplication(sys.argv)
    
    # Show splash screen
    splash = SplashScreen()
    splash.show()
    
    # Create main window but don't show it yet
    window = PCOMMMainFrame()
    
    # Timer to close splash and show main window after 3 seconds
    def show_main_window():
        splash.close()
        window.showMaximized()  # Show maximized
        # Set equal dock sizes after window is shown and maximized
        QTimer.singleShot(200, window.set_initial_dock_sizes)
    
    QTimer.singleShot(3000, show_main_window)  # 3000ms = 3 seconds
    
    sys.exit(app.exec())